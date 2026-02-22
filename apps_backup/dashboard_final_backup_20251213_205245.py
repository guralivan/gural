# -*- coding: utf-8 -*-
import os
import json
import base64
import re
from io import BytesIO
from collections import Counter
import urllib.parse as _urlparse
import locale

import pandas as pd
import numpy as np
import streamlit as st
import requests
import plotly.express as px
import plotly.graph_objects as go
from datetime import datetime, timedelta
import warnings
warnings.filterwarnings('ignore')

# Импорт Prophet с обработкой ошибок
try:
    from prophet import Prophet
    from prophet.plot import plot_plotly, plot_components_plotly
    PROPHET_AVAILABLE = True
except ImportError:
    PROPHET_AVAILABLE = False
    # Prophet не установлен - функциональность прогнозирования недоступна

# Импорт OpenAI с обработкой ошибок
try:
    import openai
    OPENAI_AVAILABLE = True
except ImportError:
    OPENAI_AVAILABLE = False
    # OpenAI не установлен - функциональность анализа изображений через нейросеть недоступна

# Импорт python-docx для чтения файлов WGSN
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    # python-docx не установлен - функциональность чтения WGSN файлов недоступна

# Настройка локали для правильного отображения чисел с пробелами
try:
    locale.setlocale(locale.LC_ALL, 'ru_RU.UTF-8')
except:
    try:
        locale.setlocale(locale.LC_ALL, 'Russian_Russia.1251')
    except:
        pass

try:
    from PIL import Image
except Exception:
    Image = None

st.set_page_config(page_title="WB Dashboard — Анализ товаров", layout="wide")

# CSS стили для улучшения отображения таблиц
st.markdown("""
<style>
/* Убираем только ограничения по максимальной высоте */
.stDataFrame {
    max-height: none !important;
}
.stDataFrame > div {
    max-height: none !important;
}
div[data-testid="stDataFrame"] {
    max-height: none !important;
}
</style>
""", unsafe_allow_html=True)

# ================= ФУНКЦИИ ДЛЯ РАБОТЫ С ПАРАМЕТРАМИ =================

def is_valid_param_name(param_name: str) -> bool:
    """
    Проверяет, является ли имя параметра валидным.
    Исключает артикулы и другие недопустимые значения.
    """
    if not param_name or not isinstance(param_name, str):
        return False
    
    param_name_clean = param_name.strip()
    
    # Исключаем пустые значения
    if not param_name_clean:
        return False
    
    # Исключаем артикулы (чисто числовые значения длиной более 6 цифр)
    if param_name_clean.isdigit() and len(param_name_clean) >= 6:
        return False
    
    # Исключаем служебные поля
    excluded_names = ["Артикул", "URL", "Статус", "Параметры", "sku", "SKU"]
    if param_name_clean in excluded_names:
        return False
    
    return True

def save_param_value(sku: str, param: str, value: str, save_history: bool = True):
    """Сохраняет значение параметра для товара с сохранением истории изменений"""
    # Проверяем валидность имени параметра
    if not is_valid_param_name(param):
        return
    
    # Проверяем, не является ли параметр удаленным
    if "deleted_params" not in st.session_state:
        st.session_state["deleted_params"] = load_deleted_params_from_file()
    
    deleted_params = st.session_state.get("deleted_params", set())
    if param in deleted_params:
        # Параметр был удален, не сохраняем его
        return
    
    if "param_values" not in st.session_state:
        st.session_state["param_values"] = {}
    if param not in st.session_state["param_values"]:
        st.session_state["param_values"][param] = {}
    
    # Проверяем, есть ли старое значение
    old_value = st.session_state["param_values"][param].get(sku, None)
    
    # Сохраняем новое значение в правильной структуре: param_values[param][sku] = value
    st.session_state["param_values"][param][sku] = value
    
    # Если значение изменилось и нужно сохранить историю
    if save_history and old_value is not None and old_value != value:
        add_param_history(sku, param, old_value, value)

def add_param_history(sku: str, param: str, old_value: str, new_value: str):
    """Добавляет запись в историю изменений параметра (без сохранения в файл - сохраняется пакетами)"""
    if "param_history" not in st.session_state:
        st.session_state["param_history"] = []
    
    history_entry = {
        "timestamp": pd.Timestamp.now().strftime("%Y-%m-%d %H:%M:%S"),
        "sku": str(sku),
        "param": param,
        "old_value": old_value,
        "new_value": new_value
    }
    
    st.session_state["param_history"].append(history_entry)
    
    # НЕ сохраняем историю в файл здесь - сохраняется пакетами после обработки товара

def get_param_history(sku: str = None, param: str = None):
    """Получает историю изменений параметров с фильтрацией"""
    history = st.session_state.get("param_history", [])
    
    if sku:
        history = [h for h in history if h["sku"] == str(sku)]
    if param:
        history = [h for h in history if h["param"] == param]
    
    # Сортируем по дате (новые сверху)
    history.sort(key=lambda x: x["timestamp"], reverse=True)
    
    return history

def save_param_history_to_file():
    """Сохраняет историю изменений параметров в файл"""
    history = st.session_state.get("param_history", [])
    current_file = st.session_state.get("cached_file_name", None)
    
    if history and current_file:
        try:
            file_hash = hash(current_file) % 1000000
            history_file = f"param_history_{file_hash}.json"
            
            history_data = {
                "file_name": current_file,
                "history": history,
                "last_updated": pd.Timestamp.now().strftime("%Y-%m-%d %H:%M:%S")
            }
            
            with open(history_file, "w", encoding="utf-8") as f:
                json.dump(history_data, f, ensure_ascii=False, indent=2)
            
            return True
        except Exception as e:
            return False
    
    return False

def load_param_history_from_file():
    """Загружает историю изменений параметров из файла"""
    current_file = st.session_state.get("cached_file_name", None)
    
    if current_file:
        try:
            file_hash = hash(current_file) % 1000000
            history_file = f"param_history_{file_hash}.json"
            
            if os.path.exists(history_file):
                with open(history_file, "r", encoding="utf-8") as f:
                    history_data = json.load(f)
                
                # Проверяем, что файл соответствует текущему файлу
                if history_data.get("file_name") == current_file:
                    st.session_state["param_history"] = history_data.get("history", [])
                    return True
        except Exception as e:
            pass
    
    # Если файл не найден или не соответствует, инициализируем пустую историю
    if "param_history" not in st.session_state:
        st.session_state["param_history"] = []
    
    return False

def get_param_values():
    """Получает все сохраненные значения параметров, исключая удаленные параметры"""
    param_values = st.session_state.get("param_values", {})
    
    # Загружаем список удаленных параметров
    if "deleted_params" not in st.session_state:
        st.session_state["deleted_params"] = load_deleted_params_from_file()
    
    deleted_params = st.session_state.get("deleted_params", set())
    
    # Фильтруем удаленные параметры
    if deleted_params:
        param_values = {k: v for k, v in param_values.items() if k not in deleted_params}
    
    return param_values

def get_hierarchy_params():
    """Получает список иерархических параметров из session_state или возвращает значения по умолчанию"""
    if "hierarchy_params" in st.session_state:
        return st.session_state["hierarchy_params"]
    # Значения по умолчанию
    return ["Тип", "Комплект", "Рукав", "Ворот"]

def get_subtype_params():
    """Получает список параметров подтипа из session_state или возвращает значения по умолчанию"""
    if "subtype_params" in st.session_state:
        return st.session_state["subtype_params"]
    # Значения по умолчанию
    return ["Рукав", "Ворот"]

def get_visual_params():
    """Получает список визуальных параметров из session_state или возвращает значения по умолчанию"""
    if "visual_params" in st.session_state:
        return st.session_state["visual_params"]
    # Значения по умолчанию
    return ["Цвет", "Принт", "Логотип", "Строчка", "Шов", "Строчка шов"]

def save_hierarchy_config():
    """Сохраняет конфигурацию иерархии параметров в файл"""
    try:
        hierarchy_config = {
            "hierarchy_params": st.session_state.get("hierarchy_params", ["Тип", "Комплект", "Рукав", "Ворот"]),
            "subtype_params": st.session_state.get("subtype_params", ["Рукав", "Ворот"]),
            "visual_params": st.session_state.get("visual_params", ["Цвет", "Принт", "Логотип", "Строчка", "Шов", "Строчка шов"])
        }
        with open("hierarchy_config.json", "w", encoding="utf-8") as f:
            json.dump(hierarchy_config, f, ensure_ascii=False, indent=2)
        return True
    except Exception as e:
        return False

def load_hierarchy_config():
    """Загружает конфигурацию иерархии параметров из файла"""
    try:
        if os.path.exists("hierarchy_config.json"):
            with open("hierarchy_config.json", "r", encoding="utf-8") as f:
                hierarchy_config = json.load(f)
                st.session_state["hierarchy_params"] = hierarchy_config.get("hierarchy_params", ["Тип", "Комплект", "Рукав", "Ворот"])
                st.session_state["subtype_params"] = hierarchy_config.get("subtype_params", ["Рукав", "Ворот"])
                st.session_state["visual_params"] = hierarchy_config.get("visual_params", ["Цвет", "Принт", "Логотип", "Строчка", "Шов", "Строчка шов"])
                return True
    except Exception as e:
        pass
    return False

def save_mass_analysis_progress(skus_to_process: list, processed_skus: set, results: list, settings: dict):
    """Сохраняет прогресс массового анализа"""
    try:
        progress_data = {
            "skus_to_process": skus_to_process,
            "processed_skus": list(processed_skus),
            "results": results,
            "settings": settings,
            "timestamp": pd.Timestamp.now().strftime("%Y-%m-%d %H:%M:%S")
        }
        with open("mass_analysis_progress.json", "w", encoding="utf-8") as f:
            json.dump(progress_data, f, ensure_ascii=False, indent=2)
        return True
    except Exception as e:
        return False

def load_mass_analysis_progress():
    """Загружает сохраненный прогресс массового анализа"""
    try:
        if os.path.exists("mass_analysis_progress.json"):
            with open("mass_analysis_progress.json", "r", encoding="utf-8") as f:
                progress_data = json.load(f)
            return progress_data
    except Exception as e:
        pass
    return None

def clear_mass_analysis_progress():
    """Удаляет сохраненный прогресс массового анализа"""
    try:
        if os.path.exists("mass_analysis_progress.json"):
            os.remove("mass_analysis_progress.json")
        return True
    except Exception as e:
        return False

def save_excluded_params_settings():
    """Сохраняет настройки исключения параметров в файл"""
    try:
        excluded_params = st.session_state.get("excluded_params", [])
        excluded_param_values = st.session_state.get("excluded_param_values", {})
        
        settings = {
            "excluded_params": excluded_params,
            "excluded_param_values": excluded_param_values,
            "timestamp": pd.Timestamp.now().strftime("%Y-%m-%d %H:%M:%S")
        }
        
        with open("excluded_params_settings.json", "w", encoding="utf-8") as f:
            json.dump(settings, f, ensure_ascii=False, indent=2)
        
        return True
    except Exception as e:
        st.error(f"Ошибка сохранения настроек исключения параметров: {e}")
        return False

def load_excluded_params_settings():
    """Загружает настройки исключения параметров из файла"""
    try:
        if os.path.exists("excluded_params_settings.json"):
            with open("excluded_params_settings.json", "r", encoding="utf-8") as f:
                settings = json.load(f)
                
            if "excluded_params" in settings:
                st.session_state["excluded_params"] = settings["excluded_params"]
            else:
                st.session_state["excluded_params"] = []
                
            if "excluded_param_values" in settings:
                st.session_state["excluded_param_values"] = settings["excluded_param_values"]
            else:
                st.session_state["excluded_param_values"] = {}
                
            return True
    except Exception as e:
        # Не показываем ошибку, просто используем значения по умолчанию
        pass
    return False

def save_mass_analysis_results(results: list):
    """Сохраняет результаты массового анализа в отдельный файл для постоянного хранения"""
    try:
        # Очищаем удаленные параметры перед сохранением
        cleaned_results = remove_deleted_params_from_mass_results(results)
        with open("mass_analysis_results.json", "w", encoding="utf-8") as f:
            json.dump(cleaned_results, f, ensure_ascii=False, indent=2)
        return True
    except Exception as e:
        return False

def load_mass_analysis_results():
    """Загружает сохраненные результаты массового анализа"""
    try:
        if os.path.exists("mass_analysis_results.json"):
            with open("mass_analysis_results.json", "r", encoding="utf-8") as f:
                results = json.load(f)
            results = results if isinstance(results, list) else []
            # Очищаем удаленные параметры из результатов
            results = remove_deleted_params_from_mass_results(results)
            return results
    except Exception as e:
        pass
    return []

def remove_deleted_params_from_mass_results(results: list) -> list:
    """Удаляет удаленные параметры из результатов массового анализа"""
    if not results:
        return results
    
    # Получаем список удаленных параметров
    deleted_params = set()
    if "deleted_params" in st.session_state:
        deleted_params = st.session_state["deleted_params"]
    else:
        deleted_params = load_deleted_params_from_file()
        # Сохраняем в session_state для будущих использований
        if deleted_params:
            st.session_state["deleted_params"] = deleted_params
    
    if not deleted_params:
        return results
    
    # Список обязательных полей, которые не являются параметрами
    required_fields = {"Артикул", "URL", "Статус", "Параметры"}
    
    cleaned_results = []
    for result in results:
        # Создаем копию результата без удаленных параметров
        # Сохраняем обязательные поля и параметры, которые не были удалены
        cleaned_result = {
            k: v for k, v in result.items() 
            if k in required_fields or k not in deleted_params
        }
        cleaned_results.append(cleaned_result)
    
    return cleaned_results

def export_mass_analysis_results_to_csv(results: list):
    """Экспортирует результаты массового анализа в CSV"""
    try:
        if not results:
            return None
        
        df = pd.DataFrame(results)
        
        # Убеждаемся, что колонка Артикул есть и в правильном формате
        if "Артикул" in df.columns:
            df["Артикул"] = df["Артикул"].astype(str).str.replace(".0", "", regex=False)
        
        # Сортируем колонки: сначала основные, потом параметры
        main_cols = ["Артикул", "URL", "Статус", "Параметры"]
        param_cols = [col for col in df.columns if col not in main_cols]
        column_order = [col for col in main_cols if col in df.columns] + sorted(param_cols)
        
        if column_order:
            df = df[column_order]
        
        csv_data = df.to_csv(encoding='utf-8-sig', index=False)
        return csv_data
    except Exception as e:
        st.error(f"Ошибка экспорта в CSV: {e}")
        return None

def import_mass_analysis_results_from_csv(csv_data: str):
    """Импортирует результаты массового анализа из CSV и применяет параметры"""
    try:
        import io
        df = pd.read_csv(io.StringIO(csv_data), encoding='utf-8-sig')
        
        results = []
        applied_count = 0
        
        # Инициализируем param_values если нет
        if "param_values" not in st.session_state:
            st.session_state["param_values"] = {}
        if "param_options" not in st.session_state:
            st.session_state["param_options"] = {}
        
        for _, row in df.iterrows():
            # Получаем артикул
            sku_raw = row.get("Артикул", "")
            if pd.isna(sku_raw) or not sku_raw:
                continue
            
            sku = str(int(sku_raw)) if isinstance(sku_raw, (int, float)) else str(sku_raw).replace(".0", "")
            
            # Создаем запись результата
            result_row = {
                "Артикул": sku,
                "URL": row.get("URL", ""),
                "Статус": row.get("Статус", "✅ Импортировано"),
                "Параметры": row.get("Параметры", "")
            }
            
            # Извлекаем параметры из строки (если есть колонка Параметры)
            params_dict = {}
            if "Параметры" in row and pd.notna(row["Параметры"]):
                params_str = str(row["Параметры"])
                for param_pair in params_str.split(", "):
                    if ":" in param_pair:
                        param_name, param_value = param_pair.split(":", 1)
                        param_name = param_name.strip()
                        param_value = param_value.strip()
                        if param_name and param_value:
                            params_dict[param_name] = param_value
            
            # Также проверяем отдельные колонки параметров
            param_cols = [col for col in df.columns if col not in ["Артикул", "URL", "Статус", "Параметры"]]
            for param_name in param_cols:
                param_value = row.get(param_name)
                if pd.notna(param_value) and str(param_value).strip():
                    params_dict[param_name] = str(param_value).strip()
            
            # Применяем параметры к param_values
            if params_dict:
                # Загружаем список удаленных параметров
                if "deleted_params" not in st.session_state:
                    st.session_state["deleted_params"] = load_deleted_params_from_file()
                
                deleted_params = st.session_state.get("deleted_params", set())
                
                for param_name, param_value in params_dict.items():
                    # Пропускаем невалидные параметры
                    if not is_valid_param_name(param_name):
                        continue
                    
                    # Пропускаем удаленные параметры
                    if param_name in deleted_params:
                        continue
                    
                    # Сохраняем параметр
                    save_param_value(sku, param_name, str(param_value), save_history=True)
                    
                    # Обновляем param_options
                    if param_name not in st.session_state["param_options"]:
                        st.session_state["param_options"][param_name] = []
                    if param_value and str(param_value) not in st.session_state["param_options"][param_name]:
                        st.session_state["param_options"][param_name].append(str(param_value))
                    
                    applied_count += 1
            
            # Добавляем параметры в результат
            result_row.update(params_dict)
            results.append(result_row)
        
        # Сохраняем результаты в session_state
        if "mass_analysis_results" not in st.session_state:
            st.session_state["mass_analysis_results"] = []
        
        # Объединяем с существующими результатами (обновляем по артикулу)
        existing_results = {str(r.get("Артикул", "")).replace(".0", ""): r for r in st.session_state["mass_analysis_results"]}
        for result in results:
            sku = str(result.get("Артикул", "")).replace(".0", "")
            existing_results[sku] = result
        
        st.session_state["mass_analysis_results"] = list(existing_results.values())
        
        # Сохраняем результаты и параметры в файлы
        save_mass_analysis_results(st.session_state["mass_analysis_results"])
        save_param_values_to_file()
        
        return len(results), applied_count
    except Exception as e:
        st.error(f"Ошибка импорта из CSV: {e}")
        return 0, 0

def apply_mass_analysis_results_to_params(results: list):
    """Автоматически применяет параметры из результатов массового анализа к param_values"""
    try:
        applied_count = 0
        
        if "param_values" not in st.session_state:
            st.session_state["param_values"] = {}
        if "param_options" not in st.session_state:
            st.session_state["param_options"] = {}
        
        for result in results:
            # Пропускаем неуспешные результаты
            status = result.get("Статус", "")
            if not status.startswith("✅"):
                continue
            
            sku_raw = result.get("Артикул", "")
            if not sku_raw:
                continue
            
            sku = str(int(sku_raw)) if isinstance(sku_raw, (int, float)) else str(sku_raw).replace(".0", "")
            
            # Извлекаем параметры из результата
            params_dict = {}
            
            # Получаем параметры из колонки "Параметры" если есть
            if "Параметры" in result and result["Параметры"]:
                params_str = str(result["Параметры"])
                for param_pair in params_str.split(", "):
                    if ":" in param_pair:
                        param_name, param_value = param_pair.split(":", 1)
                        param_name = param_name.strip()
                        param_value = param_value.strip()
                        if param_name and param_value:
                            params_dict[param_name] = param_value
            
            # Также берем параметры из всех ключей результата (кроме служебных)
            excluded_keys = ["Артикул", "URL", "Статус", "Параметры"]
            for key, value in result.items():
                if key not in excluded_keys and value and str(value).strip():
                    params_dict[key] = str(value).strip()
            
            # Применяем параметры
            if params_dict:
                # Загружаем список удаленных параметров
                if "deleted_params" not in st.session_state:
                    st.session_state["deleted_params"] = load_deleted_params_from_file()
                
                deleted_params = st.session_state.get("deleted_params", set())
                
                for param_name, param_value in params_dict.items():
                    # Пропускаем невалидные параметры
                    if not is_valid_param_name(param_name):
                        continue
                    
                    # Пропускаем удаленные параметры
                    if param_name in deleted_params:
                        continue
                    
                    # Сохраняем параметр с историей
                    save_param_value(sku, param_name, str(param_value), save_history=True)
                    
                    # Обновляем param_options
                    if param_name not in st.session_state["param_options"]:
                        st.session_state["param_options"][param_name] = []
                    if param_value and str(param_value) not in st.session_state["param_options"][param_name]:
                        st.session_state["param_options"][param_name].append(str(param_value))
                    
                    applied_count += 1
        
        return applied_count
    except Exception as e:
        st.error(f"Ошибка применения параметров: {e}")
        return 0

def save_param_values_to_file():
    """Сохраняет параметры в файл, привязанные к текущему файлу"""
    param_values = get_param_values()
    current_file = st.session_state.get("cached_file_name", None)
    
    # Инициализируем deleted_params если еще не инициализирован
    if "deleted_params" not in st.session_state:
        st.session_state["deleted_params"] = load_deleted_params_from_file()
    
    # Удаляем удаленные параметры из param_values перед сохранением
    deleted_params = st.session_state.get("deleted_params", set())
    if deleted_params:
        param_values = {k: v for k, v in param_values.items() if k not in deleted_params}
    
    # Получаем param_options для сохранения
    param_options = st.session_state.get("param_options", {})
    # Также удаляем удаленные параметры из param_options
    if deleted_params:
        param_options = {k: v for k, v in param_options.items() if k not in deleted_params}
    
    if param_values and current_file:
        try:
            # Создаем имя файла на основе текущего файла
            file_hash = hash(current_file) % 1000000  # Простой хеш для имени файла
            param_file = f"param_values_{file_hash}.json"
            
            # Сохраняем параметры с информацией о файле, включая param_options
            param_data = {
                "file_name": current_file,
                "param_values": param_values,
                "param_options": param_options,  # Добавляем варианты параметров
                "timestamp": pd.Timestamp.now().strftime("%Y-%m-%d %H:%M:%S")
            }
            
            with open(param_file, "w", encoding="utf-8") as f:
                json.dump(param_data, f, ensure_ascii=False, indent=2)
            
            # Также сохраняем в общий реестр файлов
            update_file_params_registry(current_file, param_file)
            return True
        except Exception as e:
            st.error(f"Ошибка сохранения параметров: {e}")
            return False
    elif param_values and not current_file:
        # Fallback для случаев, когда нет текущего файла
        try:
            # Сохраняем и param_values, и param_options в глобальный файл
            global_data = {
                "param_values": param_values,
                "param_options": param_options
            }
            with open("param_values_global.json", "w", encoding="utf-8") as f:
                json.dump(global_data, f, ensure_ascii=False, indent=2)
            return True
        except Exception:
            return False
    elif param_options and not param_values:
        # Сохраняем только param_options, если нет param_values
        try:
            if current_file:
                file_hash = hash(current_file) % 1000000
                param_file = f"param_values_{file_hash}.json"
                param_data = {
                    "file_name": current_file,
                    "param_values": {},
                    "param_options": param_options,
                    "timestamp": pd.Timestamp.now().strftime("%Y-%m-%d %H:%M:%S")
                }
                with open(param_file, "w", encoding="utf-8") as f:
                    json.dump(param_data, f, ensure_ascii=False, indent=2)
                update_file_params_registry(current_file, param_file)
                return True
            else:
                global_data = {"param_values": {}, "param_options": param_options}
                with open("param_values_global.json", "w", encoding="utf-8") as f:
                    json.dump(global_data, f, ensure_ascii=False, indent=2)
            return True
        except Exception:
            return False
    return False

def update_file_params_registry(file_name, param_file):
    """Обновляет реестр файлов и их параметров"""
    try:
        registry_file = "file_params_registry.json"
        registry = {}
        
        if os.path.exists(registry_file):
            with open(registry_file, "r", encoding="utf-8") as f:
                registry = json.load(f)
        
        registry[file_name] = {
            "param_file": param_file,
            "last_updated": pd.Timestamp.now().strftime("%Y-%m-%d %H:%M:%S")
        }
        
        with open(registry_file, "w", encoding="utf-8") as f:
            json.dump(registry, f, ensure_ascii=False, indent=2)
            
    except Exception as e:
        st.error(f"Ошибка обновления реестра: {e}")

def save_deleted_params_to_file():
    """Сохраняет список удаленных параметров в файл и очищает все файлы параметров от удаленных параметров"""
    try:
        # Получаем текущие удаленные параметры из session_state
        current_deleted = st.session_state.get("deleted_params", set())
        
        # Убеждаемся, что это set
        if not isinstance(current_deleted, set):
            current_deleted = set(current_deleted) if current_deleted else set()
            st.session_state["deleted_params"] = current_deleted
        
        # Загружаем существующие удаленные параметры из файла для объединения
        existing_deleted = load_deleted_params_from_file()
        
        # Убеждаемся, что это set
        if not isinstance(existing_deleted, set):
            existing_deleted = set(existing_deleted) if existing_deleted else set()
        
        # Объединяем оба набора (приоритет у текущих из session_state)
        deleted_params = existing_deleted.union(current_deleted)
        
        # Обновляем session_state
        st.session_state["deleted_params"] = deleted_params
        
        # Конвертируем set в list для JSON
        deleted_params_list = sorted(list(deleted_params))
        
        # Сохраняем в файл
        with open("deleted_params.json", "w", encoding="utf-8") as f:
            json.dump(deleted_params_list, f, ensure_ascii=False, indent=2)
        
        # Очищаем все файлы параметров от удаленных параметров
        if deleted_params:
            clean_deleted_params_from_all_files(deleted_params)
        
        return True
    except Exception as e:
        import traceback
        if hasattr(st, 'error'):
            st.error(f"Ошибка сохранения deleted_params: {e}\n{traceback.format_exc()}")
        else:
            print(f"Ошибка сохранения deleted_params: {e}\n{traceback.format_exc()}")
        return False

def clean_deleted_params_from_all_files(deleted_params: set):
    """Удаляет удаленные параметры из всех файлов параметров"""
    if not deleted_params:
        return
    
    import glob
    
    # Очищаем все файлы param_values*.json (включая param_values.json без подчеркивания)
    param_files = glob.glob("param_values*.json")
    cleaned_count = 0
    for param_file in param_files:
        try:
            with open(param_file, "r", encoding="utf-8") as f:
                param_data = json.load(f)
            
            # Поддерживаем оба формата: новый (с file_name) и старый (прямой словарь)
            if isinstance(param_data, dict) and "param_values" in param_data:
                # Новый формат
                file_params = param_data.get("param_values", {})
                file_options = param_data.get("param_options", {})
                
                # Удаляем удаленные параметры
                cleaned_params = {k: v for k, v in file_params.items() if k not in deleted_params}
                cleaned_options = {k: v for k, v in file_options.items() if k not in deleted_params}
                
                # Обновляем данные
                param_data["param_values"] = cleaned_params
                param_data["param_options"] = cleaned_options
                
                # Проверяем, были ли изменения
                if cleaned_params != file_params or cleaned_options != file_options:
                    # Сохраняем обратно
                    with open(param_file, "w", encoding="utf-8") as f:
                        json.dump(param_data, f, ensure_ascii=False, indent=2)
                    cleaned_count += 1
            else:
                # Старый формат - прямой словарь
                if isinstance(param_data, dict):
                    cleaned_params = {k: v for k, v in param_data.items() if k not in deleted_params}
                    if cleaned_params != param_data:
                        with open(param_file, "w", encoding="utf-8") as f:
                            json.dump(cleaned_params, f, ensure_ascii=False, indent=2)
                        cleaned_count += 1
        except Exception as e:
            # Пропускаем файлы с ошибками, но не прерываем процесс
            continue
    
    # Очищаем глобальный файл
    if os.path.exists("param_values_global.json"):
        try:
            with open("param_values_global.json", "r", encoding="utf-8") as f:
                global_data = json.load(f)
            
            # Поддерживаем оба формата
            if isinstance(global_data, dict) and "param_values" in global_data:
                original_params = global_data.get("param_values", {})
                original_options = global_data.get("param_options", {})
                
                cleaned_params = {k: v for k, v in original_params.items() if k not in deleted_params}
                cleaned_options = {k: v for k, v in original_options.items() if k not in deleted_params}
                
                if cleaned_params != original_params or cleaned_options != original_options:
                    global_data["param_values"] = cleaned_params
                    global_data["param_options"] = cleaned_options
                    with open("param_values_global.json", "w", encoding="utf-8") as f:
                        json.dump(global_data, f, ensure_ascii=False, indent=2)
                    cleaned_count += 1
            else:
                # Старый формат
                if isinstance(global_data, dict):
                    cleaned_global = {k: v for k, v in global_data.items() if k not in deleted_params}
                    if cleaned_global != global_data:
                        with open("param_values_global.json", "w", encoding="utf-8") as f:
                            json.dump(cleaned_global, f, ensure_ascii=False, indent=2)
                        cleaned_count += 1
        except Exception:
            pass
    
    # Очищаем table_cache.json
    if os.path.exists("table_cache.json"):
        try:
            with open("table_cache.json", "r", encoding="utf-8") as f:
                table_cache_data = json.load(f)
            
            if isinstance(table_cache_data, dict):
                changed = False
                # Очищаем param_values
                if "param_values" in table_cache_data:
                    original = table_cache_data["param_values"]
                    table_cache_data["param_values"] = {
                        k: v for k, v in original.items() 
                        if k not in deleted_params
                    }
                    if table_cache_data["param_values"] != original:
                        changed = True
                
                # Очищаем param_options
                if "param_options" in table_cache_data:
                    original = table_cache_data["param_options"]
                    table_cache_data["param_options"] = {
                        k: v for k, v in original.items() 
                        if k not in deleted_params
                    }
                    if table_cache_data["param_options"] != original:
                        changed = True
                
                if changed:
                    with open("table_cache.json", "w", encoding="utf-8") as f:
                        json.dump(table_cache_data, f, ensure_ascii=False, indent=2)
                    cleaned_count += 1
        except Exception:
            pass
    
    # Возвращаем количество очищенных файлов
    return cleaned_count
    
    # Возвращаем количество очищенных файлов
    return cleaned_count
    
    # Очищаем глобальный файл
    if os.path.exists("param_values_global.json"):
        try:
            with open("param_values_global.json", "r", encoding="utf-8") as f:
                global_data = json.load(f)
            
            # Поддерживаем оба формата
            if isinstance(global_data, dict) and "param_values" in global_data:
                cleaned_params = {k: v for k, v in global_data.get("param_values", {}).items() if k not in deleted_params}
                cleaned_options = {k: v for k, v in global_data.get("param_options", {}).items() if k not in deleted_params}
                global_data["param_values"] = cleaned_params
                global_data["param_options"] = cleaned_options
            else:
                # Старый формат
                global_data = {k: v for k, v in global_data.items() if k not in deleted_params} if isinstance(global_data, dict) else {}
            
            with open("param_values_global.json", "w", encoding="utf-8") as f:
                json.dump(global_data, f, ensure_ascii=False, indent=2)
        except Exception:
            pass
    
    # Очищаем table_cache.json
    if os.path.exists("table_cache.json"):
        try:
            with open("table_cache.json", "r", encoding="utf-8") as f:
                table_cache_data = json.load(f)
            
            if isinstance(table_cache_data, dict):
                # Очищаем param_values
                if "param_values" in table_cache_data:
                    table_cache_data["param_values"] = {
                        k: v for k, v in table_cache_data["param_values"].items() 
                        if k not in deleted_params
                    }
                
                # Очищаем param_options
                if "param_options" in table_cache_data:
                    table_cache_data["param_options"] = {
                        k: v for k, v in table_cache_data["param_options"].items() 
                        if k not in deleted_params
                    }
                
                # Сохраняем обратно
                with open("table_cache.json", "w", encoding="utf-8") as f:
                    json.dump(table_cache_data, f, ensure_ascii=False, indent=2)
        except Exception:
            pass

def load_deleted_params_from_file():
    """Загружает список удаленных параметров из файла"""
    try:
        if os.path.exists("deleted_params.json"):
            with open("deleted_params.json", "r", encoding="utf-8") as f:
                deleted_params_list = json.load(f)
            # Конвертируем list обратно в set
            return set(deleted_params_list) if isinstance(deleted_params_list, list) else set()
    except Exception:
        pass
    return set()

def load_param_values_from_file():
    """Загружает параметры из всех файлов param_values_*.json и объединяет их"""
    import glob
    
    # Загружаем список удаленных параметров ПЕРЕД загрузкой параметров
    if "deleted_params" not in st.session_state:
        st.session_state["deleted_params"] = load_deleted_params_from_file()
    
    # Получаем список удаленных параметров, чтобы не загружать их
    deleted_params = st.session_state.get("deleted_params", set())
    
    all_param_values = {}
    all_param_options = {}
    
    # Загружаем параметры из всех файлов param_values_*.json
    # Очищаем все файлы param_values*.json (включая param_values.json без подчеркивания)
    param_files = glob.glob("param_values*.json")
    
    for param_file in param_files:
        try:
            with open(param_file, "r", encoding="utf-8") as f:
                param_data = json.load(f)
                
            # Поддерживаем оба формата: новый (с file_name) и старый (прямой словарь)
            if isinstance(param_data, dict) and "param_values" in param_data:
                # Новый формат
                file_params = param_data.get("param_values", {})
                # Загружаем param_options если есть в файле
                file_options = param_data.get("param_options", {})
            else:
                # Старый формат
                file_params = param_data if isinstance(param_data, dict) else {}
                file_options = {}
            
            # Объединяем параметры из всех файлов (пропускаем удаленные)
            for param_name, sku_values in file_params.items():
                # Пропускаем удаленные параметры
                if param_name in deleted_params:
                    continue
                    
                if param_name and isinstance(sku_values, dict):
                    if param_name not in all_param_values:
                        all_param_values[param_name] = {}
                    # Объединяем значения по артикулам (поздние значения перезаписывают ранние)
                    all_param_values[param_name].update(sku_values)
                    
                    # Загружаем варианты из file_options если есть, иначе собираем из значений
                    if param_name in file_options and isinstance(file_options[param_name], list):
                        # Используем сохраненные варианты из файла
                        if param_name not in all_param_options:
                            all_param_options[param_name] = []
                        for option in file_options[param_name]:
                            if option and str(option) not in all_param_options[param_name]:
                                all_param_options[param_name].append(str(option))
                    else:
                        # Собираем все уникальные значения для param_options из значений
                        if param_name not in all_param_options:
                            all_param_options[param_name] = []
                        for value in sku_values.values():
                            if value and str(value) not in all_param_options[param_name]:
                                all_param_options[param_name].append(str(value))
        except Exception as e:
            continue
    
    # Также проверяем глобальный файл
    if os.path.exists("param_values_global.json"):
        try:
            with open("param_values_global.json", "r", encoding="utf-8") as f:
                global_data = json.load(f)
                # Поддерживаем оба формата: новый (с param_values) и старый (прямой словарь)
                if isinstance(global_data, dict) and "param_values" in global_data:
                    global_params = global_data.get("param_values", {})
                    global_options = global_data.get("param_options", {})
                else:
                    global_params = global_data if isinstance(global_data, dict) else {}
                    global_options = {}
                
                if isinstance(global_params, dict):
                    for param_name, sku_values in global_params.items():
                        # Пропускаем удаленные параметры
                        if param_name in deleted_params:
                            continue
                            
                        if param_name and isinstance(sku_values, dict):
                            if param_name not in all_param_values:
                                all_param_values[param_name] = {}
                            all_param_values[param_name].update(sku_values)
                            
                            # Загружаем варианты из global_options если есть, иначе собираем из значений
                            if param_name in global_options and isinstance(global_options[param_name], list):
                                if param_name not in all_param_options:
                                    all_param_options[param_name] = []
                                for option in global_options[param_name]:
                                    if option and str(option) not in all_param_options[param_name]:
                                        all_param_options[param_name].append(str(option))
                            else:
                                if param_name not in all_param_options:
                                    all_param_options[param_name] = []
                                for value in sku_values.values():
                                    if value and str(value) not in all_param_options[param_name]:
                                        all_param_options[param_name].append(str(value))
        except Exception:
            pass
    
    # Загружаем параметры для текущего файла (если есть) - приоритет выше
    current_file = st.session_state.get("cached_file_name", None)
    if current_file:
        try:
            registry_file = "file_params_registry.json"
            if os.path.exists(registry_file):
                with open(registry_file, "r", encoding="utf-8") as f:
                    registry = json.load(f)
                
                if current_file in registry:
                    param_file = registry[current_file]["param_file"]
                    if os.path.exists(param_file):
                        with open(param_file, "r", encoding="utf-8") as f:
                            param_data = json.load(f)
                            file_params = param_data.get("param_values", {}) if isinstance(param_data, dict) else {}
                            file_options = param_data.get("param_options", {}) if isinstance(param_data, dict) else {}
                            
                            # Объединяем с уже загруженными (приоритет текущего файла)
                            for param_name, sku_values in file_params.items():
                                # Пропускаем удаленные параметры
                                if param_name in deleted_params:
                                    continue
                                    
                                if param_name and isinstance(sku_values, dict):
                                    if param_name not in all_param_values:
                                        all_param_values[param_name] = {}
                                    all_param_values[param_name].update(sku_values)
                                    
                                    # Загружаем варианты из file_options если есть, иначе собираем из значений
                                    if param_name in file_options and isinstance(file_options[param_name], list):
                                        # Используем сохраненные варианты из файла (приоритет выше)
                                        all_param_options[param_name] = list(file_options[param_name])
                                    else:
                                        if param_name not in all_param_options:
                                            all_param_options[param_name] = []
                                        for value in sku_values.values():
                                            if value and str(value) not in all_param_options[param_name]:
                                                all_param_options[param_name].append(str(value))
        except Exception:
            pass
    
    # Инициализируем deleted_params если еще не инициализирован
    if "deleted_params" not in st.session_state:
        st.session_state["deleted_params"] = load_deleted_params_from_file()
    
    # Устанавливаем объединенные параметры
    if all_param_values or all_param_options:
        if all_param_values:
            st.session_state["param_values"] = all_param_values
        if all_param_options:
            # Объединяем с существующими param_options если есть
            # НО исключаем удаленные параметры
            existing_options = st.session_state.get("param_options", {})
            for param_name, options in all_param_options.items():
                # Пропускаем удаленные параметры
                if param_name in deleted_params:
                    continue
                if param_name not in existing_options:
                    existing_options[param_name] = options
                else:
                    # Объединяем варианты, сохраняя порядок из файла
                    combined = list(existing_options[param_name])
                    for option in options:
                        if option and str(option) not in combined:
                            combined.append(str(option))
                    existing_options[param_name] = combined
            st.session_state["param_options"] = existing_options
        # Создаем param_options из значений только если их еще нет в файлах
        # Но сначала удаляем удаленные параметры из param_values в session_state
        if deleted_params and "param_values" in st.session_state:
            params_to_remove = [p for p in st.session_state["param_values"].keys() if p in deleted_params]
            for param_name in params_to_remove:
                del st.session_state["param_values"][param_name]
        
        if not all_param_options:
            create_param_options_from_values()
        return True
    else:
        st.session_state["param_values"] = {}
        st.session_state["param_options"] = {}
        return False
                
def create_param_options_from_values():
    """Создает param_options на основе загруженных данных"""
    if "param_options" not in st.session_state:
        st.session_state["param_options"] = {}
    
    # Получаем список удаленных параметров
    deleted_params = st.session_state.get("deleted_params", load_deleted_params_from_file())
    if deleted_params:
        st.session_state["deleted_params"] = deleted_params
        # Убеждаемся, что это set
        if not isinstance(deleted_params, set):
            deleted_params = set(deleted_params)
            st.session_state["deleted_params"] = deleted_params
    
    # Получаем param_values через функцию, которая фильтрует удаленные параметры
    param_values = get_param_values()
    
    # Дополнительная проверка - удаляем удаленные параметры еще раз на всякий случай
    if deleted_params:
        params_to_remove = [p for p in param_values.keys() if p in deleted_params]
        for param_name in params_to_remove:
            if param_name in param_values:
                del param_values[param_name]
        if params_to_remove:
            st.session_state["param_values"] = param_values
    for param_name, param_data in param_values.items():
        # Пропускаем невалидные параметры (артикулы и т.д.)
        if not is_valid_param_name(param_name):
            continue
        
        # Пропускаем удаленные параметры
        if param_name in deleted_params:
            continue
        
        # Собираем все уникальные значения для каждого параметра
        unique_values = list(set([v for v in param_data.values() if v and v.strip()]))
        if unique_values:
            # Если параметр уже есть в param_options, объединяем значения
            if param_name in st.session_state["param_options"]:
                existing_values = st.session_state["param_options"][param_name]
                combined_values = list(set(existing_values + unique_values))
                st.session_state["param_options"][param_name] = sorted(combined_values)
            else:
                st.session_state["param_options"][param_name] = sorted(unique_values)
        elif param_name == "Крой":
            # Для параметра "Крой" добавляем стандартные варианты, если нет данных
            st.session_state["param_options"][param_name] = ["Классический", "Приталенный", "Свободный", "Оверсайз"]
        else:
            # Если параметр есть в param_values, но нет вариантов, добавляем его в param_options с пустым списком
            if param_name not in st.session_state["param_options"]:
                st.session_state["param_options"][param_name] = []

# ================= ФУНКЦИИ ДЛЯ ПРОГНОЗИРОВАНИЯ С PROPHET =================

def prepare_data_for_prophet(df, metric_column, date_column=None):
    """Подготавливает данные для Prophet"""
    if not PROPHET_AVAILABLE:
        return None
    
    # Если нет колонки с датами, создаем искусственную временную последовательность
    if date_column is None or date_column not in df.columns:
        # Создаем даты на основе индекса
        start_date = datetime.now() - timedelta(days=len(df)-1)
        dates = [start_date + timedelta(days=i) for i in range(len(df))]
        df_prophet = pd.DataFrame({
            'ds': dates,
            'y': df[metric_column].values
        })
    else:
        # Используем существующую колонку с датами
        df_prophet = pd.DataFrame({
            'ds': pd.to_datetime(df[date_column]),
            'y': df[metric_column].values
        })
    
    # Удаляем строки с NaN значениями
    df_prophet = df_prophet.dropna()
    
    return df_prophet

def create_prophet_forecast(df_prophet, periods=30, seasonality_mode='additive'):
    """Создает прогноз с помощью Prophet"""
    if not PROPHET_AVAILABLE or df_prophet is None or len(df_prophet) < 2:
        return None, None, None
    
    try:
        # Создаем модель Prophet
        model = Prophet(
            seasonality_mode=seasonality_mode,
            daily_seasonality=True,
            weekly_seasonality=True,
            yearly_seasonality=True,
            changepoint_prior_scale=0.05
        )
        
        # Обучаем модель
        model.fit(df_prophet)
        
        # Создаем будущие даты
        future = model.make_future_dataframe(periods=periods)
        
        # Делаем прогноз
        forecast = model.predict(future)
        
        return model, forecast, future
        
    except Exception as e:
        st.error(f"Ошибка при создании прогноза: {e}")
        return None, None, None

def plot_prophet_forecast(model, forecast, title="Прогноз Prophet"):
    """Создает график прогноза с помощью plotly"""
    if not PROPHET_AVAILABLE or model is None or forecast is None:
        return None
    
    try:
        # Создаем график с помощью Prophet
        fig = plot_plotly(model, forecast)
        
        # Обновляем заголовок
        fig.update_layout(
            title=title,
            xaxis_title="Дата",
            yaxis_title="Значение",
            width=1000,
            height=600
        )
        
        return fig
        
    except Exception as e:
        st.error(f"Ошибка при создании графика: {e}")
        return None

def plot_prophet_components(model, forecast, title="Компоненты прогноза"):
    """Создает график компонентов прогноза"""
    if not PROPHET_AVAILABLE or model is None or forecast is None:
        return None
    
    try:
        # Создаем график компонентов
        fig = plot_components_plotly(model, forecast)
        
        # Обновляем заголовок
        fig.update_layout(title=title)
        
        return fig
        
    except Exception as e:
        st.error(f"Ошибка при создании графика компонентов: {e}")
        return None

def save_main_page_data_to_file():
    """Сохраняет данные главной страницы в файл, привязанные к текущему файлу"""
    current_file = st.session_state.get("cached_file_name", None)
    
    try:
        main_page_data = {
            "search": st.session_state.get("search", ""),
            "spp": st.session_state.get("spp", 25),
            "buyout_pct": st.session_state.get("buyout_pct", 25),
            "revenue_min": st.session_state.get("revenue_min", 0),
            "revenue_max": st.session_state.get("revenue_max", 1000000),
            "price_min": st.session_state.get("price_min", 0),
            "price_max": st.session_state.get("price_max", 10000),
            "show_images": st.session_state.get("show_images", False),
            "sort_column": st.session_state.get("sort_column", "Выручка"),
            "sort_descending": st.session_state.get("sort_descending", True)
        }
        
        if current_file:
            # Сохраняем настройки для конкретного файла
            file_hash = hash(current_file) % 1000000
            settings_file = f"main_page_settings_{file_hash}.json"
            
            settings_data = {
                "file_name": current_file,
                "settings": main_page_data,
                "timestamp": pd.Timestamp.now().strftime("%Y-%m-%d %H:%M:%S")
            }
            
            with open(settings_file, "w", encoding="utf-8") as f:
                json.dump(settings_data, f, ensure_ascii=False, indent=2)
            
            # Обновляем реестр настроек файлов
            update_file_settings_registry(current_file, settings_file)
        else:
            # Fallback для случаев, когда нет текущего файла
            with open("main_page_data_global.json", "w", encoding="utf-8") as f:
                json.dump(main_page_data, f, ensure_ascii=False, indent=2)
        
        return True
    except Exception as e:
        st.error(f"Ошибка сохранения настроек: {e}")
        return False

def update_file_settings_registry(file_name, settings_file):
    """Обновляет реестр файлов и их настроек"""
    try:
        registry_file = "file_settings_registry.json"
        registry = {}
        
        if os.path.exists(registry_file):
            with open(registry_file, "r", encoding="utf-8") as f:
                registry = json.load(f)
        
        registry[file_name] = {
            "settings_file": settings_file,
            "last_updated": pd.Timestamp.now().strftime("%Y-%m-%d %H:%M:%S")
        }
        
        with open(registry_file, "w", encoding="utf-8") as f:
            json.dump(registry, f, ensure_ascii=False, indent=2)
            
    except Exception as e:
        st.error(f"Ошибка обновления реестра настроек: {e}")

def load_main_page_data_from_file():
    """Загружает данные главной страницы из файла, привязанные к текущему файлу"""
    current_file = st.session_state.get("cached_file_name", None)
    
    if current_file:
        try:
            # Пытаемся загрузить настройки для текущего файла
            registry_file = "file_settings_registry.json"
            if os.path.exists(registry_file):
                with open(registry_file, "r", encoding="utf-8") as f:
                    registry = json.load(f)
                
                if current_file in registry:
                    settings_file = registry[current_file]["settings_file"]
                    if os.path.exists(settings_file):
                        with open(settings_file, "r", encoding="utf-8") as f:
                            settings_data = json.load(f)
                            data = settings_data.get("settings", {})
                            
                            # Загружаем данные в session_state
                            st.session_state["search"] = data.get("search", "")
                            st.session_state["spp"] = data.get("spp", 25)
                            st.session_state["buyout_pct"] = data.get("buyout_pct", 25)
                            st.session_state["revenue_min"] = data.get("revenue_min", 0)
                            st.session_state["revenue_max"] = data.get("revenue_max", 1000000)
                            st.session_state["price_min"] = data.get("price_min", 0)
                            st.session_state["price_max"] = data.get("price_max", 10000)
                            st.session_state["show_images"] = data.get("show_images", False)
                            st.session_state["sort_column"] = data.get("sort_column", "Выручка")
                            st.session_state["sort_descending"] = data.get("sort_descending", True)
                            return True
            
            # Если настройки для файла не найдены, инициализируем значения по умолчанию
            st.session_state["search"] = ""
            st.session_state["spp"] = 25
            st.session_state["buyout_pct"] = 25
            st.session_state["revenue_min"] = 0
            st.session_state["revenue_max"] = 1000000
            st.session_state["price_min"] = 0
            st.session_state["price_max"] = 10000
            st.session_state["show_images"] = False
            st.session_state["sort_column"] = "Выручка"
            st.session_state["sort_descending"] = True
            return True
            
        except Exception as e:
            st.error(f"Ошибка загрузки настроек: {e}")
            return False
    else:
        # Fallback для случаев, когда нет текущего файла
        try:
            if os.path.exists("main_page_data_global.json"):
                with open("main_page_data_global.json", "r", encoding="utf-8") as f:
                    data = json.load(f)
                    
                    # Загружаем данные в session_state
                    st.session_state["search"] = data.get("search", "")
                    st.session_state["spp"] = data.get("spp", 25)
                    st.session_state["buyout_pct"] = data.get("buyout_pct", 25)
                    st.session_state["revenue_min"] = data.get("revenue_min", 0)
                    st.session_state["revenue_max"] = data.get("revenue_max", 1000000)
                    st.session_state["price_min"] = data.get("price_min", 0)
                    st.session_state["price_max"] = data.get("price_max", 10000)
                    st.session_state["show_images"] = data.get("show_images", False)
                    st.session_state["sort_column"] = data.get("sort_column", "Выручка")
                    st.session_state["sort_descending"] = data.get("sort_descending", True)
                    return True
        except Exception:
            pass
    return False

# Автоматическая загрузка параметров при запуске приложения
# Сначала загружаем deleted_params, чтобы не загружать удаленные параметры
if "deleted_params" not in st.session_state:
    st.session_state["deleted_params"] = load_deleted_params_from_file()
    # Убеждаемся, что это set
    if not isinstance(st.session_state["deleted_params"], set):
        st.session_state["deleted_params"] = set(st.session_state["deleted_params"])

# Очищаем все файлы от удаленных параметров при старте (один раз)
if "files_cleaned_on_startup" not in st.session_state:
    deleted_params = st.session_state.get("deleted_params", set())
    if deleted_params:
        # Принудительно очищаем все файлы от удаленных параметров
        cleaned_count = clean_deleted_params_from_all_files(deleted_params)
        if cleaned_count > 0:
            # Файлы были изменены, нужно перезагрузить параметры
            pass
        # Также удаляем удаленные параметры из session_state
        if "param_values" in st.session_state:
            for param_name in list(st.session_state["param_values"].keys()):
                if param_name in deleted_params:
                    del st.session_state["param_values"][param_name]
        if "param_options" in st.session_state:
            for param_name in list(st.session_state["param_options"].keys()):
                if param_name in deleted_params:
                    del st.session_state["param_options"][param_name]
    st.session_state["files_cleaned_on_startup"] = True

# Загружаем параметры ТОЛЬКО после очистки файлов
if "param_values" not in st.session_state:
    load_param_values_from_file()
    load_param_history_from_file()
    
    # После загрузки еще раз удаляем удаленные параметры из session_state
    deleted_params = st.session_state.get("deleted_params", set())
    if deleted_params:
        if "param_values" in st.session_state:
            for param_name in list(st.session_state["param_values"].keys()):
                if param_name in deleted_params:
                    del st.session_state["param_values"][param_name]
        if "param_options" in st.session_state:
            for param_name in list(st.session_state["param_options"].keys()):
                if param_name in deleted_params:
                    del st.session_state["param_options"][param_name]

# КРИТИЧЕСКИ ВАЖНО: ВСЕГДА фильтруем удаленные параметры при каждом rerun
# Это последняя линия защиты от восстановления удаленных параметров
# Загружаем deleted_params из файла, чтобы быть уверенными в актуальности
if "deleted_params" in st.session_state:
    # Объединяем с данными из файла на случай, если они были обновлены
    file_deleted_params = load_deleted_params_from_file()
    current_deleted_params = st.session_state.get("deleted_params", set())
    if isinstance(file_deleted_params, set):
        deleted_params = file_deleted_params.union(current_deleted_params if isinstance(current_deleted_params, set) else set())
    else:
        deleted_params = current_deleted_params if isinstance(current_deleted_params, set) else set()
    
    # Убеждаемся, что это set
    if not isinstance(deleted_params, set):
        deleted_params = set(deleted_params) if deleted_params else set()
    
    # Обновляем session_state
    st.session_state["deleted_params"] = deleted_params
else:
    deleted_params = load_deleted_params_from_file()
    if not isinstance(deleted_params, set):
        deleted_params = set(deleted_params) if deleted_params else set()
    st.session_state["deleted_params"] = deleted_params

# Удаляем удаленные параметры из всех мест в session_state
if deleted_params:
    # Удаляем из param_values
    if "param_values" in st.session_state:
        params_to_remove = [p for p in st.session_state["param_values"].keys() if p in deleted_params]
        for param_name in params_to_remove:
            del st.session_state["param_values"][param_name]
    
    # Удаляем из param_options
    if "param_options" in st.session_state:
        params_to_remove = [p for p in st.session_state["param_options"].keys() if p in deleted_params]
        for param_name in params_to_remove:
            del st.session_state["param_options"][param_name]

# Автоматическая загрузка настроек исключения параметров при запуске
if "excluded_params_loaded" not in st.session_state:
    load_excluded_params_settings()
    st.session_state["excluded_params_loaded"] = True

# Автоматическая загрузка результатов массового анализа при запуске
# Инициализируем пустым списком, если еще не загружено
if "mass_analysis_results" not in st.session_state:
    st.session_state["mass_analysis_results"] = []

# Загружаем результаты из файлов при каждом запуске (если еще не загружены)
if not st.session_state.get("mass_analysis_results", []):
    # Сначала пытаемся загрузить из сохраненных результатов
    saved_results = load_mass_analysis_results()
    if saved_results:
        st.session_state["mass_analysis_results"] = saved_results
    else:
        # Если нет сохраненных результатов, пытаемся загрузить из прогресса
        saved_progress = load_mass_analysis_progress()
        if saved_progress and saved_progress.get("results"):
            st.session_state["mass_analysis_results"] = saved_progress["results"]
            # Сохраняем результаты в отдельный файл для постоянного хранения
            save_mass_analysis_results(saved_progress["results"])

# Автоматическая загрузка данных главной страницы при запуске приложения
if "main_page_data_loaded" not in st.session_state:
    load_main_page_data_from_file()
    st.session_state["main_page_data_loaded"] = True

def _cache_root():
    return os.path.join(os.path.dirname(os.path.abspath(__file__)), "wb_cache")
def _cache_dir():
    d = _cache_root()
    os.makedirs(d, exist_ok=True)
    os.makedirs(os.path.join(d, "imgs"), exist_ok=True)
    return d
def _url_cache_path():
    return os.path.join(_cache_dir(), "image_cache.json")
def load_url_cache():
    p = _url_cache_path()
    if os.path.exists(p):
        try:
            with open(p, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            return {}
    return {}
def save_url_cache(m: dict):
    try:
        with open(_url_cache_path(), "w", encoding="utf-8") as f:
            json.dump(m, f, ensure_ascii=False, indent=2)
    except Exception:
        pass
def get_url_cache():
    if "img_url_cache" not in st.session_state:
        st.session_state["img_url_cache"] = load_url_cache()
    return st.session_state["img_url_cache"]

def img_path_for(nm: str, fmt: str = "JPEG"):
    nm = str(nm).replace(".0", "")
    ext = "jpg" if fmt.upper() == "JPEG" else "png"
    return os.path.join(_cache_dir(), "imgs", f"{nm}.{ext}")
def get_cached_image_path(nm: str):
    """Проверяет, есть ли изображение в кеше и возвращает путь к нему"""
    nm = str(nm).replace(".0", "")
    
    # Проверяем различные форматы изображений
    for ext in ("jpg", "png", "jpeg", "webp", "JPG", "PNG", "JPEG", "WEBP"):
        p = os.path.join(_cache_dir(), "imgs", f"{nm}.{ext}")
        if os.path.exists(p) and os.path.getsize(p) > 0:  # Проверяем, что файл не пустой
            return p
    
    return ""
def ensure_image_cached(nm: str, url: str, fmt: str = "JPEG", timeout: int = 25) -> str:
    """Скачивает изображение по URL и сохраняет в кеш"""
    try:
        # Сначала проверяем, есть ли уже изображение в кеше
        p_exist = get_cached_image_path(nm)
        if p_exist:
            return p_exist
        
        if not url:
            return ""
        
        # Создаем путь для сохранения
        path = img_path_for(nm, fmt)
        
        # Убеждаемся, что директория существует
        os.makedirs(os.path.dirname(path), exist_ok=True)
        
        # Скачиваем изображение
        headers = {
            "User-Agent": "WB-Dashboard/1.0",
            "Accept": "image/webp,image/apng,image/*,*/*;q=0.8"
        }
        
        with requests.get(url, headers=headers, timeout=timeout, stream=True) as r:
            if r.status_code != 200:
                return ""
            
            # Проверяем, что это действительно изображение
            content_type = r.headers.get('content-type', '').lower()
            if not content_type.startswith('image/'):
                return ""
            
            # Сохраняем файл
            with open(path, "wb") as f:
                for chunk in r.iter_content(8192):
                    if chunk:
                        f.write(chunk)
            
            # Проверяем, что файл создался и не пустой
            if os.path.exists(path) and os.path.getsize(path) > 0:
                return path
            else:
                # Удаляем пустой файл
                if os.path.exists(path):
                    os.remove(path)
                return ""
                
    except Exception as e:
        # В случае ошибки удаляем частично скачанный файл
        try:
            path = img_path_for(nm, fmt)
            if os.path.exists(path):
                os.remove(path)
        except:
            pass
        return ""
def get_cache_status():
    """Возвращает информацию о состоянии кеша изображений"""
    cache_dir = os.path.join(_cache_dir(), "imgs")
    if not os.path.exists(cache_dir):
        return {"count": 0, "size": 0, "files": []}
    
    files = []
    total_size = 0
    
    for filename in os.listdir(cache_dir):
        if filename.lower().endswith(('.jpg', '.jpeg', '.png', '.webp')):
            file_path = os.path.join(cache_dir, filename)
            file_size = os.path.getsize(file_path)
            files.append({
                "name": filename,
                "size": file_size,
                "path": file_path
            })
            total_size += file_size
    
    return {
        "count": len(files),
        "size": total_size,
        "files": files
    }

def load_image_bytes(path: str, max_w: int | None = None) -> bytes:
    if not path or not os.path.exists(path):
        return b""
    if Image is None or max_w is None:
        try:
            with open(path, "rb") as f:
                return f.read()
        except Exception:
            return b""
    try:
        im = Image.open(path)
        if im.mode not in ("RGB", "RGBA"):
            im = im.convert("RGB")
        if max_w and im.width > max_w:
            ratio = max_w / float(im.width)
            im = im.resize((int(im.width * ratio), int(im.height * ratio)))
        bio = BytesIO()
        im.save(bio, format="JPEG", quality=85)
        return bio.getvalue()
    except Exception:
        try:
            with open(path, "rb") as f:
                return f.read()
        except Exception:
            return b""
def img_data_uri(nm: str, max_w: int | None = None) -> str:
    p = get_cached_image_path(nm)
    if not p:
        return ""
    try:
        data = load_image_bytes(p, max_w=max_w)
        if not data:
            return ""
        b64 = base64.b64encode(data).decode("ascii")
        return f"data:image/jpeg;base64,{b64}"
    except Exception:
        return ""

def build_wb_product_url(nm, host="https://global.wildberries.ru"):
    return f"{host.rstrip('/')}/catalog/{str(nm).replace('.0','')}/detail.aspx"

def extract_sku_from_url(url: str) -> str | None:
    """
    Извлекает артикул (SKU) из URL товара Wildberries.
    
    Поддерживаемые форматы:
    - https://www.wildberries.ru/catalog/{sku}/detail.aspx
    - https://wildberries.ru/catalog/{sku}/detail.aspx
    - https://global.wildberries.ru/catalog/{sku}/detail.aspx
    - https://www.wildberries.ru/catalog/{sku}/detail
    - https://www.wildberries.ru/catalog/{sku}/
    
    Args:
        url: URL товара Wildberries
        
    Returns:
        Артикул товара или None, если не удалось извлечь
    """
    if not url or not isinstance(url, str):
        return None
    
    try:
        # Убираем пробелы и лишние символы
        url = url.strip()
        
        # Паттерны для поиска артикула в URL
        # Паттерн 1: /catalog/{sku}/detail
        pattern1 = r'/catalog/(\d+)/detail'
        match = re.search(pattern1, url)
        if match:
            return match.group(1)
        
        # Паттерн 2: /catalog/{sku}/
        pattern2 = r'/catalog/(\d+)/?'
        match = re.search(pattern2, url)
        if match:
            return match.group(1)
        
        # Паттерн 3: /catalog/{sku} (без слеша в конце)
        pattern3 = r'/catalog/(\d+)(?:\?|$)'
        match = re.search(pattern3, url)
        if match:
            return match.group(1)
        
        return None
    except Exception:
        return None

def get_cached_images_for_sku(sku: str, max_images: int = 3) -> list:
    """
    Получает пути к кешированным изображениям товара по артикулу.
    Использует уже скачанные изображения из кеша.
    
    Args:
        sku: Артикул товара
        max_images: Максимальное количество изображений (по умолчанию 3)
        
    Returns:
        Список путей к изображениям
    """
    images = []
    sku_clean = str(sku).replace(".0", "")
    
    # Проверяем основное изображение (обычно это первое фото товара)
    main_image = get_cached_image_path(sku_clean)
    if main_image and os.path.exists(main_image):
        images.append(main_image)
    
    # Если нужно больше изображений, можно искать дополнительные
    # (например, с суффиксами _1, _2 и т.д.)
    if len(images) < max_images:
        cache_dir = os.path.join(_cache_dir(), "imgs")
        if os.path.exists(cache_dir):
            # Ищем дополнительные изображения с суффиксами
            for suffix in range(1, max_images):
                for ext in ("jpg", "png", "jpeg", "webp", "JPG", "PNG", "JPEG", "WEBP"):
                    additional_path = os.path.join(cache_dir, f"{sku_clean}_{suffix}.{ext}")
                    if os.path.exists(additional_path) and additional_path not in images:
                        images.append(additional_path)
                        break
    
    return images[:max_images]

def analyze_image_with_ai(image_path_or_url: str, api_key: str = None, selected_params: list = None, product_name: str = None, sku: str = None) -> dict:
    """
    Анализирует изображение товара через нейросеть (OpenAI Vision API).
    Работает как с локальными файлами, так и с URL.
    
    Args:
        image_path_or_url: Путь к локальному файлу изображения или URL
        api_key: API ключ OpenAI (если None, берется из secrets.toml или session_state)
        selected_params: Список параметров для анализа
        product_name: Название товара (если доступно)
        sku: Артикул товара (для получения названия, если product_name не указан)
        
    Returns:
        Словарь с определенными параметрами
    """
    params = {}
    
    # Проверяем наличие библиотеки openai (используем глобальную проверку)
    if not OPENAI_AVAILABLE:
        # Не показываем ошибку здесь, так как она может показываться многократно
        # Вместо этого вернем пустой словарь
            return params
    
    try:
        
        # Получаем API ключ
        if not api_key:
            # Пытаемся получить из st.secrets (Streamlit автоматически загружает secrets.toml)
            try:
                api_key = st.secrets.get('openai_api_key', '')
            except:
                pass
            
            # Если не нашли в secrets, пытаемся из session_state
            if not api_key:
                api_key = st.session_state.get('openai_api_key')
        
        if not api_key:
            # Если ключа нет, возвращаем пустой словарь
            return params
        
        # Используем OpenAI Vision API с увеличенным timeout
        client = openai.OpenAI(api_key=api_key, timeout=120.0)  # 120 секунд timeout
        
        # Получаем существующие параметры и их варианты
        param_options = st.session_state.get("param_options", {})
        
        # Если указаны выбранные параметры, фильтруем только их
        if selected_params and len(selected_params) > 0:
            param_options = {k: v for k, v in param_options.items() if k in selected_params}
        
        # Определяем иерархию сегментов
        # Получаем иерархию из session_state или используем значения по умолчанию
        hierarchy_params = get_hierarchy_params()
        subtype_params = get_subtype_params()  # Параметры, которые находятся внутри "Подтип"
        visual_params = get_visual_params()
        
        # Разделяем параметры на иерархические и визуальные
        hierarchical_params_dict = {}
        visual_params_dict = {}
        other_params_dict = {}
        
        for param_name, options in param_options.items():
            if param_name in hierarchy_params:
                hierarchical_params_dict[param_name] = options
            elif param_name in visual_params or any(vp in param_name for vp in visual_params):
                visual_params_dict[param_name] = options
            else:
                other_params_dict[param_name] = options
        
        # Получаем название товара из API WB (если доступно)
        if not product_name and sku:
            try:
                product_name = get_product_name_from_wb(sku)
            except:
                pass
        
        # Если название все еще не получено, пытаемся извлечь SKU из пути
        if not product_name:
            try:
                filename = os.path.basename(image_path_or_url)
                # Извлекаем SKU из имени файла (формат: SKU_screenshotapi_0.png)
                if "_screenshotapi_" in filename:
                    sku_from_path = filename.split("_screenshotapi_")[0]
                    if sku_from_path and sku_from_path.isdigit():
                        product_name = get_product_name_from_wb(sku_from_path)
            except:
                pass
        
        # Формируем промпт с учетом иерархии сегментов
        prompt_parts = ["""Ты - помощник для анализа товаров электронной коммерции. Проанализируй это изображение товара (одежды или аксессуара) с Wildberries.

ВАЖНО: 
- Это изображение товара для продажи в интернет-магазине
- Ты анализируешь ТОВАР (одежду, предмет), а не людей
- Игнорируй людей на изображении полностью - они просто демонстрируют товар
- Это коммерческий анализ товара для каталога, не анализ людей
- На скриншоте может быть видно НАЗВАНИЕ ТОВАРА - обязательно учитывай его при анализе!
- Если на скриншоте видно название товара, используй его для более точного определения параметров

АНАЛИЗ ДОЛЖЕН ПРОХОДИТЬ ПО ИЕРАРХИИ СЕГМЕНТОВ:

ШАГ 1: Определи ТИП товара (верхний уровень сегментации)
ШАГ 2: Внутри определенного типа определи КОМПЛЕКТ
ШАГ 3: Внутри определенного комплекта определи ПОДТИП (через параметры РУКАВ и ВОРОТ)
ШАГ 4: Внутри определенного подтипа определи РУКАВ
ШАГ 5: Внутри определенного рукава определи ВОРОТ
ШАГ 6: ТОЛЬКО после определения всех иерархических параметров (Тип → Комплект → Подтип (Рукав, Ворот)) анализируй ВИЗУАЛЬНЫЕ параметры (Цвет, Принт, Логотип, Строчка шов и т.д.)

Визуальные параметры анализируются только в контексте определенной иерархии сегментов!"""]
        
        # Добавляем информацию о названии товара, если оно доступно
        if product_name:
            prompt_parts.append(f"\nВАЖНО: Название товара из каталога: {product_name}")
            prompt_parts.append("Используй это название для более точного определения параметров товара!")
            prompt_parts.append("Если на скриншоте видно название товара, сравни его с указанным выше и учитывай при анализе.")
        
        # Добавляем информацию о существующих параметрах с учетом иерархии
        if param_options:
            prompt_parts.append("\nКРИТИЧЕСКИ ВАЖНО: Используй ТОЛЬКО параметры из списка ниже! НЕ добавляй параметры, которых нет в списке!")
            
            # Сначала добавляем иерархические параметры
            if hierarchical_params_dict:
                prompt_parts.append("\n=== ИЕРАРХИЧЕСКИЕ ПАРАМЕТРЫ (определяй в указанном порядке) ===")
                for param_name in hierarchy_params:
                    if param_name in hierarchical_params_dict:
                        options = hierarchical_params_dict[param_name]
                        if isinstance(options, list) and options:
                            prompt_parts.append(f"\n{param_name} (варианты): {', '.join(options)}")
                        else:
                            prompt_parts.append(f"\n{param_name}: (определи значение самостоятельно на основе изображения)")
            
            # Затем добавляем визуальные параметры
            if visual_params_dict:
                prompt_parts.append("\n=== ВИЗУАЛЬНЫЕ ПАРАМЕТРЫ (анализируй ТОЛЬКО после определения иерархии) ===")
                for param_name, options in visual_params_dict.items():
                    # Включаем параметр в промпт, даже если у него нет вариантов
                    if isinstance(options, list):
                        if options:  # Если есть варианты
                            if param_name == "Цвет":
                                # Разделяем основные цвета и сочетания
                                single_colors = [opt for opt in options if "-" not in opt]
                                combined_colors = [opt for opt in options if "-" in opt]
                                
                                prompt_parts.append(f"\n- {param_name}:")
                                prompt_parts.append(f"  * Основные цвета (используй когда товар одного цвета): {', '.join(single_colors)}")
                                if combined_colors:
                                    prompt_parts.append(f"  * Сочетания двух цветов (используй ТОЛЬКО когда видно два разных цвета на товаре): {', '.join(combined_colors)}")
                                    prompt_parts.append(f"  ВАЖНО: Сочетания типа 'Черно-белый' используй ТОЛЬКО если товар действительно имеет два разных цвета!")
                            else:
                                prompt_parts.append(f"- {param_name}: {', '.join(options)}")
                        else:  # Если вариантов нет, все равно упоминаем параметр
                            prompt_parts.append(f"- {param_name}: (определи значение самостоятельно на основе изображения)")
                    elif options:  # Если options не список, но есть значение (строка)
                        prompt_parts.append(f"- {param_name}: {options}")
                    else:  # Если параметр есть, но без вариантов
                        prompt_parts.append(f"- {param_name}: (определи значение самостоятельно на основе изображения)")
            
            # Добавляем остальные параметры (если есть)
            if other_params_dict:
                prompt_parts.append("\n=== ДОПОЛНИТЕЛЬНЫЕ ПАРАМЕТРЫ ===")
                for param_name, options in other_params_dict.items():
                    if isinstance(options, list) and options:
                        prompt_parts.append(f"- {param_name}: {', '.join(options)}")
                    else:
                        prompt_parts.append(f"- {param_name}: (определи значение самостоятельно на основе изображения)")
            
            # Формируем JSON структуру на основе существующих параметров
            json_keys = []
            for param_name in param_options.keys():
                json_keys.append(f'  "{param_name}": "значение из вариантов выше или пустая строка"')
            
            prompt_parts.append(f"""
Ответь в формате JSON, используя ТОЛЬКО параметры из списка выше:
{{
{chr(10).join(json_keys)}
}}

КРИТИЧЕСКИ ВАЖНО - СТРОГО СЛЕДУЙ ИЕРАРХИИ: 
1. СНАЧАЛА определи иерархические параметры в строгом порядке:
   - Сначала ТИП товара (футболка, лонгслив, комплект, худи, свитшот и т.д.)
   - Затем КОМПЛЕКТ (внутри определенного типа)
   - Затем ПОДТИП (определяется через параметры РУКАВ и ВОРОТ внутри определенного комплекта)
   - Затем РУКАВ (внутри определенного подтипа)
   - Затем ВОРОТ (внутри определенного рукава)

2. ТОЛЬКО ПОСЛЕ определения всех иерархических параметров анализируй ВИЗУАЛЬНЫЕ параметры:
   - Цвет (анализируй цвет товара в контексте определенной иерархии)
   - Принт (есть ли принт на товаре)
   - Логотип (есть ли логотип)
   - Строчка/Шов (цветная строчка, особенности швов)
   - И другие визуальные параметры

3. ОБЯЗАТЕЛЬНО включи ВСЕ параметры из списка выше в ответ JSON
4. Используй ТОЛЬКО параметры из списка выше
5. НЕ добавляй параметры, которых нет в списке
6. Для параметра "Цвет": если товар имеет ОДИН цвет - используй основной цвет (Белый, Черный, Синий и т.д.)
7. Для параметра "Цвет": если товар имеет ДВА разных цвета - используй сочетание (Черно-белый, Сине-белый и т.д.)
8. НЕ используй сочетания цветов если товар однотонный!
9. Если параметр определить невозможно после тщательного анализа, укажи пустую строку ""
10. Если есть существующие варианты для параметра, ОБЯЗАТЕЛЬНО выбери один из них
11. Будь внимателен и строго следуй иерархии: Тип → Комплект → Подтип (Рукав, Ворот) → Визуальные параметры""")
        else:
            # Если параметров нет, используем базовые
            prompt_parts.append("""
Определи следующие параметры товара:
1. Цвет товара - укажи цвет из доступных вариантов
2. Другие видимые характеристики товара

Ответь в формате JSON с параметрами, которые ты видишь на изображении.""")
        
        prompt_parts.append("""

ВАЖНО: 
- Анализируй только характеристики ТОВАРА (цвет, фасон, детали товара)
- Не анализируй людей на изображении
- Фокусируйся на визуальных характеристиках самого товара""")
        
        prompt = "\n".join(prompt_parts)
        
        # Определяем, это локальный файл или URL
        is_local_file = False
        if image_path_or_url:
            is_local_file = os.path.exists(image_path_or_url) and os.path.isfile(image_path_or_url)
        
        if is_local_file:
            # Проверяем размер файла (OpenAI имеет ограничения)
            file_size = os.path.getsize(image_path_or_url)
            max_size = 20 * 1024 * 1024  # 20 MB
            
            if file_size > max_size:
                st.warning(f"⚠️ Изображение слишком большое ({file_size / 1024 / 1024:.1f} MB). Максимальный размер: 20 MB")
                return params
            
            # Локальный файл - конвертируем в base64
            try:
                with open(image_path_or_url, 'rb') as image_file:
                    image_bytes = image_file.read()
                    image_data = base64.b64encode(image_bytes).decode('utf-8')
                    # Определяем MIME тип
                    ext = os.path.splitext(image_path_or_url)[1].lower()
                    mime_type = {
                        '.jpg': 'image/jpeg',
                        '.jpeg': 'image/jpeg',
                        '.png': 'image/png',
                        '.webp': 'image/webp'
                    }.get(ext, 'image/jpeg')
                    
                    image_content = f"data:{mime_type};base64,{image_data}"
            except Exception as e:
                st.error(f"❌ Не удалось прочитать файл изображения: {str(e)[:100]}")
                return params
        else:
            # URL - используем напрямую
            image_content = image_path_or_url
        
        # Отправляем запрос в OpenAI Vision API с retry логикой для таймаутов
        max_retries = 3
        retry_count = 0
        response = None
        
        while retry_count < max_retries:
            try:
                response = client.chat.completions.create(
                    model="gpt-4o",  # Модель с поддержкой vision
                    messages=[
                        {
                            "role": "user",
                            "content": [
                                {"type": "text", "text": prompt},
                                {
                                    "type": "image_url",
                                    "image_url": {
                                        "url": image_content
                                    }
                                }
                            ]
                        }
                    ],
                    max_tokens=300
                )
                # Если запрос успешен, выходим из цикла
                break
            except openai.APITimeoutError as timeout_err:
                retry_count += 1
                if retry_count < max_retries:
                    # Экспоненциальная задержка: 5, 10, 20 секунд
                    wait_time = (2 ** (retry_count - 1)) * 5
                    import time
                    time.sleep(wait_time)
                    continue
                else:
                    # Превышено количество попыток
                    raise Exception(f"APITimeoutError: Превышено время ожидания после {max_retries} попыток: {str(timeout_err)}")
            except openai.RateLimitError as e:
                # Пробрасываем ошибку rate limit для обработки на верхнем уровне
                raise Exception(f"Rate limit error: {str(e)}")
            except openai.APIError as e:
                # Проверяем, является ли это ошибкой rate limit по коду статуса
                error_str = str(e)
                if "429" in error_str or "rate limit" in error_str.lower() or "too many requests" in error_str.lower():
                    raise Exception(f"Rate limit error: {error_str}")
                # Пробуем альтернативную модель
                try:
                    response = client.chat.completions.create(
                        model="gpt-4-turbo",
                        messages=[
                            {
                                "role": "user",
                                "content": [
                                    {"type": "text", "text": prompt},
                                    {
                                        "type": "image_url",
                                        "image_url": {
                                            "url": image_content
                                        }
                                    }
                                ]
                            }
                        ],
                        max_tokens=300
                    )
                    # Если запрос успешен, выходим из цикла
                    break
                except openai.APITimeoutError as timeout_err:
                    retry_count += 1
                    if retry_count < max_retries:
                        wait_time = (2 ** (retry_count - 1)) * 5
                        import time
                        time.sleep(wait_time)
                        continue
                    else:
                        raise Exception(f"APITimeoutError: Превышено время ожидания после {max_retries} попыток: {str(timeout_err)}")
                except openai.RateLimitError as rate_err:
                    raise Exception(f"Rate limit error: {str(rate_err)}")
                except openai.APIError as api_err:
                    error_str = str(api_err)
                    if "429" in error_str or "rate limit" in error_str.lower() or "too many requests" in error_str.lower():
                        raise Exception(f"Rate limit error: {error_str}")
                    raise e
                except Exception:
                    raise e
        
        # Проверяем, что получили ответ
        if response is None:
            raise Exception("Не удалось получить ответ от API после всех попыток")
        
        # Парсим ответ
        response_text = response.choices[0].message.content
        
        # Для отладки показываем сырой ответ (можно убрать в продакшене)
        debug_mode = st.session_state.get('debug_ai_analysis', False)
        if debug_mode:
            with st.expander("🔍 Отладочная информация (ответ API)", expanded=False):
                st.code(response_text)
        
        # Пытаемся извлечь JSON из ответа
        import json
        # Ищем JSON в ответе
        json_start = response_text.find('{')
        json_end = response_text.rfind('}') + 1
        
        if json_start >= 0 and json_end > json_start:
            json_str = response_text[json_start:json_end]
            try:
                parsed_params = json.loads(json_str)
                
                # Сначала проверяем, что все параметры из param_options присутствуют в ответе
                if param_options:
                    for param_name in param_options.keys():
                        if param_name not in parsed_params:
                            # Если параметр не был возвращен нейросетью, добавляем его как пустую строку
                            # Это позволит отследить, какие параметры не были определены
                            parsed_params[param_name] = ""
                            if debug_mode:
                                st.info(f"ℹ️ Параметр '{param_name}' не был возвращен нейросетью, добавляем как пустое значение")
                
                # Обрабатываем параметры с учетом существующих вариантов
                # Игнорируем параметры, которых нет в существующих
                for key, value in parsed_params.items():
                    # Пропускаем параметры, которых нет в существующих вариантах
                    if param_options and len(param_options) > 0 and key not in param_options:
                        if debug_mode:
                            st.info(f"ℹ️ Параметр '{key}' пропущен, так как его нет в существующих параметрах")
                        continue
                    
                    # Обрабатываем значение (включая пустые строки)
                    if value is not None:
                        value_clean = str(value).strip()
                    else:
                        value_clean = ""
                    
                    # Сохраняем значение, даже если оно пустое (для отслеживания)
                    # Но в финальный результат включаем только непустые значения
                    if value_clean:
                        
                        # Проверяем, есть ли существующие варианты для этого параметра
                        if param_options and key in param_options and param_options[key]:
                            # Ищем точное совпадение
                            exact_match = None
                            for option in param_options[key]:
                                if option.lower() == value_clean.lower():
                                    exact_match = option
                                    break
                            
                            if exact_match:
                                params[key] = exact_match
                            else:
                                # Ищем частичное совпадение (содержит или содержится)
                                partial_match = None
                                value_lower = value_clean.lower()
                                for option in param_options[key]:
                                    option_lower = option.lower()
                                    if value_lower in option_lower or option_lower in value_lower:
                                        partial_match = option
                                        break
                                
                                if partial_match:
                                    params[key] = partial_match
                                else:
                                    # Если не нашли совпадение, но есть варианты - все равно сохраняем значение
                                    # Пользователь может решить, использовать его или нет
                                    params[key] = value_clean
                                    if debug_mode:
                                        st.warning(f"⚠️ Значение '{value_clean}' для параметра '{key}' не найдено в существующих вариантах: {param_options[key][:5]}")
                        else:
                            # Если нет существующих вариантов, используем значение как есть
                            params[key] = value_clean
            except json.JSONDecodeError as e:
                if debug_mode:
                    st.error(f"❌ Ошибка парсинга JSON: {e}")
                    st.code(json_str)
                # Пытаемся извлечь параметры вручную через регулярные выражения
                # Используем только существующие параметры
                import re
                param_names_to_extract = list(param_options.keys()) if param_options else ["Цвет", "Стиль", "Материал", "Длина"]
                for param_name in param_names_to_extract:
                    pattern = f'"{param_name}"\\s*:\\s*"([^"]+)"'
                    match = re.search(pattern, response_text, re.IGNORECASE)
                    if match:
                        value = match.group(1).strip()
                        if value:
                            # Проверяем соответствие существующим вариантам
                            if param_name in param_options and param_options[param_name]:
                                # Ищем совпадение
                                for option in param_options[param_name]:
                                    if option.lower() == value.lower():
                                        params[param_name] = option
                                        break
                            else:
                                params[param_name] = value
        else:
            # Если не нашли JSON, пытаемся извлечь параметры другим способом
            if debug_mode:
                st.warning("⚠️ JSON не найден в ответе. Пытаемся извлечь параметры другим способом...")
                st.code(response_text[:500])
            
            # Пытаемся найти параметры в тексте
            # Используем только существующие параметры
            import re
            param_names_to_extract = list(param_options.keys()) if param_options else ["Цвет", "Стиль", "Материал", "Длина"]
            for param_name in param_names_to_extract:
                # Ищем паттерны типа "Цвет: Красный" или "Цвет": "Красный"
                patterns = [
                    f'{param_name}\\s*[:=]\\s*"([^"]+)"',
                    f'{param_name}\\s*[:=]\\s*([А-Яа-яA-Za-z\\-]+)',
                ]
                for pattern in patterns:
                    match = re.search(pattern, response_text, re.IGNORECASE)
                    if match:
                        value = match.group(1).strip()
                        if value and len(value) < 50:  # Ограничиваем длину значения
                            # Проверяем соответствие существующим вариантам
                            if param_name in param_options and param_options[param_name]:
                                # Ищем совпадение
                                for option in param_options[param_name]:
                                    if option.lower() == value.lower() or value.lower() in option.lower():
                                        params[param_name] = option
                                        break
                            else:
                                params[param_name] = value
                            break
        
    except Exception as e:
        # Логируем ошибку для отладки
        error_msg = str(e)
        error_type = type(e).__name__
        
        # Проверяем тип ошибки и показываем соответствующее сообщение
        if "api_key" in error_msg.lower() or "authentication" in error_msg.lower() or "401" in error_msg:
            st.error(f"❌ Ошибка аутентификации: Проверьте правильность API ключа OpenAI")
        elif "rate limit" in error_msg.lower() or "quota" in error_msg.lower() or "429" in error_msg:
            st.error(f"❌ Превышен лимит запросов к API. Попробуйте позже.")
        elif "invalid" in error_msg.lower() and ("image" in error_msg.lower() or "format" in error_msg.lower()):
            st.error(f"❌ Не удалось обработать изображение. Проверьте формат файла.")
        elif "timeout" in error_msg.lower() or "APITimeoutError" in error_type:
            st.error(f"❌ Превышено время ожидания ответа от API. Попробуйте позже.")
        elif "connection" in error_msg.lower() or "network" in error_msg.lower():
            st.error(f"❌ Ошибка подключения к API. Проверьте интернет-соединение.")
        else:
            # Для отладки показываем краткую информацию об ошибке
            st.warning(f"⚠️ Ошибка при анализе изображения: {error_type}")
            # В режиме отладки можно показать полную ошибку
            if st.session_state.get('debug_mode', False):
                st.error(f"Детали ошибки: {error_msg[:200]}")
        
        return params
    
    return params

def analyze_image_for_params(image_url: str) -> dict:
    """
    Анализирует изображение товара и определяет параметры (цвет, стиль и т.д.).
    Использует нейросеть для анализа (не скачивает изображение).
    
    Args:
        image_url: URL изображения товара
        
    Returns:
        Словарь с определенными параметрами
    """
    # Используем нейросеть для анализа
    return analyze_image_with_ai(image_url)

def extract_dominant_colors_from_image(image, num_colors: int = 5) -> list:
    """
    Извлекает доминирующие цвета из изображения.
    Упрощенная версия без sklearn для совместимости.
    """
    try:
        # Преобразуем изображение в массив
        img_array = np.array(image)
        
        # Изменяем форму для анализа
        pixels = img_array.reshape(-1, 3)
        
        # Упрощенный алгоритм: берем случайную выборку пикселей
        # и находим наиболее частые цвета
        sample_size = min(10000, len(pixels))
        if len(pixels) > sample_size:
            sample_indices = np.random.choice(len(pixels), sample_size, replace=False)
            sample_pixels = pixels[sample_indices]
        else:
            sample_pixels = pixels
        
        # Квантуем цвета (округление до ближайших значений)
        quantized = (sample_pixels / 32).astype(int) * 32
        quantized = np.clip(quantized, 0, 255)
        
        # Подсчитываем частоту цветов
        color_counts = Counter(tuple(c) for c in quantized)
        
        # Берем топ-N цветов
        top_colors = color_counts.most_common(num_colors)
        
        # Возвращаем RGB значения
        return [list(color[0]) for color in top_colors]
        
    except Exception:
        # Fallback: просто берем средний цвет
        try:
            img_array = np.array(image)
            avg_color = img_array.mean(axis=(0, 1)).astype(int)
            return [list(avg_color)]
        except:
            return []

def get_color_name_russian(rgb: list) -> str:
    """
    Преобразует RGB в русское название цвета.
    """
    try:
        r, g, b = rgb[0], rgb[1], rgb[2]
        
        # Простая эвристика для определения цвета
        # Определяем доминирующий канал
        max_val = max(r, g, b)
        min_val = min(r, g, b)
        diff = max_val - min_val
        
        # Серый/белый/черный
        if diff < 30:
            if max_val > 200:
                return "Белый"
            elif max_val < 50:
                return "Черный"
            else:
                return "Серый"
        
        # Определяем цвет по доминирующему каналу
        if r > g and r > b:
            if r > 200 and g < 100 and b < 100:
                return "Красный"
            elif r > 150 and g > 100:
                return "Оранжевый"
            elif r > 150 and b > 100:
                return "Розовый"
            else:
                return "Красный"
        elif g > r and g > b:
            if g > 200 and r < 100 and b < 100:
                return "Зеленый"
            elif g > 150 and b > 100:
                return "Бирюзовый"
            else:
                return "Зеленый"
        elif b > r and b > g:
            if b > 200 and r < 100 and g < 100:
                return "Синий"
            elif b > 150 and r > 100:
                return "Фиолетовый"
            else:
                return "Синий"
        elif r > 200 and g > 200:
            return "Желтый"
        elif r > 150 and g > 150 and b > 150:
            return "Бежевый"
        else:
            return "Разноцветный"
            
    except Exception:
        return None

def analyze_style_from_image(image) -> dict:
    """
    Анализирует изображение для определения стиля товара.
    Базовая эвристика на основе цветов и контраста.
    """
    params = {}
    
    try:
        img_array = np.array(image)
        
        # Вычисляем среднюю яркость
        brightness = img_array.mean()
        
        # Вычисляем контраст (стандартное отклонение)
        contrast = img_array.std()
        
        # Определяем стиль на основе характеристик
        if brightness > 200:
            params['Стиль'] = 'Светлый'
        elif brightness < 80:
            params['Стиль'] = 'Темный'
        else:
            params['Стиль'] = 'Средний'
        
        # Определяем тип по контрасту
        if contrast > 60:
            if 'Стиль' in params:
                params['Стиль'] += ', Контрастный'
        
    except Exception:
        pass
    
    return params

def get_product_image_urls_from_wb_api(sku: str, max_images: int = 3) -> list:
    """
    Получает URL изображений товара из API Wildberries.
    
    Args:
        sku: Артикул товара
        max_images: Максимальное количество изображений
        
    Returns:
        Список URL изображений
    """
    image_urls = []
    
    try:
        # API Wildberries для получения данных о товаре
        api_url = f"https://card.wb.ru/cards/v1/detail?appType=1&curr=rub&dest=-1257786&spp=30&nm={sku}"
        
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
            'Accept': 'application/json',
            'Accept-Language': 'ru-RU,ru;q=0.9,en-US;q=0.8,en;q=0.7',
        }
        
        response = requests.get(api_url, headers=headers, timeout=10)
        
        if response.status_code == 200:
            data = response.json()
            
            if 'data' in data and 'products' in data['data'] and len(data['data']['products']) > 0:
                product = data['data']['products'][0]
                
                # Получаем изображения из поля 'pics'
                if 'pics' in product and product['pics']:
                    # Определяем параметры для формирования URL
                    # Формат: https://basket-{basket_id}.wbbasket.ru/vol{vol_id}/part{part_id}/{sku}/images/big/{pic_id}.jpg
                    sku_int = int(sku) if sku.isdigit() else 0
                    
                    if sku_int > 0:
                        # Определяем корзину (basket_id) на основе диапазонов артикулов
                        # WB использует разные корзины для разных диапазонов артикулов
                        basket_id = 1  # По умолчанию
                        if sku_int >= 14300000 and sku_int < 28800000:
                            basket_id = 2
                        elif sku_int >= 28800000 and sku_int < 43500000:
                            basket_id = 3
                        elif sku_int >= 43500000 and sku_int < 58000000:
                            basket_id = 4
                        elif sku_int >= 58000000 and sku_int < 72500000:
                            basket_id = 5
                        elif sku_int >= 72500000 and sku_int < 87000000:
                            basket_id = 6
                        elif sku_int >= 87000000 and sku_int < 101500000:
                            basket_id = 7
                        elif sku_int >= 101500000 and sku_int < 116000000:
                            basket_id = 8
                        elif sku_int >= 116000000 and sku_int < 130500000:
                            basket_id = 9
                        elif sku_int >= 130500000 and sku_int < 145000000:
                            basket_id = 10
                        elif sku_int >= 145000000 and sku_int < 159500000:
                            basket_id = 11
                        elif sku_int >= 159500000 and sku_int < 174000000:
                            basket_id = 12
                        elif sku_int >= 174000000 and sku_int < 188500000:
                            basket_id = 13
                        elif sku_int >= 188500000 and sku_int < 203000000:
                            basket_id = 14
                        elif sku_int >= 203000000 and sku_int < 217500000:
                            basket_id = 15
                        elif sku_int >= 217500000 and sku_int < 232000000:
                            basket_id = 16
                        elif sku_int >= 232000000:
                            basket_id = 17
                        
                        # Определяем vol и part из артикула
                        # vol = первые 3 цифры (или первые 2 для старых артикулов)
                        # part = первые 4 цифры
                        vol_id = sku_int // 100000 if sku_int >= 100000 else sku_int // 10000
                        part_id = sku_int // 1000 if sku_int >= 1000 else sku_int // 100
                        
                        # Формируем URL для каждого изображения
                        for pic_id in product['pics'][:max_images]:
                            image_url = f"https://basket-{basket_id:02d}.wbbasket.ru/vol{vol_id}/part{part_id}/{sku}/images/big/{pic_id}.jpg"
                            image_urls.append(image_url)
                
                # Альтернативный способ: если есть поле 'photos' или 'images'
                elif 'photos' in product and product['photos'] and not image_urls:
                    for photo_url in product['photos'][:max_images]:
                        if isinstance(photo_url, str) and photo_url.startswith('http'):
                            image_urls.append(photo_url)
        
    except Exception as e:
        # В случае ошибки возвращаем пустой список
        pass
    
    return image_urls[:max_images]

def get_product_name_from_wb(sku: str) -> str:
    """
    Получает название товара из API Wildberries.
    
    Args:
        sku: Артикул товара
        
    Returns:
        Название товара или пустая строка
    """
    try:
        api_url = f"https://card.wb.ru/cards/v1/detail?appType=1&curr=rub&dest=-1257786&spp=30&nm={sku}"
        
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
            'Accept': 'application/json',
            'Accept-Language': 'ru-RU,ru;q=0.9,en-US;q=0.8,en;q=0.7',
        }
        
        response = requests.get(api_url, headers=headers, timeout=10)
        
        if response.status_code == 200:
            data = response.json()
            
            if 'data' in data and 'products' in data['data'] and len(data['data']['products']) > 0:
                product = data['data']['products'][0]
                if 'name' in product:
                    return product['name']
    except Exception as e:
        pass
    
    return ""

def get_screenshotapi_image_urls(sku: str, token: str = None, max_images: int = 5) -> list:
    """
    Генерирует URL-ы для получения изображений товара через screenshotapi.net.
    
    Args:
        sku: Артикул товара
        token: API токен screenshotapi.net (если None, берется из session_state или secrets)
        max_images: Максимальное количество изображений (по умолчанию 5)
        
    Returns:
        Список URL-ов изображений
    """
    image_urls = []
    
    try:
        # Получаем токен если не передан
        if not token:
            try:
                token = st.secrets.get('screenshotapi_token', '')
            except:
                pass
            
            if not token:
                token = st.session_state.get('screenshotapi_token', '')
        
        if not token:
            # Если токена нет, возвращаем пустой список
            return []
        
        # Формируем URL страницы товара
        product_url = build_wb_product_url(sku, host="https://global.wildberries.ru")
        encoded_url = _urlparse.quote(product_url, safe="")
        
        # Базовый URL для screenshotapi.net
        base_url = "https://shot.screenshotapi.net/v3/screenshot"
        
        # Первое изображение - основное (из главного слайдера)
        # Увеличиваем timeout и добавляем больше времени на загрузку
        first_image_url = (
            f"{base_url}?token={token}"
            f"&fresh=true"
            f"&url={encoded_url}"
            f"&width=1370"
            f"&height=544"
            f"&output=image"
            f"&file_type=png"
            f"&wait_for_event=load"
            f"&wait_for=2000"
            f"&selector_to_click=.mainSlider--JxK8z%20img"
        )
        image_urls.append(first_image_url)
        
        # Остальные 4 изображения - из вертикального слайдера
        for i in range(2, min(6, max_images + 1)):  # nth-child от 2 до 5
            additional_image_url = (
                f"{base_url}?token={token}"
                f"&fresh=true"
                f"&url={encoded_url}"
                f"&output=image"
                f"&file_type=png"
                f"&wait_for_event=load"
                f"&selector_to_click=.swiper-vertical%20.swiper-slide%3Anth-child({i})%20img"
            )
            image_urls.append(additional_image_url)
        
    except Exception as e:
        # В случае ошибки возвращаем пустой список
        pass
    
    return image_urls[:max_images]

def has_downloaded_images(sku: str, min_images: int = 1) -> bool:
    """
    Проверяет, есть ли у товара скачанные изображения через screenshotapi.
    
    Args:
        sku: Артикул товара
        min_images: Минимальное количество изображений для считания, что фото есть (по умолчанию 1)
        
    Returns:
        True, если есть хотя бы min_images скачанных изображений
    """
    try:
        sku_clean = str(sku).replace(".0", "")
        cached_count = 0
        
        # Проверяем наличие изображений в кеше (формат: {sku}_screenshotapi_{index})
        for idx in range(5):  # Проверяем до 5 изображений
            cache_key = f"{sku_clean}_screenshotapi_{idx}"
            cached_path = get_cached_image_path(cache_key)
            if cached_path and os.path.exists(cached_path):
                cached_count += 1
                if cached_count >= min_images:
                    return True
        
        return False
    except Exception:
        return False

def get_screenshotapi_images_with_status(sku: str, token: str = None, max_images: int = 5) -> list:
    """
    Получает изображения товара через screenshotapi.net с информацией о статусе (в кеше или скачивается).
    
    Args:
        sku: Артикул товара
        token: API токен screenshotapi.net (если None, берется из session_state или secrets)
        max_images: Максимальное количество изображений (по умолчанию 5)
        
    Returns:
        Список словарей с информацией о каждом изображении:
        [{"index": 1, "path": "...", "status": "cached"|"downloading"|"downloaded"|"error", "url": "..."}, ...]
    """
    images_info = []
    
    try:
        # Получаем URL изображений
        image_urls = get_screenshotapi_image_urls(sku, token=token, max_images=max_images)
        
        if not image_urls:
            return []
        
        # Проверяем каждое изображение
        for idx, img_url in enumerate(image_urls):
            cache_key = f"{sku}_screenshotapi_{idx}"
            
            # Проверяем, есть ли изображение в кеше
            cached_path = get_cached_image_path(cache_key)
            if cached_path and os.path.exists(cached_path):
                images_info.append({
                    "index": idx + 1,
                    "path": cached_path,
                    "status": "cached",
                    "url": img_url
                })
            else:
                # Изображение нужно скачать
                images_info.append({
                    "index": idx + 1,
                    "path": None,
                    "status": "downloading",
                    "url": img_url,
                    "cache_key": cache_key
                })
        
    except Exception as e:
        pass
    
    return images_info

def analyze_combination_products_with_ai(combo_products_df, combination_key: str, category: str = "Рашрашд мужской (компрессионная одежда)") -> dict:
    """
    Анализирует товары в комбинации с помощью ИИ для получения рекомендаций по улучшению.
    
    Args:
        combo_products_df: DataFrame с товарами комбинации
        combination_key: Ключ комбинации (параметры)
        category: Категория товаров (по умолчанию "Рашрашд мужской (компрессионная одежда)")
        
    Returns:
        Словарь с анализом и рекомендациями
    """
    if not OPENAI_AVAILABLE:
        return {
            "error": "OpenAI не доступен. Установите библиотеку openai для использования этой функции."
        }
    
    try:
        # Получаем API ключ
        api_key = st.session_state.get('openai_api_key', '')
        if not api_key:
            try:
                api_key = st.secrets.get('openai_api_key', '')
            except:
                pass
        
        if not api_key:
            return {
                "error": "API ключ OpenAI не найден. Укажите ключ в настройках."
            }
        
        client = openai.OpenAI(api_key=api_key)
        
        # Собираем изображения товаров (максимум 5 товаров для анализа)
        image_paths = []
        product_info = []
        
        for idx, row in combo_products_df.head(5).iterrows():
            sku = str(row.get("Артикул", "")).replace(".0", "")
            if not sku:
                continue
            
            # Получаем путь к изображению
            image_path = get_cached_image_path(sku)
            if image_path and os.path.exists(image_path):
                image_paths.append(image_path)
                
                # Собираем информацию о товаре
                product_data = {
                    "sku": sku,
                    "revenue": row.get("Выручка", 0),
                    "orders": row.get("Заказы", 0),
                    "price": row.get("Средняя цена", 0),
                    "position": row.get("Позиция в выдаче", 0)
                }
                product_info.append(product_data)
        
        if not image_paths:
            return {
                "error": "Не найдено изображений товаров для анализа. Убедитесь, что изображения скачаны."
            }
        
        # Подготавливаем изображения для отправки в API
        image_contents = []
        for img_path in image_paths:
            try:
                with open(img_path, 'rb') as image_file:
                    image_bytes = image_file.read()
                    image_data = base64.b64encode(image_bytes).decode('utf-8')
                    ext = os.path.splitext(img_path)[1].lower()
                    mime_type = {
                        '.jpg': 'image/jpeg',
                        '.jpeg': 'image/jpeg',
                        '.png': 'image/png',
                        '.webp': 'image/webp'
                    }.get(ext, 'image/jpeg')
                    image_contents.append({
                        "type": "image_url",
                        "image_url": {
                            "url": f"data:{mime_type};base64,{image_data}"
                        }
                    })
            except Exception as e:
                continue
        
        if not image_contents:
            return {
                "error": "Не удалось подготовить изображения для анализа."
            }
        
        # Формируем промпт для анализа
        avg_revenue = combo_products_df["Выручка"].mean() if "Выручка" in combo_products_df.columns else 0
        avg_position = combo_products_df["Позиция в выдаче"].mean() if "Позиция в выдаче" in combo_products_df.columns else 0
        total_products = len(combo_products_df)
        
        prompt = f"""Ты эксперт по анализу товаров и маркетингу. Проанализируй товары в комбинации параметров и дай рекомендации по улучшению.

КОНТЕКСТ:
- Категория: {category}
- Комбинация параметров: {combination_key}
- Количество товаров в комбинации: {total_products}
- Средняя выручка на товар: {avg_revenue:,.0f} руб
- Средняя позиция в выдаче: {avg_position:.1f}

ЗАДАЧА:
1. Проанализируй изображения товаров в этой комбинации
2. Определи общие характеристики и паттерны
3. Выяви слабые места и возможности для улучшения
4. Дай конкретные рекомендации по изменению товаров для получения конкурентного преимущества

ФОРМАТ ОТВЕТА (JSON):
{{
    "common_characteristics": "Общие характеристики товаров в комбинации",
    "strengths": ["Сильная сторона 1", "Сильная сторона 2", ...],
    "weaknesses": ["Слабая сторона 1", "Слабая сторона 2", ...],
    "recommendations": [
        {{
            "priority": "high|medium|low",
            "category": "Дизайн|Материал|Цвет|Принт|Детали",
            "recommendation": "Конкретная рекомендация по улучшению",
            "expected_impact": "Ожидаемый эффект от изменения"
        }},
        ...
    ],
    "competitive_advantages": ["Преимущество 1", "Преимущество 2", ...],
    "market_insights": "Анализ рыночной ситуации и трендов для этой категории"
}}

ВАЖНО:
- Будь конкретным и практичным в рекомендациях
- Учитывай специфику категории "{category}"
- Фокусируйся на изменениях, которые дадут реальное конкурентное преимущество
- Учитывай текущие показатели (выручка, позиция) при рекомендациях"""

        # Отправляем запрос в OpenAI
        messages = [
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": prompt}
                ] + image_contents
            }
        ]
        
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=messages,
            max_tokens=2000,
            temperature=0.7
        )
        
        response_text = response.choices[0].message.content
        
        # Парсим JSON из ответа
        import json
        json_start = response_text.find('{')
        json_end = response_text.rfind('}') + 1
        
        if json_start >= 0 and json_end > json_start:
            json_str = response_text[json_start:json_end]
            try:
                analysis_result = json.loads(json_str)
                analysis_result["raw_response"] = response_text
                return analysis_result
            except json.JSONDecodeError:
                return {
                    "error": "Не удалось распарсить ответ ИИ",
                    "raw_response": response_text
                }
        else:
            return {
                "error": "ИИ не вернул JSON формат",
                "raw_response": response_text
            }
            
    except openai.RateLimitError:
        return {
            "error": "Превышен лимит запросов к OpenAI API. Попробуйте позже."
        }
    except openai.APITimeoutError:
        return {
            "error": "Превышено время ожидания ответа от OpenAI API. Попробуйте позже."
        }
    except Exception as e:
        return {
            "error": f"Ошибка при анализе: {str(e)}"
        }

def get_category_cagr_analysis(category: str = "Рашрашд мужской (компрессионная одежда)") -> dict:
    """
    Возвращает информацию о CAGR (Compound Annual Growth Rate) для категории.
    
    Args:
        category: Категория товаров
        
    Returns:
        Словарь с информацией о CAGR и анализом рынка
    """
    # Данные о CAGR для компрессионной одежды (примерные значения на основе рыночных данных)
    cagr_data = {
        "Рашрашд мужской (компрессионная одежда)": {
            "cagr_5_years": 8.5,  # Примерный CAGR за 5 лет
            "market_size_2024": "2.5 млрд руб",  # Размер рынка в 2024
            "projected_growth_2025": 9.2,  # Прогнозируемый рост в 2025
            "key_drivers": [
                "Рост популярности фитнеса и здорового образа жизни",
                "Увеличение спроса на функциональную одежду",
                "Развитие e-commerce и онлайн-продаж",
                "Расширение ассортимента и улучшение качества продукции"
            ],
            "trends": [
                "Фокус на экологичных материалах",
                "Технологичные ткани с улучшенными свойствами",
                "Персонализация и кастомизация",
                "Интеграция умных технологий"
            ],
            "competition_level": "Высокий",
            "market_maturity": "Растущий рынок"
        }
    }
    
    # Если категория не найдена, возвращаем общие данные
    if category not in cagr_data:
        return {
            "cagr_5_years": 7.0,
            "market_size_2024": "N/A",
            "projected_growth_2025": 7.5,
            "key_drivers": ["Общие рыночные факторы"],
            "trends": ["Общие тренды рынка"],
            "competition_level": "Средний",
            "market_maturity": "Стабильный"
        }
    
    return cagr_data[category]

def read_wgsn_files() -> dict:
    """
    Читает все файлы из папки WGSN и возвращает их содержимое.
    
    Returns:
        Словарь с содержимым файлов: {"filename": "content", ...}
    """
    wgsn_data = {}
    
    if not DOCX_AVAILABLE:
        return {
            "error": "Библиотека python-docx не установлена. Установите её командой: pip install python-docx"
        }
    
    try:
        wgsn_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "WGSN")
        
        if not os.path.exists(wgsn_dir):
            return {
                "error": f"Папка WGSN не найдена: {wgsn_dir}"
            }
        
        # Читаем все .docx файлы из папки WGSN
        for filename in os.listdir(wgsn_dir):
            if filename.endswith('.docx'):
                file_path = os.path.join(wgsn_dir, filename)
                try:
                    doc = Document(file_path)
                    # Извлекаем весь текст из документа
                    full_text = []
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            full_text.append(paragraph.text.strip())
                    
                    # Также извлекаем текст из таблиц
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = []
                            for cell in row.cells:
                                if cell.text.strip():
                                    row_text.append(cell.text.strip())
                            if row_text:
                                full_text.append(" | ".join(row_text))
                    
                    wgsn_data[filename] = "\n".join(full_text)
                except Exception as e:
                    wgsn_data[filename] = f"Ошибка при чтении файла: {str(e)}"
        
        if not wgsn_data:
            return {
                "error": "Не найдено .docx файлов в папке WGSN"
            }
        
        return wgsn_data
        
    except Exception as e:
        return {
            "error": f"Ошибка при чтении файлов WGSN: {str(e)}"
        }

def analyze_wgsn_trends_with_ai(wgsn_content: dict, category: str = "Рашрашд мужской (компрессионная одежда)", combination_key: str = "") -> dict:
    """
    Анализирует содержимое файлов WGSN с помощью ИИ для получения трендов и рекомендаций.
    
    Args:
        wgsn_content: Словарь с содержимым файлов WGSN
        category: Категория товаров
        combination_key: Ключ комбинации параметров
        
    Returns:
        Словарь с анализом трендов WGSN
    """
    if not OPENAI_AVAILABLE:
        return {
            "error": "OpenAI не доступен. Установите библиотеку openai для использования этой функции."
        }
    
    if "error" in wgsn_content:
        return wgsn_content
    
    try:
        # Получаем API ключ
        api_key = st.session_state.get('openai_api_key', '')
        if not api_key:
            try:
                api_key = st.secrets.get('openai_api_key', '')
            except:
                pass
        
        if not api_key:
            return {
                "error": "API ключ OpenAI не найден. Укажите ключ в настройках."
            }
        
        client = openai.OpenAI(api_key=api_key)
        
        # Объединяем содержимое всех файлов WGSN
        combined_content = ""
        file_names = []
        for filename, content in wgsn_content.items():
            # Увеличиваем размер для более детального анализа
            combined_content += f"\n\n=== Файл: {filename} ===\n{content[:10000]}"
            file_names.append(filename)
        
        # Если содержимое слишком большое, берем больше контента из каждого файла
        if len(combined_content) > 50000:
            combined_content = ""
            for filename, content in wgsn_content.items():
                combined_content += f"\n\n=== Файл: {filename} ===\n{content[:15000]}"
        
        # Парсим параметры комбинации для детального анализа
        combination_params = {}
        if combination_key:
            parts = [p.strip() for p in combination_key.split("|")]
            for part in parts:
                if ":" in part:
                    param_name, param_value = part.split(":", 1)
                    combination_params[param_name.strip()] = param_value.strip()
        
        # Определяем специфику категории
        category_specifics = ""
        if "рашград" in category.lower() or "компрессион" in category.lower() or "футболка" in category.lower():
            category_specifics = """
СПЕЦИФИКА КАТЕГОРИИ:
- Это мужская компрессионная/спортивная одежда (рашград)
- Товары предназначены для спорта, фитнеса, активного образа жизни
- Важны: функциональность, комфорт, технологичность материалов
- Целевая аудитория: мужчины, занимающиеся спортом, фитнесом, ведущие активный образ жизни
- Ключевые требования: влагоотведение, воздухопроницаемость, эластичность, долговечность
- Конкурентная среда: высокая конкуренция, важно выделяться через дизайн, технологии, качество
"""
        
        # Формируем детальный промпт для анализа
        prompt = f"""Ты эксперт по анализу модных трендов WGSN (Worth Global Style Network) и специалист по спортивной/компрессионной одежде. 
Проведи глубокий и детальный анализ файлов WGSN и дай конкретные, развернутые рекомендации для товаров в указанной комбинации параметров.

КОНТЕКСТ АНАЛИЗА:
- Категория товаров: {category}
{category_specifics}
- Текущая комбинация параметров: {combination_key}
- Детали комбинации: {json.dumps(combination_params, ensure_ascii=False, indent=2) if combination_params else "Параметры не указаны"}
- Файлы WGSN для анализа: {', '.join(file_names)}

СОДЕРЖИМОЕ ФАЙЛОВ WGSN:
{combined_content}

ДЕТАЛЬНАЯ ЗАДАЧА АНАЛИЗА:

1. ГЛУБОКИЙ АНАЛИЗ ТРЕНДОВ WGSN:
   - Выдели ВСЕ релевантные тренды из файлов WGSN для категории "{category}"
   - Определи приоритетность трендов (что актуально сейчас, что будет актуально в ближайшем будущем)
   - Найди тренды, которые напрямую связаны с компрессионной/спортивной одеждой, футболками, мужской спортивной одеждой
   - Выяви тренды по цветам, материалам, технологиям, дизайну, которые применимы к данной категории

2. КОНКРЕТНОЕ СОПОСТАВЛЕНИЕ С КОМБИНАЦИЕЙ:
   - Проанализируй КАЖДЫЙ параметр из комбинации: {combination_key}
   - Для каждого параметра найди соответствующие тренды из WGSN
   - Определи, какие параметры соответствуют актуальным трендам, а какие устарели
   - Выяви параметры, которые можно улучшить на основе трендов WGSN

3. РАЗВЕРНУТЫЕ РЕКОМЕНДАЦИИ:
   - Дай КОНКРЕТНЫЕ рекомендации по изменению КАЖДОГО параметра комбинации на основе трендов WGSN
   - Укажи, какие конкретные значения параметров будут актуальны согласно WGSN
   - Объясни, ПОЧЕМУ эти изменения дадут конкурентное преимущество
   - Предложи новые комбинации параметров, которые будут соответствовать трендам WGSN

4. АНАЛИЗ ПО КАТЕГОРИЯМ:
   - ЦВЕТА: какие цвета актуальны для спортивной/компрессионной одежды по WGSN, какие устарели
   - МАТЕРИАЛЫ: какие материалы и технологии рекомендуются WGSN для данной категории
   - ДИЗАЙН: какие дизайнерские решения, принты, детали актуальны
   - ТЕХНОЛОГИИ: какие технологические решения (влагозащита, терморегуляция и т.д.) актуальны

5. КОНКУРЕНТНОЕ ПРЕИМУЩЕСТВО:
   - Определи, какие тренды WGSN еще не широко используются в категории "{category}"
   - Предложи, как можно выделиться на рынке, применяя эти тренды
   - Укажи конкретные шаги для получения конкурентного преимущества

ФОРМАТ ОТВЕТА (JSON):
{{
    "wgsn_trends_summary": "Развернутое резюме ключевых трендов из WGSN, применимых к категории {category}. Минимум 200 слов.",
    "category_analysis": {{
        "current_trends": "Актуальные тренды для {category} на основе WGSN",
        "emerging_trends": "Появляющиеся тренды, которые станут актуальными",
        "declining_trends": "Устаревающие тренды, от которых стоит отойти"
    }},
    "combination_parameter_analysis": [
        {{
            "parameter": "Название параметра из комбинации",
            "current_value": "Текущее значение",
            "wgsn_trend_status": "Соответствует|Частично соответствует|Не соответствует трендам WGSN",
            "wgsn_recommendation": "Конкретная рекомендация WGSN для этого параметра",
            "suggested_values": ["Рекомендуемое значение 1", "Рекомендуемое значение 2", ...],
            "rationale": "Обоснование рекомендации на основе WGSN"
        }},
        ...
    ],
    "relevant_trends": [
        {{
            "trend_name": "Название тренда",
            "source_file": "Имя файла WGSN",
            "trend_type": "Цвет|Материал|Дизайн|Технология|Стиль|Функциональность",
            "description": "Детальное описание тренда (минимум 100 слов)",
            "relevance_to_category": "Почему тренд критически важен для {category}",
            "application_to_combination": "Конкретное применение тренда к комбинации {combination_key}",
            "implementation_steps": ["Шаг 1", "Шаг 2", "Шаг 3"],
            "expected_market_impact": "Ожидаемое влияние на рынок и продажи"
        }},
        ...
    ],
    "trend_recommendations": [
        {{
            "priority": "high|medium|low",
            "trend_category": "Цвет|Дизайн|Материал|Технология|Стиль|Функциональность",
            "recommendation": "Детальная рекомендация на основе трендов WGSN (минимум 150 слов)",
            "specific_changes": "Конкретные изменения в параметрах комбинации",
            "expected_impact": "Детальное описание ожидаемого эффекта от применения тренда",
            "wgsn_source": "Конкретная цитата или ссылка на источник из файлов WGSN",
            "timeline": "Когда лучше внедрить (сейчас|ближайшие 3 месяца|6 месяцев|год)"
        }},
        ...
    ],
    "new_combination_suggestions": [
        {{
            "combination": "Новая комбинация параметров на основе трендов WGSN",
            "rationale": "Обоснование новой комбинации",
            "wgsn_trends_used": ["Тренд 1", "Тренд 2", ...],
            "competitive_advantage": "Какое конкурентное преимущество даст эта комбинация"
        }},
        ...
    ],
    "future_trends": [
        {{
            "trend_name": "Название будущего тренда",
            "timeframe": "Когда станет актуальным",
            "description": "Описание тренда",
            "preparation_steps": ["Шаг подготовки 1", "Шаг подготовки 2", ...]
        }},
        ...
    ],
    "competitive_advantages_from_trends": [
        {{
            "advantage": "Конкретное конкурентное преимущество",
            "wgsn_trend_basis": "На каком тренде WGSN основано",
            "implementation": "Как реализовать",
            "market_positioning": "Как позиционировать на рынке"
        }},
        ...
    ],
    "market_insights": "Развернутый анализ рыночной ситуации на основе трендов WGSN для категории {category}. Минимум 300 слов.",
    "action_plan": {{
        "immediate_actions": ["Действие 1", "Действие 2", ...],
        "short_term_actions": ["Действие 1 (3 месяца)", "Действие 2 (3 месяца)", ...],
        "long_term_actions": ["Действие 1 (6-12 месяцев)", "Действие 2 (6-12 месяцев)", ...]
    }}
}}

КРИТИЧЕСКИ ВАЖНО:
- Будь МАКСИМАЛЬНО конкретным и детальным в рекомендациях
- Ссылайся на КОНКРЕТНЫЕ тренды из файлов WGSN с указанием источника
- Анализируй КАЖДЫЙ параметр комбинации отдельно
- Давай РАЗВЕРНУТЫЕ объяснения (минимум 100-150 слов для каждого важного пункта)
- Учитывай специфику категории "{category}" - это спортивная/компрессионная одежда для мужчин
- Фокусируйся на трендах, которые дадут РЕАЛЬНОЕ конкурентное преимущество
- Предлагай КОНКРЕТНЫЕ значения параметров, а не общие рекомендации
- Объясняй ПОЧЕМУ каждое изменение важно с точки зрения трендов WGSN"""

        # Отправляем запрос в OpenAI
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            max_tokens=4000,
            temperature=0.7
        )
        
        response_text = response.choices[0].message.content
        
        # Парсим JSON из ответа
        import json
        json_start = response_text.find('{')
        json_end = response_text.rfind('}') + 1
        
        if json_start >= 0 and json_end > json_start:
            json_str = response_text[json_start:json_end]
            try:
                analysis_result = json.loads(json_str)
                analysis_result["raw_response"] = response_text
                analysis_result["source_files"] = file_names
                return analysis_result
            except json.JSONDecodeError:
                return {
                    "error": "Не удалось распарсить ответ ИИ",
                    "raw_response": response_text,
                    "source_files": file_names
                }
        else:
            return {
                "error": "ИИ не вернул JSON формат",
                "raw_response": response_text,
                "source_files": file_names
            }
            
    except openai.RateLimitError:
        return {
            "error": "Превышен лимит запросов к OpenAI API. Попробуйте позже."
        }
    except openai.APITimeoutError:
        return {
            "error": "Превышено время ожидания ответа от OpenAI API. Попробуйте позже."
        }
    except Exception as e:
        return {
            "error": f"Ошибка при анализе WGSN: {str(e)}"
        }

def get_product_params_from_images(url: str, api_key: str = None, max_images: int = 5, selected_params: list = None) -> dict:
    """
    Определяет параметры товара по фотографиям через нейросеть.
    Использует screenshotapi.net для получения изображений товара.
    Оптимизировано для использования кешированных изображений - не скачивает повторно уже имеющиеся.
    
    Args:
        url: URL товара Wildberries
        api_key: API ключ OpenAI (если None, берется из session_state)
        max_images: Максимальное количество изображений для анализа (по умолчанию 5)
        selected_params: Список параметров для анализа (если None, анализируются все)
        
    Returns:
        Словарь с определенными параметрами
    """
    sku = extract_sku_from_url(url)
    if not sku:
        return {}
    
    all_params = {}
    
    try:
        import time
        
        # Получаем API ключ если не передан
        if not api_key:
            api_key = st.session_state.get('openai_api_key', '')
        
        if not api_key:
            return {}
        
        # Сначала проверяем, какие изображения уже есть в кеше
        sku_clean = str(sku).replace(".0", "")
        cached_images = []
        images_to_download = []  # Список индексов изображений, которые нужно скачать
        
        # Проверяем наличие изображений в кеше
        for idx in range(max_images):
            cache_key = f"{sku_clean}_screenshotapi_{idx}"
            cached_path = get_cached_image_path(cache_key)
            
            if cached_path and os.path.exists(cached_path):
                # Изображение уже в кеше - используем его
                cached_images.append((idx, cached_path))
            else:
                # Изображение нужно скачать
                images_to_download.append(idx)
        
        # Если все изображения уже в кеше, используем их без дополнительных запросов
        if not images_to_download:
            # Все изображения в кеше - используем их
            image_paths = [path for _, path in sorted(cached_images)]
        else:
            # Нужно скачать некоторые изображения
            # Получаем URL изображений через screenshotapi.net только для недостающих
            image_urls = get_screenshotapi_image_urls(sku, max_images=max_images)
            
            if not image_urls:
                # Если не удалось получить URL, используем только кешированные
                image_paths = [path for _, path in sorted(cached_images)]
            else:
                # Скачиваем только недостающие изображения
                image_paths = [None] * max_images
                
                # Заполняем пути для уже кешированных изображений
                for idx, path in cached_images:
                    image_paths[idx] = path
                
                # Скачиваем недостающие изображения
                for idx in images_to_download:
                    if idx < len(image_urls):
                        cache_key = f"{sku_clean}_screenshotapi_{idx}"
                        img_url = image_urls[idx]
                        
                        # Скачиваем изображение (ensure_image_cached проверит кеш внутри себя)
                        cached_path = ensure_image_cached(cache_key, img_url, "PNG", timeout=30)
                        if cached_path and os.path.exists(cached_path):
                            image_paths[idx] = cached_path
                        
                        # Задержка между запросами только для новых скачиваний
                        if idx < len(images_to_download) - 1:  # Не делаем задержку после последнего
                            time.sleep(0.5)  # 500ms задержка между запросами
                
                # Убираем None значения
                image_paths = [path for path in image_paths if path is not None]
        
        if not image_paths:
            return {}
        
        # Анализируем каждое изображение через нейросеть
        image_params_list = []
        for img_path in image_paths:
            img_params = analyze_image_with_ai(img_path, api_key, selected_params=selected_params)
            if img_params:
                image_params_list.append(img_params)
        
        # Объединяем результаты всех изображений
        if image_params_list:
            # Для каждого параметра берем наиболее частый
            param_keys = set()
            for img_params in image_params_list:
                param_keys.update(img_params.keys())
            
            for key in param_keys:
                values = [p.get(key) for p in image_params_list if p.get(key)]
                if values:
                    value_counts = Counter(values)
                    all_params[key] = value_counts.most_common(1)[0][0]
        
    except Exception as e:
        pass
    
    return all_params

def determine_completeness(sku: str) -> str:
    """
    Определяет, является ли товар комплектом или одной вещью.
    Анализирует характеристики и описание товара через API Wildberries.
    
    Args:
        sku: Артикул товара
        
    Returns:
        "Комплект" или "Один"
    """
    try:
        api_url = f"https://card.wb.ru/cards/v1/detail?appType=1&curr=rub&dest=-1257786&spp=30&nm={sku}"
        
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
            'Accept': 'application/json',
        }
        
        response = requests.get(api_url, headers=headers, timeout=10)
        
        if response.status_code == 200:
            data = response.json()
            
            if 'data' in data and 'products' in data['data'] and len(data['data']['products']) > 0:
                product = data['data']['products'][0]
                
                # Собираем текст для анализа
                text_to_analyze = []
                
                # Получаем название товара
                if 'name' in product:
                    text_to_analyze.append(product['name'].lower())
                
                # Получаем описание
                if 'description' in product:
                    text_to_analyze.append(product['description'].lower())
                
                # Получаем характеристики
                if 'characteristics' in product:
                    for char in product['characteristics']:
                        if 'name' in char and 'value' in char:
                            text_to_analyze.append(f"{char['name']} {char['value']}".lower())
                
                # Объединяем весь текст
                full_text = " ".join(text_to_analyze)
                
                # Ключевые слова, указывающие на комплект
                kit_keywords = [
                    'комплект', 'набор', 'set', 'комплектация', 'в комплекте',
                    '2 шт', '3 шт', '4 шт', '5 шт', '6 шт', '7 шт', '8 шт', '9 шт', '10 шт',
                    'две', 'три', 'четыре', 'пять', 'шесть', 'семь', 'восемь', 'девять', 'десять',
                    'пара', 'пары', 'шт.', 'штук', 'штуки',
                    'включает', 'состоит из', 'содержит', 'в наборе',
                    'комплект из', 'набор из', 'комплектация из'
                ]
                
                # Проверяем наличие ключевых слов
                for keyword in kit_keywords:
                    if keyword in full_text:
                        return "Комплект"
                
                # Дополнительная проверка: ищем числа перед словами "шт", "штук", "предмет"
                import re
                quantity_patterns = [
                    r'\d+\s*(шт|штук|предмет|вещь|изделие)',
                    r'(две|три|четыре|пять|шесть|семь|восемь|девять|десять)\s*(шт|штук|предмет|вещь)',
                ]
                
                for pattern in quantity_patterns:
                    if re.search(pattern, full_text):
                        return "Комплект"
                
    except Exception as e:
        # В случае ошибки возвращаем "Один" по умолчанию
        pass
    
    # По умолчанию считаем, что это одна вещь
    return "Один"

def get_product_params_from_url(url: str) -> dict:
    """
    Пытается получить параметры товара по ссылке на Wildberries.
    
    ВАЖНО: Wildberries не предоставляет публичного API для получения параметров товара.
    Эта функция может работать только если:
    1. Используется официальный API Wildberries (требует токен)
    2. Используется парсинг HTML (может нарушать ToS Wildberries)
    
    Args:
        url: URL товара Wildberries
        
    Returns:
        Словарь с параметрами товара или пустой словарь
    """
    sku = extract_sku_from_url(url)
    if not sku:
        return {}
    
    params = {}
    
    try:
        # Попытка получить данные через API Wildberries (если доступен)
        # Для этого нужен токен API, который обычно хранится в secrets.toml
        api_url = f"https://card.wb.ru/cards/v1/detail?appType=1&curr=rub&dest=-1257786&spp=30&nm={sku}"
        
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
            'Accept': 'application/json',
        }
        
        response = requests.get(api_url, headers=headers, timeout=10)
        
        if response.status_code == 200:
            data = response.json()
            
            # Извлекаем параметры из ответа API
            if 'data' in data and 'products' in data['data'] and len(data['data']['products']) > 0:
                product = data['data']['products'][0]
                
                # Извлекаем характеристики товара
                if 'characteristics' in product:
                    for char in product['characteristics']:
                        if 'name' in char and 'value' in char:
                            param_name = char['name'].strip()
                            param_value = char['value'].strip()
                            if param_name and param_value:
                                params[param_name] = param_value
                
                # Извлекаем дополнительные данные
                if 'brand' in product:
                    params['Бренд'] = product['brand']
                if 'name' in product:
                    params['Название'] = product['name']
                if 'colors' in product and product['colors']:
                    params['Цвет'] = ', '.join([c.get('name', '') for c in product['colors'] if c.get('name')])
                
    except Exception as e:
        # Если API недоступен или произошла ошибка, возвращаем пустой словарь
        pass
    
    return params

def build_screenshot_url(page_url: str, key: str,
                         w: int = 400, h: int = 600,
                         fmt: str = "JPEG", profile: str = "D4",
                         base: str = "https://api.s-shot.ru"):
    q = _urlparse.quote(page_url, safe="")
    return f"{base.rstrip('/')}/{int(w)}x{int(h)}/{fmt}/{key}/{profile}/?{q}"
def screenshot_for_article(nm, conf):
    if not conf.get("key"):
        return ""
    page = build_wb_product_url(nm, conf.get("wb_host", "https://global.wildberries.ru"))
    return build_screenshot_url(
        page, conf.get("key", ""), conf.get("w", 400), conf.get("h", 600),
        conf.get("fmt", "JPEG"), conf.get("profile", "D4"), conf.get("base", "https://api.s-shot.ru")
    )

@st.cache_data(show_spinner=False)
def read_table(file_bytes: bytes, filename: str):
    try:
        if filename.lower().endswith((".xlsx", ".xls")):
            df_raw = pd.read_excel(BytesIO(file_bytes), sheet_name=0, header=None)
        else:
            df_raw = pd.read_csv(BytesIO(file_bytes), header=None, sep=None, engine="python")
    except Exception as e:
        st.error(f"Ошибка чтения файла: {e}")
        return None, None, {}
    key_candidates = ["Артикул", "Выручка", "Заказы", "Название"]
    header_row = None
    for i in range(min(30, len(df_raw))):
        vals = df_raw.iloc[i].astype(str).str.strip().tolist()
        if any(k in vals for k in key_candidates):
            header_row = i
            break
    if header_row is None:
        header_row = 0
    if filename.lower().endswith((".xlsx", ".xls")):
        df = pd.read_excel(BytesIO(file_bytes), sheet_name=0, header=header_row)
    else:
        df = pd.read_csv(BytesIO(file_bytes), header=header_row, sep=None, engine="python")
    df = df.loc[:, ~df.columns.astype(str).str.startswith("Unnamed")]
    df = df.loc[:, df.columns.notna()]
    df.columns = [str(c).strip() for c in df.columns]
    rename_map = {
        "Средняя цена без СПП": "Средняя цена",
        "Средняя цена без СПП, ₽": "Средняя цена",
        "Цена": "Средняя цена",
        "Выручка, ₽": "Выручка",
        "Orders": "Заказы",
        "Brand": "Бренд",
        "Supplier": "Поставщик",
        "Subject": "Предмет",
        "Creation date": "Дата создания",
        "Дата": "Дата создания",
        "Позиция": "Позиция в выдаче",
        "CPM": "Стоимость за 1000 показов",
        "Упущенная выручка, ₽": "Упущенная выручка",
    }
    df = df.rename(columns={k: v for k, v in rename_map.items() if k in df.columns})
    num_cols = ["Выручка","Заказы","Средняя цена","Упущенная выручка",
                "Позиция в выдаче","Стоимость за 1000 показов","Буст на позицию","Буст с позиции"]
    for c in num_cols:
        if c in df.columns:
            df[c] = pd.to_numeric(
                df[c].astype(str).str.replace(r"[^\d,.-]", "", regex=True).str.replace(",", ".", regex=False),
                errors="coerce",
            )
    if "Дата создания" in df.columns:
        df["Дата создания"] = pd.to_datetime(df["Дата создания"], errors="coerce")
    if "Тип рекламы" in df.columns:
        df["Тип рекламы"] = df["Тип рекламы"].replace({"b": "Поиск", "c": "Автомат"})
    if ("Буст на позицию" in df.columns) and ("Буст с позиции" in df.columns) and ("Дельта" not in df.columns):
        df["Дельта"] = df["Буст с позиции"] - df["Буст на позицию"]
    return df, df_raw, {"header_row": header_row, "columns": list(df.columns)}

def format_thousands(x, decimals=0):
    if x is None or (isinstance(x, float) and np.isnan(x)):
        return ""
    try:
        xf = float(x)
    except Exception:
        return str(x) if x is not None else ""
    if decimals == 0:
        return f"{int(round(xf))}"
    return f"{xf:.{decimals}f}"

def format_thousands_with_spaces(x, decimals=0):
    if x is None or (isinstance(x, float) and np.isnan(x)):
        return ""
    try:
        xf = float(x)
    except Exception:
        return str(x) if x is not None else ""
    if decimals == 0:
        return f"{int(round(xf)):,}".replace(",", " ")
    return f"{xf:,.{decimals}f}".replace(",", " ").replace(".", ",")
def fmt_rub(x, decimals=0):
    s = format_thousands(x, decimals=decimals)
    return (s + " ₽") if s != "" else ""
def fmt_units(x, unit="шт."):
    s = format_thousands(x, decimals=0)
    return (s + f" {unit}") if s != "" else ""

def fmt_rub_kpi(x, decimals=0):
    s = format_thousands_with_spaces(x, decimals=decimals)
    return (s + " ₽") if s != "" else ""
def fmt_units_kpi(x, unit="шт."):
    s = format_thousands_with_spaces(x, decimals=0)
    return (s + f" {unit}") if s != "" else ""
def fmt_date(d):
    if d is None or (isinstance(d, float) and np.isnan(d)):
        return ""
    try:
        dt = pd.to_datetime(d)
        # Русские названия месяцев
        months = {
            1: "января", 2: "февраля", 3: "марта", 4: "апреля", 5: "мая", 6: "июня",
            7: "июля", 8: "августа", 9: "сентября", 10: "октября", 11: "ноября", 12: "декабря"
        }
        return f"{dt.day} {months[dt.month]} {dt.year}"
    except Exception:
        return str(d) if d is not None else ""
def parse_thousands_input(s, default_val):
    if s is None or str(s).strip() == "":
        return default_val
    try:
        cleaned = (str(s).replace("\\xa0"," ").replace("\\u00a0"," ").replace(" ", " "))
        cleaned = cleaned.replace(" ", "").replace(",", "").strip()
        return int(float(cleaned))
    except Exception:
        return default_val
def sort_df(df, col, asc):
    if col not in df.columns:
        return df
    if pd.api.types.is_numeric_dtype(df[col]):
        return df.sort_values(by=col, ascending=asc, na_position="last", kind="mergesort")
    return df.sort_values(by=col, ascending=asc, na_position="last", kind="mergesort",
                          key=lambda s: s.astype(str).str.lower())

def get_param_schemas():
    if "param_schemas" not in st.session_state:
        st.session_state["param_schemas"] = {}
    return st.session_state["param_schemas"]
def get_param_values():
    """Получает param_values из session_state, исключая удаленные параметры"""
    if "param_values" not in st.session_state:
        st.session_state["param_values"] = {}
    
    param_values = st.session_state["param_values"]
    
    # Загружаем список удаленных параметров
    if "deleted_params" not in st.session_state:
        st.session_state["deleted_params"] = load_deleted_params_from_file()
    
    deleted_params = st.session_state.get("deleted_params", set())
    
    # Фильтруем удаленные параметры
    if deleted_params:
        param_values = {k: v for k, v in param_values.items() if k not in deleted_params}
    
    return param_values

def kpi_row(df):
    total_rev = float(df["Выручка"].sum()) if "Выручка" in df.columns else float('nan')
    total_orders = df["Заказы"].sum() if "Заказы" in df.columns else np.nan
    avg_check = (df["Выручка"].sum() / df["Заказы"].sum()) if ("Выручка" in df.columns and "Заказы" in df.columns and df["Заказы"].sum() > 0) else np.nan
    lost_rev = df["Упущенная выручка"].sum() if "Упущенная выручка" in df.columns else np.nan
    sku_count = (df["Артикул"].nunique() if "Артикул" in df.columns else len(df)) if len(df) > 0 else 0
    rev_per_sku = (total_rev / sku_count) if (isinstance(total_rev, (int,float,np.floating)) and not pd.isna(total_rev) and sku_count > 0) else np.nan
    k1, k2, k3, k4, k5, k6 = st.columns(6)
    k1.metric("Выручка (в выборке)", fmt_rub_kpi(total_rev))
    k2.metric("Заказы (в выборке)", fmt_units_kpi(total_orders, "шт."))
    k3.metric("Средний чек", fmt_rub_kpi(avg_check))
    k4.metric("Упущенная выручка", fmt_rub_kpi(lost_rev))
    k5.metric("Выручка / Кол-во товаров", fmt_rub_kpi(rev_per_sku))
    k6.metric("Количество артикулов", fmt_units_kpi(sku_count, "шт."))

# Параметры уже загружены выше при инициализации

# Автоматическая загрузка конфигурации иерархии параметров
if "hierarchy_params" not in st.session_state:
    load_hierarchy_config()

# Автоматическая загрузка OpenAI API ключа из secrets.toml при запуске
if "openai_api_key" not in st.session_state:
    try:
        default_key = st.secrets.get('openai_api_key', '')
        if default_key:
            st.session_state['openai_api_key'] = default_key
    except:
        pass

# Автоматическая загрузка ScreenshotAPI токена из secrets.toml при запуске
if "screenshotapi_token" not in st.session_state:
    try:
        default_token = st.secrets.get('screenshotapi_token', '')
        if default_token:
            st.session_state['screenshotapi_token'] = default_token
        else:
            # Устанавливаем токен по умолчанию
            st.session_state['screenshotapi_token'] = 'MDA94V3-4C14RXS-KSE2KMK-T08BD9M'
    except:
        # Устанавливаем токен по умолчанию
        st.session_state['screenshotapi_token'] = 'MDA94V3-4C14RXS-KSE2KMK-T08BD9M'

# Автоматическая загрузка последней таблицы параметров при запуске
if "table_loaded" not in st.session_state:
    try:
        import json
        import os
        # Загружаем список удаленных параметров
        deleted_params = load_deleted_params_from_file()
        
        if os.path.exists("table_cache.json"):
            with open("table_cache.json", "r", encoding="utf-8") as f:
                table_cache_data = json.load(f)
            
            # Восстанавливаем данные
            param_values_raw = table_cache_data.get("param_values", {})
            param_options_raw = table_cache_data.get("param_options", {})
            
            # Фильтруем невалидные и удаленные параметры
            # Убеждаемся, что deleted_params это set
            if not isinstance(deleted_params, set):
                deleted_params = set(deleted_params) if deleted_params else set()
            
            st.session_state["param_values"] = {
                k: v for k, v in param_values_raw.items() 
                if is_valid_param_name(k) and k not in deleted_params
            }
            st.session_state["param_options"] = {
                k: v for k, v in param_options_raw.items() 
                if is_valid_param_name(k) and k not in deleted_params
            }
            
            # Дополнительная проверка - удаляем удаленные параметры еще раз на всякий случай
            if deleted_params:
                for param_name in list(st.session_state["param_values"].keys()):
                    if param_name in deleted_params:
                        del st.session_state["param_values"][param_name]
                for param_name in list(st.session_state["param_options"].keys()):
                    if param_name in deleted_params:
                        del st.session_state["param_options"][param_name]
            st.session_state["param_ratings"] = table_cache_data.get("param_ratings", {})
            
            # Сохраняем deleted_params в session_state
            st.session_state["deleted_params"] = deleted_params
            
            # Отмечаем, что таблица загружена
            st.session_state["table_loaded"] = True
            
            # Показываем уведомление о загрузке
            st.sidebar.success(f"📂 Автозагрузка: таблица параметров восстановлена")
    except Exception as e:
        st.session_state["table_loaded"] = True  # Отмечаем, что попытка загрузки была

# Автоматическая загрузка последнего файла при запуске
if "file_auto_loaded" not in st.session_state:
    try:
        import json
        import os
        
        # Проверяем наличие кешированного файла
        if os.path.exists("file_cache_meta.json"):
            with open("file_cache_meta.json", "r", encoding="utf-8") as f:
                meta_data = json.load(f)
            
            filename = meta_data.get("filename")
            cache_path = os.path.join("file_cache", filename)
            
            # Если файл существует, устанавливаем флаг для автозагрузки
            if os.path.exists(cache_path):
                st.session_state["auto_load_file"] = True
                st.session_state["file_auto_loaded"] = True
                
                # Показываем уведомление
                st.sidebar.info(f"📂 Найден кешированный файл: {filename}")
            else:
                st.session_state["file_auto_loaded"] = True
        else:
            st.session_state["file_auto_loaded"] = True
    except Exception as e:
        st.session_state["file_auto_loaded"] = True  # Отмечаем, что попытка загрузки была

# Инициализация session_state
if "schemas" not in st.session_state:
    st.session_state["schemas"] = {}

# Функции для работы с кешем файлов
def save_file_cache(file_data, filename):
    """Сохраняет файл в кеш"""
    try:
        import os
        cache_dir = "file_cache"
        if not os.path.exists(cache_dir):
            os.makedirs(cache_dir)
        
        # Сохраняем файл
        cache_path = os.path.join(cache_dir, filename)
        with open(cache_path, "wb") as f:
            f.write(file_data)
        
        # Сохраняем метаданные
        import json
        meta_data = {
            "filename": filename,
            "timestamp": pd.Timestamp.now().strftime("%Y-%m-%d %H:%M:%S"),
            "size": len(file_data)
        }
        
        with open("file_cache_meta.json", "w", encoding="utf-8") as f:
            json.dump(meta_data, f, ensure_ascii=False, indent=2)
        
        return True
    except Exception as e:
        st.error(f"Ошибка сохранения файла в кеш: {e}")
        return False

def get_analysis_period(df, df_raw=None, header_row=None):
    """Извлекает анализируемый период из заголовка таблицы"""
    try:
        # Сначала пытаемся найти период в заголовке таблицы
        if df_raw is not None and header_row is not None:
            # Ищем строку с "Анализируемый период" в заголовке
            for i in range(max(0, header_row - 5), min(len(df_raw), header_row + 1)):
                row_values = df_raw.iloc[i].astype(str).str.strip().tolist()
                
                # Ищем "Анализируемый период" в строке
                for j, cell_value in enumerate(row_values):
                    if "анализируемый период" in cell_value.lower():
                        # Ищем даты в соседних ячейках
                        for k in range(j + 1, min(len(row_values), j + 5)):
                            period_value = row_values[k]
                            if period_value and period_value != "nan":
                                # Пытаемся извлечь даты из строки вида "01.01.2025 - 30.04.2025"
                                import re
                                date_pattern = r'(\d{2}\.\d{2}\.\d{4})\s*-\s*(\d{2}\.\d{2}\.\d{4})'
                                match = re.search(date_pattern, period_value)
                                
                                if match:
                                    start_date_str = match.group(1)
                                    end_date_str = match.group(2)
                                    
                                    try:
                                        # Парсим даты
                                        start_date = pd.to_datetime(start_date_str, format='%d.%m.%Y')
                                        end_date = pd.to_datetime(end_date_str, format='%d.%m.%Y')
                                        
                                        # Форматируем даты для отображения
                                        months = {
                                            1: "января", 2: "февраля", 3: "марта", 4: "апреля", 5: "мая", 6: "июня",
                                            7: "июля", 8: "августа", 9: "сентября", 10: "октября", 11: "ноября", 12: "декабря"
                                        }
                                        
                                        start_date_formatted = f"{start_date.day} {months[start_date.month]} {start_date.year}"
                                        end_date_formatted = f"{end_date.day} {months[end_date.month]} {end_date.year}"
                                        
                                        # Вычисляем количество дней
                                        days_diff = (end_date - start_date).days
                                        
                                        return {
                                            "start_date": start_date,
                                            "end_date": end_date,
                                            "start_date_str": start_date_formatted,
                                            "end_date_str": end_date_formatted,
                                            "days_count": days_diff + 1,
                                            "period_str": f"{start_date_formatted} - {end_date_formatted} ({days_diff + 1} дней)",
                                            "source": "header"
                                        }
                                    except Exception as e:
                                        st.warning(f"Не удалось распарсить даты из заголовка: {e}")
                                        break
                                break
                        break
        
        # Если не нашли в заголовке, пытаемся извлечь из колонки "Дата создания"
        if "Дата создания" in df.columns and not df["Дата создания"].isna().all():
            # Получаем минимальную и максимальную даты
            min_date = df["Дата создания"].dropna().min()
            max_date = df["Дата создания"].dropna().max()
            
            if pd.notna(min_date) and pd.notna(max_date):
                # Форматируем даты для отображения
                months = {
                    1: "января", 2: "февраля", 3: "марта", 4: "апреля", 5: "мая", 6: "июня",
                    7: "июля", 8: "августа", 9: "сентября", 10: "октября", 11: "ноября", 12: "декабря"
                }
                
                min_date_str = f"{min_date.day} {months[min_date.month]} {min_date.year}"
                max_date_str = f"{max_date.day} {months[max_date.month]} {max_date.year}"
                
                # Вычисляем количество дней
                days_diff = (max_date - min_date).days
                
                return {
                    "start_date": min_date,
                    "end_date": max_date,
                    "start_date_str": min_date_str,
                    "end_date_str": max_date_str,
                    "days_count": days_diff + 1,
                    "period_str": f"{min_date_str} - {max_date_str} ({days_diff + 1} дней)",
                    "source": "date_column"
                }
        
        return None
    except Exception as e:
        st.error(f"Ошибка при определении периода анализа: {e}")
        return None

def get_file_statistics(df):
    """Получает статистику по загруженному файлу"""
    try:
        stats = {
            "total_rows": len(df),
            "total_products": df["Артикул"].nunique() if "Артикул" in df.columns else len(df),
            "total_revenue": df["Выручка"].sum() if "Выручка" in df.columns else 0,
            "total_orders": df["Заказы"].sum() if "Заказы" in df.columns else 0,
            "avg_price": df["Средняя цена"].mean() if "Средняя цена" in df.columns else 0,
            "columns_count": len(df.columns)
        }
        
        # Добавляем информацию о колонках
        stats["available_columns"] = list(df.columns)
        
        return stats
    except Exception as e:
        st.error(f"Ошибка при получении статистики файла: {e}")
        return None

def load_file_cache():
    """Загружает файл из кеша"""
    try:
        import json
        import os
        
        # Проверяем наличие метаданных
        if not os.path.exists("file_cache_meta.json"):
            return None, None
        
        # Загружаем метаданные
        with open("file_cache_meta.json", "r", encoding="utf-8") as f:
            meta_data = json.load(f)
        
        filename = meta_data.get("filename")
        cache_path = os.path.join("file_cache", filename)
        
        # Проверяем наличие файла
        if not os.path.exists(cache_path):
            return None, None
        
        # Загружаем файл
        with open(cache_path, "rb") as f:
            file_data = f.read()
        
        return file_data, meta_data
    except Exception as e:
        return None, None

def get_file_cache_info():
    """Получает информацию о кешированном файле"""
    try:
        import json
        import os
        
        if os.path.exists("file_cache_meta.json"):
            with open("file_cache_meta.json", "r", encoding="utf-8") as f:
                meta_data = json.load(f)
            return meta_data
        return None
    except:
        return None

def get_all_cached_files():
    """Получает список всех кешированных файлов"""
    try:
        import os
        cache_dir = "file_cache"
        if not os.path.exists(cache_dir):
            return []
        
        cached_files = []
        for filename in os.listdir(cache_dir):
            if filename.endswith(('.xlsx', '.xls', '.csv')):
                file_path = os.path.join(cache_dir, filename)
                file_size = os.path.getsize(file_path)
                file_time = os.path.getmtime(file_path)
                
                cached_files.append({
                    "filename": filename,
                    "size": file_size,
                    "timestamp": pd.Timestamp.fromtimestamp(file_time).strftime("%Y-%m-%d %H:%M:%S"),
                    "path": file_path
                })
        
        # Сортируем по времени изменения (новые сначала)
        cached_files.sort(key=lambda x: x["timestamp"], reverse=True)
        return cached_files
    except Exception as e:
        st.error(f"Ошибка при получении списка кешированных файлов: {e}")
        return []

def save_file_to_cache(file_data, filename):
    """Сохраняет файл в кеш с улучшенной системой"""
    try:
        import os
        cache_dir = "file_cache"
        if not os.path.exists(cache_dir):
            os.makedirs(cache_dir)
        
        # Сохраняем файл
        cache_path = os.path.join(cache_dir, filename)
        with open(cache_path, "wb") as f:
            f.write(file_data)
        
        # Обновляем метаданные
        meta_data = {
            "filename": filename,
            "timestamp": pd.Timestamp.now().strftime("%Y-%m-%d %H:%M:%S"),
            "size": len(file_data),
            "last_used": pd.Timestamp.now().strftime("%Y-%m-%d %H:%M:%S")
        }
        
        with open("file_cache_meta.json", "w", encoding="utf-8") as f:
            json.dump(meta_data, f, ensure_ascii=False, indent=2)
        
        return True
    except Exception as e:
        st.error(f"Ошибка сохранения файла в кеш: {e}")
        return False

# --- UI (урезанный пример, ключевые места с прибыль и миниатюрами) ---
with st.sidebar.expander("Загрузка файла", expanded=True):
    # Получаем список всех кешированных файлов
    cached_files = get_all_cached_files()
    
    if cached_files:
        st.write("**📂 Кешированные файлы:**")
        
        # Показываем последний использованный файл
        current_file_info = get_file_cache_info()
        if current_file_info:
            st.info(f"🔄 Текущий: {current_file_info['filename']}\n🕒 {current_file_info['timestamp']}")
        
        # Список всех файлов
        for i, file_info in enumerate(cached_files[:5]):  # Показываем только 5 последних
            col_file, col_load, col_del = st.columns([3, 1, 1])
            
            with col_file:
                file_size_mb = file_info["size"] / (1024 * 1024)
                st.caption(f"📄 {file_info['filename']}\n💾 {file_size_mb:.1f} MB • {file_info['timestamp']}")
            
            with col_load:
                if st.button("📂", key=f"load_{i}"):
                    # Загружаем выбранный файл
                    try:
                        with open(file_info["path"], "rb") as f:
                            file_data = f.read()
                        
                        # Обновляем метаданные
                        meta_data = {
                            "filename": file_info["filename"],
                            "timestamp": file_info["timestamp"],
                            "size": file_info["size"],
                            "last_used": pd.Timestamp.now().strftime("%Y-%m-%d %H:%M:%S")
                        }
                        
                        with open("file_cache_meta.json", "w", encoding="utf-8") as f:
                            json.dump(meta_data, f, ensure_ascii=False, indent=2)
                        
                        # Сохраняем в session_state
                        st.session_state["cached_file_data"] = file_data
                        st.session_state["cached_file_name"] = file_info["filename"]
                        st.session_state["file_loaded_from_cache"] = True
                        st.session_state["load_from_cache"] = True
                        
                        # Загружаем параметры и настройки для нового файла
                        load_param_values_from_file()
                        load_main_page_data_from_file()
                        st.rerun()
                    except Exception as e:
                        st.error(f"❌ Ошибка загрузки: {e}")
            
            with col_del:
                if st.button("🗑️", key=f"del_{i}"):
                    try:
                        import os
                        os.remove(file_info["path"])
                        st.success(f"✅ Файл {file_info['filename']} удален из кеша")
                        st.rerun()
                    except Exception as e:
                        st.error(f"❌ Ошибка удаления: {e}")
        
        # Кнопка очистки всего кеша
        if st.button("🗑️ Очистить весь кеш файлов", type="secondary"):
            try:
                import os
                import shutil
                
                # Удаляем все файлы кеша
                cache_dir = "file_cache"
                if os.path.exists(cache_dir):
                    shutil.rmtree(cache_dir)
                
                # Удаляем метаданные
                if os.path.exists("file_cache_meta.json"):
                    os.remove("file_cache_meta.json")
                
                # Очищаем session_state
                if "cached_file_data" in st.session_state:
                    del st.session_state["cached_file_data"]
                if "cached_file_name" in st.session_state:
                    del st.session_state["cached_file_name"]
                if "file_loaded_from_cache" in st.session_state:
                    del st.session_state["file_loaded_from_cache"]
                
                st.success("✅ Весь кеш файлов очищен!")
                st.rerun()
            except Exception as e:
                st.error(f"❌ Ошибка очистки кеша: {e}")
    else:
        st.info("📂 Кеш файлов пуст")
    
    # Кнопка для сброса текущего файла (если файл загружен)
    if st.session_state.get("file_loaded_from_cache", False):
        if st.button("🔄 Загрузить новый файл", type="secondary"):
            # Очищаем session_state
            if "cached_file_data" in st.session_state:
                del st.session_state["cached_file_data"]
            if "cached_file_name" in st.session_state:
                del st.session_state["cached_file_name"]
            if "file_loaded_from_cache" in st.session_state:
                del st.session_state["file_loaded_from_cache"]
            st.rerun()
    
    # Всегда показываем кнопку загрузки файла
    st.markdown("---")
    st.markdown("**📤 Загрузить новый файл:**")
    uploaded = st.file_uploader("Excel/CSV с товарами", type=["xlsx","xls","csv"], key="main_uploader")
    
    # Автоматическое сохранение загруженного файла в кеш
    if uploaded is not None:
        file_data = uploaded.read()
        # Возвращаем указатель в начало для дальнейшего чтения
        uploaded.seek(0)
        
        # Сохраняем в кеш с улучшенной системой
        if save_file_to_cache(file_data, uploaded.name):
            st.success(f"💾 Файл автоматически сохранен в кеш: {uploaded.name}")
            
            # Сохраняем информацию в session_state
            st.session_state["cached_file_data"] = file_data
            st.session_state["cached_file_name"] = uploaded.name
            st.session_state["file_loaded_from_cache"] = True
            st.session_state["uploaded_file"] = uploaded
            
            # Загружаем параметры и настройки для нового файла
            load_param_values_from_file()
            load_main_page_data_from_file()
            st.rerun()
    
    # Секция для загрузки товаров по ссылкам Wildberries
    st.markdown("---")
    with st.expander("🔗 Загрузка товаров по ссылкам Wildberries", expanded=False):
        st.info("💡 **Альтернативный способ:** Загрузите товары по ссылкам Wildberries. Система автоматически получит изображения и определит параметры через нейросеть OpenAI.")
        
        urls_input = st.text_area(
            "Введите ссылки на товары Wildberries (по одной на строку):",
            height=150,
            help="Каждая ссылка должна быть на отдельной строке.\nПример:\nhttps://www.wildberries.ru/catalog/12345678/detail.aspx\nhttps://global.wildberries.ru/catalog/87654321/detail.aspx",
            key="wb_urls_input"
        )
        
        # Настройки для загрузки по ссылкам
        col_delay1, col_delay2 = st.columns(2)
        with col_delay1:
            wb_rate_limit_delay = st.slider(
                "Задержка между запросами к WB (сек):",
                min_value=0.5,
                max_value=5.0,
                value=1.0,
                step=0.1,
                help="Помогает избежать rate limit при парсинге Wildberries",
                key="wb_delay_slider"
            )
        with col_delay2:
            max_images_per_url = st.slider(
                "Изображений на товар:",
                min_value=1,
                max_value=5,
                value=5,
                help="Количество изображений для анализа (рекомендуется 5)",
                key="max_images_url_slider"
            )
        
        if st.button("📥 Загрузить товары по ссылкам", type="primary", key="load_urls_btn"):
            if urls_input:
                urls = [url.strip() for url in urls_input.split('\n') if url.strip()]
                
                if urls:
                    progress_bar = st.progress(0)
                    status_text = st.empty()
                    
                    results = []
                    api_key = st.session_state.get('openai_api_key', '')
                    
                    if not api_key:
                        st.error("❌ Необходимо указать API ключ OpenAI в настройках (вкладка '⚙️ Установка параметров')")
                    else:
                        total_urls = len(urls)
                        success_count = 0
                        error_count = 0
                        
                        for idx, url in enumerate(urls):
                            status_text.text(f"Обработка товара {idx + 1}/{total_urls}: {url[:60]}...")
                            
                            # Извлекаем артикул
                            sku = extract_sku_from_url(url)
                            if not sku:
                                results.append({
                                    "URL": url,
                                    "Артикул": "Не найден",
                                    "Статус": "❌ Ошибка извлечения артикула"
                                })
                                error_count += 1
                                # Ограничиваем значение progress до максимум 1.0
                                progress_value = min((idx + 1) / total_urls, 1.0)
                                progress_bar.progress(progress_value)
                                continue
                            
                            # Получаем параметры через анализ изображений
                            try:
                                params = get_product_params_from_images(url, api_key, max_images=max_images_per_url)
                                
                                if params:
                                    result_row = {
                                        "URL": url,
                                        "Артикул": sku,
                                        "Статус": "✅ Успешно",
                                        **params
                                    }
                                    results.append(result_row)
                                    success_count += 1
                                    
                                    # Сохраняем параметры в session_state
                                    if "param_values" not in st.session_state:
                                        st.session_state["param_values"] = {}
                                    
                                    for param_name, param_value in params.items():
                                        if param_name not in st.session_state["param_values"]:
                                            st.session_state["param_values"][param_name] = {}
                                        st.session_state["param_values"][param_name][sku] = param_value
                                    
                                else:
                                    results.append({
                                        "URL": url,
                                        "Артикул": sku,
                                        "Статус": "⚠️ Параметры не определены"
                                    })
                                    error_count += 1
                            except Exception as e:
                                results.append({
                                    "URL": url,
                                    "Артикул": sku,
                                    "Статус": f"❌ Ошибка: {str(e)[:50]}"
                                })
                                error_count += 1
                            
                            # Обновляем прогресс
                            progress_bar.progress((idx + 1) / total_urls)
                            
                            # Задержка между товарами для избежания rate limit
                            import time
                            if idx < total_urls - 1:  # Не ждем после последнего товара
                                time.sleep(wb_rate_limit_delay)
                        
                        # Очищаем прогресс
                        progress_bar.empty()
                        status_text.empty()
                        
                        # Сохраняем результаты
                        if results:
                            results_df = pd.DataFrame(results)
                            st.session_state['url_loaded_products'] = results_df
                            
                            # Показываем статистику
                            st.success(f"✅ Обработка завершена!")
                            st.info(f"📊 **Статистика:** Успешно: {success_count} | Ошибок: {error_count} | Всего: {total_urls}")
                            
                            # Показываем результаты
                            st.dataframe(results_df, use_container_width=True)
                            
                            # Сохраняем параметры в файл
                            if success_count > 0:
                                if save_param_values_to_file():
                                    st.success("💾 Параметры сохранены в файл!")
                                else:
                                    st.warning("⚠️ Параметры сохранены в памяти, но не в файл")
                                
                                st.info("💡 **Совет:** Параметры товаров автоматически сохранены. Вы можете использовать их в массовом анализе или просмотреть в редакторе параметров.")
            else:
                st.warning("⚠️ Введите хотя бы одну ссылку на товар Wildberries")
    
    # Обработка загрузки из кеша (ручная или автоматическая)
    if st.session_state.get("load_from_cache", False) or st.session_state.get("auto_load_file", False):
        file_data, meta_data = load_file_cache()
        if file_data and meta_data:
            # Создаем объект, похожий на uploaded file
            class CachedFile:
                def __init__(self, data, name):
                    self.data = data
                    self.name = name
                
                def read(self):
                    return self.data
                
                def seek(self, pos):
                    pass  # Заглушка для совместимости
            
            uploaded = CachedFile(file_data, meta_data["filename"])
            
            # Сохраняем информацию о загруженном файле в session_state
            st.session_state["cached_file_data"] = file_data
            st.session_state["cached_file_name"] = meta_data["filename"]
            st.session_state["file_loaded_from_cache"] = True
            
            # Загружаем параметры и настройки для файла
            load_param_values_from_file()
            load_main_page_data_from_file()
            
            if st.session_state.get("auto_load_file", False):
                st.success(f"🔄 Автозагрузка: файл восстановлен из кеша - {meta_data['filename']}")
                st.session_state["auto_load_file"] = False
            else:
                st.success(f"✅ Файл загружен из кеша: {meta_data['filename']}")
            
            # Сбрасываем флаги
            st.session_state["load_from_cache"] = False
        else:
            st.error("❌ Не удалось загрузить файл из кеша")
            st.session_state["load_from_cache"] = False
            st.session_state["auto_load_file"] = False
            uploaded = None
    else:
        # Проверяем, есть ли сохраненный файл в session_state
        if st.session_state.get("file_loaded_from_cache", False) and st.session_state.get("cached_file_data"):
            # Восстанавливаем файл из session_state
            class CachedFile:
                def __init__(self, data, name):
                    self.data = data
                    self.name = name
                
                def read(self):
                    return self.data
                
                def seek(self, pos):
                    pass  # Заглушка для совместимости
            
            uploaded = CachedFile(st.session_state["cached_file_data"], st.session_state["cached_file_name"])
        else:
            # Используем файл, загруженный через основную кнопку загрузки
            uploaded = st.session_state.get("uploaded_file", None)
    
    # Автоматическая загрузка последнего файла из кеша, если файл не загружен
    if uploaded is None and not st.session_state.get("auto_load_attempted", False):
        # Пытаемся загрузить последний файл из кеша
        file_data, meta_data = load_file_cache()
        if file_data and meta_data:
            # Создаем объект, похожий на uploaded file
            class CachedFile:
                def __init__(self, data, name):
                    self.data = data
                    self.name = name
                
                def read(self):
                    return self.data
                
                def seek(self, pos):
                    pass  # Заглушка для совместимости
            
            uploaded = CachedFile(file_data, meta_data["filename"])
            
            # Сохраняем информацию о загруженном файле в session_state
            st.session_state["cached_file_data"] = file_data
            st.session_state["cached_file_name"] = meta_data["filename"]
            st.session_state["file_loaded_from_cache"] = True
            st.session_state["auto_load_attempted"] = True
            
            # Загружаем параметры и настройки для файла
            load_param_values_from_file()
            load_main_page_data_from_file()
            # Перезагружаем страницу для отображения данных
            st.rerun()
        else:
            # Если файла в кеше нет, пробуем найти последний файл в папке file_cache
            cached_files = get_all_cached_files()
            if cached_files:
                # Берем самый последний файл
                latest_file = cached_files[0]
                try:
                    with open(latest_file["path"], "rb") as f:
                        file_data = f.read()
                    
                    # Обновляем метаданные
                    meta_data = {
                        "filename": latest_file["filename"],
                        "timestamp": latest_file["timestamp"],
                        "size": latest_file["size"],
                        "last_used": pd.Timestamp.now().strftime("%Y-%m-%d %H:%M:%S")
                    }
                    
                    with open("file_cache_meta.json", "w", encoding="utf-8") as f:
                        json.dump(meta_data, f, ensure_ascii=False, indent=2)
                    
                    # Создаем объект CachedFile
                    class CachedFile:
                        def __init__(self, data, name):
                            self.data = data
                            self.name = name
                        
                        def read(self):
                            return self.data
                        
                        def seek(self, pos):
                            pass
                    
                    uploaded = CachedFile(file_data, latest_file["filename"])
                    
                    # Сохраняем информацию в session_state
                    st.session_state["cached_file_data"] = file_data
                    st.session_state["cached_file_name"] = latest_file["filename"]
                    st.session_state["file_loaded_from_cache"] = True
                    st.session_state["auto_load_attempted"] = True
                    
                    # Загружаем параметры и настройки
                    load_param_values_from_file()
                    load_main_page_data_from_file()
                    # Перезагружаем страницу для отображения данных
                    st.rerun()
                except Exception as e:
                    st.session_state["auto_load_attempted"] = True
            else:
                st.session_state["auto_load_attempted"] = True

with st.sidebar.expander("Скриншоты страниц (s-shot.ru)", expanded=True):
    sc_key = st.text_input("Ключ s-shot", value="KEYSV7S9IWCFGI50SA8")
    sc_base = st.text_input("Базовый URL", value="https://api.s-shot.ru")
    sc_host = st.text_input("Домен карточки WB", value="https://global.wildberries.ru")
    sc_w = st.number_input("Ширина", 100, 2000, 400, 10)
    sc_h = st.number_input("Высота", 100, 2000, 600, 10)
    sc_fmt = st.selectbox("Формат", ["JPEG","PNG"], 0)
    sc_profile = st.text_input("Профиль", value="D4")
    
    # Информация о кеше
    url_cache = get_url_cache()
    cached_count = len(url_cache)
    st.info(f"📦 В кеше: {cached_count} изображений")
    
    # Кнопки управления кешем
    col_cache1, col_cache2 = st.columns(2)
    if col_cache1.button("🗑️ Очистить кеш"):
        st.session_state["img_url_cache"] = {}
        save_url_cache({})
        st.rerun()
    
    if col_cache2.button("💾 Сохранить параметры"):
        if save_param_values_to_file():
            st.success("✅ Параметры сохранены!")
        else:
            st.error("❌ Ошибка сохранения параметров")
    
    # Кнопка сохранения данных главной страницы
    if st.button("💾 Сохранить настройки главной страницы", use_container_width=True):
        # Сохраняем текущие значения в session_state
        st.session_state["search"] = st.session_state.get("search_input", "")
        st.session_state["spp"] = st.session_state.get("spp_input", 25)
        st.session_state["buyout_pct"] = st.session_state.get("buyout_input", 25)
        st.session_state["revenue_min"] = st.session_state.get("revenue_min_input", 0)
        st.session_state["revenue_max"] = st.session_state.get("revenue_max_input", 1000000)
        st.session_state["price_min"] = st.session_state.get("price_min_input", 0)
        st.session_state["price_max"] = st.session_state.get("price_max_input", 10000)
        
        if save_main_page_data_to_file():
            st.success("✅ Настройки главной страницы сохранены!")
        else:
            st.error("❌ Ошибка сохранения настроек")
    
    st.divider()
    
    # Управление таблицей параметров
    st.write("**Управление таблицей параметров:**")
    
    # Кнопка сохранения
    if st.button("💾 Сохранить таблицу в кеш", use_container_width=True):
            # Получаем удаленные параметры для фильтрации
            deleted_params = st.session_state.get("deleted_params", load_deleted_params_from_file())
            param_values = st.session_state.get("param_values", {})
            param_options = st.session_state.get("param_options", {})
            
            # Фильтруем удаленные параметры перед сохранением
            filtered_param_values = {
                k: v for k, v in param_values.items() 
                if k not in deleted_params
            } if deleted_params else param_values
            
            filtered_param_options = {
                k: v for k, v in param_options.items() 
                if k not in deleted_params
            } if deleted_params else param_options
            
            # Сохраняем текущую таблицу параметров
            table_cache_data = {
                "param_values": filtered_param_values,
                "param_options": filtered_param_options,
                "timestamp": pd.Timestamp.now().strftime("%Y-%m-%d %H:%M:%S")
            }
            
            # Сохраняем в файл
            try:
                import json
                with open("table_cache.json", "w", encoding="utf-8") as f:
                    json.dump(table_cache_data, f, ensure_ascii=False, indent=2)
                st.success("✅ Таблица сохранена в кеш!")
            except Exception as e:
                st.error(f"❌ Ошибка сохранения: {e}")
    
    # Кнопка загрузки
    if st.button("📂 Загрузить таблицу из кеша", use_container_width=True):
            # Загружаем последнюю сохраненную таблицу
            try:
                import json
                import os
                if os.path.exists("table_cache.json"):
                    with open("table_cache.json", "r", encoding="utf-8") as f:
                        table_cache_data = json.load(f)
                    
                    # Загружаем список удаленных параметров
                    deleted_params = load_deleted_params_from_file()
                    
                    # Восстанавливаем данные
                    param_values_raw = table_cache_data.get("param_values", {})
                    param_options_raw = table_cache_data.get("param_options", {})
                    
                    # Фильтруем невалидные и удаленные параметры
                    st.session_state["param_values"] = {
                        k: v for k, v in param_values_raw.items() 
                        if is_valid_param_name(k) and k not in deleted_params
                    }
                    st.session_state["param_options"] = {
                        k: v for k, v in param_options_raw.items() 
                        if is_valid_param_name(k) and k not in deleted_params
                    }
                    st.session_state["param_ratings"] = table_cache_data.get("param_ratings", {})
                    
                    # Сохраняем deleted_params в session_state
                    st.session_state["deleted_params"] = deleted_params
                    
                    timestamp = table_cache_data.get("timestamp", "неизвестно")
                    st.success(f"✅ Таблица загружена! (сохранена: {timestamp})")
                    st.rerun()
                else:
                    st.warning("Кеш таблицы не найден")
            except Exception as e:
                st.error(f"❌ Ошибка загрузки: {e}")
    
    # Кнопка очистки
    if st.button("🗑️ Очистить кеш таблицы", use_container_width=True):
            # Очищаем кеш таблицы параметров
            try:
                import os
                if os.path.exists("table_cache.json"):
                    os.remove("table_cache.json")
                    st.success("✅ Кеш таблицы очищен!")
                else:
                    st.warning("Кеш таблицы не найден")
                st.rerun()
            except Exception as e:
                st.error(f"❌ Ошибка очистки кеша: {e}")
    
    # Информация о кеше таблицы
    try:
        import json
        import os
        if os.path.exists("table_cache.json"):
            with open("table_cache.json", "r", encoding="utf-8") as f:
                table_cache_data = json.load(f)
            timestamp = table_cache_data.get("timestamp", "неизвестно")
            param_count = len(table_cache_data.get("param_options", {}))
            st.info(f"📦 Кеш таблицы: {param_count} параметров, сохранен {timestamp}")
    except:
        pass

if uploaded is None:
    st.info("Загрузите файл с данными.")
else:
    df, raw, meta = read_table(uploaded.read(), uploaded.name)
    if df is None or df.empty:
        st.error("Не удалось прочитать таблицу.")
    else:
        # Сохраняем df в session_state для использования в других вкладках
        st.session_state["df"] = df
        st.title("📊 Дашборд WB")
        
        # Получаем анализируемый период (для внутреннего использования)
        analysis_period = get_analysis_period(df, raw, meta.get("header_row"))
        
        # Отображаем информацию о периоде анализа
        if analysis_period:
            source_text = "из заголовка таблицы" if analysis_period.get("source") == "header" else "из колонки 'Дата создания'"
            st.success(f"📅 **Анализируемый период:** {analysis_period['period_str']} ({source_text})")
        
        # Импортируем модуль анализа сезонности
        try:
            from seasonality_module import (
                load_seasonality_data, clean_seasonality_data, create_seasonality_graph,
                get_status_stats, style_dataframe, load_custom_data, create_manual_entry_data
            )
            seasonality_available = True
        except ImportError:
            seasonality_available = False
        
        # Создаем вкладки
        if seasonality_available:
            if PROPHET_AVAILABLE:
                tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs(["📊 Анализ данных", "⚙️ Установка параметров", "📈 Аналитика по параметрам", "📅 Анализ сезонности", "🔮 Прогнозирование", "📜 История изменений"])
            else:
                tab1, tab2, tab3, tab4, tab5 = st.tabs(["📊 Анализ данных", "⚙️ Установка параметров", "📈 Аналитика по параметрам", "📅 Анализ сезонности", "📜 История изменений"])
        else:
            if PROPHET_AVAILABLE:
                tab1, tab2, tab3, tab4, tab5 = st.tabs(["📊 Анализ данных", "⚙️ Установка параметров", "📈 Аналитика по параметрам", "🔮 Прогнозирование", "📜 История изменений"])
            else:
                tab1, tab2, tab3, tab4 = st.tabs(["📊 Анализ данных", "⚙️ Установка параметров", "📈 Аналитика по параметрам", "📜 История изменений"])
        
        with tab1:
            # Основные фильтры
            col1, col2, col3, col4 = st.columns(4)
        
            search = col1.text_input("🔍 Поиск", value=st.session_state.get("search", ""), key="search_input")
            spp = col2.number_input("💰 СПП, %", 0, 100, st.session_state.get("spp", 25), 1, key="spp_input")
            buyout_pct = col3.number_input("📈 Процент выкупа, %", 1, 100, st.session_state.get("buyout_pct", 25), 1, key="buyout_input")
            
            # Сохраняем значения в session_state
            st.session_state["search"] = search
            st.session_state["spp"] = spp
            st.session_state["buyout_pct"] = buyout_pct
            
            # Автоматически сохраняем настройки при изменении
            save_main_page_data_to_file()
            
            # Кнопка обновления данных
            col4.markdown("🔄 Обновить данные")
            if col4.button("Обновить", type="primary"):
                st.rerun()
            
            # Фильтр по предмету на отдельной строчке
            if "Предмет" in df.columns:
                subjects = sorted(df["Предмет"].dropna().unique())
                selected_subjects = st.multiselect("📦 Предмет", subjects, default=subjects)
            else:
                selected_subjects = []
            
            # Фильтры по параметрам товаров
            param_values = get_param_values()
            selected_param_filters = {}
            
            if param_values:
                # Переключатель для включения/выключения фильтров по параметрам
                enable_param_filters = st.checkbox(
                    "🎨 Включить фильтры по параметрам", 
                    value=False,
                    help="Включить фильтрацию товаров по их параметрам (цвет, материал и т.д.)"
                )
                
                if enable_param_filters:
                    st.subheader("🎨 Фильтры по параметрам")
                    
                    # Создаем колонки для фильтров параметров
                    param_cols = st.columns(min(len(param_values), 4))  # Максимум 4 колонки
                    
                    for i, (param_name, param_data) in enumerate(param_values.items()):
                        col_idx = i % 4
                        
                        with param_cols[col_idx]:
                            # Получаем уникальные значения для этого параметра (убираем дубликаты)
                            unique_values = sorted(list(set([v for v in param_data.values() if v and v.strip()])))
                            
                            if unique_values:
                                # Создаем multiselect для каждого параметра
                                selected_values = st.multiselect(
                                    f"🎨 {param_name}",
                                    unique_values,
                                    default=[],  # По умолчанию ничего не выбрано
                                    help=f"Выберите значения для параметра '{param_name}'"
                                )
                                selected_param_filters[param_name] = selected_values
                            else:
                                selected_param_filters[param_name] = []
            
            # Настройки отображения
            # Инициализируем img_size по умолчанию (гарантируем, что это число)
            if "img_size" not in st.session_state:
                st.session_state["img_size"] = 200
            img_size = int(st.session_state.get("img_size", 200))
            col_img1, col_img2, col_img3 = st.columns(3)
            show_images = col_img1.checkbox("🖼️ Показывать изображения", value=False)
            if show_images:
                img_size = int(col_img2.number_input("📏 Размер миниатюр (px)", min_value=50, max_value=300, value=int(img_size), step=10))
                st.session_state["img_size"] = img_size
                if col_img3.button("🔄 Обновить кеш изображений", type="secondary"):
                    # Очищаем кеш URL
                    st.session_state["img_url_cache"] = {}
                    save_url_cache({})
                    st.rerun()
                
                # Показываем информацию о кеше изображений
                cache_status = get_cache_status()
                col_cache_info, col_cache_clear = st.columns([2, 1])
                
                with col_cache_info:
                    if cache_status["count"] > 0:
                        size_mb = cache_status["size"] / (1024 * 1024)
                        st.info(f"📦 Кеш изображений: {cache_status['count']} файлов ({size_mb:.1f} МБ)")
                    else:
                        st.info("📦 Кеш изображений пуст")
                
                with col_cache_clear:
                    if cache_status["count"] > 0:
                        if st.button("🗑️ Очистить кеш", key="clear_image_cache"):
                            try:
                                cache_dir = os.path.join(_cache_dir(), "imgs")
                                if os.path.exists(cache_dir):
                                    for file in cache_status["files"]:
                                        if os.path.exists(file["path"]):
                                            os.remove(file["path"])
                                st.success(f"✅ Удалено {cache_status['count']} изображений из кеша")
                                st.rerun()
                            except Exception as e:
                                st.error(f"❌ Ошибка очистки кеша: {e}")
                if not sc_key:
                    st.info("💡 Для отображения изображений товаров введите API ключ s-shot.ru в боковой панели")
            
            # Фильтр сортировки
            col_sort1, col_sort2 = st.columns(2)
            
            # Определяем доступные колонки для сортировки
            sortable_columns = []
            if "Выручка" in df.columns:
                sortable_columns.append("Выручка")
            if "Средняя цена" in df.columns:
                sortable_columns.append("Средняя цена")
            if "Заказы" in df.columns:
                sortable_columns.append("Заказы")
            if "Дата создания" in df.columns:
                sortable_columns.append("Дата создания")
            if "Прибыль" in df.columns:
                sortable_columns.append("Прибыль")
            if "Упущенная выручка" in df.columns:
                sortable_columns.append("Упущенная выручка")
            if "Позиция в выдаче" in df.columns:
                sortable_columns.append("Позиция в выдаче")
            
            # Добавляем опцию "Без сортировки"
            sortable_columns.insert(0, "Без сортировки")
            
            # Находим индекс "Выручка" для установки по умолчанию
            default_index = 0  # По умолчанию "Без сортировки"
            if "Выручка" in sortable_columns:
                default_index = sortable_columns.index("Выручка")
            
            sort_column = col_sort1.selectbox("📊 Сортировка по", sortable_columns, index=default_index)
            sort_ascending = col_sort2.selectbox("🔽 Направление", ["По убыванию", "По возрастанию"], index=0) == "По возрастанию"
            
            st.divider()
            
            # Выручка
            col7, col8 = st.columns(2)
            
            if "Выручка" in df.columns:
                default_revenue_max = int(df["Выручка"].max()) if not df["Выручка"].isna().all() else 1000000
                revenue_min = col7.number_input("Выручка от", min_value=0, value=st.session_state.get("revenue_min", 0), step=1000, key="revenue_min_input")
                revenue_max = col8.number_input("Выручка до", min_value=0, value=st.session_state.get("revenue_max", default_revenue_max), step=1000, key="revenue_max_input")
                
                # Сохраняем значения в session_state
                st.session_state["revenue_min"] = revenue_min
                st.session_state["revenue_max"] = revenue_max
                
                # Автоматически сохраняем настройки при изменении
                save_main_page_data_to_file()
            else:
                revenue_min = st.session_state.get("revenue_min", 0)
                revenue_max = st.session_state.get("revenue_max", 1000000)
            
            # Цена
            col9, col10 = st.columns(2)
            
            if "Средняя цена" in df.columns:
                default_price_max = int(df["Средняя цена"].max()) if not df["Средняя цена"].isna().all() else 10000
                price_min = col9.number_input("Цена (до СПП) от", min_value=0, value=st.session_state.get("price_min", 0), step=100, key="price_min_input")
                price_max = col10.number_input("Цена (до СПП) до", min_value=0, value=st.session_state.get("price_max", default_price_max), step=100, key="price_max_input")
                
                # Сохраняем значения в session_state
                st.session_state["price_min"] = price_min
                st.session_state["price_max"] = price_max
                
                # Автоматически сохраняем настройки при изменении
                save_main_page_data_to_file()
            else:
                price_min = st.session_state.get("price_min", 0)
                price_max = st.session_state.get("price_max", 10000)
            
            # Дата создания
            
            # Определяем диапазон дат для фильтрации
            # Всегда используем полный диапазон дат создания для фильтра
            if "Дата создания" in df.columns:
                # Используем даты из колонки "Дата создания"
                date_range = df["Дата создания"].dropna()
                if not date_range.empty:
                    min_date = date_range.min().date()
                    max_date = date_range.max().date()
                    date_source = "колонка 'Дата создания'"
                else:
                    min_date = pd.Timestamp.now().date()
                    max_date = pd.Timestamp.now().date()
                    date_source = "не найдены"
                    st.warning("⚠️ В данных нет информации о датах")
            else:
                min_date = pd.Timestamp.now().date()
                max_date = pd.Timestamp.now().date()
                date_source = "не найдены"
                st.warning("⚠️ Колонка 'Дата создания' не найдена в данных")
            

            
            # Создаем ползунок для выбора диапазона дат
            date_range_days = (max_date - min_date).days
            if date_range_days > 0:
                date_slider = st.slider(
                    "📅 Выберите диапазон дат для фильтрации",
                    min_value=min_date,
                    max_value=max_date,
                    value=(min_date, max_date),
                    help="Выберите период для анализа. По умолчанию показан полный диапазон дат создания."
                )
                date_min, date_max = date_slider
            else:
                date_min = min_date
                date_max = max_date
                st.info(f"📅 Данные за один день: {min_date.strftime('%d.%m.%Y')}")
            fdf = df.copy()
            
            # Применяем фильтры
            if search:
                mask = fdf.apply(lambda x: x.astype(str).str.contains(search, case=False, na=False)).any(axis=1)
                fdf = fdf[mask]
            
            if selected_subjects and "Предмет" in fdf.columns:
                fdf = fdf[fdf["Предмет"].isin(selected_subjects)]
            
            # Фильтр по выручке
            if "Выручка" in fdf.columns:
                fdf = fdf[(fdf["Выручка"] >= revenue_min) & (fdf["Выручка"] <= revenue_max)]
            
            # Фильтр по цене
            if "Средняя цена" in fdf.columns:
                fdf = fdf[(fdf["Средняя цена"] >= price_min) & (fdf["Средняя цена"] <= price_max)]
            
            # Фильтр по дате создания
            if "Дата создания" in fdf.columns:
                fdf = fdf[(fdf["Дата создания"].dt.date >= date_min) & (fdf["Дата создания"].dt.date <= date_max)]
            
            # Фильтры по параметрам товаров
            if param_values and selected_param_filters and enable_param_filters:
                for param_name, selected_values in selected_param_filters.items():
                    if selected_values:  # Если выбраны значения для фильтрации
                        # Находим артикулы, которые соответствуют выбранным значениям параметра
                        matching_skus = []
                        if param_name in param_values:
                            for sku, value in param_values[param_name].items():
                                if value in selected_values:
                                    matching_skus.append(sku)
                        
                        if matching_skus:
                            # Фильтруем данные по найденным артикулам
                            mask = fdf["Артикул"].astype(str).str.replace(".0", "").isin(matching_skus)
                            fdf = fdf[mask]
            
            if "Средняя цена" in fdf.columns:
                fdf["Цена (с СПП)"] = fdf["Средняя цена"] * (1 - float(spp)/100.0)
            buyout_k = float(buyout_pct)/100.0 if buyout_pct else 0.0
            if "Заказы" in fdf.columns:
                fdf["Выкупы"] = pd.to_numeric(fdf["Заказы"], errors="coerce") * buyout_k
            else:
                fdf["Выкупы"] = np.nan
            # === FIX: Прибыль = Выручка * (процент выкупа) ===
            if "Выручка" in fdf.columns and buyout_k > 0:
                fdf["Прибыль"] = pd.to_numeric(fdf["Выручка"], errors="coerce") * buyout_k
            else:
                fdf["Прибыль"] = np.nan
            
            # Применяем сортировку
            if sort_column and sort_column != "Без сортировки" and sort_column in fdf.columns:
                fdf = sort_df(fdf, sort_column, sort_ascending)
            
            kpi_row(fdf)
            st.divider()

            # Миниатюры кэш
            url_cache = get_url_cache()
            
            # Создаем копию для отображения
            display_df = fdf.copy()
            # Сохраняем display_df в session_state для использования в других вкладках
            st.session_state["display_df"] = display_df
            
            # Добавляем изображения
            if show_images and "Артикул" in display_df.columns:
                imgs = []
                loaded_count = 0
                cached_count = 0
                total_items = len(display_df)
                
                # Создаем прогресс-бар если много товаров
                if total_items > 10:
                    progress_bar = st.progress(0)
                    status_text = st.empty()
                
                for i, a in enumerate(display_df["Артикул"].astype(str)):
                    k = a.replace(".0","")
                    url = url_cache.get(k, "")
                    if not url and sc_key:
                        url = screenshot_for_article(k, {"key": sc_key,"w": sc_w,"h": sc_h,"fmt": sc_fmt,"profile": sc_profile,"base": sc_base,"wb_host": sc_host})
                        if url:
                            url_cache[k] = url
                            save_url_cache(url_cache)
                    
                    # Проверяем кеш изображений - сначала ищем уже скачанные
                    cached_path = get_cached_image_path(k)
                    if cached_path and os.path.exists(cached_path):
                        # Изображение уже в кеше, используем его
                        path = cached_path
                        cached_count += 1
                    elif url:
                        # Изображения нет в кеше, но есть URL - скачиваем
                        path = ensure_image_cached(k, url, sc_fmt)
                    else:
                        # Нет ни кеша, ни URL
                        path = ""
                    
                    if path and os.path.exists(path):
                        # Создаем data URI для Streamlit
                        img_bytes = load_image_bytes(path, img_size)
                        if img_bytes:
                            b64_data = base64.b64encode(img_bytes).decode()
                            data_uri = f"data:image/jpeg;base64,{b64_data}"
                            imgs.append(data_uri)
                            loaded_count += 1
                        else:
                            imgs.append("")
                    else:
                        imgs.append("")
                    
                    # Обновляем прогресс
                    if total_items > 10:
                        # Ограничиваем значение progress до максимум 1.0
                        progress = min((i + 1) / total_items, 1.0)
                        progress_bar.progress(progress)
                        status_text.text(f"Загружаем изображения: {i + 1}/{total_items}")
                
                # Очищаем прогресс-бар
                if total_items > 10:
                    progress_bar.empty()
                    status_text.empty()
                
                display_df.insert(1, "Изображение", imgs)
                
                # Показываем статистику загрузки
                if cached_count > 0:
                    st.success(f"📊 Загружено изображений: {loaded_count} из {len(display_df)} товаров (из кеша: {cached_count})")
                else:
                    st.success(f"📊 Загружено изображений: {loaded_count} из {len(display_df)} товаров")
            
            # Добавляем отдельный столбец со ссылками на артикулы
            if "Артикул" in display_df.columns:
                # Создаем специальный формат для отображения "Открыть" в Streamlit
                display_df["Ссылка"] = "Открыть"
                # Но фактические ссылки будут настроены через column_config
            
            # Форматирование даты для Streamlit таблицы
            if "Дата создания" in display_df.columns:
                display_df["Дата создания"] = display_df["Дата создания"].apply(fmt_date)
            
            # Оставляем числовые данные как есть для корректной сортировки в Streamlit таблице
            
            # Изменение порядка столбцов
            desired_order = [
                "Артикул", "Ссылка", "Дата создания", "Выручка", "Заказы", "Выкупы", 
                "Средняя цена", "Цена (с СПП)", "Упущенная выручка", "Прибыль",
                "Предмет", "Позиция в выдаче", "Стоимость за 1000 показов", 
                "Тип рекламы", "Буст на позицию", "Буст с позиции", "Дельта",
                "Название", "Поставщик", "Бренд"
            ]
            
            # Добавляем изображения в начало если нужно
            if show_images and "Изображение" in display_df.columns:
                desired_order.insert(1, "Изображение")
            
            # Информация о загруженных параметрах
            if param_values:
                total_params = sum(len(param_data) for param_data in param_values.values())
                st.info(f"📊 Загружено параметров: {list(param_values.keys())} ({total_params} значений)")
            
            # Получаем список всех параметров (исключаем "крой")
            all_params = list(param_values.keys())
            if "param_options" in st.session_state:
                all_params.extend([p for p in st.session_state["param_options"].keys() if p not in all_params])
            
            # Исключаем столбец "крой" из отображения
            all_params = [param for param in all_params if param.lower() != "крой"]
            
            # Добавляем столбцы параметров в DataFrame
            for param in all_params:
                param_column_data = []
                for sku in display_df["Артикул"].astype(str):
                    sku_clean = sku.replace(".0", "")
                    param_value = param_values.get(param, {}).get(sku_clean, "")
                    param_column_data.append(param_value)
                display_df[param] = param_column_data
            
            # Переупорядочиваем столбцы - параметры после "Дата создания"
            existing_cols = [col for col in desired_order if col in display_df.columns]
            
            # Находим позицию "Дата создания" в desired_order
            date_creation_index = -1
            if "Дата создания" in desired_order:
                date_creation_index = desired_order.index("Дата создания")
            
            # Разделяем столбцы на основные и параметры
            main_cols = [col for col in existing_cols if col not in all_params]
            param_cols = [col for col in all_params if col in display_df.columns]
            other_cols = [col for col in display_df.columns if col not in existing_cols and col not in all_params]
            
            # Если "Дата создания" найдена, вставляем параметры после неё
            if date_creation_index >= 0 and "Дата создания" in main_cols:
                date_index = main_cols.index("Дата создания")
                # Вставляем параметры после "Дата создания"
                final_order = main_cols[:date_index+1] + param_cols + main_cols[date_index+1:] + other_cols
            else:
                # Если "Дата создания" не найдена, добавляем параметры в конец основных столбцов
                final_order = main_cols + param_cols + other_cols
            
            display_df = display_df[final_order]

            from streamlit import column_config as cc
            
            # Настройка конфигурации столбцов для лучшего отображения
            col_cfg = {}
            
            # Конфигурация для изображений
            if "Изображение" in display_df.columns:
                # Убеждаемся, что img_size - это число (гарантируем целое число)
                try:
                    img_width = int(float(img_size)) + 20
                except (ValueError, TypeError):
                    img_width = 220
                col_cfg["Изображение"] = cc.ImageColumn("Изображение", width=int(img_width))
            
            # Конфигурация для артикула (числовой тип)
            if "Артикул" in display_df.columns:
                col_cfg["Артикул"] = cc.NumberColumn("Артикул", format="%.0f", width=int(120))
            
            # Конфигурация для ссылки на товар с динамическими URL
            if "Ссылка" in display_df.columns and "Артикул" in display_df.columns:
                # Создаем ссылки на основе артикулов
                links_data = []
                for sku in display_df["Артикул"].astype(str):
                    sku_clean = sku.replace(".0", "")
                    links_data.append(f"https://global.wildberries.ru/catalog/{sku_clean}/detail.aspx")
                display_df["Ссылка"] = links_data
                # Явно указываем тип данных как строковый
                display_df["Ссылка"] = display_df["Ссылка"].astype(str)
                col_cfg["Ссылка"] = cc.LinkColumn("Ссылка", display_text="🔗", width=int(60))
            
            # Конфигурация для числовых столбцов (NumberColumn для корректной сортировки)
            money_columns = ["Выручка", "Средняя цена", "Цена (с СПП)", "Упущенная выручка", "Прибыль"]
            for col in money_columns:
                if col in display_df.columns:
                    col_cfg[col] = cc.NumberColumn(col, format="%.0f ₽", width=int(120))
            
            # Конфигурация для числовых столбцов с единицами
            if "Заказы" in display_df.columns:
                col_cfg["Заказы"] = cc.NumberColumn("Заказы", format="%.0f шт.", width=int(120))
            if "Выкупы" in display_df.columns:
                col_cfg["Выкупы"] = cc.NumberColumn("Выкупы", format="%.0f шт.", width=int(120))
            
            # Конфигурация для даты - отключаем редактирование чтобы избежать проблем с типами
            if "Дата создания" in display_df.columns:
                col_cfg["Дата создания"] = cc.TextColumn("Дата создания", width=int(150), disabled=True)
            
    # Конфигурация для параметров товаров
    for param in all_params:
        if param in display_df.columns:
            if param in st.session_state.get("param_options", {}):
                # Selectbox для параметров с вариантами
                options = [""] + st.session_state["param_options"][param]
                col_cfg[param] = cc.SelectboxColumn(
                    param, 
                    options=options, 
                    width=int(150)
                )
            else:
                # Проверяем тип данных в колонке для правильной настройки
                column_data = display_df[param]
                if column_data.dtype in ['int64', 'float64'] or pd.api.types.is_numeric_dtype(column_data):
                    # Для числовых данных используем NumberColumn
                    col_cfg[param] = cc.NumberColumn(param, width=int(150), format="%.0f")
                else:
                    # Для текстовых данных используем TextColumn
                    col_cfg[param] = cc.TextColumn(param, width=int(150))
            
    # Отображаем редактируемую таблицу с возможностью сортировки
    # Вычисляем оптимальную высоту на основе количества строк
    num_rows = len(display_df)
    # Минимум 400px, максимум 5000px для очень больших таблиц
    optimal_height = max(400, min(5000, num_rows * 30 + 100))
    
    # Показываем информацию о таблице
    col_info1, col_info2, col_info3 = st.columns(3)
    with col_info1:
        st.metric("📊 Строк в таблице", f"{num_rows:,}")
    with col_info2:
        st.metric("📏 Высота таблицы", f"{optimal_height}px")
    with col_info3:
        if num_rows > 50:
            st.metric("📈 Статус", "Большая таблица")
        elif num_rows > 20:
            st.metric("📈 Статус", "Средняя таблица")
        else:
            st.metric("📈 Статус", "Компактная таблица")
    
    # Дополнительная проверка col_cfg - убеждаемся, что все width - это числа
    if col_cfg:
        for col_name, col_config in col_cfg.items():
            if hasattr(col_config, 'width') and col_config.width is not None:
                try:
                    col_config.width = int(col_config.width)
                except (ValueError, TypeError):
                    # Если width не может быть преобразован в int, убираем его
                    col_config.width = None
    
    # Поле поиска по артикулу
    if "Артикул" in display_df.columns:
        search_col1, search_col2 = st.columns([3, 1])
        with search_col1:
            sku_search = st.text_input(
                "🔍 Поиск по артикулу:",
                value=st.session_state.get("sku_search", ""),
                placeholder="Введите артикул для поиска...",
                key="sku_search_input"
            )
            st.session_state["sku_search"] = sku_search
        with search_col2:
            if st.button("🔄 Сбросить", key="reset_sku_search"):
                st.session_state["sku_search"] = ""
                sku_search = ""
                st.rerun()
        
        # Фильтруем таблицу по поисковому запросу
        if sku_search and sku_search.strip():
            search_term = sku_search.strip()
            # Ищем артикулы, содержащие поисковый запрос
            mask = display_df["Артикул"].astype(str).str.contains(search_term, case=False, na=False)
            filtered_display_df = display_df[mask].copy()
            
            if len(filtered_display_df) > 0:
                st.info(f"🔍 Найдено товаров: {len(filtered_display_df)} из {len(display_df)}")
            else:
                st.warning(f"⚠️ Товары с артикулом, содержащим '{search_term}', не найдены")
        else:
            filtered_display_df = display_df.copy()
    else:
        filtered_display_df = display_df.copy()
    
    edited_df = st.data_editor(
        filtered_display_df, 
        hide_index=True, 
        column_config=col_cfg if col_cfg else None,
        column_order=None,  # Позволяет пользователю переупорядочивать столбцы
        num_rows="fixed",
        use_container_width=True,
        key="main_table_editor"
    )
            
    # Сохраняем изменения параметров обратно в param_values
    if all_params:
        changes_made = False
        for index, row in edited_df.iterrows():
            # Обрабатываем артикул как число, но конвертируем в строку для ключей
            sku_raw = row["Артикул"]
            if pd.isna(sku_raw):
                continue
            sku = str(int(sku_raw)) if isinstance(sku_raw, (int, float)) else str(sku_raw)
            
            for param in all_params:
                if param in row and row[param]:
                    if param not in param_values:
                        param_values[param] = {}
                    if sku not in param_values[param] or param_values[param][sku] != str(row[param]):
                        param_values[param][sku] = str(row[param])
                        changes_made = True
                elif param in param_values and sku in param_values[param]:
                    # Удаляем пустые значения
                    if not row.get(param):
                        del param_values[param][sku]
                        changes_made = True
        
        # Показываем уведомление об изменениях (без автоматического сохранения)
        if changes_made:
            st.success("✅ Изменения параметров сохранены в памяти!")
                    
    # Кнопка для сохранения в файл
    col_save1, col_save2 = st.columns([1, 4])
    with col_save1:
        if st.button("💾 Сохранить в файл", type="primary"):
            if save_param_values_to_file():
                st.success("✅ Параметры сохранены в файл!")
            else:
                st.error("❌ Ошибка сохранения в файл")
        
        with tab2:
            st.subheader("⚙️ Установка параметров товаров")
            
            # ========== РЕДАКТОР ПАРАМЕТРОВ ==========
            with st.expander("📝 Редактор параметров", expanded=False):
                st.markdown("### 📝 Редактор параметров")
                st.markdown("Управление параметрами товаров: добавление, редактирование, удаление параметров и их вариантов")
                
                # Инициализация param_options если нет
                if "param_options" not in st.session_state:
                    st.session_state["param_options"] = {}
                
                # Инициализация списка удаленных параметров
                if "deleted_params" not in st.session_state:
                    st.session_state["deleted_params"] = load_deleted_params_from_file()
                
                param_options = st.session_state.get("param_options", {})
                param_values = get_param_values()
                deleted_params = st.session_state.get("deleted_params", set())
                
                # Синхронизируем param_options с param_values - добавляем параметры из param_values, которых нет в param_options
                # НО исключаем параметры, которые были явно удалены пользователем
                # Также удаляем удаленные параметры из param_values и param_options, если они там есть
                if param_values:
                    # Сначала удаляем удаленные параметры напрямую из session_state
                    if deleted_params:
                        # Удаляем из param_options
                        if "param_options" in st.session_state:
                            params_to_remove_from_options = [p for p in st.session_state["param_options"].keys() if p in deleted_params]
                            for param_name in params_to_remove_from_options:
                                del st.session_state["param_options"][param_name]
                                if param_name in param_options:
                                    del param_options[param_name]
                        
                        # Удаляем из param_values
                        if "param_values" in st.session_state:
                            params_to_remove_from_values = [p for p in st.session_state["param_values"].keys() if p in deleted_params]
                            for param_name in params_to_remove_from_values:
                                del st.session_state["param_values"][param_name]
                    
                    # Перезагружаем param_values через функцию (она фильтрует удаленные)
                    param_values = get_param_values()
                    
                    # Теперь работаем только с отфильтрованными данными
                    params_to_remove = []
                    for param_name in list(param_values.keys()):
                        if param_name in deleted_params:
                            params_to_remove.append(param_name)
                    
                    for param_name in params_to_remove:
                        if param_name in param_values:
                            del param_values[param_name]
                        if param_name in param_options:
                            del param_options[param_name]
                    
                    if params_to_remove:
                        st.session_state["param_values"] = param_values
                        st.session_state["param_options"] = param_options
                    
                    # Теперь синхронизируем - добавляем параметры из param_values, которых нет в param_options
                    # param_values уже отфильтрован от удаленных параметров
                    for param_name in param_values.keys():
                        # Просто проверяем, что параметра нет в param_options
                        if param_name not in param_options:
                            # Добавляем параметр в param_options с пустым списком вариантов
                            param_options[param_name] = []
                            st.session_state["param_options"] = param_options
                
                # Обновляем локальную переменную после синхронизации
                param_options = st.session_state.get("param_options", {})
                # Фильтруем удаленные параметры из param_options для отображения
                param_options = {k: v for k, v in param_options.items() if k not in deleted_params}
                
                # Статистика
                col_stat1, col_stat2, col_stat3, col_stat4 = st.columns(4)
                with col_stat1:
                    st.metric("Всего параметров", len(param_options))
                with col_stat2:
                    total_options = sum(len(opts) if isinstance(opts, list) else 1 for opts in param_options.values())
                    st.metric("Всего вариантов", total_options)
                with col_stat3:
                    params_with_values = sum(1 for param_name in param_options.keys() if param_name in param_values and param_values[param_name])
                    st.metric("Параметров с данными", params_with_values)
                with col_stat4:
                    # Кнопка сохранения параметров
                    if st.button("💾 Сохранить все", type="primary", use_container_width=True):
                        # Сначала сохраняем deleted_params (это очистит все файлы)
                        if save_deleted_params_to_file():
                            # Затем сохраняем param_options и param_values
                            if save_param_values_to_file():
                                # Очищаем результаты массового анализа от удаленных параметров
                                if "mass_analysis_results" in st.session_state:
                                    st.session_state["mass_analysis_results"] = remove_deleted_params_from_mass_results(
                                        st.session_state["mass_analysis_results"]
                                    )
                                    save_mass_analysis_results(st.session_state["mass_analysis_results"])
                                st.success("✅ Все параметры и настройки сохранены!")
                                st.rerun()
                            else:
                                st.error("❌ Ошибка сохранения параметров")
                        else:
                            st.error("❌ Ошибка сохранения списка удаленных параметров")
                
                st.divider()
                
                # Режимы работы редактора
                editor_mode = st.radio(
                    "Режим работы:",
                    ["📋 Просмотр и редактирование", "➕ Добавить новый параметр", "🔗 Управление иерархией", "📥 Импорт / 📤 Экспорт"],
                    horizontal=True,
                    key="editor_mode"
                )
                
                if editor_mode == "📋 Просмотр и редактирование":
                    if param_options:
                        # Поиск параметров
                        search_param = st.text_input("🔍 Поиск параметра:", placeholder="Введите название параметра...", key="search_param_editor")
                        
                        # Фильтрация параметров
                        filtered_params = param_options.copy()
                        if search_param:
                            filtered_params = {k: v for k, v in param_options.items() if search_param.lower() in k.lower()}
                        
                        if filtered_params:
                            # Сортировка параметров
                            sort_order = st.selectbox("Сортировка:", ["По алфавиту", "По количеству вариантов", "По дате добавления"], key="sort_params")
                            
                            if sort_order == "По количеству вариантов":
                                filtered_params = dict(sorted(filtered_params.items(), key=lambda x: len(x[1]) if isinstance(x[1], list) else 0, reverse=True))
                            elif sort_order == "По алфавиту":
                                filtered_params = dict(sorted(filtered_params.items()))
                            
                            # Определяем иерархические параметры для выделения
                            # Иерархия параметров: Подтип - это логическая группировка для Рукав и Ворот
                            hierarchy_params = get_hierarchy_params()
                            subtype_params = get_subtype_params()  # Параметры, которые находятся внутри "Подтип"
                            visual_params = get_visual_params()
                            
                            # Разделяем параметры на категории для отображения
                            hierarchical_params_dict = {k: v for k, v in filtered_params.items() if k in hierarchy_params}
                            visual_params_dict = {k: v for k, v in filtered_params.items() if k in visual_params or any(vp in k for vp in visual_params)}
                            other_params_dict = {k: v for k, v in filtered_params.items() if k not in hierarchy_params and k not in visual_params_dict}
                            
                            # Сначала показываем иерархические параметры с выделением
                            if hierarchical_params_dict:
                                st.markdown("### 📊 Иерархические параметры (сегментация)")
                                st.info("💡 Эти параметры определяют иерархию сегментов: Тип → Комплект → Подтип (Рукав, Ворот)")
                                
                                for param_name, options in hierarchical_params_dict.items():
                                    with st.container():
                                        # Выделяем иерархические параметры цветом
                                        st.markdown(f'<div style="background-color: #e3f2fd; padding: 10px; border-radius: 5px; border-left: 4px solid #2196F3; margin-bottom: 10px;">', unsafe_allow_html=True)
                                        col_header1, col_header2, col_header3 = st.columns([3, 1, 1])
                                        
                                        with col_header1:
                                            # Редактирование названия параметра
                                            if st.session_state.get(f"editing_param_name_{param_name}", False):
                                                new_name_col1, new_name_col2, new_name_col3 = st.columns([3, 1, 1])
                                                with new_name_col1:
                                                    new_param_name = st.text_input(
                                                        "Новое название:",
                                                        value=param_name,
                                                        key=f"new_name_input_{param_name}"
                                                    )
                                                with new_name_col2:
                                                    if st.button("💾", key=f"save_name_{param_name}"):
                                                        if new_param_name and new_param_name.strip() and new_param_name != param_name:
                                                            # Переименовываем параметр
                                                            if new_param_name not in param_options:
                                                                param_options[new_param_name] = param_options.pop(param_name)
                                                                # Обновляем в param_values
                                                                if param_name in param_values:
                                                                    param_values[new_param_name] = param_values.pop(param_name)
                                                                st.session_state["param_options"] = param_options
                                                                st.session_state["param_values"] = param_values
                                                                st.success(f"✅ Параметр переименован: '{param_name}' → '{new_param_name}'")
                                                                st.rerun()
                                                            else:
                                                                st.error(f"❌ Параметр '{new_param_name}' уже существует!")
                                                        else:
                                                            st.session_state[f"editing_param_name_{param_name}"] = False
                                                            st.rerun()
                                                with new_name_col3:
                                                    if st.button("❌", key=f"cancel_name_{param_name}"):
                                                        st.session_state[f"editing_param_name_{param_name}"] = False
                                                        st.rerun()
                                            else:
                                                options_count = len(options) if isinstance(options, list) else 0
                                                # Определяем категорию параметра
                                                if param_name in subtype_params:
                                                    category = "📊 Иерархический (Подтип)"
                                                elif param_name in hierarchy_params:
                                                    category = "📊 Иерархический"
                                                elif param_name in visual_params or any(vp in param_name for vp in visual_params):
                                                    category = "🎨 Визуальный"
                                                else:
                                                    category = "📋 Другой"
                                                st.markdown(f"### {category} - {param_name} ({options_count} вариантов)")
                                        
                                        with col_header2:
                                            if st.button("✏️ Редактировать", key=f"edit_param_{param_name}"):
                                                st.session_state[f"editing_param_name_{param_name}"] = True
                                                st.rerun()
                                        
                                        with col_header3:
                                            if st.button("🗑️ Удалить", key=f"delete_param_{param_name}", type="secondary"):
                                                # Удаляем параметр напрямую из session_state
                                                if param_name in st.session_state.get("param_options", {}):
                                                    del st.session_state["param_options"][param_name]
                                                if param_name in st.session_state.get("param_values", {}):
                                                    del st.session_state["param_values"][param_name]
                                                
                                                # Также удаляем из локальных переменных
                                                if param_name in param_options:
                                                    del param_options[param_name]
                                                if param_name in param_values:
                                                    del param_values[param_name]
                                                
                                                # Добавляем в список удаленных параметров
                                                if "deleted_params" not in st.session_state:
                                                    st.session_state["deleted_params"] = load_deleted_params_from_file()
                                                
                                                # Убеждаемся, что это set
                                                if not isinstance(st.session_state["deleted_params"], set):
                                                    st.session_state["deleted_params"] = set(st.session_state["deleted_params"])
                                                
                                                # Добавляем параметр в список удаленных
                                                st.session_state["deleted_params"].add(param_name)
                                                
                                                # КРИТИЧЕСКИ ВАЖНО: Сохраняем список удаленных параметров в файл СРАЗУ
                                                # Это нужно сделать ДО очистки файлов, чтобы deleted_params.json был актуален
                                                deleted_params_list = sorted(list(st.session_state["deleted_params"]))
                                                try:
                                                    with open("deleted_params.json", "w", encoding="utf-8") as f:
                                                        json.dump(deleted_params_list, f, ensure_ascii=False, indent=2)
                                                except Exception as e:
                                                    st.error(f"❌ Ошибка сохранения deleted_params.json: {e}")
                                                
                                                # Очищаем все файлы от удаленных параметров
                                                deleted_params_set = st.session_state.get("deleted_params", set())
                                                if deleted_params_set:
                                                    cleaned_files = clean_deleted_params_from_all_files(deleted_params_set)
                                                    if cleaned_files > 0:
                                                        st.success(f"✅ Очищено {cleaned_files} файлов, удален параметр '{param_name}' из всех товаров")
                                                    else:
                                                        st.info(f"ℹ️ Параметр '{param_name}' удален. Файлы уже были очищены ранее.")
                                                
                                                # Сохраняем список удаленных параметров в файл еще раз (на всякий случай)
                                                if save_deleted_params_to_file():
                                                    # Проверяем, что параметр действительно сохранен
                                                    saved_params = load_deleted_params_from_file()
                                                    if param_name in saved_params:
                                                        st.success(f"✅ Параметр '{param_name}' удален и сохранен в deleted_params.json")
                                                    else:
                                                        st.error(f"❌ Параметр '{param_name}' не был сохранен в deleted_params.json!")
                                                    
                                                    # Удаляем параметр из результатов массового анализа
                                                    if "mass_analysis_results" in st.session_state:
                                                        st.session_state["mass_analysis_results"] = remove_deleted_params_from_mass_results(
                                                            st.session_state["mass_analysis_results"]
                                                        )
                                                        # Сохраняем очищенные результаты
                                                        save_mass_analysis_results(st.session_state["mass_analysis_results"])
                                                    
                                                    # Сохраняем очищенные параметры в файлы (это удалит удаленные параметры из сохраненных файлов)
                                                if save_param_values_to_file():
                                                        st.info(f"ℹ️ Параметры сохранены в файлы")
                                                else:
                                                    st.error(f"❌ Ошибка сохранения списка удаленных параметров! Попробуйте нажать кнопку '💾 Сохранить все'.")
                                                
                                                # Принудительно удаляем параметр из session_state еще раз (на всякий случай)
                                                if param_name in st.session_state.get("param_options", {}):
                                                    del st.session_state["param_options"][param_name]
                                                if param_name in st.session_state.get("param_values", {}):
                                                    del st.session_state["param_values"][param_name]
                                                
                                                st.rerun()
                                    
                                    # Отображение вариантов
                                    if isinstance(options, list) and options:
                                        st.markdown("**Варианты значений:**")
                                        
                                        # Редактируемая таблица вариантов
                                        for idx, option in enumerate(options):
                                            col_opt1, col_opt2, col_opt3, col_opt4 = st.columns([4, 1, 1, 1])
                                            
                                            with col_opt1:
                                                if st.session_state.get(f"editing_option_{param_name}_{idx}", False):
                                                    edited_opt = st.text_input(
                                                        "Редактировать:",
                                                        value=option,
                                                        key=f"edit_opt_input_{param_name}_{idx}"
                                                    )
                                                else:
                                                    st.write(f"• {option}")
                                            
                                            with col_opt2:
                                                if st.session_state.get(f"editing_option_{param_name}_{idx}", False):
                                                    if st.button("💾", key=f"save_opt_{param_name}_{idx}"):
                                                        if edited_opt and edited_opt.strip():
                                                            options[idx] = edited_opt.strip()
                                                            st.session_state["param_options"] = param_options
                                                            st.session_state[f"editing_option_{param_name}_{idx}"] = False
                                                            if save_param_values_to_file():
                                                                st.success("✅ Вариант обновлен!")
                                                            st.rerun()
                                                else:
                                                    if st.button("✏️", key=f"edit_opt_{param_name}_{idx}"):
                                                        st.session_state[f"editing_option_{param_name}_{idx}"] = True
                                                        st.rerun()
                                            
                                            with col_opt3:
                                                if st.session_state.get(f"editing_option_{param_name}_{idx}", False):
                                                    if st.button("❌", key=f"cancel_opt_{param_name}_{idx}"):
                                                        st.session_state[f"editing_option_{param_name}_{idx}"] = False
                                                        st.rerun()
                                            
                                            with col_opt4:
                                                if st.button("🗑️", key=f"del_opt_{param_name}_{idx}"):
                                                    options.pop(idx)
                                                    if not options:
                                                        # Если вариантов не осталось, удаляем параметр
                                                        if param_name in param_options:
                                                            del param_options[param_name]
                                                        if param_name in param_values:
                                                            del param_values[param_name]
                                                        
                                                        # Добавляем в список удаленных параметров
                                                        if "deleted_params" not in st.session_state:
                                                            st.session_state["deleted_params"] = load_deleted_params_from_file()
                                                        
                                                        # Убеждаемся, что это set
                                                        if not isinstance(st.session_state["deleted_params"], set):
                                                            st.session_state["deleted_params"] = set(st.session_state["deleted_params"])
                                                        
                                                        # Добавляем параметр в список удаленных
                                                        st.session_state["deleted_params"].add(param_name)
                                                        
                                                        # Очищаем все файлы от удаленных параметров перед сохранением
                                                        deleted_params_set = st.session_state.get("deleted_params", set())
                                                        if deleted_params_set:
                                                            cleaned_files = clean_deleted_params_from_all_files(deleted_params_set)
                                                            if cleaned_files > 0:
                                                                st.info(f"✅ Очищено {cleaned_files} файлов, удален параметр '{param_name}' из всех товаров")
                                                        
                                                        # Сохраняем список удаленных параметров в файл
                                                        if save_deleted_params_to_file():
                                                            # Удаляем параметр из результатов массового анализа
                                                            if "mass_analysis_results" in st.session_state:
                                                                st.session_state["mass_analysis_results"] = remove_deleted_params_from_mass_results(
                                                                    st.session_state["mass_analysis_results"]
                                                                )
                                                                # Сохраняем очищенные результаты
                                                                save_mass_analysis_results(st.session_state["mass_analysis_results"])
                                                        
                                                            # Сохраняем очищенные параметры в файлы
                                                            if save_param_values_to_file():
                                                                st.success("✅ Параметр удален (варианты закончились)!")
                                                            else:
                                                                st.warning("⚠️ Параметр удален, но не сохранен в файл")
                                                        else:
                                                            st.error("❌ Ошибка сохранения списка удаленных параметров!")
                                                    else:
                                                        # Вариантов еще есть, просто обновляем список
                                                        st.session_state["param_options"] = param_options
                                                        st.session_state["param_values"] = param_values
                                                        if save_param_values_to_file():
                                                            st.success("✅ Вариант удален!")
                                                        st.rerun()
                                        
                                        # Добавление нового варианта
                                        col_add1, col_add2 = st.columns([4, 1])
                                        with col_add1:
                                            new_opt = st.text_input(
                                                f"Добавить вариант в '{param_name}':",
                                                placeholder="Введите новый вариант",
                                                key=f"new_opt_{param_name}"
                                            )
                                        with col_add2:
                                            if st.button("➕", key=f"add_opt_{param_name}"):
                                                if new_opt and new_opt.strip():
                                                    if new_opt.strip() not in options:
                                                        options.append(new_opt.strip())
                                                        st.session_state["param_options"] = param_options
                                                        if save_param_values_to_file():
                                                            st.success(f"✅ Вариант '{new_opt.strip()}' добавлен!")
                                                        st.rerun()
                                                    else:
                                                        st.warning("Этот вариант уже существует")
                                    else:
                                        st.info("ℹ️ Варианты не указаны - нейросеть определит значение самостоятельно")
                                    
                                    st.divider()
                                    st.markdown('</div>', unsafe_allow_html=True)
                            
                            # Затем показываем визуальные параметры
                            if visual_params_dict:
                                st.markdown("### 🎨 Визуальные параметры (анализируются в контексте иерархии)")
                                st.info("💡 Эти параметры анализируются только после определения иерархии сегментов")
                                
                                for param_name, options in visual_params_dict.items():
                                    with st.container():
                                        st.markdown(f'<div style="background-color: #fff3e0; padding: 10px; border-radius: 5px; border-left: 4px solid #FF9800; margin-bottom: 10px;">', unsafe_allow_html=True)
                                        col_header1, col_header2, col_header3 = st.columns([3, 1, 1])
                                        
                                        with col_header1:
                                            options_count = len(options) if isinstance(options, list) else 0
                                            category = "🎨 Визуальный"
                                            st.markdown(f"### {category} - {param_name} ({options_count} вариантов)")
                                        
                                        with col_header2:
                                            if st.button("✏️ Редактировать", key=f"edit_param_{param_name}"):
                                                st.session_state[f"editing_param_name_{param_name}"] = True
                                                st.rerun()
                                        
                                        with col_header3:
                                            if st.button("🗑️ Удалить", key=f"delete_param_{param_name}", type="secondary"):
                                                if param_name in param_options:
                                                    del param_options[param_name]
                                                if param_name in param_values:
                                                    del param_values[param_name]
                                                st.session_state["param_options"] = param_options
                                                st.session_state["param_values"] = param_values
                                                if save_param_values_to_file():
                                                    st.success(f"✅ Параметр '{param_name}' удален!")
                                                st.rerun()
                                        
                                        # Отображение вариантов (копируем логику из иерархических параметров)
                                        if isinstance(options, list) and options:
                                            st.markdown("**Варианты значений:**")
                                            for idx, option in enumerate(options):
                                                col_opt1, col_opt2, col_opt3, col_opt4 = st.columns([4, 1, 1, 1])
                                                
                                                with col_opt1:
                                                    if st.session_state.get(f"editing_option_{param_name}_{idx}", False):
                                                        edited_opt = st.text_input(
                                                            "Редактировать:",
                                                            value=option,
                                                            key=f"edit_opt_input_{param_name}_{idx}"
                                                        )
                                                    else:
                                                        st.write(f"• {option}")
                                                
                                                with col_opt2:
                                                    if st.session_state.get(f"editing_option_{param_name}_{idx}", False):
                                                        if st.button("💾", key=f"save_opt_{param_name}_{idx}"):
                                                            if edited_opt and edited_opt.strip():
                                                                options[idx] = edited_opt.strip()
                                                                st.session_state["param_options"] = param_options
                                                                st.session_state[f"editing_option_{param_name}_{idx}"] = False
                                                                if save_param_values_to_file():
                                                                    st.success("✅ Вариант обновлен!")
                                                                st.rerun()
                                                    else:
                                                        if st.button("✏️", key=f"edit_opt_{param_name}_{idx}"):
                                                            st.session_state[f"editing_option_{param_name}_{idx}"] = True
                                                            st.rerun()
                                                
                                                with col_opt3:
                                                    if st.session_state.get(f"editing_option_{param_name}_{idx}", False):
                                                        if st.button("❌", key=f"cancel_opt_{param_name}_{idx}"):
                                                            st.session_state[f"editing_option_{param_name}_{idx}"] = False
                                                            st.rerun()
                                                
                                                with col_opt4:
                                                    if st.button("🗑️", key=f"del_opt_{param_name}_{idx}"):
                                                        options.pop(idx)
                                                        if not options:
                                                            if param_name in param_options:
                                                                del param_options[param_name]
                                                            if param_name in param_values:
                                                                del param_values[param_name]
                                                        st.session_state["param_options"] = param_options
                                                        st.session_state["param_values"] = param_values
                                                        if save_param_values_to_file():
                                                            st.success("✅ Вариант удален!")
                                                        st.rerun()
                                            
                                            # Добавление нового варианта
                                            col_add1, col_add2 = st.columns([4, 1])
                                            with col_add1:
                                                new_opt = st.text_input(
                                                    f"Добавить вариант в '{param_name}':",
                                                    placeholder="Введите новый вариант",
                                                    key=f"new_opt_{param_name}"
                                                )
                                            with col_add2:
                                                if st.button("➕", key=f"add_opt_{param_name}"):
                                                    if new_opt and new_opt.strip():
                                                        if new_opt.strip() not in options:
                                                            options.append(new_opt.strip())
                                                            st.session_state["param_options"] = param_options
                                                            if save_param_values_to_file():
                                                                st.success(f"✅ Вариант '{new_opt.strip()}' добавлен!")
                                                            st.rerun()
                                                        else:
                                                            st.warning("Этот вариант уже существует")
                                        else:
                                            st.info("ℹ️ Варианты не указаны - нейросеть определит значение самостоятельно")
                                        
                                        st.divider()
                                        st.markdown('</div>', unsafe_allow_html=True)
                            
                            # Затем показываем остальные параметры
                            if other_params_dict:
                                st.markdown("### 📋 Дополнительные параметры")
                                
                                for param_name, options in other_params_dict.items():
                                    with st.container():
                                        col_header1, col_header2, col_header3 = st.columns([3, 1, 1])
                                        
                                        with col_header1:
                                            options_count = len(options) if isinstance(options, list) else 0
                                            category = "📋 Другой"
                                            st.markdown(f"### {category} - {param_name} ({options_count} вариантов)")
                                        
                                        with col_header2:
                                            if st.button("✏️ Редактировать", key=f"edit_param_{param_name}"):
                                                st.session_state[f"editing_param_name_{param_name}"] = True
                                                st.rerun()
                                        
                                        with col_header3:
                                            if st.button("🗑️ Удалить", key=f"delete_param_{param_name}", type="secondary"):
                                                if param_name in param_options:
                                                    del param_options[param_name]
                                                if param_name in param_values:
                                                    del param_values[param_name]
                                                st.session_state["param_options"] = param_options
                                                st.session_state["param_values"] = param_values
                                                if save_param_values_to_file():
                                                    st.success(f"✅ Параметр '{param_name}' удален!")
                                                st.rerun()
                                        
                                        # Отображение вариантов (аналогично выше)
                                        if isinstance(options, list) and options:
                                            st.markdown("**Варианты значений:**")
                                            for idx, option in enumerate(options):
                                                col_opt1, col_opt2, col_opt3, col_opt4 = st.columns([4, 1, 1, 1])
                                                
                                                with col_opt1:
                                                    if st.session_state.get(f"editing_option_{param_name}_{idx}", False):
                                                        edited_opt = st.text_input(
                                                            "Редактировать:",
                                                            value=option,
                                                            key=f"edit_opt_input_{param_name}_{idx}"
                                                        )
                                                    else:
                                                        st.write(f"• {option}")
                                                
                                                with col_opt2:
                                                    if st.session_state.get(f"editing_option_{param_name}_{idx}", False):
                                                        if st.button("💾", key=f"save_opt_{param_name}_{idx}"):
                                                            if edited_opt and edited_opt.strip():
                                                                options[idx] = edited_opt.strip()
                                                                st.session_state["param_options"] = param_options
                                                                st.session_state[f"editing_option_{param_name}_{idx}"] = False
                                                                if save_param_values_to_file():
                                                                    st.success("✅ Вариант обновлен!")
                                                                st.rerun()
                                                    else:
                                                        if st.button("✏️", key=f"edit_opt_{param_name}_{idx}"):
                                                            st.session_state[f"editing_option_{param_name}_{idx}"] = True
                                                            st.rerun()
                                                
                                                with col_opt3:
                                                    if st.session_state.get(f"editing_option_{param_name}_{idx}", False):
                                                        if st.button("❌", key=f"cancel_opt_{param_name}_{idx}"):
                                                            st.session_state[f"editing_option_{param_name}_{idx}"] = False
                                                            st.rerun()
                                                
                                                with col_opt4:
                                                    if st.button("🗑️", key=f"del_opt_{param_name}_{idx}"):
                                                        options.pop(idx)
                                                        if not options:
                                                            if param_name in param_options:
                                                                del param_options[param_name]
                                                            if param_name in param_values:
                                                                del param_values[param_name]
                                                        st.session_state["param_options"] = param_options
                                                        st.session_state["param_values"] = param_values
                                                        if save_param_values_to_file():
                                                            st.success("✅ Вариант удален!")
                                                        st.rerun()
                                            
                                            # Добавление нового варианта
                                            col_add1, col_add2 = st.columns([4, 1])
                                            with col_add1:
                                                new_opt = st.text_input(
                                                    f"Добавить вариант в '{param_name}':",
                                                    placeholder="Введите новый вариант",
                                                    key=f"new_opt_{param_name}"
                                                )
                                            with col_add2:
                                                if st.button("➕", key=f"add_opt_{param_name}"):
                                                    if new_opt and new_opt.strip():
                                                        if new_opt.strip() not in options:
                                                            options.append(new_opt.strip())
                                                            st.session_state["param_options"] = param_options
                                                            if save_param_values_to_file():
                                                                st.success(f"✅ Вариант '{new_opt.strip()}' добавлен!")
                                                            st.rerun()
                                                        else:
                                                            st.warning("Этот вариант уже существует")
                                        else:
                                            st.info("ℹ️ Варианты не указаны - нейросеть определит значение самостоятельно")
                                        
                                        st.divider()
                        else:
                            st.info("🔍 Параметры не найдены по запросу")
                    else:
                        st.info("📝 Параметры еще не добавлены. Используйте режим '➕ Добавить новый параметр'")
                
                elif editor_mode == "➕ Добавить новый параметр":
                    st.markdown("### ➕ Добавить новый параметр")
                    
                    # Готовые шаблоны параметров
                    param_templates = {
                        "Комплект": ["Комплект", "Один"],
                        "Цвет": [
                            "Белый", "Черный", "Серый", "Синий", "Красный", "Зеленый", "Розовый", "Бежевый",
                            "Черно-белый", "Бело-черный", "Сине-белый", "Бело-синий",
                            "Красно-белый", "Бело-красный", "Черно-серый", "Серо-черный"
                        ],
                        "Длина": ["Короткая", "Средняя", "Длинная", "Мини", "Макси", "Миди", "Анкл"],
                        "Пуговицы": ["С пуговицами", "Без пуговиц", "На молнии", "На липучке", "На кнопках", "На шнурке"],
                        "Материал": ["Хлопок", "Полиэстер", "Шерсть", "Лен", "Джинс", "Кожа", "Замша", "Трикотаж"],
                        "Размер": ["XS", "S", "M", "L", "XL", "XXL", "XXXL"],
                        "Сезон": ["Лето", "Зима", "Весна", "Осень", "Демисезон"],
                        "Стиль": ["Классический", "Спортивный", "Повседневный", "Деловой", "Вечерний", "Романтический"],
                        "Тип": []  # Без вариантов
                    }
                    
                    col_template1, col_template2 = st.columns([2, 1])
                    with col_template1:
                        template_choice = st.selectbox(
                            "Выберите шаблон или создайте новый:",
                            ["Создать новый параметр"] + list(param_templates.keys()),
                            key="template_choice_editor"
                        )
                    
                    with col_template2:
                        if template_choice != "Создать новый параметр":
                            if st.button("📋 Применить шаблон"):
                                st.session_state["temp_param_name_editor"] = template_choice
                                st.session_state["temp_param_options_editor"] = " / ".join(param_templates[template_choice]) if param_templates[template_choice] else ""
                    
                    col_name, col_options = st.columns([1, 2])
                    with col_name:
                        param_name_new = st.text_input(
                            "Название параметра:",
                            value=st.session_state.get("temp_param_name_editor", ""),
                            placeholder="Например: Тип, Цвет, Материал",
                            key="param_name_new_editor"
                        )
                    
                    with col_options:
                        param_options_new = st.text_area(
                            "Варианты (через слэш /, можно оставить пустым):",
                            value=st.session_state.get("temp_param_options_editor", ""),
                            placeholder="Красный / Синий / Зеленый\n(или оставьте пустым для свободного ввода)",
                            height=100,
                            key="param_options_new_editor"
                        )
                    
                    col_add_btn, col_clear_btn = st.columns([1, 1])
                    with col_add_btn:
                        if st.button("➕ Добавить параметр", type="primary", use_container_width=True):
                            if param_name_new and param_name_new.strip():
                                # Очищаем и разбираем варианты
                                if param_options_new:
                                    options = [opt.strip() for opt in param_options_new.split("/") if opt.strip()]
                                else:
                                    options = []
                                
                                # Проверяем, не существует ли уже такой параметр
                                if param_name_new.strip() in param_options:
                                    st.error(f"❌ Параметр '{param_name_new.strip()}' уже существует!")
                                else:
                                    # Добавляем параметр
                                    param_options[param_name_new.strip()] = options
                                    if param_name_new.strip() not in param_values:
                                        param_values[param_name_new.strip()] = {}
                                    
                                    st.session_state["param_options"] = param_options
                                    st.session_state["param_values"] = param_values
                                    
                                    if save_param_values_to_file():
                                        st.success(f"✅ Параметр '{param_name_new.strip()}' добавлен! ({len(options)} вариантов)")
                                        # Очищаем поля
                                        if "temp_param_name_editor" in st.session_state:
                                            del st.session_state["temp_param_name_editor"]
                                        if "temp_param_options_editor" in st.session_state:
                                            del st.session_state["temp_param_options_editor"]
                                        st.rerun()
                            else:
                                st.warning("⚠️ Введите название параметра")
                    
                    with col_clear_btn:
                        if st.button("🗑️ Очистить", use_container_width=True):
                            if "temp_param_name_editor" in st.session_state:
                                del st.session_state["temp_param_name_editor"]
                            if "temp_param_options_editor" in st.session_state:
                                del st.session_state["temp_param_options_editor"]
                            st.rerun()
                
                elif editor_mode == "🔗 Управление иерархией":
                    st.markdown("### 🔗 Управление иерархией параметров")
                    st.info("💡 Настройте иерархию параметров для правильной структуризации анализа. Иерархические параметры влияют на порядок определения параметров нейросетью.")
                    
                    # Загружаем конфигурацию иерархии при первом запуске
                    if "hierarchy_params" not in st.session_state:
                        load_hierarchy_config()
                    
                    # Получаем текущие значения
                    current_hierarchy = get_hierarchy_params()
                    current_subtype = get_subtype_params()
                    current_visual = get_visual_params()
                    
                    # Доступные параметры для выбора
                    available_params_list = list(param_options.keys())
                    
                    col_hier1, col_hier2 = st.columns(2)
                    
                    with col_hier1:
                        st.markdown("#### 📊 Иерархические параметры")
                        st.write("Параметры, которые определяются первыми и влияют на структуру анализа:")
                        
                        hierarchy_params_selected = st.multiselect(
                            "Выберите иерархические параметры (порядок важен):",
                            options=available_params_list,
                            default=current_hierarchy,
                            help="Порядок параметров в списке определяет порядок в иерархии. Первый параметр - самый верхний уровень.",
                            key="hierarchy_params_multiselect"
                        )
                        
                        # Показываем текущую иерархию
                        if hierarchy_params_selected:
                            st.success(f"✅ Иерархия: {' → '.join(hierarchy_params_selected)}")
                        else:
                            st.warning("⚠️ Не выбрано ни одного иерархического параметра")
                    
                    with col_hier2:
                        st.markdown("#### 🔗 Параметры подтипа")
                        st.write("Параметры, которые группируются внутри 'Подтип':")
                        
                        # Фильтруем только те параметры, которые есть в иерархии
                        subtype_candidates = [p for p in available_params_list if p in hierarchy_params_selected]
                        
                        if subtype_candidates:
                            subtype_params_selected = st.multiselect(
                                "Выберите параметры подтипа:",
                                options=subtype_candidates,
                                default=[p for p in current_subtype if p in subtype_candidates],
                                help="Эти параметры будут сгруппированы как 'Подтип' в иерархии",
                                key="subtype_params_multiselect"
                            )
                        else:
                            st.info("ℹ️ Сначала выберите иерархические параметры")
                            subtype_params_selected = []
                    
                    st.divider()
                    
                    st.markdown("#### 🎨 Визуальные параметры")
                    st.write("Параметры, описывающие внешний вид товара (цвет, принт, строчка и т.д.):")
                    
                    visual_params_selected = st.multiselect(
                        "Выберите визуальные параметры:",
                        options=available_params_list,
                        default=current_visual,
                        help="Визуальные параметры определяются после иерархических",
                        key="visual_params_multiselect"
                    )
                    
                    st.divider()
                    
                    # Кнопки сохранения
                    col_save_hier1, col_save_hier2 = st.columns([1, 1])
                    with col_save_hier1:
                        if st.button("💾 Сохранить иерархию", type="primary", use_container_width=True):
                            st.session_state["hierarchy_params"] = hierarchy_params_selected
                            st.session_state["subtype_params"] = subtype_params_selected
                            st.session_state["visual_params"] = visual_params_selected
                            
                            if save_hierarchy_config():
                                st.success("✅ Иерархия параметров сохранена!")
                                st.rerun()
                            else:
                                st.warning("⚠️ Иерархия сохранена в памяти, но не в файл")
                    
                    with col_save_hier2:
                        if st.button("🔄 Сбросить к значениям по умолчанию", use_container_width=True):
                            st.session_state["hierarchy_params"] = ["Тип", "Комплект", "Рукав", "Ворот"]
                            st.session_state["subtype_params"] = ["Рукав", "Ворот"]
                            st.session_state["visual_params"] = ["Цвет", "Принт", "Логотип", "Строчка", "Шов", "Строчка шов"]
                            
                            if save_hierarchy_config():
                                st.success("✅ Иерархия сброшена к значениям по умолчанию!")
                                st.rerun()
                    
                    # Показываем текущую конфигурацию
                    st.markdown("#### 📋 Текущая конфигурация иерархии")
                    col_curr1, col_curr2, col_curr3 = st.columns(3)
                    with col_curr1:
                        st.write("**Иерархические:**")
                        if hierarchy_params_selected:
                            for i, param in enumerate(hierarchy_params_selected, 1):
                                st.write(f"{i}. {param}")
                        else:
                            st.write("Не выбрано")
                    
                    with col_curr2:
                        st.write("**Подтип:**")
                        if subtype_params_selected:
                            for param in subtype_params_selected:
                                st.write(f"• {param}")
                        else:
                            st.write("Не выбрано")
                    
                    with col_curr3:
                        st.write("**Визуальные:**")
                        if visual_params_selected:
                            for param in visual_params_selected:
                                st.write(f"• {param}")
                        else:
                            st.write("Не выбрано")
                
                elif editor_mode == "📥 Импорт / 📤 Экспорт":
                    st.markdown("### 📥 Импорт / 📤 Экспорт параметров")
                    
                    # Кнопка сохранения параметров
                    st.divider()
                    col_save_editor1, col_save_editor2 = st.columns([1, 3])
                    with col_save_editor1:
                        if st.button("💾 Сохранить все параметры", type="primary", use_container_width=True):
                            # Сначала сохраняем deleted_params (это очистит все файлы)
                            if save_deleted_params_to_file():
                                # Затем сохраняем param_options и param_values
                                if save_param_values_to_file():
                                    # Очищаем результаты массового анализа от удаленных параметров
                                    if "mass_analysis_results" in st.session_state:
                                        st.session_state["mass_analysis_results"] = remove_deleted_params_from_mass_results(
                                            st.session_state["mass_analysis_results"]
                                        )
                                        save_mass_analysis_results(st.session_state["mass_analysis_results"])
                                    st.success("✅ Все параметры и настройки сохранены!")
                                    st.rerun()
                                else:
                                    st.error("❌ Ошибка сохранения параметров")
                            else:
                                st.error("❌ Ошибка сохранения списка удаленных параметров")
                    
                    with col_save_editor2:
                        st.info("💡 Сохраняет все параметры, варианты и список удаленных параметров в файлы")
                        
                        # Показываем текущие удаленные параметры для отладки
                        deleted_params = st.session_state.get("deleted_params", set())
                        if deleted_params:
                            st.caption(f"🗑️ Удаленные параметры: {', '.join(sorted(deleted_params))}")
                        else:
                            st.caption("ℹ️ Нет удаленных параметров")
                    
                    st.divider()
                    
                    col_export, col_import = st.columns(2)
                    
                    with col_export:
                        st.markdown("#### 📤 Экспорт параметров")
                        if st.button("💾 Экспортировать в JSON", use_container_width=True):
                            export_data = {
                                "param_options": param_options,
                                "param_values": param_values,
                                "export_date": pd.Timestamp.now().strftime("%Y-%m-%d %H:%M:%S")
                            }
                            st.download_button(
                                label="⬇️ Скачать файл",
                                data=json.dumps(export_data, ensure_ascii=False, indent=2),
                                file_name=f"parameters_export_{pd.Timestamp.now().strftime('%Y%m%d_%H%M%S')}.json",
                                mime="application/json",
                                use_container_width=True
                            )
                    
                    with col_import:
                        st.markdown("#### 📥 Импорт параметров")
                        uploaded_file = st.file_uploader(
                            "Загрузить JSON файл:",
                            type=["json"],
                            key="import_params_file"
                        )
                        if uploaded_file:
                            try:
                                import_data = json.load(uploaded_file)
                                if "param_options" in import_data:
                                    st.success(f"✅ Файл загружен: {len(import_data.get('param_options', {}))} параметров")
                                    
                                    if st.button("📥 Импортировать параметры", type="primary", use_container_width=True):
                                        # Объединяем с существующими параметрами
                                        imported_options = import_data.get("param_options", {})
                                        imported_values = import_data.get("param_values", {})
                                        
                                        # Обновляем param_options
                                        for param_name, options in imported_options.items():
                                            param_options[param_name] = options
                                        
                                        # Обновляем param_values
                                        for param_name, values in imported_values.items():
                                            if param_name not in param_values:
                                                param_values[param_name] = {}
                                            param_values[param_name].update(values)
                                        
                                        st.session_state["param_options"] = param_options
                                        st.session_state["param_values"] = param_values
                                        
                                        if save_param_values_to_file():
                                            st.success(f"✅ Импортировано {len(imported_options)} параметров!")
                                            st.rerun()
                            except Exception as e:
                                st.error(f"❌ Ошибка при импорте: {e}")
            
            st.divider()
            # ========== КОНЕЦ РЕДАКТОРА ПАРАМЕТРОВ ==========
            
            # Получаем данные
            param_values = get_param_values()
            
            # Инициализация session_state для автосохранения (отключено)
            if "last_autosave" not in st.session_state:
                st.session_state["last_autosave"] = 0
            
            # Автосохранение включено
            import time
            current_time = time.time()
            if current_time - st.session_state["last_autosave"] > 30:  # 30 секунд
                if save_param_values_to_file():
                    st.session_state["last_autosave"] = current_time
                    st.info("🔄 Автосохранение выполнено")
            
            # ========== СЕКЦИЯ: Определение параметров по ссылке (через анализ изображений) ==========
            st.divider()
            st.subheader("🔗 Определение параметров по ссылке (анализ через нейросеть)")
            st.write("Вставьте ссылку на товар Wildberries. Система получит до 5 фотографий товара через screenshotapi.net и отправит их в нейросеть для анализа:")
            
            # Поле для API ключа OpenAI
            # Загружаем ключ из secrets.toml при первом запуске
            if 'openai_api_key' not in st.session_state:
                try:
                    # Пытаемся загрузить из secrets.toml
                    default_key = st.secrets.get('openai_api_key', '')
                    if default_key:
                        st.session_state['openai_api_key'] = default_key
                except:
                    pass
            
            with st.expander("⚙️ Настройки API", expanded=False):
                # Проверяем наличие библиотеки openai
                if not OPENAI_AVAILABLE:
                    st.warning("⚠️ Для анализа изображений через нейросеть необходимо установить библиотеку: `pip install openai`")
                    st.code("pip install openai", language="bash")
                else:
                    # Режим отладки
                    debug_mode = st.checkbox(
                        "🔍 Режим отладки",
                        value=st.session_state.get('debug_ai_analysis', False),
                        help="Показывать детальную информацию об анализе изображений",
                        key="debug_ai_checkbox"
                    )
                    if debug_mode:
                        st.session_state['debug_ai_analysis'] = True
                    else:
                        st.session_state['debug_ai_analysis'] = False
                default_key_value = st.session_state.get('openai_api_key', '')
                openai_api_key = st.text_input(
                    "🔑 OpenAI API ключ",
                    value=default_key_value,
                    type="password",
                    help="Введите ваш OpenAI API ключ для анализа изображений. Ключ сохраняется только в текущей сессии.",
                    key="openai_api_key_input"
                )
                if openai_api_key:
                    st.session_state['openai_api_key'] = openai_api_key
                    if openai_api_key == default_key_value and default_key_value:
                        st.success("✅ Используется стандартный API ключ")
                    else:
                        st.success("✅ API ключ сохранен")
                else:
                    if default_key_value:
                        st.info("ℹ️ Используется стандартный API ключ из настроек")
                    else:
                        st.info("ℹ️ Для работы с нейросетью необходим OpenAI API ключ. Получить можно на https://platform.openai.com/api-keys")
                
                # Поле для токена screenshotapi.net
                st.divider()
                st.markdown("**📸 ScreenshotAPI.net (для получения изображений товаров)**")
                
                # Загружаем токен из secrets.toml при первом запуске
                if 'screenshotapi_token' not in st.session_state:
                    try:
                        default_token = st.secrets.get('screenshotapi_token', '')
                        if default_token:
                            st.session_state['screenshotapi_token'] = default_token
                        else:
                            # Устанавливаем токен по умолчанию
                            st.session_state['screenshotapi_token'] = 'MDA94V3-4C14RXS-KSE2KMK-T08BD9M'
                    except:
                        # Устанавливаем токен по умолчанию
                        st.session_state['screenshotapi_token'] = 'MDA94V3-4C14RXS-KSE2KMK-T08BD9M'
                
                default_token_value = st.session_state.get('screenshotapi_token', 'MDA94V3-4C14RXS-KSE2KMK-T08BD9M')
                screenshotapi_token = st.text_input(
                    "🔑 ScreenshotAPI.net токен",
                    value=default_token_value,
                    type="password",
                    help="Введите ваш токен screenshotapi.net для получения изображений товаров. Токен сохраняется только в текущей сессии.",
                    key="screenshotapi_token_input"
                )
                if screenshotapi_token:
                    st.session_state['screenshotapi_token'] = screenshotapi_token
                    if screenshotapi_token == default_token_value and default_token_value:
                        st.success("✅ Используется стандартный токен")
                    else:
                        st.success("✅ Токен сохранен")
                else:
                    if default_token_value:
                        st.info("ℹ️ Используется стандартный токен из настроек")
                    else:
                        st.info("ℹ️ Для получения изображений через screenshotapi.net необходим токен. Получить можно на https://screenshotapi.net")
            
            # Переключатель между одиночным и массовым анализом
            analysis_mode = st.radio(
                "Режим анализа:",
                ["Массовый анализ", "Одиночный анализ"],
                horizontal=True,
                index=0,
                key="analysis_mode_radio"
            )
            
            if analysis_mode == "Одиночный анализ":
                col_url1, col_url2 = st.columns([3, 1])
                with col_url1:
                    product_url = st.text_input(
                        "Ссылка на товар Wildberries",
                        placeholder="https://www.wildberries.ru/catalog/12345678/detail.aspx",
                        key="product_url_input"
                    )
                
                with col_url2:
                    st.write("")  # Отступ
                    st.write("")  # Отступ
                    analyze_btn = st.button("🤖 Анализировать через нейросеть", type="primary", key="extract_params_btn")
                
                if analyze_btn:
                    if product_url:
                        with st.spinner("⏳ Анализ фотографий товара..."):
                            # Извлекаем артикул
                            sku = extract_sku_from_url(product_url)
                            
                            if sku:
                                st.success(f"✅ Артикул найден: **{sku}**")
                                
                                # Получаем и анализируем изображения
                                progress_bar = st.progress(0)
                                status_text = st.empty()
                                
                                # ПЕРВЫЙ ЭТАП: Определение комплектности по характеристикам
                                status_text.text("🔍 Анализ характеристик товара для определения комплектности...")
                                progress_bar.progress(10)
                                
                                completeness = determine_completeness(sku)
                                initial_params = {}
                                if completeness:
                                    initial_params["Комплект"] = completeness
                                    st.info(f"📦 **Комплектность:** {completeness}")
                                
                                # ВТОРОЙ ЭТАП: Получение изображений через screenshotapi.net
                                status_text.text("📸 Получение изображений через screenshotapi.net...")
                                progress_bar.progress(20)
                                
                                # Получаем информацию о статусе изображений
                                images_info = get_screenshotapi_images_with_status(sku, max_images=5)
                                
                                if not images_info:
                                    st.error("❌ Не удалось получить URL-ы изображений. Проверьте токен screenshotapi.net.")
                                else:
                                    # Разделяем изображения на те, что в кеше, и те, что нужно скачать
                                    cached_images = [img for img in images_info if img["status"] == "cached"]
                                    downloading_images = [img for img in images_info if img["status"] == "downloading"]
                                    
                                    # Показываем информацию о статусе изображений
                                    st.markdown("### 📸 Статус изображений (5 шт.)")
                                    
                                    # Создаем контейнеры для отображения статуса
                                    status_cols = st.columns(2)
                                    
                                    with status_cols[0]:
                                        st.markdown(f"**✅ В кеше:** {len(cached_images)} из 5")
                                        if cached_images:
                                            for img_info in cached_images:
                                                st.success(f"📷 Фото {img_info['index']}: В кеше")
                                    
                                    with status_cols[1]:
                                        st.markdown(f"**⬇️ Скачивается:** {len(downloading_images)} из 5")
                                        if downloading_images:
                                            for img_info in downloading_images:
                                                st.info(f"📷 Фото {img_info['index']}: Скачивается...")
                                    
                                    # Скачиваем изображения, которых нет в кеше
                                    if downloading_images:
                                        status_text.text(f"⬇️ Скачивание {len(downloading_images)} изображений...")
                                        progress_bar.progress(30)
                                        
                                        downloaded_count = 0
                                        download_status_placeholder = st.empty()
                                        
                                        for img_idx, img_info in enumerate(downloading_images):
                                            status_text.text(f"⬇️ Скачивание фото {img_info['index']}/5...")
                                            progress_bar.progress(30 + int((img_idx + 1) / len(downloading_images) * 20))
                                            
                                            # Увеличиваем timeout для первого изображения (оно может загружаться дольше)
                                            timeout_value = 45 if img_info["index"] == 1 else 30
                                            cached_path = ensure_image_cached(
                                                img_info["cache_key"], 
                                                img_info["url"], 
                                                "PNG", 
                                                timeout=timeout_value
                                            )
                                            if cached_path and os.path.exists(cached_path):
                                                img_info["path"] = cached_path
                                                img_info["status"] = "downloaded"
                                                downloaded_count += 1
                                                # Обновляем статус в реальном времени
                                                with download_status_placeholder.container():
                                                    st.success(f"✅ Фото {img_info['index']}: Скачано")
                                            else:
                                                img_info["status"] = "error"
                                                with download_status_placeholder.container():
                                                    st.error(f"❌ Фото {img_info['index']}: Ошибка скачивания")
                                            
                                            # Задержка между запросами
                                            import time
                                            time.sleep(0.5)
                                        
                                        if downloaded_count > 0:
                                            st.success(f"✅ Скачано изображений: {downloaded_count} из {len(downloading_images)}")
                                        elif len(downloading_images) > 0:
                                            st.warning(f"⚠️ Не удалось скачать изображения. Проверьте токен screenshotapi.net и интернет-соединение.")
                                    
                                    # Собираем все успешно загруженные изображения
                                    all_loaded_images = []
                                    for img_info in images_info:
                                        if img_info["status"] in ["cached", "downloaded"] and img_info.get("path"):
                                            all_loaded_images.append(img_info["path"])
                                    
                                    if not all_loaded_images:
                                        st.error("❌ Не удалось загрузить ни одного изображения")
                                    else:
                                        st.success(f"✅ Готово к анализу: {len(all_loaded_images)} из 5 изображений")
                                        
                                        # Показываем все загруженные изображения
                                        st.markdown("### 📸 Загруженные изображения")
                                        # Создаем колонки для отображения изображений (по 3 в ряд)
                                        num_cols = min(3, len(all_loaded_images))
                                        img_cols = st.columns(num_cols)
                                        
                                        for idx, img_path in enumerate(all_loaded_images):
                                            col_idx = idx % num_cols
                                            with img_cols[col_idx]:
                                                try:
                                                    # Находим номер изображения из images_info
                                                    img_index = None
                                                    for img_info in images_info:
                                                        if img_info.get("path") == img_path:
                                                            img_index = img_info.get("index", idx + 1)
                                                            break
                                                    
                                                    if img_index is None:
                                                        img_index = idx + 1
                                                    
                                                    # Определяем статус изображения
                                                    status_icon = "✅" if any(img_info.get("path") == img_path and img_info.get("status") == "cached" for img_info in images_info) else "⬇️"
                                                    status_text_img = "В кеше" if any(img_info.get("path") == img_path and img_info.get("status") == "cached" for img_info in images_info) else "Скачано"
                                                    
                                                    st.image(img_path, use_container_width=True, caption=f"{status_icon} Фото {img_index} ({status_text_img})")
                                                except Exception as e:
                                                    st.write(f"📷 Фото {idx+1}: {os.path.basename(img_path)}")
                                        
                                        st.divider()
                                        
                                        # Проверяем наличие библиотеки и API ключа
                                        if not OPENAI_AVAILABLE:
                                            st.error("⚠️ Для анализа изображений через нейросеть необходимо установить библиотеку: `pip install openai`")
                                        else:
                                            api_key = st.session_state.get('openai_api_key', '')
                                            if not api_key:
                                                st.error("❌ Необходимо указать OpenAI API ключ в настройках выше")
                                            else:
                                                # Показываем информацию о существующих параметрах
                                                param_options = st.session_state.get("param_options", {})
                                                if param_options:
                                                    with st.expander("📋 Используемые параметры для анализа", expanded=False):
                                                        for param_name, options in param_options.items():
                                                            if options:
                                                                st.write(f"**{param_name}:** {', '.join(options[:10])}{'...' if len(options) > 10 else ''}")
                                                
                                                status_text.text("🤖 Отправка изображений в нейросеть для анализа...")
                                                progress_bar.progress(50)
                                                
                                                # Анализируем изображения через нейросеть
                                                all_params = {}
                                                
                                                errors_count = 0
                                                successful_analyses = 0
                                                total_images = len(all_loaded_images)
                                                for idx, img_path in enumerate(all_loaded_images):
                                                    status_text.text(f"🤖 Анализ изображения {idx+1}/{total_images}...")
                                                    progress_bar.progress(50 + int((idx + 1) / total_images * 40))
                                                    try:
                                                        # Проверяем, что файл существует и доступен
                                                        if not os.path.exists(img_path):
                                                            st.error(f"❌ Файл изображения не найден: {img_path}")
                                                            errors_count += 1
                                                            continue
                                                        
                                                        # Получаем название товара для анализа
                                                        product_name = get_product_name_from_wb(sku)
                                                        img_params = analyze_image_with_ai(img_path, api_key, selected_params=None, product_name=product_name, sku=sku)
                                                        if img_params and len(img_params) > 0:
                                                            successful_analyses += 1
                                                            # Объединяем параметры
                                                            for key, value in img_params.items():
                                                                if key not in all_params:
                                                                    all_params[key] = []
                                                                all_params[key].append(value)
                                                            
                                                            # Показываем результат анализа для каждого изображения в режиме отладки
                                                            if st.session_state.get('debug_ai_analysis', False):
                                                                with st.expander(f"✅ Результат анализа изображения {idx+1}", expanded=False):
                                                                    for k, v in img_params.items():
                                                                        st.write(f"**{k}:** {v}")
                                                        else:
                                                            # Если параметры не определены, это не обязательно ошибка
                                                            # (нейросеть могла просто не найти параметры на изображении)
                                                            if st.session_state.get('debug_ai_analysis', False):
                                                                st.info(f"ℹ️ Изображение {idx+1}: параметры не определены нейросетью")
                                                    except Exception as e:
                                                        errors_count += 1
                                                        error_msg = str(e)
                                                        st.warning(f"⚠️ Ошибка при анализе изображения {idx+1}: {error_msg[:150]}")
                                                        if st.session_state.get('debug_ai_analysis', False):
                                                            import traceback
                                                            st.code(traceback.format_exc())
                                                
                                                # Определяем финальные параметры (берем наиболее частые значения)
                                                final_params = {}
                                                for key, values in all_params.items():
                                                    if values:
                                                        # Берем наиболее частое значение
                                                        value_counts = Counter(values)
                                                        most_common = value_counts.most_common(1)[0][0]
                                                        final_params[key] = most_common
                                                
                                                # Добавляем параметр комплектности, определенный на первом этапе
                                                if initial_params:
                                                    final_params.update(initial_params)
                                                
                                                progress_bar.progress(100)
                                                status_text.empty()
                                                
                                                if final_params:
                                                    st.success(f"✅ Найдено параметров: {len(final_params)}")
                                                    
                                                    # Показываем найденные параметры
                                                    with st.expander("📊 Определенные параметры", expanded=True):
                                                        for param_name, param_value in final_params.items():
                                                            st.write(f"**{param_name}:** {param_value}")
                                                    
                                                    # Предлагаем сохранить параметры
                                                    if st.button("💾 Сохранить параметры для этого товара", key="save_extracted_params"):
                                                        # Проверяем, есть ли этот артикул в текущей таблице
                                                        display_df = st.session_state.get("display_df", pd.DataFrame())
                                                        
                                                        if "Артикул" in display_df.columns:
                                                            sku_in_table = str(sku) in [str(s).replace(".0", "") for s in display_df["Артикул"].dropna().unique()]
                                                            
                                                            if sku_in_table:
                                                                # Сохраняем параметры с историей
                                                                for param_name, param_value in final_params.items():
                                                                    if param_name not in param_values:
                                                                        param_values[param_name] = {}
                                                                    # Используем функцию с сохранением истории
                                                                    save_param_value(str(sku), param_name, str(param_value), save_history=True)
                                                                
                                                                # Обновляем param_options
                                                                if "param_options" not in st.session_state:
                                                                    st.session_state["param_options"] = {}
                                                                
                                                                for param_name, param_value in final_params.items():
                                                                    if param_name not in st.session_state["param_options"]:
                                                                        st.session_state["param_options"][param_name] = []
                                                                    if param_value not in st.session_state["param_options"][param_name]:
                                                                        st.session_state["param_options"][param_name].append(param_value)
                                                                
                                                                # Сохраняем в файл
                                                                if save_param_values_to_file():
                                                                    st.success(f"✅ Параметры сохранены для артикула {sku}!")
                                                                    st.rerun()
                                                                else:
                                                                    st.warning("⚠️ Параметры сохранены в памяти, но не в файл")
                                                            else:
                                                                st.warning(f"⚠️ Артикул {sku} не найден в текущей таблице. Загрузите файл с этим артикулом.")
                                                        else:
                                                            st.warning("⚠️ Таблица не загружена. Загрузите файл с данными.")
                                                else:
                                                    if errors_count > 0:
                                                        st.error(f"❌ Ошибки при анализе {errors_count} изображений. Проверьте:")
                                                        st.write("1. Правильность API ключа OpenAI")
                                                        st.write("2. Доступность изображений в кеше")
                                                        st.write("3. Формат изображений (должны быть JPG, PNG или WebP)")
                                                        st.write("4. Размер изображений (не должны быть слишком большими)")
                                                        st.write("5. Баланс на счету OpenAI API")
                                                    elif successful_analyses == 0:
                                                        st.info("ℹ️ Нейросеть не смогла определить параметры из изображений.")
                                                        st.write("**Возможные причины:**")
                                                        st.write("- Изображения не содержат достаточной информации о товаре")
                                                        st.write("- Товар на изображении не виден четко")
                                                        st.write("- Параметры не соответствуют существующим вариантам")
                                                        st.write("- Попробуйте включить режим отладки для детальной информации")
                                                    else:
                                                        st.warning(f"⚠️ Проанализировано {successful_analyses} изображений, но параметры не были определены.")
                                                        st.write("Возможно, найденные параметры не соответствуют существующим вариантам.")
                                                    
                                                    if st.session_state.get('debug_ai_analysis', False):
                                                        st.write(f"**Найденные параметры:** {all_params}")
                            else:
                                st.warning("⚠️ Изображения товара не найдены. Не удалось загрузить изображения через screenshotapi.net.")
                                st.info("💡 **Совет:** Проверьте токен screenshotapi.net и интернет-соединение.")
                    else:
                        st.warning("⚠️ Введите ссылку на товар")
            else:
                # Массовый анализ
                st.markdown("### 📋 Массовый анализ товаров")
                
                is_analysis_running = st.session_state.get("mass_analysis_btn_clicked", False)
                
                # Показываем таблицу результатов массового анализа со статусами (ВСЕГДА)
                st.markdown("#### 📊 Результаты массового анализа")
                
                # Загружаем результаты из файлов (всегда проверяем, чтобы не потерять данные)
                mass_analysis_results = st.session_state.get("mass_analysis_results", [])
                
                # Если результатов нет в session_state, пытаемся загрузить из файлов
                if not mass_analysis_results:
                    # Сначала пытаемся загрузить из постоянного файла результатов
                    saved_results = load_mass_analysis_results()
                    if saved_results:
                        mass_analysis_results = saved_results
                        st.session_state["mass_analysis_results"] = saved_results
                        # НЕ применяем автоматически - пользователь должен применить вручную
                        # apply_mass_analysis_results_to_params(saved_results)
                        # save_param_values_to_file()
                    else:
                        # Если нет постоянного файла, пытаемся загрузить из прогресса
                        saved_progress = load_mass_analysis_progress()
                        if saved_progress and saved_progress.get("results"):
                            mass_analysis_results = saved_progress["results"]
                            st.session_state["mass_analysis_results"] = saved_progress["results"]
                            # Сохраняем в постоянный файл для будущих запусков
                            save_mass_analysis_results(saved_progress["results"])
                            # НЕ применяем автоматически - пользователь должен применить вручную
                            # apply_mass_analysis_results_to_params(saved_progress["results"])
                            # save_param_values_to_file()
                
                # Получаем все артикулы из основной таблицы для отображения всех товаров
                df = st.session_state.get("df", None)
                display_df = st.session_state.get("display_df", pd.DataFrame())
                
                # Используем основной df, если он есть, иначе display_df
                source_df = None
                if df is not None and not df.empty and "Артикул" in df.columns:
                    source_df = df
                elif not display_df.empty and "Артикул" in display_df.columns:
                    source_df = display_df
                
                # Создаем полный список результатов: обработанные + необработанные товары
                all_results = []
                
                # Создаем словарь обработанных товаров для быстрого поиска
                processed_skus = {str(r.get("Артикул", "")).replace(".0", "") for r in mass_analysis_results}
                
                # Добавляем обработанные товары
                all_results.extend(mass_analysis_results)
                
                # Добавляем необработанные товары из основной таблицы
                if source_df is not None and not source_df.empty:
                    for _, row in source_df.iterrows():
                        sku = str(row.get("Артикул", "")).replace(".0", "")
                        if sku and sku not in processed_skus:
                            # Создаем запись для необработанного товара
                            wb_url = build_wb_product_url(sku)
                            
                            # Получаем параметры из param_values, если они есть
                            param_values = get_param_values()
                            unprocessed_row = {
                                "Артикул": sku,
                                "URL": wb_url,
                                "Статус": "⏳ Необработан"
                            }
                            
                            # Добавляем параметры, если они есть (исключаем удаленные)
                            deleted_params = st.session_state.get("deleted_params", load_deleted_params_from_file())
                            for param_name, sku_values in param_values.items():
                                if param_name not in deleted_params and sku in sku_values:
                                    unprocessed_row[param_name] = sku_values[sku]
                            
                            all_results.append(unprocessed_row)
                
                # Обновляем mass_analysis_results для использования в остальном коде
                mass_analysis_results = all_results
                
                # Очищаем удаленные параметры из результатов перед отображением
                mass_analysis_results = remove_deleted_params_from_mass_results(mass_analysis_results)
                st.session_state["mass_analysis_results"] = mass_analysis_results
                
                if mass_analysis_results:
                    results_df = pd.DataFrame(mass_analysis_results)
                    
                    # Фильтруем удаленные параметры из столбцов DataFrame
                    deleted_params = st.session_state.get("deleted_params", set())
                    if not isinstance(deleted_params, set):
                        deleted_params = set(deleted_params) if deleted_params else set()
                    
                    # Удаляем столбцы, которые являются удаленными параметрами
                    cols_to_keep = [col for col in results_df.columns if col not in deleted_params]
                    if cols_to_keep:
                        results_df = results_df[cols_to_keep]
                    
                    # Категоризируем результаты
                    success_results = [r for r in mass_analysis_results if r.get("Статус", "").startswith("✅")]
                    error_results = [r for r in mass_analysis_results if "❌" in r.get("Статус", "") or "Ошибка" in r.get("Статус", "")]
                    undefined_results = [r for r in mass_analysis_results if "⚠️ Параметры не определены" in r.get("Статус", "")]
                    skipped_results = [r for r in mass_analysis_results if "⏭️" in r.get("Статус", "")]
                    unprocessed_results = [r for r in mass_analysis_results if "⏳ Необработан" in r.get("Статус", "")]
                    
                    # Определяем товары с неполными параметрами
                    param_values = get_param_values()
                    all_param_names = list(st.session_state.get("param_options", {}).keys())
                    incomplete_params_results = []
                    
                    if all_param_names and param_values:
                        for result in success_results:
                            sku_raw = result.get("Артикул", "")
                            if not sku_raw:
                                continue
                            sku_str = str(sku_raw).replace(".0", "")
                            
                            # Подсчитываем заполненные параметры
                            filled_params_count = 0
                            for param_name in all_param_names:
                                # Проверяем в param_values
                                if param_name in param_values and sku_str in param_values[param_name]:
                                    value = param_values[param_name][sku_str]
                                    if value and str(value).strip():
                                        filled_params_count += 1
                                # Также проверяем в самом результате
                                elif param_name in result and result[param_name] and str(result[param_name]).strip():
                                    filled_params_count += 1
                            
                            # Если заполнены не все параметры, но хотя бы один - это неполные параметры
                            if 0 < filled_params_count < len(all_param_names):
                                incomplete_params_results.append(result)
                    
                    # Статистика по результатам
                    success_count = len(success_results)
                    error_count = len(error_results)
                    undefined_count = len(undefined_results)
                    skipped_count = len(skipped_results)
                    unprocessed_count = len(unprocessed_results)
                    incomplete_count = len(incomplete_params_results)
                    total_count = len(mass_analysis_results)
                    
                    col_stats1, col_stats2, col_stats3, col_stats4, col_stats5, col_stats6, col_stats7 = st.columns(7)
                    with col_stats1:
                        st.metric("✅ Успешно", success_count, delta=f"{success_count/total_count*100:.1f}%" if total_count > 0 else "0%")
                    with col_stats2:
                        st.metric("⚠️ Неопределено", undefined_count, delta=f"{undefined_count/total_count*100:.1f}%" if total_count > 0 else "0%")
                    with col_stats3:
                        st.metric("❌ Ошибки", error_count, delta=f"{error_count/total_count*100:.1f}%" if total_count > 0 else "0%")
                    with col_stats4:
                        st.metric("⏭️ Пропущено", skipped_count, delta=f"{skipped_count/total_count*100:.1f}%" if total_count > 0 else "0%")
                    with col_stats5:
                        st.metric("⏳ Необработан", unprocessed_count, delta=f"{unprocessed_count/total_count*100:.1f}%" if total_count > 0 else "0%")
                    with col_stats6:
                        st.metric("📝 Неполные", incomplete_count, delta=f"{incomplete_count/total_count*100:.1f}%" if total_count > 0 else "0%")
                    with col_stats7:
                        st.metric("📦 Всего", total_count)
                    
                    # Фильтр для таблицы результатов
                    st.markdown("**🔍 Фильтр результатов:**")
                    filter_options = ["Все результаты", "✅ Успешно", "⚠️ Неопределено", "❌ Ошибки", "⏭️ Пропущено", "⏳ Необработан", "📝 Неполные параметры"]
                    selected_filter = st.radio(
                        "Показать:",
                        options=filter_options,
                        horizontal=True,
                        key="results_filter_radio"
                    )
                    
                    # Применяем фильтр
                    if selected_filter == "✅ Успешно":
                        filtered_results = success_results
                    elif selected_filter == "⚠️ Неопределено":
                        filtered_results = undefined_results
                    elif selected_filter == "❌ Ошибки":
                        filtered_results = error_results
                    elif selected_filter == "⏭️ Пропущено":
                        filtered_results = skipped_results
                    elif selected_filter == "⏳ Необработан":
                        filtered_results = unprocessed_results
                    elif selected_filter == "📝 Неполные параметры":
                        filtered_results = incomplete_params_results
                    else:
                        filtered_results = mass_analysis_results
                    
                    # Отображаем отфильтрованную таблицу (ВСЕГДА)
                    if filtered_results:
                        filtered_df = pd.DataFrame(filtered_results)
                        
                        # Фильтруем удаленные параметры из столбцов DataFrame
                        deleted_params = st.session_state.get("deleted_params", set())
                        if not isinstance(deleted_params, set):
                            deleted_params = set(deleted_params) if deleted_params else set()
                        
                        # Удаляем столбцы, которые являются удаленными параметрами
                        cols_to_keep = [col for col in filtered_df.columns if col not in deleted_params]
                        filtered_df = filtered_df[cols_to_keep] if cols_to_keep else filtered_df
                        
                        # Убеждаемся, что все колонки отображаются
                        # Переупорядочиваем колонки: сначала основные, потом параметры
                        main_cols = ["Артикул", "URL", "Статус", "Параметры"]
                        param_cols = [col for col in filtered_df.columns if col not in main_cols]
                        column_order = [col for col in main_cols if col in filtered_df.columns] + param_cols
                        
                        # Переупорядочиваем DataFrame
                        if column_order:
                            filtered_df = filtered_df[column_order]
                        
                        st.dataframe(
                            filtered_df, 
                            use_container_width=True,
                            height=400,
                            key="mass_analysis_results_table"
                        )
                    else:
                        # Показываем пустую таблицу с заголовками
                        empty_cols = ["Артикул", "URL", "Статус"]
                        # Добавляем колонки параметров, если они есть
                        if mass_analysis_results:
                            # Берем все уникальные ключи из результатов
                            sample_result = mass_analysis_results[0] if mass_analysis_results else {}
                            for key in sample_result.keys():
                                if key not in empty_cols and key not in ["Артикул", "URL", "Статус", "Параметры"]:
                                    empty_cols.append(key)
                        empty_df = pd.DataFrame(columns=empty_cols)
                        st.dataframe(
                            empty_df,
                            use_container_width=True,
                            key="mass_analysis_results_table_empty"
                        )
                        st.info(f"ℹ️ Нет результатов для категории '{selected_filter}'")
                    
                    # Кнопка для применения параметров в общую таблицу
                    st.divider()
                    col_apply1, col_apply2 = st.columns([1, 2])
                    with col_apply1:
                        if st.button("💾 Применить параметры в общую таблицу", type="primary", use_container_width=True, key="apply_mass_results_to_table"):
                            # Параметры уже сохранены в session_state через save_param_value во время массового анализа
                            # Но нужно обновить основную таблицу, чтобы отобразить новые параметры
                            if save_param_values_to_file():
                                st.success(f"✅ Параметры применены в общую таблицу! Обновлено {success_count} товаров.")
                                # Обновляем таблицу для отображения новых параметров
                                if "df" in st.session_state:
                                    # Перезагружаем данные, чтобы показать новые параметры
                                    st.session_state["data_loaded"] = False
                                st.rerun()
                            else:
                                st.error("❌ Ошибка при сохранении параметров")
                    with col_apply2:
                        st.info("💡 Параметры из успешных результатов массового анализа будут применены в основную таблицу. Существующие параметры будут обновлены.")
                    
                    # Экспорт/Импорт результатов в CSV
                    st.divider()
                    st.markdown("#### 💾 Экспорт / Импорт результатов")
                    col_export1, col_export2 = st.columns(2)
                    
                    with col_export1:
                        # Экспорт в CSV
                        if st.button("📥 Экспорт результатов в CSV", type="secondary", use_container_width=True, key="export_mass_results_csv"):
                            csv_data = export_mass_analysis_results_to_csv(mass_analysis_results)
                            if csv_data:
                                timestamp = pd.Timestamp.now().strftime("%Y%m%d_%H%M%S")
                                st.download_button(
                                    label="💾 Скачать CSV файл",
                                    data=csv_data,
                                    file_name=f"mass_analysis_results_{timestamp}.csv",
                                    mime="text/csv",
                                    key="download_mass_results_csv"
                                )
                                st.success(f"✅ Результаты подготовлены к экспорту! Всего записей: {len(mass_analysis_results)}")
                            else:
                                st.error("❌ Ошибка при экспорте результатов")
                    
                    with col_export2:
                        # Импорт из CSV
                        st.markdown("**Импорт из CSV:**")
                        uploaded_file = st.file_uploader(
                            "Загрузите CSV файл с результатами",
                            type=['csv'],
                            key="import_mass_results_csv"
                        )
                        if uploaded_file is not None:
                            try:
                                csv_data = uploaded_file.read().decode('utf-8-sig')
                                imported_count, applied_count = import_mass_analysis_results_from_csv(csv_data)
                                if imported_count > 0:
                                    st.success(f"✅ Импортировано записей: {imported_count}, применено параметров: {applied_count}")
                                    st.info("💡 Параметры автоматически применены. Обновите страницу, чтобы увидеть изменения.")
                                    st.rerun()
                                else:
                                    st.warning("⚠️ Не удалось импортировать данные из CSV файла")
                            except Exception as e:
                                st.error(f"❌ Ошибка при импорте CSV: {e}")
                    
                    st.divider()
                    
                    # Анализируем результаты для определения товаров, требующих повторной обработки
                    error_results = [r for r in mass_analysis_results if not r.get("Статус", "").startswith("✅")]
                    undefined_results = [r for r in mass_analysis_results if "⚠️ Параметры не определены" in r.get("Статус", "")]
                    
                    # Проверяем товары с незаполненными параметрами
                    param_values = get_param_values()
                    all_param_names = list(st.session_state.get("param_options", {}).keys())
                    skus_with_empty_params = []
                    
                    if all_param_names and param_values:
                        # Получаем все артикулы из таблицы
                        source_df = st.session_state.get("df")
                        if source_df is None:
                            source_df = st.session_state.get("display_df")
                        if source_df is not None and "Артикул" in source_df.columns:
                            for sku in source_df["Артикул"].dropna().unique():
                                sku_str = str(sku).replace(".0", "")
                                # Проверяем, заполнены ли все параметры для этого товара
                                filled_params_count = 0
                                for param_name in all_param_names:
                                    if param_name in param_values and sku_str in param_values[param_name]:
                                        value = param_values[param_name][sku_str]
                                        if value and str(value).strip():
                                            filled_params_count += 1
                                
                                # Если ни один параметр не заполнен, добавляем в список
                                if filled_params_count == 0:
                                    skus_with_empty_params.append(sku_str)
                    
                    # Показываем кнопки для повторной обработки
                    if error_results or undefined_results or skus_with_empty_params:
                        st.markdown("#### 🔄 Повторная обработка товаров")
                        
                        col_retry1, col_retry2, col_retry3 = st.columns(3)
                        
                        with col_retry1:
                            if error_results:
                                if st.button(f"🔄 Обработать товары с ошибками ({len(error_results)})", type="primary", use_container_width=True, key="retry_errors_from_results_top"):
                                    retry_skus = [str(r["Артикул"]) for r in error_results]
                                    success_results = [r for r in mass_analysis_results if r.get("Статус", "").startswith("✅")]
                                    
                                    retry_settings = {
                                        "max_images": st.session_state.get("mass_max_images_slider", 5),
                                        "delay": st.session_state.get("mass_delay_slider", 2.0),
                                        "selected_params": st.session_state.get("mass_selected_params_multiselect", [])
                                    }
                                    
                                    retry_progress = {
                                        "skus_to_process": retry_skus,
                                        "processed_skus": [],
                                        "results": success_results,
                                        "settings": retry_settings,
                                        "timestamp": pd.Timestamp.now().strftime("%Y-%m-%d %H:%M:%S"),
                                        "is_retry": True
                                    }
                                    
                                    save_mass_analysis_progress(
                                        retry_skus,
                                        set(),
                                        success_results,
                                        retry_settings
                                    )
                                    
                                    st.session_state["resume_mass_analysis"] = True
                                    st.session_state["saved_progress"] = retry_progress
                                    st.success(f"✅ Запланирована повторная обработка {len(retry_skus)} товаров с ошибками")
                                    st.rerun()
                            else:
                                st.info("✅ Нет товаров с ошибками")
                        
                        with col_retry2:
                            if undefined_results:
                                if st.button(f"🔄 Обработать товары с неопределенными параметрами ({len(undefined_results)})", type="primary", use_container_width=True, key="retry_undefined_from_results_top"):
                                    retry_skus = [str(r["Артикул"]) for r in undefined_results]
                                    success_results = [r for r in mass_analysis_results if r.get("Статус", "").startswith("✅")]
                                    
                                    retry_settings = {
                                        "max_images": st.session_state.get("mass_max_images_slider", 5),
                                        "delay": st.session_state.get("mass_delay_slider", 2.0),
                                        "selected_params": st.session_state.get("mass_selected_params_multiselect", [])
                                    }
                                    
                                    retry_progress = {
                                        "skus_to_process": retry_skus,
                                        "processed_skus": [],
                                        "results": success_results,
                                        "settings": retry_settings,
                                        "timestamp": pd.Timestamp.now().strftime("%Y-%m-%d %H:%M:%S"),
                                        "is_retry": True
                                    }
                                    
                                    save_mass_analysis_progress(
                                        retry_skus,
                                        set(),
                                        success_results,
                                        retry_settings
                                    )
                                    
                                    st.session_state["resume_mass_analysis"] = True
                                    st.session_state["saved_progress"] = retry_progress
                                    st.success(f"✅ Запланирована повторная обработка {len(retry_skus)} товаров с неопределенными параметрами")
                                    st.rerun()
                            else:
                                st.info("✅ Нет товаров с неопределенными параметрами")
                        
                        with col_retry3:
                            if skus_with_empty_params:
                                if st.button(f"🔄 Обработать товары с незаполненными параметрами ({len(skus_with_empty_params)})", type="primary", use_container_width=True, key="retry_empty_params_from_results_top"):
                                    success_results = [r for r in mass_analysis_results if r.get("Статус", "").startswith("✅")]
                                    
                                    retry_settings = {
                                        "max_images": st.session_state.get("mass_max_images_slider", 5),
                                        "delay": st.session_state.get("mass_delay_slider", 2.0),
                                        "selected_params": st.session_state.get("mass_selected_params_multiselect", [])
                                    }
                                    
                                    retry_progress = {
                                        "skus_to_process": skus_with_empty_params,
                                        "processed_skus": [],
                                        "results": success_results,
                                        "settings": retry_settings,
                                        "timestamp": pd.Timestamp.now().strftime("%Y-%m-%d %H:%M:%S"),
                                        "is_retry": True
                                    }
                                    
                                    save_mass_analysis_progress(
                                        skus_with_empty_params,
                                        set(),
                                        success_results,
                                        retry_settings
                                    )
                                    
                                    st.session_state["resume_mass_analysis"] = True
                                    st.session_state["saved_progress"] = retry_progress
                                    st.success(f"✅ Запланирована обработка {len(skus_with_empty_params)} товаров с незаполненными параметрами")
                                    st.rerun()
                            else:
                                st.info("✅ Нет товаров с незаполненными параметрами")
                        
                        # Общая кнопка для обработки всех проблемных товаров
                        if (error_results or undefined_results or skus_with_empty_params):
                            st.divider()
                            all_problem_skus = set()
                            if error_results:
                                all_problem_skus.update([str(r["Артикул"]) for r in error_results])
                            if undefined_results:
                                all_problem_skus.update([str(r["Артикул"]) for r in undefined_results])
                            if skus_with_empty_params:
                                all_problem_skus.update(skus_with_empty_params)
                            
                            if all_problem_skus:
                                if st.button(f"🔄 Обработать ВСЕ проблемные товары ({len(all_problem_skus)})", type="primary", use_container_width=True, key="retry_all_problems_from_results_top"):
                                    success_results = [r for r in mass_analysis_results if r.get("Статус", "").startswith("✅")]
                                    
                                    retry_settings = {
                                        "max_images": st.session_state.get("mass_max_images_slider", 5),
                                        "delay": st.session_state.get("mass_delay_slider", 2.0),
                                        "selected_params": st.session_state.get("mass_selected_params_multiselect", [])
                                    }
                                    
                                    retry_progress = {
                                        "skus_to_process": list(all_problem_skus),
                                        "processed_skus": [],
                                        "results": success_results,
                                        "settings": retry_settings,
                                        "timestamp": pd.Timestamp.now().strftime("%Y-%m-%d %H:%M:%S"),
                                        "is_retry": True
                                    }
                                    
                                    save_mass_analysis_progress(
                                        list(all_problem_skus),
                                        set(),
                                        success_results,
                                        retry_settings
                                    )
                                    
                                    st.session_state["resume_mass_analysis"] = True
                                    st.session_state["saved_progress"] = retry_progress
                                    st.success(f"✅ Запланирована обработка {len(all_problem_skus)} проблемных товаров")
                                    st.rerun()
                else:
                    # Если результатов нет, показываем пустую таблицу и сообщение
                    # Но сначала еще раз пытаемся загрузить из файлов (на случай, если файл появился)
                    saved_results = load_mass_analysis_results()
                    if saved_results:
                        mass_analysis_results = saved_results
                        st.session_state["mass_analysis_results"] = saved_results
                        st.success(f"✅ Загружено {len(saved_results)} результатов из сохраненного файла")
                        st.rerun()
                    else:
                        saved_progress = load_mass_analysis_progress()
                        if saved_progress and saved_progress.get("results"):
                            mass_analysis_results = saved_progress["results"]
                            st.session_state["mass_analysis_results"] = saved_progress["results"]
                            save_mass_analysis_results(saved_progress["results"])
                            st.success(f"✅ Загружено {len(saved_progress['results'])} результатов из прогресса")
                            st.rerun()
                        else:
                            st.info("ℹ️ Результаты массового анализа будут отображаться здесь после запуска анализа")
                            # Создаем пустую таблицу с основными колонками
                            empty_df = pd.DataFrame(columns=["Артикул", "URL", "Статус"])
                            st.dataframe(
                                empty_df,
                                use_container_width=True,
                                key="mass_analysis_results_table_no_data"
                            )
                    
                    st.divider()
                
                st.info("💡 Система автоматически проанализирует все товары из таблицы ниже. Для каждого товара будут получены изображения через screenshotapi.net и проанализированы через нейросеть.")
                
                # Проверяем наличие таблицы - используем основной df из session_state
                df = st.session_state.get("df", None)
                display_df = st.session_state.get("display_df", pd.DataFrame())
                
                # Используем основной df, если он есть, иначе display_df
                if df is not None and not df.empty and "Артикул" in df.columns:
                    source_df = df
                elif not display_df.empty and "Артикул" in display_df.columns:
                    source_df = display_df
                else:
                    source_df = None
                
                if source_df is None or source_df.empty or "Артикул" not in source_df.columns:
                    # Проверяем, есть ли вообще данные в session_state
                    if st.session_state.get("df") is None:
                        st.warning("⚠️ Файл не загружен. Загрузите файл с данными товаров в первой вкладке.")
                    else:
                        st.warning("⚠️ В таблице нет данных с артикулами. Убедитесь, что файл содержит колонку 'Артикул'.")
                else:
                    # Получаем список артикулов из таблицы в порядке их появления в таблице
                    # Порядок соответствует порядку строк в исходной таблице (сверху вниз)
                    skus = source_df["Артикул"].dropna().unique()
                    skus = [str(sku).replace(".0", "") for sku in skus if str(sku).strip()]
                    
                    if not skus:
                        st.warning("⚠️ В таблице нет артикулов для анализа.")
                    else:
                        st.success(f"✅ Найдено товаров в таблице: **{len(skus)}**")
                        
                        # Выбор параметров для анализа (скрыто в expander)
                        all_param_names = list(st.session_state.get("param_options", {}).keys())
                        mass_selected_params = []
                        
                        if all_param_names:
                            with st.expander("⚙️ Настройки анализа (опционально)", expanded=False):
                                st.markdown("**📋 Выберите параметры для анализа:**")
                                st.info("💡 Выберите конкретные параметры, которые нужно определить через нейросеть. Если не выбрано, будут анализироваться все доступные параметры.")
                                mass_selected_params = st.multiselect(
                                    "Параметры для анализа:",
                                    options=all_param_names,
                                    default=[],  # По умолчанию пусто - будут анализироваться все
                                    help="Оставьте пустым, чтобы анализировать все параметры. Или выберите конкретные параметры.",
                                    key="mass_selected_params_multiselect"
                                )
                        else:
                            st.warning("⚠️ Нет доступных параметров. Добавьте параметры в разделе 'Управление параметрами'.")
                        
                        st.divider()
                        
                        col_mass1, col_mass2, col_mass3 = st.columns(3)
                        with col_mass1:
                            mass_max_images = st.slider(
                                "Изображений на товар:",
                                min_value=1,
                                max_value=5,
                                value=5,
                                help="Количество изображений для анализа каждого товара",
                                key="mass_max_images_slider"
                            )
                        
                        with col_mass2:
                            mass_delay = st.slider(
                                "Задержка между товарами (сек):",
                                min_value=1.0,
                                max_value=10.0,
                                value=2.0,
                                step=0.5,
                                help="Задержка между обработкой товаров для избежания rate limit",
                                key="mass_delay_slider"
                            )
                        
                        with col_mass3:
                            mass_limit = st.number_input(
                                "Ограничение количества товаров:",
                                min_value=0,
                                max_value=len(skus),
                                value=len(skus),
                                help="0 = все товары",
                                key="mass_limit_input"
                            )
                        
                        # Опция фильтрации товаров без фото
                        st.divider()
                        filter_no_images = st.checkbox(
                            "🔍 Обрабатывать только товары БЕЗ скачанных фото",
                            value=False,
                            help="Если включено, будут обработаны только товары, у которых нет скачанных изображений через screenshotapi",
                            key="filter_no_images_checkbox"
                        )
                        
                        if filter_no_images:
                            # Проверяем, сколько товаров без фото
                            skus_without_images = []
                            with st.spinner("Проверка наличия фото..."):
                                for sku in skus:
                                    if not has_downloaded_images(sku, min_images=1):
                                        skus_without_images.append(sku)
                            
                            if skus_without_images:
                                st.success(f"✅ Найдено товаров без фото: **{len(skus_without_images)}** из {len(skus)}")
                                with st.expander("👀 Превью товаров без фото", expanded=False):
                                    for idx, sku in enumerate(skus_without_images[:10], 1):
                                        wb_url = build_wb_product_url(sku)
                                        st.write(f"{idx}. Артикул: {sku} - [Открыть на WB]({wb_url})")
                                    if len(skus_without_images) > 10:
                                        st.write(f"... и еще {len(skus_without_images) - 10} товаров")
                            else:
                                st.info("ℹ️ У всех товаров уже есть скачанные фото")
                        
                        # Проверяем наличие сохраненного прогресса
                        saved_progress = load_mass_analysis_progress()
                        if saved_progress:
                            processed_count = len(saved_progress.get('processed_skus', []))
                            total_count = len(saved_progress.get('skus_to_process', []))
                            remaining_count = total_count - processed_count
                            
                            with st.expander("📥 Найден сохраненный прогресс", expanded=False):
                                st.info(f"**Сохранено:** {saved_progress.get('timestamp', 'неизвестно')}")
                                st.write(f"**Всего товаров:** {total_count}")
                                st.write(f"**Обработано:** {processed_count}")
                                st.write(f"**Осталось:** {remaining_count}")
                                
                                # Показываем статистику по ошибкам
                                saved_results = saved_progress.get('results', [])
                                if saved_results:
                                    error_results = [r for r in saved_results if not r.get("Статус", "").startswith("✅")]
                                    success_results = [r for r in saved_results if r.get("Статус", "").startswith("✅")]
                                    st.write(f"**Успешно:** {len(success_results)}")
                                    st.write(f"**С ошибками:** {len(error_results)}")
                                st.info("💡 Используйте кнопку '▶️ Продолжить' ниже для возобновления анализа")
                                
                                if st.button("🗑️ Удалить сохраненный прогресс", key="clear_progress_btn"):
                                    clear_mass_analysis_progress()
                                    st.success("✅ Прогресс удален")
                                    st.rerun()
                        
                        # Автоматическое возобновление, если есть сохраненный прогресс и не было явного запуска нового анализа
                        auto_resume = saved_progress and not st.session_state.get("mass_analysis_btn_clicked", False)
                        
                        # Показываем кнопки управления анализом
                        col_start, col_resume, col_stop = st.columns([1, 1, 1])
                        with col_start:
                            if st.button("🚀 Запустить массовый анализ", type="primary", key="mass_analysis_btn", disabled=is_analysis_running, use_container_width=True):
                                st.session_state["mass_analysis_btn_clicked"] = True
                                # Сбрасываем флаг продолжения при новом запуске
                                if "resume_mass_analysis" in st.session_state:
                                    del st.session_state["resume_mass_analysis"]
                                st.rerun()
                        with col_resume:
                            if saved_progress and not is_analysis_running:
                                if st.button("▶️ Продолжить", type="primary", key="continue_analysis_btn", use_container_width=True):
                                    st.session_state["resume_mass_analysis"] = True
                                    st.session_state["saved_progress"] = saved_progress
                                    st.session_state["mass_analysis_btn_clicked"] = True
                                    st.rerun()
                            else:
                                st.write("")  # Пустое место для выравнивания
                        with col_stop:
                            if is_analysis_running:
                                if st.button("⏹️ Остановить анализ", type="secondary", key="stop_mass_analysis_btn", use_container_width=True):
                                    st.session_state["mass_analysis_btn_clicked"] = False
                                    st.session_state["mass_analysis_stopped"] = True
                                    st.warning("⚠️ Анализ остановлен пользователем")
                                    st.rerun()
                            else:
                                st.write("")  # Пустое место для выравнивания
                        
                        # Запускаем анализ, если кнопка была нажата
                        if st.session_state.get("mass_analysis_btn_clicked", False) and not st.session_state.get("mass_analysis_stopped", False):
                            
                            # Фильтруем товары без фото, если включена опция
                            if filter_no_images:
                                skus_filtered = []
                                with st.spinner("Фильтрация товаров без фото..."):
                                    for sku in skus:
                                        if not has_downloaded_images(sku, min_images=1):
                                            skus_filtered.append(sku)
                                skus = skus_filtered
                                
                                if not skus:
                                    st.warning("⚠️ Нет товаров без фото для обработки")
                                    st.stop()
                            
                            # Ограничиваем количество товаров, если указано
                            if mass_limit > 0:
                                skus_to_process = skus[:mass_limit]
                            else:
                                skus_to_process = skus
                            
                            if skus_to_process:
                                # Проверяем, нужно ли возобновить с сохраненного прогресса
                                resume_analysis = st.session_state.get("resume_mass_analysis", False) or auto_resume
                                saved_progress_data = st.session_state.get("saved_progress", None) or saved_progress
                                
                                # Начинаем массовый анализ
                                mass_progress_bar = st.progress(0)
                                mass_status_text = st.empty()
                                mass_results_container = st.container()
                                
                                api_key = st.session_state.get('openai_api_key', '')
                                
                                if not api_key:
                                    st.error("❌ Необходимо указать API ключ OpenAI в настройках выше")
                                else:
                                    # Если возобновляем, используем сохраненные данные
                                    if resume_analysis and saved_progress_data:
                                        skus_to_process = saved_progress_data.get("skus_to_process", skus_to_process)
                                        processed_skus = set(saved_progress_data.get("processed_skus", []))
                                        mass_results = saved_progress_data.get("results", [])
                                        settings = saved_progress_data.get("settings", {})
                                        mass_max_images = settings.get("max_images", mass_max_images)
                                        mass_delay = settings.get("delay", mass_delay)
                                        # Восстанавливаем выбранные параметры из сохраненного прогресса
                                        mass_selected_params = settings.get("selected_params", [])
                                        
                                        # Проверяем, это повторная обработка ошибок
                                        is_retry = saved_progress_data.get("is_retry", False)
                                        
                                        # Показываем информацию о возобновлении
                                        remaining_count = len(skus_to_process) - len(processed_skus)
                                        if is_retry:
                                            st.info(f"🔄 Повторная обработка товаров с ошибками: {len(skus_to_process)} товаров. Успешные результаты ({len(mass_results)}) будут сохранены.")
                                        else:
                                            st.info(f"🔄 Возобновление анализа: обработано {len(processed_skus)} из {len(skus_to_process)}, осталось {remaining_count}")
                                        
                                        st.session_state["resume_mass_analysis"] = False
                                        st.session_state["saved_progress"] = None
                                    else:
                                        processed_skus = set()
                                        mass_results = []
                                        # Используем выбранные параметры или все доступные
                                        selected_params_for_analysis = mass_selected_params if mass_selected_params else None
                                        settings = {
                                            "max_images": mass_max_images,
                                            "delay": mass_delay,
                                            "selected_params": selected_params_for_analysis
                                        }
                                    
                                    total_skus = len(skus_to_process)
                                    mass_success_count = len([r for r in mass_results if r.get("Статус", "").startswith("✅")])
                                    mass_error_count = len([r for r in mass_results if not r.get("Статус", "").startswith("✅")])
                                    
                                    with mass_results_container:
                                        st.markdown("### 📊 Результаты массового анализа")
                                        results_table_placeholder = st.empty()
                                        apply_button_placeholder = st.empty()
                                        
                                        # Показываем уже обработанные результаты сразу при возобновлении
                                        if mass_results:
                                            # Пересчитываем счетчики из актуальных результатов
                                            current_success_count = len([r for r in mass_results if r.get("Статус", "").startswith("✅")])
                                            current_error_count = len([r for r in mass_results if not r.get("Статус", "").startswith("✅")])
                                            
                                            results_df = pd.DataFrame(mass_results)
                                            
                                            # Фильтруем удаленные параметры из столбцов DataFrame
                                            deleted_params = st.session_state.get("deleted_params", set())
                                            if not isinstance(deleted_params, set):
                                                deleted_params = set(deleted_params) if deleted_params else set()
                                            
                                            # Удаляем столбцы, которые являются удаленными параметрами
                                            cols_to_keep = [col for col in results_df.columns if col not in deleted_params]
                                            if cols_to_keep:
                                                results_df = results_df[cols_to_keep]
                                            
                                            with results_table_placeholder.container():
                                                st.markdown(f"**📋 Уже обработано: {len(mass_results)} товаров** (✅ Успешно: {current_success_count}, ❌ Ошибок: {current_error_count})")
                                                st.dataframe(results_df, use_container_width=True, height=400)
                                            
                                            # Кнопка для применения параметров - доступна сразу во время обработки
                                            if current_success_count > 0:
                                                with apply_button_placeholder.container():
                                                    st.divider()
                                                    col_apply1, col_apply2 = st.columns([1, 2])
                                                    with col_apply1:
                                                        if st.button("💾 Применить параметры в общую таблицу", type="primary", use_container_width=True, key="apply_mass_results_final"):
                                                            # Сохраняем параметры в файл
                                                            if save_param_values_to_file():
                                                                if save_param_history_to_file():
                                                                    # Перезагружаем параметры из файла, чтобы убедиться, что они применены
                                                                    load_param_values_from_file()
                                                                    current_success_count = len([r for r in mass_results if r.get("Статус", "").startswith("✅")])
                                                                    st.success(f"✅ Параметры применены в общую таблицу! Обновлено {current_success_count} товаров.")
                                                                    # Обновляем таблицу для отображения новых параметров
                                                                    if "df" in st.session_state:
                                                                        # Перезагружаем данные, чтобы показать новые параметры
                                                                        st.session_state["data_loaded"] = False
                                                                        # Очищаем display_df, чтобы он пересоздался с новыми параметрами
                                                                        if "display_df" in st.session_state:
                                                                            del st.session_state["display_df"]
                                                                    st.rerun()
                                                                else:
                                                                    st.error("❌ Ошибка при сохранении истории изменений")
                                                            else:
                                                                st.error("❌ Ошибка при сохранении параметров")
                                                    with col_apply2:
                                                        st.info("💡 Параметры из успешных результатов будут применены в основную таблицу. Можно нажимать во время обработки.")
                                    
                                    # Проверяем, не был ли анализ остановлен
                                    if st.session_state.get("mass_analysis_stopped", False):
                                        st.warning("⚠️ Анализ был остановлен пользователем")
                                        st.session_state["mass_analysis_stopped"] = False
                                        st.session_state["mass_analysis_btn_clicked"] = False
                                        # Сохраняем текущие результаты
                                        if mass_results:
                                            st.session_state["mass_analysis_results"] = mass_results
                                            save_mass_analysis_results(mass_results)
                                        st.stop()
                                    
                                    # Обрабатываем только необработанные товары
                                    remaining_skus = [sku for sku in skus_to_process if str(sku) not in processed_skus]
                                    
                                    # Если все товары уже обработаны, показываем финальные результаты
                                    if not remaining_skus and mass_results:
                                        st.info("ℹ️ Все товары уже обработаны. Результаты показаны выше.")
                                        mass_progress_bar.progress(1.0)
                                        mass_status_text.empty()
                                        
                                        # Показываем финальные результаты
                                        st.markdown("---")
                                        st.markdown("### 📊 Итоговый отчет массового анализа")
                                        
                                        # Пересчитываем счетчики из актуальных результатов
                                        final_success_count = len([r for r in mass_results if r.get("Статус", "").startswith("✅")])
                                        final_error_count = len([r for r in mass_results if not r.get("Статус", "").startswith("✅")])
                                        
                                        col_success, col_error, col_total = st.columns(3)
                                        with col_success:
                                            st.metric("✅ Успешно", final_success_count)
                                        with col_error:
                                            st.metric("❌ Ошибок", final_error_count)
                                        with col_total:
                                            st.metric("📦 Всего", total_skus)
                                        
                                        # Удаляем сохраненный прогресс
                                        clear_mass_analysis_progress()
                                        
                                        # Сохраняем результаты в session_state
                                        st.session_state["mass_analysis_results"] = mass_results
                                        save_mass_analysis_results(mass_results)
                                        
                                        # Автоматически применяем параметры из результатов
                                        applied_params_count = apply_mass_analysis_results_to_params(mass_results)
                                        if applied_params_count > 0:
                                            # Сохраняем параметры в файл
                                            save_param_values_to_file()
                                        
                                        # Сбрасываем флаг запуска анализа, чтобы старые результаты снова показывались
                                        st.session_state["mass_analysis_btn_clicked"] = False
                                        
                                        # Таблица уже показана выше в mass_results_container, не дублируем
                                        # Кнопка для применения параметров уже доступна выше, но показываем здесь тоже для итогового отчета
                                        # Пересчитываем счетчики из актуальных результатов
                                        final_success_count = len([r for r in mass_results if r.get("Статус", "").startswith("✅")])
                                        if final_success_count > 0:
                                            # Кнопка уже есть в mass_results_container, здесь просто подтверждаем
                                            st.success(f"✅ Анализ завершен! Параметры можно применить кнопкой выше.")
                                        else:
                                            # Сохраняем параметры в файл даже если нет успешных результатов (для истории)
                                            if save_param_values_to_file():
                                                st.info("💾 Параметры сохранены в файл!")
                                            if save_param_history_to_file():
                                                st.info("💾 История изменений сохранена!")
                                        
                                        # Определяем проблемные товары (всегда, не только при mass_error_count > 0)
                                        error_results = [r for r in mass_results if not r.get("Статус", "").startswith("✅")]
                                        undefined_results = [r for r in mass_results if "⚠️ Параметры не определены" in r.get("Статус", "")]
                                        
                                        # Проверяем товары с незаполненными параметрами из таблицы
                                        param_values = get_param_values()
                                        all_param_names = list(st.session_state.get("param_options", {}).keys())
                                        skus_with_empty_params = []
                                        
                                        if all_param_names and param_values:
                                            # Получаем все артикулы из таблицы
                                            source_df = st.session_state.get("df")
                                            if source_df is None:
                                                source_df = st.session_state.get("display_df")
                                            if source_df is not None and "Артикул" in source_df.columns:
                                                for sku in source_df["Артикул"].dropna().unique():
                                                    sku_str = str(sku).replace(".0", "")
                                                    # Проверяем, заполнены ли все параметры для этого товара
                                                    filled_params_count = 0
                                                    for param_name in all_param_names:
                                                        if param_name in param_values and sku_str in param_values[param_name]:
                                                            value = param_values[param_name][sku_str]
                                                            if value and str(value).strip():
                                                                filled_params_count += 1
                                                    
                                                    # Если ни один параметр не заполнен, добавляем в список
                                                    if filled_params_count == 0:
                                                        skus_with_empty_params.append(sku_str)
                                        
                                        # Показываем детальный отчет об ошибках, если есть
                                        # Пересчитываем счетчики из актуальных результатов
                                        final_error_count = len([r for r in mass_results if not r.get("Статус", "").startswith("✅")])
                                        if final_error_count > 0:
                                            st.markdown("---")
                                            st.markdown("### ❌ Детальный отчет об ошибках")
                                            
                                            # Кнопки для обработки разных типов проблемных товаров - показываем всегда, если есть проблемы
                                            if error_results or undefined_results or skus_with_empty_params:
                                                st.markdown("#### 🔄 Повторная обработка товаров")
                                                
                                                # Определяем количество колонок в зависимости от наличия типов ошибок
                                                num_cols = sum([bool(error_results), bool(undefined_results), bool(skus_with_empty_params)])
                                                if num_cols == 1:
                                                    col_retry1, = st.columns(1)
                                                    col_retry2 = None
                                                    col_retry3 = None
                                                elif num_cols == 2:
                                                    col_retry1, col_retry2 = st.columns(2)
                                                    col_retry3 = None
                                                else:
                                                    col_retry1, col_retry2, col_retry3 = st.columns(3)
                                                
                                                # Кнопка для обработки ВСЕХ товаров с ошибками (показываем отдельно)
                                                if error_results:
                                                    st.markdown("**🔄 Обработать все товары с ошибками:**")
                                                    if st.button(f"🔄 Обработать все товары с ошибками ({len(error_results)})", type="primary", use_container_width=True, key="retry_all_errors_final_main"):
                                                        retry_skus = [str(r["Артикул"]) for r in error_results]
                                                        success_results = [r for r in mass_results if r.get("Статус", "").startswith("✅")]
                                                        
                                                        retry_settings = {
                                                            "max_images": mass_max_images,
                                                            "delay": mass_delay,
                                                            "selected_params": settings.get("selected_params", None)
                                                        }
                                                        
                                                        retry_progress = {
                                                            "skus_to_process": retry_skus,
                                                            "processed_skus": [],
                                                            "results": success_results,
                                                            "settings": retry_settings,
                                                            "timestamp": pd.Timestamp.now().strftime("%Y-%m-%d %H:%M:%S"),
                                                            "is_retry": True
                                                        }
                                                        
                                                        save_mass_analysis_progress(
                                                            retry_skus,
                                                            set(),
                                                            success_results,
                                                            retry_settings
                                                        )
                                                        
                                                        st.session_state["resume_mass_analysis"] = True
                                                        st.session_state["saved_progress"] = retry_progress
                                                        st.success(f"✅ Запланирована обработка {len(retry_skus)} товаров с ошибками")
                                                        st.rerun()
                                                    st.divider()
                                            
                                            if error_results:
                                                error_df = pd.DataFrame(error_results)
                                                error_report_cols = ["Артикул", "URL", "Статус"]
                                                if "Параметры" in error_df.columns:
                                                    error_report_cols.append("Параметры")
                                                
                                                error_report_df = error_df[error_report_cols].copy()
                                            
                                            # Товары с ошибками (исключая неопределенные параметры)
                                            real_errors = [r for r in error_results if "⚠️ Параметры не определены" not in r.get("Статус", "")] if error_results else []
                                            
                                            if col_retry1:
                                                with col_retry1:
                                                    if real_errors:
                                                        if st.button(f"🔄 Ошибки ({len(real_errors)})", type="primary", use_container_width=True, key="retry_real_errors_final"):
                                                            retry_skus = [str(r["Артикул"]) for r in real_errors]
                                                            success_results = [r for r in mass_results if r.get("Статус", "").startswith("✅")]
                                                            
                                                            retry_settings = {
                                                                "max_images": mass_max_images,
                                                                "delay": mass_delay,
                                                                "selected_params": settings.get("selected_params", None)
                                                            }
                                                            
                                                            retry_progress = {
                                                                "skus_to_process": retry_skus,
                                                                "processed_skus": [],
                                                                "results": success_results,
                                                                "settings": retry_settings,
                                                                "timestamp": pd.Timestamp.now().strftime("%Y-%m-%d %H:%M:%S"),
                                                                "is_retry": True
                                                            }
                                                            
                                                            save_mass_analysis_progress(
                                                                retry_skus,
                                                                set(),
                                                                success_results,
                                                                retry_settings
                                                            )
                                                            
                                                            st.session_state["resume_mass_analysis"] = True
                                                            st.session_state["saved_progress"] = retry_progress
                                                            st.success(f"✅ Запланирована обработка {len(retry_skus)} товаров с ошибками")
                                                            st.rerun()
                                                    elif error_results:
                                                        st.info("✅ Нет товаров с ошибками (кроме неопределенных)")
                                            
                                            if col_retry2:
                                                with col_retry2:
                                                    # Товары с неопределенными параметрами
                                                    if undefined_results:
                                                        if st.button(f"🔄 Неопределенные ({len(undefined_results)})", type="primary", use_container_width=True, key="retry_undefined_params_final"):
                                                            retry_skus = [str(r["Артикул"]) for r in undefined_results]
                                                        success_results = [r for r in mass_results if r.get("Статус", "").startswith("✅")]
                                                        
                                                        retry_settings = {
                                                            "max_images": mass_max_images,
                                                            "delay": mass_delay,
                                                            "selected_params": settings.get("selected_params", None)
                                                        }
                                                        
                                                        retry_progress = {
                                                            "skus_to_process": retry_skus,
                                                            "processed_skus": [],
                                                            "results": success_results,
                                                            "settings": retry_settings,
                                                            "timestamp": pd.Timestamp.now().strftime("%Y-%m-%d %H:%M:%S"),
                                                            "is_retry": True
                                                        }
                                                        
                                                        save_mass_analysis_progress(
                                                            retry_skus,
                                                            set(),
                                                            success_results,
                                                            retry_settings
                                                        )
                                                        
                                                        st.session_state["resume_mass_analysis"] = True
                                                        st.session_state["saved_progress"] = retry_progress
                                                        st.success(f"✅ Запланирована обработка {len(retry_skus)} товаров с неопределенными параметрами")
                                                        st.rerun()
                                                    else:
                                                        st.info("✅ Нет товаров с неопределенными параметрами")
                                            
                                            if col_retry3:
                                                with col_retry3:
                                                    # Товары с незаполненными параметрами
                                                    if skus_with_empty_params:
                                                        if st.button(f"🔄 Незаполненные ({len(skus_with_empty_params)})", type="primary", use_container_width=True, key="retry_empty_params_final"):
                                                            success_results = [r for r in mass_results if r.get("Статус", "").startswith("✅")]
                                                            
                                                            retry_settings = {
                                                                "max_images": mass_max_images,
                                                                "delay": mass_delay,
                                                                "selected_params": settings.get("selected_params", None)
                                                            }
                                                            
                                                            retry_progress = {
                                                                "skus_to_process": skus_with_empty_params,
                                                                "processed_skus": [],
                                                                "results": success_results,
                                                                "settings": retry_settings,
                                                                "timestamp": pd.Timestamp.now().strftime("%Y-%m-%d %H:%M:%S"),
                                                                "is_retry": True
                                                            }
                                                            
                                                            save_mass_analysis_progress(
                                                                skus_with_empty_params,
                                                                set(),
                                                                success_results,
                                                                retry_settings
                                                            )
                                                            
                                                            st.session_state["resume_mass_analysis"] = True
                                                            st.session_state["saved_progress"] = retry_progress
                                                            st.success(f"✅ Запланирована обработка {len(skus_with_empty_params)} товаров с незаполненными параметрами")
                                                            st.rerun()
                                            
                                            # Общая кнопка для всех проблемных товаров
                                            all_problem_skus = set()
                                            if real_errors:
                                                all_problem_skus.update([str(r["Артикул"]) for r in real_errors])
                                            if undefined_results:
                                                all_problem_skus.update([str(r["Артикул"]) for r in undefined_results])
                                            if skus_with_empty_params:
                                                all_problem_skus.update(skus_with_empty_params)
                                            
                                            if all_problem_skus:
                                                st.divider()
                                                if st.button(f"🔄 Обработать ВСЕ проблемные товары ({len(all_problem_skus)})", type="primary", use_container_width=True, key="retry_all_problems_final"):
                                                    success_results = [r for r in mass_results if r.get("Статус", "").startswith("✅")]
                                                    
                                                    retry_settings = {
                                                        "max_images": mass_max_images,
                                                        "delay": mass_delay,
                                                        "selected_params": settings.get("selected_params", None)
                                                    }
                                                    
                                                    retry_progress = {
                                                        "skus_to_process": list(all_problem_skus),
                                                        "processed_skus": [],
                                                        "results": success_results,
                                                        "settings": retry_settings,
                                                        "timestamp": pd.Timestamp.now().strftime("%Y-%m-%d %H:%M:%S"),
                                                        "is_retry": True
                                                    }
                                                    
                                                    save_mass_analysis_progress(
                                                        list(all_problem_skus),
                                                        set(),
                                                        success_results,
                                                        retry_settings
                                                    )
                                                    
                                                    st.session_state["resume_mass_analysis"] = True
                                                    st.session_state["saved_progress"] = retry_progress
                                                    st.success(f"✅ Запланирована обработка {len(all_problem_skus)} проблемных товаров")
                                                    st.rerun()
                                                
                                                st.markdown("#### 📋 Все товары с ошибками")
                                                st.dataframe(error_report_df, use_container_width=True, height=400)
                                                
                                                # Экспорт в CSV
                                                csv = error_report_df.to_csv(index=False).encode('utf-8-sig')
                                                st.download_button(
                                                    label="📥 Скачать отчет об ошибках (CSV)",
                                                    data=csv,
                                                    file_name=f"mass_analysis_errors_{pd.Timestamp.now().strftime('%Y%m%d_%H%M%S')}.csv",
                                                    mime="text/csv",
                                                    key="download_errors_csv_final"
                                                )
                                        
                                    
                                    # Обрабатываем только если есть необработанные товары
                                    if remaining_skus:
                                        import time
                                        previous_product_time = None  # Время обработки предыдущего товара
                                        
                                        for idx, sku in enumerate(remaining_skus):
                                            # Проверяем, не был ли анализ остановлен
                                            if st.session_state.get("mass_analysis_stopped", False) or not st.session_state.get("mass_analysis_btn_clicked", False):
                                                st.warning("⚠️ Анализ был остановлен")
                                                # Сохраняем текущие результаты
                                                if mass_results:
                                                    st.session_state["mass_analysis_results"] = mass_results
                                                    save_mass_analysis_results(mass_results)
                                                    # Автоматически применяем параметры из результатов
                                                    apply_mass_analysis_results_to_params(mass_results)
                                                    save_param_values_to_file()
                                                st.session_state["mass_analysis_stopped"] = False
                                                st.session_state["mass_analysis_btn_clicked"] = False
                                                break
                                            
                                            # Находим реальный индекс товара в исходном списке skus_to_process
                                            # Это важно для правильного отображения номера позиции [X/96]
                                            try:
                                                processed_idx = skus_to_process.index(str(sku))
                                            except (ValueError, AttributeError):
                                                # Если не найден (не должно быть), используем текущую логику как fallback
                                                processed_idx = len(processed_skus) + idx
                                            
                                            # Формируем URL товара
                                            wb_url = build_wb_product_url(sku)
                                            
                                            # Засекаем время начала обработки товара
                                            product_start_time = time.time()
                                            
                                            # Формируем строку статуса с информацией о времени
                                            prev_time_str = ""
                                            if previous_product_time is not None:
                                                prev_time_str = f" | ⏱️ Предыдущий: {previous_product_time:.1f}с"
                                            
                                            mass_status_text.text(f"🔄 Обработка товара {processed_idx + 1}/{total_skus}: Артикул {sku}...{prev_time_str}")
                                            # Ограничиваем значение progress до максимум 1.0
                                            progress_value = min((processed_idx) / total_skus, 1.0)
                                            mass_progress_bar.progress(progress_value)
                                            
                                            # Получаем параметры через анализ изображений
                                            try:
                                                # Передаем выбранные параметры для анализа (из настроек или из текущего выбора)
                                                if resume_analysis and saved_progress_data:
                                                    selected_params_for_analysis = settings.get("selected_params", None)
                                                else:
                                                    selected_params_for_analysis = mass_selected_params if mass_selected_params else None
                                                
                                                # Этап 1: Получение изображений
                                                elapsed = time.time() - product_start_time
                                                prev_time_str = f" | ⏱️ Предыдущий: {previous_product_time:.1f}с" if previous_product_time else ""
                                                mass_status_text.text(f"📸 [{processed_idx + 1}/{total_skus}] Артикул {sku}: Получение изображений... | ⏱️ Текущий: {elapsed:.1f}с{prev_time_str}")
                                                
                                                # Этап 2: Анализ изображений через AI
                                                elapsed = time.time() - product_start_time
                                                prev_time_str = f" | ⏱️ Предыдущий: {previous_product_time:.1f}с" if previous_product_time else ""
                                                mass_status_text.text(f"🤖 [{processed_idx + 1}/{total_skus}] Артикул {sku}: Анализ через нейросеть... | ⏱️ Текущий: {elapsed:.1f}с{prev_time_str}")
                                                
                                                params = get_product_params_from_images(wb_url, api_key, max_images=mass_max_images, selected_params=selected_params_for_analysis)
                                                
                                                # Этап 3: Обработка результатов
                                                elapsed = time.time() - product_start_time
                                                prev_time_str = f" | ⏱️ Предыдущий: {previous_product_time:.1f}с" if previous_product_time else ""
                                                mass_status_text.text(f"💾 [{processed_idx + 1}/{total_skus}] Артикул {sku}: Сохранение параметров... | ⏱️ Текущий: {elapsed:.1f}с{prev_time_str}")
                                                
                                                if params:
                                                    # Фильтруем параметры, исключая невалидные (артикулы и т.д.)
                                                    valid_params = {k: v for k, v in params.items() if is_valid_param_name(k)}
                                                    
                                                    result_row = {
                                                        "Артикул": sku,
                                                        "URL": wb_url,
                                                        "Статус": "✅ Успешно",
                                                        "Параметры": ", ".join([f"{k}: {v}" for k, v in valid_params.items()])
                                                    }
                                                    result_row.update(valid_params)
                                                    
                                                    # Проверяем, есть ли уже результат для этого артикула
                                                    sku_str = str(sku).replace(".0", "")
                                                    existing_result_idx = None
                                                    for i, existing_result in enumerate(mass_results):
                                                        existing_sku = str(existing_result.get("Артикул", "")).replace(".0", "")
                                                        if existing_sku == sku_str:
                                                            existing_result_idx = i
                                                            break
                                                    
                                                    # Если результат уже есть - обновляем, иначе добавляем
                                                    if existing_result_idx is not None:
                                                        # Обновляем существующий результат (всегда, чтобы данные были актуальными)
                                                        mass_results[existing_result_idx] = result_row
                                                    else:
                                                        # Добавляем новый результат
                                                        mass_results.append(result_row)
                                                    
                                                    # СРАЗУ синхронизируем с session_state и сохраняем в файл после каждого товара
                                                    # Очищаем удаленные параметры перед сохранением
                                                    cleaned_mass_results = remove_deleted_params_from_mass_results(mass_results.copy())
                                                    st.session_state["mass_analysis_results"] = cleaned_mass_results
                                                    save_mass_analysis_results(cleaned_mass_results)
                                                    # Обновляем локальную переменную для отображения
                                                    mass_results = cleaned_mass_results
                                                    
                                                    # Сохраняем параметры в session_state с историей
                                                    if "param_values" not in st.session_state:
                                                        st.session_state["param_values"] = {}
                                                    
                                                    # Обновляем param_options только для валидных параметров
                                                    if "param_options" not in st.session_state:
                                                        st.session_state["param_options"] = {}
                                                    
                                                    for param_name, param_value in params.items():
                                                        # Пропускаем невалидные параметры (артикулы и т.д.)
                                                        if not is_valid_param_name(param_name):
                                                            continue
                                                        
                                                        # Используем функцию с сохранением истории (без записи в файл)
                                                        save_param_value(str(sku), param_name, str(param_value), save_history=True)
                                                        
                                                        # Обновляем param_options
                                                        if "param_options" not in st.session_state:
                                                            st.session_state["param_options"] = {}
                                                        if param_name not in st.session_state["param_options"]:
                                                            st.session_state["param_options"][param_name] = []
                                                        if param_value and str(param_value) not in st.session_state["param_options"][param_name]:
                                                            st.session_state["param_options"][param_name].append(str(param_value))
                                                    
                                                    # Сохраняем историю параметров пакетом после обработки всех параметров товара
                                                    if params:
                                                        save_param_history_to_file()
                                                    
                                                else:
                                                    result_row = {
                                                        "Артикул": sku,
                                                        "URL": wb_url,
                                                        "Статус": "⚠️ Параметры не определены",
                                                        "Параметры": ""
                                                    }
                                                    
                                                    # Проверяем, есть ли уже результат для этого артикула
                                                    sku_str = str(sku).replace(".0", "")
                                                    existing_result_idx = None
                                                    for i, existing_result in enumerate(mass_results):
                                                        existing_sku = str(existing_result.get("Артикул", "")).replace(".0", "")
                                                        if existing_sku == sku_str:
                                                            existing_result_idx = i
                                                            break
                                                    
                                                    # Если результат уже есть - обновляем, иначе добавляем
                                                    if existing_result_idx is not None:
                                                        # Обновляем существующий результат (всегда, чтобы данные были актуальными)
                                                        # Если предыдущий статус был успешным, а новый ошибка - показываем ошибку
                                                        # Если предыдущий статус был ошибкой, а новый успешный - обновляем на успешный
                                                        mass_results[existing_result_idx] = result_row
                                                    else:
                                                        # Добавляем новый результат
                                                        mass_results.append(result_row)
                                                    
                                                    # СРАЗУ синхронизируем с session_state и сохраняем в файл после каждого товара
                                                    # Очищаем удаленные параметры перед сохранением
                                                    cleaned_mass_results = remove_deleted_params_from_mass_results(mass_results.copy())
                                                    st.session_state["mass_analysis_results"] = cleaned_mass_results
                                                    save_mass_analysis_results(cleaned_mass_results)
                                                    # Обновляем локальную переменную для отображения
                                                    mass_results = cleaned_mass_results
                                            
                                            except Exception as e:
                                                error_message = str(e)
                                                # Сохраняем полное сообщение об ошибке
                                                error_status = f"❌ Ошибка: {error_message[:100]}"
                                                if len(error_message) > 100:
                                                    error_status += "..."
                                                
                                                result_row = {
                                                    "Артикул": sku,
                                                    "URL": wb_url,
                                                    "Статус": error_status,
                                                    "Параметры": "",
                                                    "Ошибка (полная)": error_message  # Сохраняем полное сообщение для отчета
                                                }
                                                
                                                # Проверяем, есть ли уже результат для этого артикула
                                                sku_str = str(sku).replace(".0", "")
                                                existing_result_idx = None
                                                for i, existing_result in enumerate(mass_results):
                                                    existing_sku = str(existing_result.get("Артикул", "")).replace(".0", "")
                                                    if existing_sku == sku_str:
                                                        existing_result_idx = i
                                                        break
                                                
                                                # Если результат уже есть - обновляем, иначе добавляем
                                                if existing_result_idx is not None:
                                                    # Обновляем существующий результат (всегда, чтобы данные были актуальными)
                                                    # Если предыдущий статус был успешным, а новый ошибка - показываем ошибку
                                                    # Если предыдущий статус был ошибкой, а новый успешный - обновляем на успешный
                                                    mass_results[existing_result_idx] = result_row
                                                else:
                                                    # Добавляем новый результат
                                                    mass_results.append(result_row)
                                                
                                                # СРАЗУ синхронизируем с session_state и сохраняем в файл после каждого товара
                                                st.session_state["mass_analysis_results"] = mass_results
                                                save_mass_analysis_results(mass_results)
                                            
                                            # Вычисляем время обработки товара
                                            product_end_time = time.time()
                                            product_processing_time = product_end_time - product_start_time
                                            
                                            # Отмечаем товар как обработанный
                                            processed_skus.add(str(sku))
                                            
                                            # Определяем статус и иконку из последнего добавленного результата
                                            # Проверяем последний результат в mass_results для текущего артикула
                                            last_result_status = None
                                            for result in reversed(mass_results):
                                                if str(result.get("Артикул", "")).replace(".0", "") == str(sku).replace(".0", ""):
                                                    last_result_status = result.get("Статус", "")
                                                    break
                                            
                                            # Определяем иконку статуса
                                            if last_result_status and last_result_status.startswith("✅"):
                                                status_icon = "✅"
                                            elif last_result_status and "❌" in last_result_status:
                                                status_icon = "❌"
                                            elif last_result_status and "⚠️" in last_result_status:
                                                status_icon = "⚠️"
                                            else:
                                                status_icon = "⏳"
                                            
                                            elapsed_str = f"{product_processing_time:.1f}с"
                                            prev_time_str = f" | ⏱️ Предыдущий: {previous_product_time:.1f}с" if previous_product_time is not None else ""
                                            mass_status_text.text(f"{status_icon} [{processed_idx + 1}/{total_skus}] Артикул {sku} обработан за {elapsed_str}{prev_time_str}")
                                            
                                            # Обновляем прогресс-бар
                                            progress_value = min((processed_idx + 1) / total_skus, 1.0)
                                            mass_progress_bar.progress(progress_value)
                                            
                                            # Сохраняем время обработки текущего товара для следующего
                                            previous_product_time = product_processing_time
                                            
                                            # Вычисляем статистику (всегда актуальная)
                                            current_success_count = len([r for r in mass_results if r.get("Статус", "").startswith("✅")])
                                            current_error_count = len([r for r in mass_results if not r.get("Статус", "").startswith("✅")])
                                            
                                            # Обновляем таблицу результатов ПОСЛЕ КАЖДОГО товара
                                            # КРИТИЧЕСКИ ВАЖНО: Показываем количество товаров из mass_results, а не processed_idx
                                            # Используем актуальные данные из session_state (они уже синхронизированы выше)
                                            current_results = st.session_state.get("mass_analysis_results", mass_results)
                                            if current_results:
                                                results_df = pd.DataFrame(current_results)
                                                
                                                # Фильтруем удаленные параметры из столбцов DataFrame
                                                deleted_params = st.session_state.get("deleted_params", set())
                                                if not isinstance(deleted_params, set):
                                                    deleted_params = set(deleted_params) if deleted_params else set()
                                                
                                                # Удаляем столбцы, которые являются удаленными параметрами
                                                cols_to_keep = [col for col in results_df.columns if col not in deleted_params]
                                                if cols_to_keep:
                                                    results_df = results_df[cols_to_keep]
                                                
                                                # Пересчитываем статистику из актуальных результатов
                                                current_success_count = len([r for r in current_results if r.get("Статус", "").startswith("✅")])
                                                current_error_count = len([r for r in current_results if not r.get("Статус", "").startswith("✅")])
                                                
                                                with results_table_placeholder.container():
                                                    st.markdown(f"**📋 Обработано: {len(current_results)} товаров** (✅ Успешно: {current_success_count}, ❌ Ошибок: {current_error_count}) | 🔄 Всего попыток: {processed_idx + 1}/{total_skus}")
                                                    st.dataframe(results_df, use_container_width=True, height=400)
                                            
                                            # Сохраняем прогресс реже - каждые 5 товаров или в конце (для оптимизации)
                                            should_save_progress = (len(processed_skus) % 5 == 0) or (idx == len(remaining_skus) - 1)
                                            if should_save_progress:
                                                save_mass_analysis_progress(
                                                    skus_to_process,
                                                    processed_skus,
                                                    mass_results,
                                                    settings
                                                )
                                            
                                            # Обновляем кнопку применения параметров реже - каждые 3 товара или в конце
                                            should_update_button = (len(mass_results) % 3 == 0) or (idx == len(remaining_skus) - 1)
                                            if should_update_button:
                                                if current_success_count > 0:
                                                    with apply_button_placeholder.container():
                                                        st.divider()
                                                        col_apply1, col_apply2 = st.columns([1, 2])
                                                        with col_apply1:
                                                            # Используем уникальный ключ с индексом обработанных товаров
                                                            button_key = f"apply_mass_results_during_{len(mass_results)}"
                                                            if st.button("💾 Применить параметры в общую таблицу", type="primary", use_container_width=True, key=button_key):
                                                                # Сохраняем параметры в файл
                                                                if save_param_values_to_file():
                                                                    if save_param_history_to_file():
                                                                        # Перезагружаем параметры из файла, чтобы убедиться, что они применены
                                                                        load_param_values_from_file()
                                                                        st.success(f"✅ Параметры применены в общую таблицу! Обновлено {current_success_count} товаров.")
                                                                        # Обновляем таблицу для отображения новых параметров
                                                                        if "df" in st.session_state:
                                                                            # Перезагружаем данные, чтобы показать новые параметры
                                                                            st.session_state["data_loaded"] = False
                                                                            # Очищаем display_df, чтобы он пересоздался с новыми параметрами
                                                                            if "display_df" in st.session_state:
                                                                                del st.session_state["display_df"]
                                                                        st.rerun()
                                                                    else:
                                                                        st.error("❌ Ошибка при сохранении истории изменений")
                                                                else:
                                                                    st.error("❌ Ошибка при сохранении параметров")
                                                        with col_apply2:
                                                            st.info("💡 Параметры из успешных результатов будут применены в основную таблицу. Можно нажимать во время обработки.")
                                                else:
                                                    apply_button_placeholder.empty()
                                            
                                            # Задержка между товарами и переход к следующему
                                            if idx < len(remaining_skus) - 1:
                                                import time
                                                # Показываем, что переходим к следующему товару
                                                next_idx = processed_idx + 1
                                                mass_status_text.text(f"⏳ Переход к следующему товару... [{next_idx}/{total_skus}] | ⏱️ Предыдущий: {product_processing_time:.1f}с")
                                                time.sleep(mass_delay)
                                                # Обновляем статус перед началом обработки следующего товара
                                                mass_status_text.text(f"🔄 Начинаем обработку товара {next_idx}/{total_skus}...")
                                            else:
                                                # Последний товар обработан
                                                mass_status_text.text(f"✅ Все товары обработаны! Последний товар: {product_processing_time:.1f}с")
                                    
                                    mass_progress_bar.progress(1.0)
                                    mass_status_text.empty()
                                    
                                    # Удаляем сохраненный прогресс после успешного завершения
                                    clear_mass_analysis_progress()
                                    
                                    # Сбрасываем флаг запуска анализа, чтобы старые результаты снова показывались
                                    st.session_state["mass_analysis_btn_clicked"] = False
                                    
                                    # Финальные результаты
                                    st.markdown("---")
                                    st.markdown("### 📊 Итоговый отчет массового анализа")
                                    
                                    # Пересчитываем счетчики из актуальных результатов
                                    final_success_count = len([r for r in mass_results if r.get("Статус", "").startswith("✅")])
                                    final_error_count = len([r for r in mass_results if not r.get("Статус", "").startswith("✅")])
                                    
                                    col_success, col_error, col_total = st.columns(3)
                                    with col_success:
                                        st.metric("✅ Успешно", final_success_count)
                                    with col_error:
                                        st.metric("❌ Ошибок", final_error_count)
                                    with col_total:
                                        st.metric("📦 Всего", total_skus)
                                    
                                    # Определяем проблемные товары для кнопок повторной обработки
                                    error_results = [r for r in mass_results if not r.get("Статус", "").startswith("✅")]
                                    undefined_results = [r for r in mass_results if "⚠️ Параметры не определены" in r.get("Статус", "")]
                                    
                                    # Показываем кнопки повторной обработки сразу после итогового отчета
                                    if undefined_results or error_results:
                                        st.markdown("---")
                                        st.markdown("#### 🔄 Повторная обработка товаров")
                                        
                                        # Кнопка для товаров с неопределенными параметрами
                                        if undefined_results:
                                            if st.button(f"🔄 Обработать товары с неопределенными параметрами ({len(undefined_results)})", type="primary", use_container_width=True, key="retry_undefined_final_top"):
                                                retry_skus = [str(r["Артикул"]) for r in undefined_results]
                                                success_results = [r for r in mass_results if r.get("Статус", "").startswith("✅")]
                                                
                                                retry_settings = {
                                                    "max_images": mass_max_images if 'mass_max_images' in locals() else st.session_state.get("mass_max_images_slider", 5),
                                                    "delay": mass_delay if 'mass_delay' in locals() else st.session_state.get("mass_delay_slider", 2.0),
                                                    "selected_params": mass_selected_params if 'mass_selected_params' in locals() else st.session_state.get("mass_selected_params_multiselect", [])
                                                }
                                                
                                                retry_progress = {
                                                    "skus_to_process": retry_skus,
                                                    "processed_skus": [],
                                                    "results": success_results,
                                                    "settings": retry_settings,
                                                    "timestamp": pd.Timestamp.now().strftime("%Y-%m-%d %H:%M:%S"),
                                                    "is_retry": True
                                                }
                                                
                                                save_mass_analysis_progress(
                                                    retry_skus,
                                                    set(),
                                                    success_results,
                                                    retry_settings
                                                )
                                                
                                                st.session_state["resume_mass_analysis"] = True
                                                st.session_state["saved_progress"] = retry_progress
                                                st.success(f"✅ Запланирована обработка {len(retry_skus)} товаров с неопределенными параметрами")
                                                st.rerun()
                                        
                                        # Кнопка для всех товаров с ошибками
                                        if error_results:
                                            if st.button(f"🔄 Обработать все товары с ошибками ({len(error_results)})", type="primary", use_container_width=True, key="retry_all_errors_final_top"):
                                                retry_skus = [str(r["Артикул"]) for r in error_results]
                                                success_results = [r for r in mass_results if r.get("Статус", "").startswith("✅")]
                                                
                                                retry_settings = {
                                                    "max_images": mass_max_images if 'mass_max_images' in locals() else st.session_state.get("mass_max_images_slider", 5),
                                                    "delay": mass_delay if 'mass_delay' in locals() else st.session_state.get("mass_delay_slider", 2.0),
                                                    "selected_params": mass_selected_params if 'mass_selected_params' in locals() else st.session_state.get("mass_selected_params_multiselect", [])
                                                }
                                                
                                                retry_progress = {
                                                    "skus_to_process": retry_skus,
                                                    "processed_skus": [],
                                                    "results": success_results,
                                                    "settings": retry_settings,
                                                    "timestamp": pd.Timestamp.now().strftime("%Y-%m-%d %H:%M:%S"),
                                                    "is_retry": True
                                                }
                                                
                                                save_mass_analysis_progress(
                                                    retry_skus,
                                                    set(),
                                                    success_results,
                                                    retry_settings
                                                )
                                                
                                                st.session_state["resume_mass_analysis"] = True
                                                st.session_state["saved_progress"] = retry_progress
                                                st.success(f"✅ Запланирована обработка {len(retry_skus)} товаров с ошибками")
                                                st.rerun()
                                    
                                    # Детальный отчет об ошибках
                                    # Пересчитываем счетчики из актуальных результатов
                                    current_error_count = len([r for r in mass_results if not r.get("Статус", "").startswith("✅")])
                                    if current_error_count > 0:
                                        st.markdown("---")
                                        st.markdown("### ❌ Детальный отчет об ошибках")
                                        
                                        # Фильтруем товары с ошибками
                                        error_results = [r for r in mass_results if not r.get("Статус", "").startswith("✅")]
                                        undefined_results = [r for r in mass_results if "⚠️ Параметры не определены" in r.get("Статус", "")]
                                        
                                        # Проверяем товары с незаполненными параметрами из таблицы
                                        param_values = get_param_values()
                                        all_param_names = list(st.session_state.get("param_options", {}).keys())
                                        skus_with_empty_params = []
                                        
                                        if all_param_names and param_values:
                                            # Получаем все артикулы из таблицы
                                            source_df = st.session_state.get("df")
                                            if source_df is None:
                                                source_df = st.session_state.get("display_df")
                                            if source_df is not None and "Артикул" in source_df.columns:
                                                for sku in source_df["Артикул"].dropna().unique():
                                                    sku_str = str(sku).replace(".0", "")
                                                    # Проверяем, заполнены ли все параметры для этого товара
                                                    filled_params_count = 0
                                                    for param_name in all_param_names:
                                                        if param_name in param_values and sku_str in param_values[param_name]:
                                                            value = param_values[param_name][sku_str]
                                                            if value and str(value).strip():
                                                                filled_params_count += 1
                                                    
                                                    # Если ни один параметр не заполнен, добавляем в список
                                                    if filled_params_count == 0:
                                                        skus_with_empty_params.append(sku_str)
                                        
                                        if error_results:
                                            error_df = pd.DataFrame(error_results)
                                            # Оставляем только важные колонки для отчета об ошибках
                                            error_report_cols = ["Артикул", "URL", "Статус"]
                                            if "Параметры" in error_df.columns:
                                                error_report_cols.append("Параметры")
                                            
                                            error_report_df = error_df[error_report_cols].copy()
                                            
                                            # Кнопка для обработки всех товаров с ошибками
                                            st.markdown("#### 🔄 Повторная обработка товаров с ошибками")
                                            col_retry_all, col_retry_info = st.columns([2, 3])
                                            with col_retry_all:
                                                if st.button("🔄 Обработать все товары с ошибками снова", type="primary", key="retry_all_errors"):
                                                    # Извлекаем все артикулы с ошибками
                                                    retry_skus = [str(r["Артикул"]) for r in error_results]
                                                    
                                                    # Сохраняем настройки для повторной обработки
                                                    retry_settings = {
                                                        "max_images": mass_max_images,
                                                        "delay": mass_delay
                                                    }
                                                    
                                                    # Сохраняем текущие успешные результаты для последующего объединения
                                                    success_results = [r for r in mass_results if r.get("Статус", "").startswith("✅")]
                                                    
                                                    # Создаем новый прогресс для повторной обработки
                                                    retry_progress = {
                                                        "skus_to_process": retry_skus,
                                                        "processed_skus": [],
                                                        "results": success_results,  # Сохраняем успешные результаты
                                                        "settings": retry_settings,
                                                        "timestamp": pd.Timestamp.now().strftime("%Y-%m-%d %H:%M:%S"),
                                                        "is_retry": True,
                                                        "original_results": mass_results  # Сохраняем все исходные результаты
                                                    }
                                                    
                                                    save_mass_analysis_progress(
                                                        retry_skus,
                                                        set(),
                                                        success_results,  # Сохраняем успешные результаты
                                                        retry_settings
                                                    )
                                                    
                                                    st.session_state["resume_mass_analysis"] = True
                                                    st.session_state["saved_progress"] = retry_progress
                                                    st.success(f"✅ Запланирована повторная обработка {len(retry_skus)} товаров с ошибками")
                                                    st.rerun()
                                            with col_retry_info:
                                                # Пересчитываем счетчики из актуальных результатов
                                                current_success_count_info = len([r for r in mass_results if r.get("Статус", "").startswith("✅")])
                                                st.info(f"📊 Будет обработано **{len(error_results)}** товаров с ошибками. Успешные результаты ({current_success_count_info}) будут сохранены.")
                                            
                                            st.divider()
                                            
                                            # Группируем ошибки по типу
                                            error_types = {}
                                            for _, row in error_report_df.iterrows():
                                                status = str(row.get("Статус", ""))
                                                if "APITimeoutError" in status or "timeout" in status.lower():
                                                    error_type = "⏱️ Таймаут API"
                                                elif "Rate limit" in status or "429" in status:
                                                    error_type = "🚦 Превышен лимит запросов"
                                                elif "Параметры не определены" in status:
                                                    error_type = "⚠️ Параметры не определены"
                                                elif "Ошибка" in status:
                                                    error_type = "❌ Прочие ошибки"
                                                else:
                                                    error_type = "❓ Неизвестная ошибка"
                                                
                                                if error_type not in error_types:
                                                    error_types[error_type] = []
                                                error_types[error_type].append(row.to_dict())
                                            
                                            # Показываем ошибки по типам
                                            st.markdown("#### 📊 Ошибки по типам")
                                            for error_type, errors in error_types.items():
                                                with st.expander(f"{error_type} ({len(errors)} товаров)", expanded=False):
                                                    error_type_df = pd.DataFrame(errors)
                                                    st.dataframe(error_type_df, use_container_width=True, height=min(400, len(errors) * 50))
                                                    
                                                    # Кнопка для повторной обработки товаров с ошибками конкретного типа
                                                    if st.button(f"🔄 Повторить обработку ({len(errors)} товаров)", key=f"retry_{error_type}"):
                                                        # Извлекаем артикулы для повторной обработки
                                                        retry_skus = [str(r["Артикул"]) for r in errors]
                                                        
                                                        # Сохраняем настройки для повторной обработки
                                                        retry_settings = {
                                                            "max_images": mass_max_images,
                                                            "delay": mass_delay
                                                        }
                                                        
                                                        # Сохраняем текущие успешные результаты
                                                        success_results = [r for r in mass_results if r.get("Статус", "").startswith("✅")]
                                                        
                                                        # Создаем новый прогресс для повторной обработки
                                                        retry_progress = {
                                                            "skus_to_process": retry_skus,
                                                            "processed_skus": [],
                                                            "results": success_results,  # Сохраняем успешные результаты
                                                            "settings": retry_settings,
                                                            "timestamp": pd.Timestamp.now().strftime("%Y-%m-%d %H:%M:%S"),
                                                            "is_retry": True,
                                                            "original_results": mass_results
                                                        }
                                                        
                                                        save_mass_analysis_progress(
                                                            retry_skus,
                                                            set(),
                                                            success_results,  # Сохраняем успешные результаты
                                                            retry_settings
                                                        )
                                                        
                                                        st.session_state["resume_mass_analysis"] = True
                                                        st.session_state["saved_progress"] = retry_progress
                                                        st.success(f"✅ Запланирована повторная обработка {len(retry_skus)} товаров")
                                                        st.rerun()
                                            
                                            st.divider()
                                            
                                            # Общая таблица всех ошибок
                                            st.markdown("#### 📋 Все товары с ошибками")
                                            st.dataframe(error_report_df, use_container_width=True, height=400)
                                            
                                            # Экспорт в CSV
                                            csv = error_report_df.to_csv(index=False).encode('utf-8-sig')
                                            st.download_button(
                                                label="📥 Скачать отчет об ошибках (CSV)",
                                                data=csv,
                                                file_name=f"mass_analysis_errors_{pd.Timestamp.now().strftime('%Y%m%d_%H%M%S')}.csv",
                                                mime="text/csv",
                                                key="download_errors_csv"
                                            )
                                    
                                    # Сохраняем результаты в session_state для последующего просмотра
                                    st.session_state["mass_analysis_results"] = mass_results
                                    save_mass_analysis_results(mass_results)
                                    
                                    # Автоматически применяем параметры из результатов
                                    # Пересчитываем счетчики из актуальных результатов
                                    current_success_count = len([r for r in mass_results if r.get("Статус", "").startswith("✅")])
                                    if current_success_count > 0:
                                        applied_params_count = apply_mass_analysis_results_to_params(mass_results)
                                        if applied_params_count > 0:
                                            # Сохраняем параметры в файл
                                            save_param_values_to_file()
                                            st.success(f"✅ Анализ завершен! Параметры автоматически применены ({applied_params_count} параметров).")
                                        else:
                                            st.success(f"✅ Анализ завершен!")
                                    else:
                                        st.info("ℹ️ Нет успешных результатов для применения параметров")
                            else:
                                st.warning("⚠️ Нет товаров для анализа")
            
            st.divider()
            # ========== КОНЕЦ СЕКЦИИ: Определение параметров по ссылке ==========
            
            # Таблица с параметрами удалена - используется основная таблица в первой вкладке "📊 Анализ данных"
        
        with tab3:
            st.subheader("📈 Аналитика по параметрам")
            
            # Получаем данные параметров
            param_values = get_param_values()
            param_options = st.session_state.get("param_options", {})

            if not param_values and not param_options:
                st.warning("Сначала установите параметры товаров во второй вкладке")
            else:
                # Выбор параметра для анализа
                # Объединяем параметры из param_values и param_options, чтобы показывать все определенные параметры
                params_from_values = set(param_values.keys()) if param_values else set()
                params_from_options = set(param_options.keys()) if param_options else set()
                available_params = sorted(list(params_from_values | params_from_options))
                if available_params:
                    # Кнопка экспорта всех параметров в Excel
                    col_export_all, col_select = st.columns([1, 2])
                    
                    with col_export_all:
                        if st.button("📊 Экспорт всех параметров в Excel", type="secondary"):
                            try:
                                # Создаем Excel файл с несколькими листами
                                import io
                                from openpyxl import Workbook
                                
                                wb = Workbook()
                                # Удаляем дефолтный лист
                                wb.remove(wb.active)
                                
                                # Получаем имя загруженного файла для названия экспорта
                                base_filename = "analytics_all_parameters"
                                if hasattr(uploaded, 'name') and uploaded.name:
                                    name_without_ext = os.path.splitext(uploaded.name)[0]
                                    base_filename = f"{name_without_ext}_analytics_all_parameters"
                                
                                # Создаем лист для каждого параметра
                                for param_name in available_params:
                                    ws = wb.create_sheet(title=param_name)
                                    
                                    # Получаем данные для этого параметра
                                    param_values_set = set()
                                    if param_name in param_values:
                                        for sku, value in param_values[param_name].items():
                                            if value:
                                                param_values_set.add(value)
                                        
                                        # Создаем аналитику для этого параметра
                                        analytics_data = []
                                        
                                        for param_value in sorted(param_values_set):
                                            matching_skus = []
                                            if param_name in param_values:
                                                for sku, value in param_values[param_name].items():
                                                    if value == param_value:
                                                        matching_skus.append(sku)
                                            
                                            if matching_skus:
                                                mask = df["Артикул"].astype(str).str.replace(".0", "").isin(matching_skus)
                                                filtered_df = df[mask]
                                                
                                                if not filtered_df.empty:
                                                    total_revenue = filtered_df["Выручка"].sum() if "Выручка" in filtered_df.columns else 0
                                                    total_orders = filtered_df["Заказы"].sum() if "Заказы" in filtered_df.columns else 0
                                                    avg_price = filtered_df["Средняя цена"].mean() if "Средняя цена" in filtered_df.columns else 0
                                                    lost_revenue = filtered_df["Упущенная выручка"].sum() if "Упущенная выручка" in filtered_df.columns else 0
                                                    avg_position = filtered_df["Позиция в выдаче (средняя)"].mean() if "Позиция в выдаче (средняя)" in filtered_df.columns else 0
                                                    avg_cpm = filtered_df["Стоимость за 1000 показов на 1 артикул"].mean() if "Стоимость за 1000 показов на 1 артикул" in filtered_df.columns else 0
                                                    
                                                    analytics_data.append({
                                                        param_value: {
                                                            'Общая выручка': total_revenue,
                                                            'Количество артикулов': len(filtered_df),
                                                            'Выручка на 1 артикул': total_revenue / len(filtered_df) if len(filtered_df) > 0 else 0,
                                                            'Средняя цена без СПП': avg_price,
                                                            'Упущенная выручка': lost_revenue,
                                                            'Упущенная выручка на 1 артикул': lost_revenue / len(filtered_df) if len(filtered_df) > 0 else 0,
                                                            'Позиция в выдаче (средняя)': avg_position,
                                                            'Стоимость за 1000 показов на 1 артикул': avg_cpm
                                                        }
                                                    })
                                            
                                            # Записываем данные в лист
                                            if analytics_data:
                                                # Заголовки
                                                ws['A1'] = 'Метрика'
                                                col = 2
                                                param_values_list = sorted([list(item.keys())[0] for item in analytics_data])
                                                
                                                for param_val in param_values_list:
                                                    ws.cell(row=1, column=col, value=param_val)
                                                    col += 1
                                                
                                                # Данные
                                                metric_names = [
                                                    "Общая выручка",
                                                    "Количество артикулов", 
                                                    "Выручка на 1 артикул",
                                                    "Средняя цена без СПП",
                                                    "Упущенная выручка",
                                                    "Упущенная выручка на 1 артикул",
                                                    "Позиция в выдаче (средняя)",
                                                    "Стоимость за 1000 показов на 1 артикул"
                                                ]
                                                
                                                for row, metric in enumerate(metric_names, 2):
                                                    ws.cell(row=row, column=1, value=metric)
                                                    
                                                    col = 2
                                                    for param_val in param_values_list:
                                                        # Находим данные для этого значения параметра
                                                        for item in analytics_data:
                                                            if param_val in item:
                                                                metrics = item[param_val]
                                                                value = metrics.get(metric, 0)
                                                                if metric in ["Общая выручка", "Выручка на 1 артикул", "Средняя цена без СПП", "Упущенная выручка", "Упущенная выручка на 1 артикул"]:
                                                                    ws.cell(row=row, column=col, value=value)
                                                                else:
                                                                    ws.cell(row=row, column=col, value=value)
                                                                break
                                                    col += 1
                                
                                # Сохраняем в байты
                                excel_buffer = io.BytesIO()
                                wb.save(excel_buffer)
                                excel_buffer.seek(0)
                                
                                st.download_button(
                                    label="💾 Скачать Excel файл",
                                    data=excel_buffer.getvalue(),
                                    file_name=f"{base_filename}.xlsx",
                                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                                )
                                
                                st.success("✅ Excel файл готов к скачиванию!")
                                
                            except Exception as e:
                                st.error(f"❌ Ошибка создания Excel файла: {e}")
                    
                    with col_select:
                        # Определяем иерархию параметров
                        # Иерархия: Тип → Подтип (Рукав, Ворот) → Остальные параметры (Цвет, Строчка и т.д.)
                        hierarchy_params_list = get_hierarchy_params()
                        type_param = hierarchy_params_list[0] if hierarchy_params_list else "Тип"
                        subtype_params = get_subtype_params()  # Параметры, которые находятся внутри "Подтип"
                        visual_params = get_visual_params()
                        
                        # Разделяем параметры на категории
                        type_param_available = type_param in available_params
                        subtype_params_available = [p for p in subtype_params if p in available_params]
                        visual_params_list = [p for p in available_params if p in visual_params or any(vp in p for vp in visual_params)]
                        other_params_list = [p for p in available_params if p != type_param and p not in subtype_params and p not in visual_params_list]
                        
                        # Трехуровневая иерархия выбора
                        selected_type_value = None
                        selected_subtype_param = None
                        selected_subtype_value = None
                        selected_other_param = None
                        use_combined_analysis = False
                        combined_hierarchy_param = None
                        combined_hierarchy_value = None
                        combined_subtype_param = None
                        combined_subtype_value = None
                        
                        if type_param_available:
                            st.info(f"📊 **Иерархия сегментов**: Тип → Подтип (Рукав, Ворот) → Остальные параметры")
                            
                            # Уровень 1: Выбор параметра "Тип" и его значения
                            type_param_values = set()
                            if type_param in param_values:
                                for sku, value in param_values[type_param].items():
                                    if value and str(value).strip():
                                        type_param_values.add(str(value).strip())
                            
                            if type_param_values:
                                selected_type_values = st.multiselect(
                                    f"1️⃣ Выберите значения параметра '{type_param}' (можно выбрать несколько):",
                                    sorted(list(type_param_values)),
                                    key="analytics_type_value_selector",
                                    help="Выберите одно или несколько значений типа товара"
                                )
                                
                                if selected_type_values:
                                    if len(selected_type_values) == 1:
                                        selected_type_value = selected_type_values[0]
                                    else:
                                        selected_type_value = selected_type_values
                                    
                                    # Уровень 2: Выбор подтипа (Рукав, Ворот), если выбран тип
                                    if subtype_params_available:
                                        st.divider()
                                        selected_subtype_param = st.selectbox(
                                            "2️⃣ Выберите параметр подтипа:",
                                            ["Не выбран"] + subtype_params_available,
                                            key="analytics_subtype_param_selector",
                                            help="Выберите параметр подтипа (Рукав или Ворот)"
                                        )
                                        
                                        if selected_subtype_param != "Не выбран":
                                            # Получаем значения выбранного параметра подтипа
                                            subtype_param_values = set()
                                            if selected_subtype_param in param_values:
                                                for sku, value in param_values[selected_subtype_param].items():
                                                    if value and str(value).strip():
                                                        # Фильтруем только по выбранным типам
                                                        if type_param in param_values and sku in param_values[type_param]:
                                                            type_val = param_values[type_param][sku]
                                                            if isinstance(selected_type_value, list):
                                                                if str(type_val).strip() in [str(v).strip() for v in selected_type_value]:
                                                                    subtype_param_values.add(str(value).strip())
                                                            else:
                                                                if str(type_val).strip() == str(selected_type_value).strip():
                                                                    subtype_param_values.add(str(value).strip())
                                            
                                            if subtype_param_values:
                                                selected_subtype_values = st.multiselect(
                                                    f"Выберите значения параметра '{selected_subtype_param}' (можно выбрать несколько):",
                                                    sorted(list(subtype_param_values)),
                                                    key="analytics_subtype_value_selector",
                                                    help="Выберите одно или несколько значений подтипа"
                                                )
                                                
                                                if selected_subtype_values:
                                                    if len(selected_subtype_values) == 1:
                                                        selected_subtype_value = selected_subtype_values[0]
                                                    else:
                                                        selected_subtype_value = selected_subtype_values
                                                    combined_subtype_param = selected_subtype_param
                                                    combined_subtype_value = selected_subtype_value
                                            
                                            # Уровень 3: Выбор остальных параметров
                                            all_other_params = visual_params_list + other_params_list
                                            if all_other_params:
                                                st.divider()
                                                selected_other_param = st.selectbox(
                                                    "3️⃣ Выберите параметр для анализа:",
                                                    ["Не выбран"] + all_other_params,
                                                    key="analytics_other_param_selector",
                                                    help="Выберите параметр для анализа (Цвет, Строчка и т.д.)"
                                                )
                                                
                                                if selected_other_param != "Не выбран":
                                                    selected_param = selected_other_param
                                                    use_combined_analysis = True
                                                    combined_hierarchy_param = type_param
                                                    combined_hierarchy_value = selected_type_value
                                                else:
                                                    selected_param = selected_subtype_param
                                                    combined_hierarchy_param = type_param
                                                    combined_hierarchy_value = selected_type_value
                                            else:
                                                selected_param = selected_subtype_param
                                                combined_hierarchy_param = type_param
                                                combined_hierarchy_value = selected_type_value
                                        else:
                                            # Если подтип не выбран, можно выбрать остальные параметры
                                            all_other_params = visual_params_list + other_params_list
                                            if all_other_params:
                                                st.divider()
                                                selected_other_param = st.selectbox(
                                                    "2️⃣ Выберите параметр для анализа:",
                                                    ["Не выбран"] + all_other_params,
                                                    key="analytics_other_param_selector_no_subtype",
                                                    help="Выберите параметр для анализа (Цвет, Строчка и т.д.)"
                                                )
                                                
                                                if selected_other_param != "Не выбран":
                                                    selected_param = selected_other_param
                                                    use_combined_analysis = True
                                                    combined_hierarchy_param = type_param
                                                    combined_hierarchy_value = selected_type_value
                                                else:
                                                    selected_param = type_param
                                                    combined_hierarchy_param = None
                                                    combined_hierarchy_value = None
                                            else:
                                                selected_param = type_param
                                                combined_hierarchy_param = None
                                                combined_hierarchy_value = None
                                    else:
                                        # Если нет параметров подтипа, сразу выбираем остальные параметры
                                        all_other_params = visual_params_list + other_params_list
                                        if all_other_params:
                                            st.divider()
                                            selected_other_param = st.selectbox(
                                                "2️⃣ Выберите параметр для анализа:",
                                                ["Не выбран"] + all_other_params,
                                                key="analytics_other_param_selector_no_subtype",
                                                help="Выберите параметр для анализа (Цвет, Строчка и т.д.)"
                                            )
                                            
                                            if selected_other_param != "Не выбран":
                                                selected_param = selected_other_param
                                                use_combined_analysis = True
                                                combined_hierarchy_param = type_param
                                                combined_hierarchy_value = selected_type_value
                                            else:
                                                selected_param = type_param
                                                combined_hierarchy_param = None
                                                combined_hierarchy_value = None
                                        else:
                                            selected_param = type_param
                                            combined_hierarchy_param = None
                                            combined_hierarchy_value = None
                                else:
                                    selected_param = None
                                    combined_hierarchy_param = None
                                    combined_hierarchy_value = None
                            else:
                                if type_param in param_options:
                                    st.info(f"ℹ️ Параметр '{type_param}' определен, но еще не заполнен для товаров. Заполните его во вкладке '⚙️ Установка параметров'.")
                                else:
                                    st.warning(f"⚠️ Параметр '{type_param}' не найден. Создайте его во вкладке '⚙️ Установка параметров'.")
                                selected_param = None
                                combined_hierarchy_param = None
                                combined_hierarchy_value = None
                        else:
                            # Если нет параметра "Тип", используем обычный выбор
                            selected_param = st.selectbox(
                                "Выберите параметр для анализа",
                                available_params,
                                key="analytics_param_selector"
                            )
                            use_combined_analysis = False
                            combined_hierarchy_param = None
                            combined_hierarchy_value = None
                    
                    if selected_param and selected_param != "Не выбран":
                        # Формируем заголовок с учетом иерархии
                        title_parts = []
                        if combined_hierarchy_param and combined_hierarchy_value:
                            if isinstance(combined_hierarchy_value, list):
                                values_str = ", ".join(combined_hierarchy_value)
                                title_parts.append(f"{combined_hierarchy_param} = [{values_str}]")
                            else:
                                title_parts.append(f"{combined_hierarchy_param} = {combined_hierarchy_value}")
                        
                        if combined_subtype_param and combined_subtype_value:
                            if isinstance(combined_subtype_value, list):
                                subtype_str = ", ".join(combined_subtype_value)
                                title_parts.append(f"{combined_subtype_param} = [{subtype_str}]")
                            else:
                                title_parts.append(f"{combined_subtype_param} = {combined_subtype_value}")
                        
                        if title_parts:
                            hierarchy_str = " → ".join(title_parts)
                            st.write(f"**Аналитика по параметру: {selected_param} (в контексте: {hierarchy_str})**")
                        else:
                            st.write(f"**Аналитика по параметру: {selected_param}**")
                        
                        # Создаем DataFrame для анализа
                        analytics_data = []
                        
                        # Если используется комбинированный анализ, группируем по комбинации параметров
                        if use_combined_analysis and combined_hierarchy_param:
                            # Комбинированный анализ: группируем по комбинации иерархического и визуального параметров
                            combinations = {}
                            
                            # Собираем все комбинации значений двух параметров
                            for sku in df["Артикул"].astype(str).str.replace(".0", "").unique():
                                sku_str = str(sku)
                                hierarchy_value = None
                                visual_value = None
                                
                                if combined_hierarchy_param in param_values and sku_str in param_values[combined_hierarchy_param]:
                                    hierarchy_value = param_values[combined_hierarchy_param][sku_str]
                                
                                # Если выбрано конкретное значение(я) иерархии, фильтруем только по ним
                                if combined_hierarchy_value:
                                    if isinstance(combined_hierarchy_value, list):
                                        # Если выбрано несколько значений, проверяем вхождение в список
                                        if not hierarchy_value or str(hierarchy_value).strip() not in [str(v).strip() for v in combined_hierarchy_value]:
                                            continue
                                    else:
                                        # Если выбрано одно значение, проверяем точное совпадение
                                        if not hierarchy_value or str(hierarchy_value).strip() != str(combined_hierarchy_value).strip():
                                            continue
                                
                                # Если выбран подтип, фильтруем по нему
                                if combined_subtype_param and combined_subtype_value:
                                    if combined_subtype_param in param_values and sku_str in param_values[combined_subtype_param]:
                                        subtype_val = param_values[combined_subtype_param][sku_str]
                                        if isinstance(combined_subtype_value, list):
                                            if not subtype_val or str(subtype_val).strip() not in [str(v).strip() for v in combined_subtype_value]:
                                                continue
                                        else:
                                            if not subtype_val or str(subtype_val).strip() != str(combined_subtype_value).strip():
                                                continue
                                    else:
                                        continue
                                
                                if selected_param in param_values and sku_str in param_values[selected_param]:
                                    visual_value = param_values[selected_param][sku_str]
                                
                                if visual_value and str(visual_value).strip():
                                    # Нормализуем значения для избежания дубликатов (убираем пробелы, приводим к единому формату)
                                    # Используем strip() для удаления пробелов, но сохраняем оригинальный регистр для отображения
                                    normalized_visual = str(visual_value).strip()
                                    normalized_hierarchy = str(hierarchy_value).strip() if hierarchy_value else ""
                                    
                                    # Если выбрано конкретное значение иерархии, упрощаем ключ
                                    if combined_hierarchy_value:
                                        combo_key = f"{selected_param}: {normalized_visual}"
                                    else:
                                        combo_key = f"{combined_hierarchy_param}: {normalized_hierarchy} | {selected_param}: {normalized_visual}"
                                    
                                    # Нормализуем ключ для сравнения (приводим к нижнему регистру для избежания дубликатов)
                                    combo_key_normalized = combo_key.lower()
                                    
                                    # Используем нормализованный ключ для хранения, но сохраняем оригинальный для отображения
                                    if combo_key_normalized not in combinations:
                                        combinations[combo_key_normalized] = {
                                            "display_key": combo_key,  # Оригинальный ключ для отображения
                                            "skus": set()  # Set для хранения артикулов без дубликатов
                                        }
                                    combinations[combo_key_normalized]["skus"].add(sku_str)
                            
                            # Для каждой комбинации собираем метрики
                            for combo_key_normalized, combo_data in combinations.items():
                                combo_key = combo_data["display_key"]  # Используем оригинальный ключ для отображения
                                skus = combo_data["skus"]
                                # Преобразуем set в list для использования в isin
                                skus_list = list(skus)
                                mask = df["Артикул"].astype(str).str.replace(".0", "").isin(skus_list)
                                filtered_df = df[mask]
                                
                                if not filtered_df.empty:
                                    total_revenue = filtered_df["Выручка"].sum() if "Выручка" in filtered_df.columns else 0
                                    total_orders = filtered_df["Заказы"].sum() if "Заказы" in filtered_df.columns else 0
                                    avg_price = filtered_df["Средняя цена"].mean() if "Средняя цена" in filtered_df.columns else 0
                                    lost_revenue = filtered_df["Упущенная выручка"].sum() if "Упущенная выручка" in filtered_df.columns else 0
                                    revenue_per_product = total_revenue / len(filtered_df) if len(filtered_df) > 0 else 0
                                    lost_revenue_per_product = lost_revenue / len(filtered_df) if len(filtered_df) > 0 else 0
                                    avg_position = filtered_df["Позиция в выдаче"].mean() if "Позиция в выдаче" in filtered_df.columns else 0
                                    avg_cpm = filtered_df["Стоимость за 1000 показов"].mean() if "Стоимость за 1000 показов" in filtered_df.columns else 0
                                    
                                    analytics_data.append({
                                        "Метрика": selected_param,
                                        combo_key: {
                                            "Общая выручка": total_revenue,
                                            "Количество артикулов": len(filtered_df),
                                            "Выручка на 1 артикул": revenue_per_product,
                                            "Средняя цена без СПП": avg_price,
                                            "Упущенная выручка": lost_revenue,
                                            "Упущенная выручка на 1 артикул": lost_revenue_per_product,
                                            "Позиция в выдаче (средняя)": avg_position,
                                            "Стоимость за 1000 показов на 1 артикул": avg_cpm,
                                        }
                                    })
                        else:
                            # Обычный анализ по одному параметру
                            # Получаем уникальные значения параметра
                            param_values_set = set()
                            if selected_param in param_values:
                                for sku, value in param_values[selected_param].items():
                                    if value:
                                        # Если есть фильтр по значению(ям) иерархии, проверяем его
                                        if combined_hierarchy_value and combined_hierarchy_param:
                                            # Проверяем, что товар соответствует выбранному значению(ям) иерархии
                                            if combined_hierarchy_param in param_values and sku in param_values[combined_hierarchy_param]:
                                                hierarchy_val = param_values[combined_hierarchy_param][sku]
                                                if isinstance(combined_hierarchy_value, list):
                                                    # Если выбрано несколько значений, проверяем вхождение в список
                                                    if not hierarchy_val or str(hierarchy_val).strip() not in [str(v).strip() for v in combined_hierarchy_value]:
                                                        continue
                                                else:
                                                    # Если выбрано одно значение, проверяем точное совпадение
                                                    if not hierarchy_val or str(hierarchy_val).strip() != str(combined_hierarchy_value).strip():
                                                        continue
                                            else:
                                                continue
                                        param_values_set.add(value)
                            
                            # Для каждого значения параметра собираем метрики
                            for param_value in sorted(param_values_set):
                                # Находим артикулы с этим значением параметра
                                matching_skus = []
                                if selected_param in param_values:
                                    for sku, value in param_values[selected_param].items():
                                        if value == param_value:
                                            # Если есть фильтр по значению(ям) иерархии, проверяем его
                                            if combined_hierarchy_value and combined_hierarchy_param:
                                                if combined_hierarchy_param in param_values and sku in param_values[combined_hierarchy_param]:
                                                    hierarchy_val = param_values[combined_hierarchy_param][sku]
                                                    if isinstance(combined_hierarchy_value, list):
                                                        # Если выбрано несколько значений, проверяем вхождение в список
                                                        if not hierarchy_val or str(hierarchy_val).strip() not in [str(v).strip() for v in combined_hierarchy_value]:
                                                            continue
                                                    else:
                                                        # Если выбрано одно значение, проверяем точное совпадение
                                                        if not hierarchy_val or str(hierarchy_val).strip() != str(combined_hierarchy_value).strip():
                                                            continue
                                                else:
                                                    continue
                                        matching_skus.append(sku)
                            
                            if matching_skus:
                                # Фильтруем исходные данные по этим артикулам
                                mask = df["Артикул"].astype(str).str.replace(".0", "").isin(matching_skus)
                                filtered_df = df[mask]
                                
                                if not filtered_df.empty:
                                    # Вычисляем метрики
                                    total_revenue = filtered_df["Выручка"].sum() if "Выручка" in filtered_df.columns else 0
                                    total_orders = filtered_df["Заказы"].sum() if "Заказы" in filtered_df.columns else 0
                                    avg_price = filtered_df["Средняя цена"].mean() if "Средняя цена" in filtered_df.columns else 0
                                    lost_revenue = filtered_df["Упущенная выручка"].sum() if "Упущенная выручка" in filtered_df.columns else 0
                                    revenue_per_product = total_revenue / len(filtered_df) if len(filtered_df) > 0 else 0
                                    lost_revenue_per_product = lost_revenue / len(filtered_df) if len(filtered_df) > 0 else 0
                                    avg_position = filtered_df["Позиция в выдаче"].mean() if "Позиция в выдаче" in filtered_df.columns else 0
                                    avg_cpm = filtered_df["Стоимость за 1000 показов"].mean() if "Стоимость за 1000 показов" in filtered_df.columns else 0
                                    
                                    analytics_data.append({
                                        "Метрика": selected_param,
                                        param_value: {
                                            "Общая выручка": total_revenue,
                                            "Количество артикулов": len(filtered_df),
                                            "Выручка на 1 артикул": revenue_per_product,
                                            "Средняя цена без СПП": avg_price,
                                            "Упущенная выручка": lost_revenue,
                                            "Упущенная выручка на 1 артикул": lost_revenue_per_product,
                                            "Позиция в выдаче (средняя)": avg_position,
                                            "Стоимость за 1000 показов на 1 артикул": avg_cpm,
                                        }
                                    })
                        
                        if analytics_data:
                            # Создаем сводную таблицу
                            summary_data = {}
                            for item in analytics_data:
                                for param_val, metrics in item.items():
                                    if param_val != "Метрика":
                                        summary_data[param_val] = metrics
                            
                            # Создаем сводную таблицу в стиле как на картинке
                            if summary_data:
                                # Получаем все значения параметра (цвета) и сортируем их
                                param_values_list = sorted(summary_data.keys())
                                
                                # Создаем DataFrame с метриками по строкам и значениями параметра по столбцам
                                metric_names = [
                                    "Общая выручка",
                                    "Количество артикулов", 
                                    "Выручка на 1 артикул",
                                    "Средняя цена без СПП",
                                    "Упущенная выручка",
                                    "Упущенная выручка на 1 артикул",
                                    "Позиция в выдаче (средняя)",
                                    "Стоимость за 1000 показов на 1 артикул"
                                ]
                                
                                table_data = {"Метрика": metric_names}
                                
                                # Сохраняем числовые данные для цветового кодирования
                                numeric_data = {}
                                
                                # Добавляем данные для каждого значения параметра
                                for param_value in param_values_list:
                                    metrics = summary_data[param_value]
                                    # Сохраняем числовые значения
                                    numeric_data[param_value] = [
                                        metrics['Общая выручка'],
                                        metrics['Количество артикулов'],
                                        metrics['Выручка на 1 артикул'],
                                        metrics['Средняя цена без СПП'],
                                        metrics['Упущенная выручка'],
                                        metrics['Упущенная выручка на 1 артикул'],
                                        metrics['Позиция в выдаче (средняя)'],  # Для позиции меньше = лучше, обработаем отдельно
                                        metrics['Стоимость за 1000 показов на 1 артикул']
                                    ]
                                    
                                    # Форматированные значения для отображения
                                    table_data[param_value] = [
                                        f"₽{metrics['Общая выручка']:,.0f}".replace(",", " "),
                                        f"{metrics['Количество артикулов']:,.0f}".replace(",", " "),
                                        f"₽{metrics['Выручка на 1 артикул']:,.0f}".replace(",", " "),
                                        f"₽{metrics['Средняя цена без СПП']:,.0f}".replace(",", " "),
                                        f"₽{metrics['Упущенная выручка']:,.0f}".replace(",", " "),
                                        f"₽{metrics['Упущенная выручка на 1 артикул']:,.0f}".replace(",", " "),
                                        f"{metrics['Позиция в выдаче (средняя)']:,.0f}".replace(",", " "),
                                        f"{metrics['Стоимость за 1000 показов на 1 артикул']:,.0f}".replace(",", " ")
                                    ]
                                
                                # Создаем DataFrame
                                display_df = pd.DataFrame(table_data)
                                
                                # Функция для выделения лучших результатов
                                def highlight_best_values(values, reverse=False):
                                    """Выделяет зеленым только лучшие результаты"""
                                    if not values or all(pd.isna(v) or v == 0 for v in values):
                                        return ['background-color: white'] * len(values)
                                    
                                    # Очищаем от NaN и нулевых значений для поиска лучшего
                                    clean_values = [v for v in values if not pd.isna(v) and v != 0]
                                    if not clean_values:
                                        return ['background-color: white'] * len(values)
                                    
                                    # Находим лучшее значение
                                    if reverse:  # Для позиции: меньше = лучше
                                        best_val = min(clean_values)
                                    else:  # Для остальных метрик: больше = лучше
                                        best_val = max(clean_values)
                                    
                                    colors = []
                                    for val in values:
                                        if pd.isna(val) or val == 0:
                                            colors.append('background-color: white')
                                        elif val == best_val:
                                            colors.append('background-color: lightgreen')  # Зеленый для лучших
                                        else:
                                            colors.append('background-color: white')  # Белый для остальных
                                    
                                    return colors
                                
                                # Применяем цветовое кодирование
                                def apply_colors(df):
                                    # Создаем стили для каждой строки
                                    styles = pd.DataFrame('', index=df.index, columns=df.columns)
                                    
                                    for i, metric in enumerate(metric_names):
                                        row_values = []
                                        for param_value in param_values_list:
                                            row_values.append(numeric_data[param_value][i])
                                        
                                        # Для позиции используем обратную логику (меньше = лучше)
                                        reverse_logic = (metric == "Позиция в выдаче (средняя)")
                                        colors = highlight_best_values(row_values, reverse=reverse_logic)
                                        
                                        # Применяем цвета к соответствующим ячейкам
                                        for j, param_value in enumerate(param_values_list):
                                            styles.iloc[i, j + 1] = colors[j]  # +1 потому что первый столбец - "Метрика"
                                    
                                    return styles
                                
                                # Добавляем рейтинг
                                if "param_ratings" not in st.session_state:
                                    st.session_state["param_ratings"] = {}
                                
                                param_rating_key = f"{selected_param}_ratings"
                                if param_rating_key not in st.session_state["param_ratings"]:
                                    # Автоматически формируем рейтинг по приоритетам:
                                    # 1. Выручка на 1 артикул (больше = лучше)
                                    # 2. Средняя цена без СПП (больше = лучше) 
                                    # 3. Упущенная выручка на 1 артикул (меньше = лучше)
                                    
                                    def calculate_score(item):
                                        param_val, metrics = item
                                        # Нормализуем значения от 0 до 1
                                        revenue_per_sku = metrics["Выручка на 1 артикул"]
                                        avg_price = metrics["Средняя цена без СПП"]
                                        lost_revenue_per_sku = metrics["Упущенная выручка на 1 артикул"]
                                        
                                        # Находим мин/макс для нормализации
                                        all_revenues = [m["Выручка на 1 артикул"] for m in summary_data.values()]
                                        all_prices = [m["Средняя цена без СПП"] for m in summary_data.values()]
                                        all_lost = [m["Упущенная выручка на 1 артикул"] for m in summary_data.values()]
                                        
                                        # Нормализуем выручку (0-1, где 1 = максимум)
                                        if max(all_revenues) > min(all_revenues):
                                            norm_revenue = (revenue_per_sku - min(all_revenues)) / (max(all_revenues) - min(all_revenues))
                                        else:
                                            norm_revenue = 0.5
                                        
                                        # Нормализуем цену (0-1, где 1 = максимум)
                                        if max(all_prices) > min(all_prices):
                                            norm_price = (avg_price - min(all_prices)) / (max(all_prices) - min(all_prices))
                                        else:
                                            norm_price = 0.5
                                        
                                        # Нормализуем упущенную выручку (0-1, где 1 = минимум, т.е. лучше)
                                        if max(all_lost) > min(all_lost):
                                            norm_lost = 1 - (lost_revenue_per_sku - min(all_lost)) / (max(all_lost) - min(all_lost))
                                        else:
                                            norm_lost = 0.5
                                        
                                        # Взвешенная сумма с приоритетами
                                        score = (norm_revenue * 0.9) + (norm_price * 0.09) + (norm_lost * 0.01)
                                        return score
                                    
                                    # Сортируем по рассчитанному рейтингу
                                    sorted_by_score = sorted(
                                        summary_data.items(), 
                                        key=calculate_score,
                                        reverse=True
                                    )
                                    ratings = {param_val: i+1 for i, (param_val, _) in enumerate(sorted_by_score)}
                                    st.session_state["param_ratings"][param_rating_key] = ratings
                                
                                ratings = st.session_state["param_ratings"][param_rating_key]
                                
                                # Отображаем таблицу
                                st.write(f"**Сводная таблица по параметру: {selected_param}**")
                                st.info("💡 Цветовое выделение: 🟢 лучший результат в каждой строке. Для позиции в выдаче лучший = меньшее число. Рейтинг формируется по приоритетам: Выручка на 1 артикул (90%) → Средняя цена (9%) → Упущенная выручка на 1 артикул (1%)")
                                
                                # Создаем конфигурацию столбцов
                                column_config = {
                                    "Метрика": st.column_config.TextColumn("Метрика", width=250)
                                }
                                
                                # Настраиваем столбцы для значений параметров
                                for param_value in param_values_list:
                                    column_config[param_value] = st.column_config.TextColumn(
                                        param_value, 
                                        width=150
                                    )
                                
                                # Применяем стили и отображаем таблицу
                                styled_df = display_df.style.apply(lambda x: apply_colors(display_df), axis=None)
                                
                                # Отображаем стилизованную таблицу
                                st.dataframe(
                                    styled_df,
                                    column_config=column_config,
                                    use_container_width=True,
                                    hide_index=True
                                )
                                
                                # Добавляем строку с рейтингом отдельно (без цветового кодирования)
                                st.write("**Рейтинг:**")
                                rating_display_data = {"Метрика": ["Рейтинг"]}
                                for param_value in param_values_list:
                                    rating_display_data[param_value] = [str(ratings.get(param_value, len(param_values_list)+1))]
                                
                                rating_display_df = pd.DataFrame(rating_display_data)
                                st.dataframe(
                                    rating_display_df,
                                    column_config=column_config,
                                    use_container_width=True,
                                    hide_index=True
                                )
                                
                                st.divider()
                                
                                # Редактируемая таблица только для рейтинга
                                st.write("**Редактирование рейтинга:**")
                                
                                # Создаем DataFrame только для рейтинга
                                rating_edit_data = {}
                                for param_value in param_values_list:
                                    rating_edit_data[param_value] = [ratings.get(param_value, len(param_values_list)+1)]
                                
                                rating_edit_df = pd.DataFrame(rating_edit_data, index=["Рейтинг"])
                                
                                # Конфигурация для редактирования рейтинга
                                rating_column_config = {}
                                for param_value in param_values_list:
                                    rating_column_config[param_value] = st.column_config.NumberColumn(
                                        param_value,
                                        min_value=1,
                                        max_value=len(param_values_list),
                                        step=1,
                                        width=150
                                    )
                                
                                # Редактируемая таблица рейтинга
                                edited_rating_df = st.data_editor(
                                    rating_edit_df,
                                    column_config=rating_column_config,
                                    hide_index=False,
                                    key=f"rating_table_{selected_param}"
                                )
                                
                                # Сохраняем изменения рейтинга
                                if not edited_rating_df.equals(rating_edit_df):
                                    new_ratings = {}
                                    for param_value in param_values_list:
                                        new_ratings[param_value] = int(edited_rating_df.loc["Рейтинг", param_value])
                                    st.session_state["param_ratings"][param_rating_key] = new_ratings
                                    
                                    # Автоматически сохраняем рейтинги
                                    try:
                                        # Получаем удаленные параметры для фильтрации
                                        deleted_params = st.session_state.get("deleted_params", load_deleted_params_from_file())
                                        param_values = st.session_state.get("param_values", {})
                                        param_options = st.session_state.get("param_options", {})
                                        
                                        # Фильтруем удаленные параметры перед сохранением
                                        filtered_param_values = {
                                            k: v for k, v in param_values.items() 
                                            if k not in deleted_params
                                        } if deleted_params else param_values
                                        
                                        filtered_param_options = {
                                            k: v for k, v in param_options.items() 
                                            if k not in deleted_params
                                        } if deleted_params else param_options
                                        
                                        table_cache_data = {
                                            "param_values": filtered_param_values,
                                            "param_options": filtered_param_options,
                                            "param_ratings": st.session_state.get("param_ratings", {}),
                                            "timestamp": pd.Timestamp.now().strftime("%Y-%m-%d %H:%M:%S")
                                        }
                                        
                                        import json
                                        with open("table_cache.json", "w", encoding="utf-8") as f:
                                            json.dump(table_cache_data, f, ensure_ascii=False, indent=2)
                                        
                                        st.success("💾 Рейтинг автоматически сохранен")
                                    except Exception as e:
                                        st.error(f"❌ Ошибка сохранения рейтинга: {e}")
                                
                                # Кнопки управления
                                col_reset, col_export_analytics = st.columns(2)
                                
                                with col_reset:
                                    if st.button("🔄 Сбросить рейтинг", type="secondary"):
                                        # Пересчитываем рейтинг по приоритетам
                                        def calculate_score(item):
                                            param_val, metrics = item
                                            # Нормализуем значения от 0 до 1
                                            revenue_per_sku = metrics["Выручка на 1 артикул"]
                                            avg_price = metrics["Средняя цена без СПП"]
                                            lost_revenue_per_sku = metrics["Упущенная выручка на 1 артикул"]
                                            
                                            # Находим мин/макс для нормализации
                                            all_revenues = [m["Выручка на 1 артикул"] for m in summary_data.values()]
                                            all_prices = [m["Средняя цена без СПП"] for m in summary_data.values()]
                                            all_lost = [m["Упущенная выручка на 1 артикул"] for m in summary_data.values()]
                                            
                                            # Нормализуем выручку (0-1, где 1 = максимум)
                                            if max(all_revenues) > min(all_revenues):
                                                norm_revenue = (revenue_per_sku - min(all_revenues)) / (max(all_revenues) - min(all_revenues))
                                            else:
                                                norm_revenue = 0.5
                                            
                                            # Нормализуем цену (0-1, где 1 = максимум)
                                            if max(all_prices) > min(all_prices):
                                                norm_price = (avg_price - min(all_prices)) / (max(all_prices) - min(all_prices))
                                            else:
                                                norm_price = 0.5
                                            
                                            # Нормализуем упущенную выручку (0-1, где 1 = минимум, т.е. лучше)
                                            if max(all_lost) > min(all_lost):
                                                norm_lost = 1 - (lost_revenue_per_sku - min(all_lost)) / (max(all_lost) - min(all_lost))
                                            else:
                                                norm_lost = 0.5
                                            
                                            # Взвешенная сумма с приоритетами
                                            score = (norm_revenue * 0.9) + (norm_price * 0.09) + (norm_lost * 0.01)
                                            return score
                                        
                                        # Сортируем по рассчитанному рейтингу
                                        sorted_by_score = sorted(
                                            summary_data.items(), 
                                            key=calculate_score,
                                            reverse=True
                                        )
                                        ratings = {param_val: i+1 for i, (param_val, _) in enumerate(sorted_by_score)}
                                        st.session_state["param_ratings"][param_rating_key] = ratings
                                        st.rerun()
                                
                                with col_export_analytics:
                                    if st.button("📊 Экспорт аналитики"):
                                        # Создаем полную таблицу для экспорта (с рейтингом)
                                        export_data = display_df.copy()
                                        rating_row = ["Рейтинг"] + [str(ratings.get(param_val, len(param_values_list)+1)) for param_val in param_values_list]
                                        rating_export_df = pd.DataFrame([rating_row], columns=export_data.columns)
                                        full_export_df = pd.concat([export_data, rating_export_df], ignore_index=True)
                                        
                                        csv_data = full_export_df.to_csv(encoding='utf-8-sig', index=False)
                                        # Получаем имя загруженного файла для названия экспорта
                                        base_filename = f"analytics_{selected_param}"
                                        if hasattr(uploaded, 'name') and uploaded.name:
                                            # Убираем расширение и добавляем суффикс
                                            name_without_ext = os.path.splitext(uploaded.name)[0]
                                            base_filename = f"{name_without_ext}_analytics_{selected_param}"
                                        
                                        st.download_button(
                                            label="💾 Скачать CSV",
                                            data=csv_data,
                                            file_name=f"{base_filename}.csv",
                                            mime="text/csv"
                                        )
                                
                                # Дополнительная статистика
                                st.divider()
                                st.write("**Дополнительная статистика:**")
                                
                                col_stat1, col_stat2, col_stat3, col_stat4 = st.columns(4)
                                
                                total_products = sum(item["Количество артикулов"] for item in summary_data.values())
                                total_revenue = sum(item["Общая выручка"] for item in summary_data.values())
                                total_lost_revenue = sum(item["Упущенная выручка"] for item in summary_data.values())
                                avg_position_all = sum(item["Позиция в выдаче (средняя)"] * item["Количество артикулов"] for item in summary_data.values()) / total_products if total_products > 0 else 0
                                
                                col_stat1.metric("Всего товаров", f"{total_products} шт.")
                                col_stat2.metric("Общая выручка", f"₽{total_revenue:,.0f}".replace(",", " "))
                                col_stat3.metric("Упущенная выручка", f"₽{total_lost_revenue:,.0f}".replace(",", " "))
                                col_stat4.metric("Средняя позиция", f"{avg_position_all:.1f}")
                                
                        else:
                            st.info("Нет данных для анализа по выбранному параметру")
                
                # Добавляем анализ комбинаций иерархических параметров
                st.divider()
                st.subheader("📊 Анализ комбинаций иерархических параметров")
                
                # Определяем иерархию сегментов
                # Иерархия параметров: Подтип - это логическая группировка для Рукав и Ворот
                hierarchy_params = get_hierarchy_params()
                subtype_params = get_subtype_params()  # Параметры, которые находятся внутри "Подтип"
                visual_params = get_visual_params()
                
                # Получаем список полностью исключенных параметров
                excluded_params = set(st.session_state.get("excluded_params", []))
                
                # Фильтруем available_params, исключая полностью исключенные параметры
                active_available_params = [p for p in available_params if p not in excluded_params]
                
                # Разделяем параметры на категории (учитывая исключенные параметры)
                hierarchical_params_list = [p for p in hierarchy_params if p in active_available_params]
                visual_params_list = [p for p in active_available_params if p in visual_params or any(vp in p for vp in visual_params)]
                other_params_list = [p for p in active_available_params if p not in hierarchy_params and p not in visual_params_list]
                
                # Показываем информацию об иерархии
                if hierarchical_params_list:
                    # Формируем отображение иерархии с учетом "Подтип" как группировки
                    hierarchy_display = []
                    subtype_params_in_list = [p for p in subtype_params if p in hierarchical_params_list]
                    
                    for param in hierarchy_params:
                        if param in hierarchical_params_list:
                            if param in subtype_params:
                                # Если это первый параметр из подтипа, добавляем "Подтип" перед ним
                                if param == subtype_params[0] and subtype_params_in_list:
                                    subtype_str = " → ".join(subtype_params_in_list)
                                    hierarchy_display.append(f"Подтип ({subtype_str})")
                            else:
                                hierarchy_display.append(param)
                    
                    # Убираем дубликаты и формируем строку
                    hierarchy_str = " → ".join(hierarchy_display)
                    if not hierarchy_str:
                        hierarchy_str = " → ".join(hierarchical_params_list)
                    st.info(f"📊 **Иерархия сегментов**: {hierarchy_str} → Визуальные параметры")
                    
                    # Анализ комбинаций иерархических параметров
                    if len(hierarchical_params_list) > 1:
                        st.write("**Анализ комбинаций иерархических параметров:**")
                        
                        # Собираем все комбинации иерархических параметров
                        hierarchy_combinations = {}
                        
                        # Получаем список полностью исключенных параметров
                        excluded_params = set(st.session_state.get("excluded_params", []))
                        
                        # Фильтруем иерархические параметры, исключая полностью исключенные
                        active_hierarchical_params = [p for p in hierarchical_params_list if p not in excluded_params]
                        
                        # Получаем все артикулы с заполненными иерархическими параметрами
                        for sku in df["Артикул"].astype(str).str.replace(".0", "").unique():
                            sku_str = str(sku)
                            combination_key_parts = []
                            combination_valid = True
                            
                            # Используем только активные (не исключенные) параметры
                            for param_name in active_hierarchical_params:
                                if param_name in param_values and sku_str in param_values[param_name]:
                                    value = param_values[param_name][sku_str]
                                    if value and str(value).strip():
                                        combination_key_parts.append(f"{param_name}:{value}")
                                    else:
                                        combination_valid = False
                                        break
                                else:
                                    combination_valid = False
                                    break
                            
                            if combination_valid and combination_key_parts:
                                combination_key = " | ".join(combination_key_parts)
                                if combination_key not in hierarchy_combinations:
                                    hierarchy_combinations[combination_key] = []
                                hierarchy_combinations[combination_key].append(sku_str)
                        
                        if hierarchy_combinations:
                            # Рассчитываем метрики для каждой комбинации
                            combination_analytics = []
                            for combination_key, skus in hierarchy_combinations.items():
                                mask = df["Артикул"].astype(str).str.replace(".0", "").isin(skus)
                                filtered_df = df[mask]
                                
                                if not filtered_df.empty:
                                    total_revenue = filtered_df["Выручка"].sum() if "Выручка" in filtered_df.columns else 0
                                    total_orders = filtered_df["Заказы"].sum() if "Заказы" in filtered_df.columns else 0
                                    revenue_per_product = total_revenue / len(filtered_df) if len(filtered_df) > 0 else 0
                                    avg_price = filtered_df["Средняя цена"].mean() if "Средняя цена" in filtered_df.columns else 0
                                    
                                    combination_analytics.append({
                                        "Комбинация": combination_key,
                                        "Количество товаров": len(filtered_df),
                                        "Общая выручка": total_revenue,
                                        "Выручка на 1 артикул": revenue_per_product,
                                        "Средняя цена": avg_price
                                    })
                            
                            if combination_analytics:
                                # Сортируем по выручке на 1 артикул
                                combination_analytics.sort(key=lambda x: x["Выручка на 1 артикул"], reverse=True)
                                
                                # Показываем топ-5 комбинаций
                                st.write("**Топ-5 комбинаций иерархических параметров:**")
                                top_combinations_df = pd.DataFrame(combination_analytics[:5])
                                st.dataframe(top_combinations_df, use_container_width=True)
                                
                                # Сохраняем лучшую комбинацию
                                best_hierarchy_combination = combination_analytics[0]["Комбинация"]
                                
                                # Находим артикулы, соответствующие лучшей комбинации иерархии
                                best_combination_skus = hierarchy_combinations[best_hierarchy_combination]
                                
                                # Получаем рейтинги параметров из session_state
                                param_ratings = {}
                                # available_params уже определен выше в коде (строка 4600)
                                for param in available_params:
                                    rating_key = f"{param}_ratings"
                                    if rating_key in st.session_state.get("param_ratings", {}):
                                        param_ratings[param] = st.session_state["param_ratings"][rating_key]
                                
                                # Находим лучшие визуальные и другие параметры для этой иерархии
                                best_visual_params = {}
                                best_other_params = {}
                                
                                # Используем уже определенные списки параметров (определены выше в коде)
                                # visual_params_list и other_params_list уже определены в строках 5476-5477
                                
                                # Для визуальных параметров находим значения с рейтингом 1
                                for param_name in visual_params_list:
                                    if param_name in param_ratings:
                                        ratings = param_ratings[param_name]
                                        for value, rating in ratings.items():
                                            if rating == 1:
                                                best_visual_params[param_name] = value
                                                break
                                
                                # Для других параметров находим значения с рейтингом 1
                                for param_name in other_params_list:
                                    if param_name in param_ratings:
                                        ratings = param_ratings[param_name]
                                        for value, rating in ratings.items():
                                            if rating == 1:
                                                best_other_params[param_name] = value
                                                break
                                
                                # Формируем полную строку комбинации
                                full_combination_parts = [best_hierarchy_combination]
                                
                                # Добавляем визуальные параметры
                                if best_visual_params:
                                    visual_parts = [f"{k}:{v}" for k, v in best_visual_params.items()]
                                    full_combination_parts.extend(visual_parts)
                                
                                # Добавляем другие параметры
                                if best_other_params:
                                    other_parts = [f"{k}:{v}" for k, v in best_other_params.items()]
                                    full_combination_parts.extend(other_parts)
                                
                                # Формируем полную комбинацию
                                full_combination = " | ".join(full_combination_parts)
                                
                                st.success(f"🏆 **Лучшая комбинация иерархии**: {full_combination}")
                        else:
                            st.info("ℹ️ Не найдено товаров с заполненными всеми иерархическими параметрами")
                    else:
                        st.info("ℹ️ Для анализа комбинаций нужно минимум 2 иерархических параметра")
                
                # Добавляем анализ лучших полных комбинаций всех параметров
                st.divider()
                st.subheader("🏆 Топ-10 лучших комбинаций параметров")
                st.info("💡 Система анализирует все возможные комбинации параметров и находит лучшие по метрикам: Выручка на 1 артикул (90%) → Средняя цена (9%) → Упущенная выручка на 1 артикул (1%)")
                
                # Получаем период отчета для фильтрации новинок
                report_end_date = None
                if "Дата создания" in df.columns and not df["Дата создания"].isna().all():
                    # Используем максимальную дату из "Дата создания" как дату отчета
                    report_end_date = df["Дата создания"].dropna().max()
                else:
                    # Пытаемся получить период из анализа
                    df_raw = st.session_state.get("df_raw")
                    header_row = st.session_state.get("header_row")
                    if df_raw is not None and header_row is not None:
                        period_info = get_analysis_period(df, df_raw, header_row)
                        if period_info and "end_date" in period_info:
                            report_end_date = period_info["end_date"]
                    # Если не нашли, используем текущую дату
                    if report_end_date is None:
                        report_end_date = pd.Timestamp.now()
                
                # Добавляем выбор периода для создания комбинаций новинок
                novelty_filter_options = {
                    "Все товары": None,
                    "Создать комбинации для новинок (6 месяцев)": 6,
                    "Создать комбинации для новинок (3 месяца)": 3
                }
                
                selected_novelty_filter = st.selectbox(
                    "🔍 Комбинации для новинок:",
                    options=list(novelty_filter_options.keys()),
                    index=0,
                    help="Выберите период для создания отдельных комбинаций новинок. Товары, появившиеся в выбранный период, будут иметь дополнительный параметр 'Период появления' в комбинациях."
                )
                
                novelty_months = novelty_filter_options[selected_novelty_filter]
                
                # Определяем, какие товары являются новинками
                novelty_skus = set()
                if novelty_months is not None:
                    if "Дата создания" not in df.columns:
                        st.warning("⚠️ Для создания комбинаций новинок требуется колонка 'Дата создания'. Параметр не будет добавлен.")
                        novelty_months = None
                    elif report_end_date is None:
                        st.warning("⚠️ Не удалось определить дату отчета. Параметр новинок не будет добавлен.")
                        novelty_months = None
                    else:
                        # Вычисляем дату начала периода для новинок
                        if isinstance(report_end_date, pd.Timestamp):
                            start_date_for_novelties = report_end_date - pd.DateOffset(months=novelty_months)
                        else:
                            start_date_for_novelties = pd.to_datetime(report_end_date) - pd.DateOffset(months=novelty_months)
                        
                        # Находим товары-новинки: те, у которых "Дата создания" в пределах периода новинок
                        # Важно: проверяем, что дата создания не пустая
                        novelty_mask = (
                            (df["Дата создания"] >= start_date_for_novelties) & 
                            (df["Дата создания"] <= report_end_date) &
                            (df["Дата создания"].notna())
                        )
                        novelty_skus = set(
                            df.loc[novelty_mask, "Артикул"].astype(str).str.replace(".0", "").unique()
                        )
                        
                        if novelty_skus:
                            st.info(f"📦 Для новинок (товары с {start_date_for_novelties.strftime('%d.%m.%Y')} по {report_end_date.strftime('%d.%m.%Y')}, всего {len(novelty_skus)} товаров) будут созданы отдельные комбинации с параметром 'Период появления'")
                        else:
                            st.warning(f"⚠️ Не найдено новинок за последние {novelty_months} месяцев")
                
                # Собираем все полные комбинации параметров
                param_values = st.session_state.get("param_values", {})
                all_combinations = {}
                
                # Получаем список полностью исключенных параметров
                excluded_params = set(st.session_state.get("excluded_params", []))
                
                # Фильтруем available_params, исключая полностью исключенные параметры
                active_params = [p for p in available_params if p not in excluded_params]
                
                # Проходим по всем артикулам и собираем их полные комбинации параметров
                all_skus = df["Артикул"].astype(str).str.replace(".0", "").unique()
                
                # Разделяем товары на новинки и не-новинки для создания отдельных комбинаций
                for sku in all_skus:
                    sku_str = str(sku)
                    combination_parts = []
                    combination_valid = True
                    
                    # Собираем все параметры для этого артикула (только активные, не исключенные)
                    for param_name in active_params:
                        if param_name in param_values and sku_str in param_values[param_name]:
                            value = param_values[param_name][sku_str]
                            if value and str(value).strip():
                                combination_parts.append(f"{param_name}:{value}")
                            else:
                                # Если хотя бы один активный параметр пустой, комбинация неполная
                                combination_valid = False
                                break
                        else:
                            # Если активный параметр отсутствует, комбинация неполная
                            combination_valid = False
                            break
                    
                    if not combination_valid or not combination_parts:
                        continue
                    
                    # Если выбран период для новинок, создаем отдельные комбинации для новинок и не-новинок
                    if novelty_months is not None:
                        # Дополнительная проверка: проверяем дату создания напрямую для каждого товара
                        is_novelty = False
                        if sku_str in novelty_skus:
                            # Дополнительно проверяем дату создания напрямую в данных
                            if "Дата создания" in df.columns and report_end_date is not None:
                                # Получаем все строки для этого артикула
                                sku_mask = df["Артикул"].astype(str).str.replace(".0", "") == sku_str
                                sku_rows = df[sku_mask]
                                
                                if not sku_rows.empty:
                                    # Проверяем, есть ли хотя бы одна строка с датой создания в периоде новинок
                                    if isinstance(report_end_date, pd.Timestamp):
                                        start_date_for_novelties = report_end_date - pd.DateOffset(months=novelty_months)
                                    else:
                                        start_date_for_novelties = pd.to_datetime(report_end_date) - pd.DateOffset(months=novelty_months)
                                    
                                    # Проверяем дату создания
                                    date_mask = (
                                        (sku_rows["Дата создания"] >= start_date_for_novelties) & 
                                        (sku_rows["Дата создания"] <= report_end_date) &
                                        (sku_rows["Дата создания"].notna())
                                    )
                                    is_novelty = date_mask.any()
                            else:
                                # Если нет даты создания, используем только проверку по novelty_skus
                                is_novelty = True
                        else:
                            is_novelty = False
                        
                        if is_novelty:
                            # Для новинок: добавляем параметр "Период появления" и создаем отдельную комбинацию
                            novelty_combination_parts = combination_parts.copy()
                            novelty_combination_parts.append(f"Период появления:Новинка ({novelty_months} месяцев)")
                            combination_key = " | ".join(sorted(novelty_combination_parts))
                            if combination_key not in all_combinations:
                                all_combinations[combination_key] = []
                            all_combinations[combination_key].append(sku_str)
                        else:
                            # Для не-новинок: создаем обычную комбинацию без параметра "Период появления"
                            combination_key = " | ".join(sorted(combination_parts))
                            if combination_key not in all_combinations:
                                all_combinations[combination_key] = []
                            all_combinations[combination_key].append(sku_str)
                    else:
                        # Если фильтр новинок не выбран, создаем обычные комбинации для всех товаров
                        combination_key = " | ".join(sorted(combination_parts))
                        if combination_key not in all_combinations:
                            all_combinations[combination_key] = []
                        all_combinations[combination_key].append(sku_str)
                
                if all_combinations:
                    # Рассчитываем метрики для каждой комбинации
                    combination_analytics = []
                    
                    for combination_key, skus in all_combinations.items():
                        # Проверяем, является ли это комбинацией новинок
                        is_novelty_combination = novelty_months is not None and "Период появления:Новинка" in combination_key
                        
                        # Создаем маску для фильтрации товаров
                        mask = df["Артикул"].astype(str).str.replace(".0", "").isin(skus)
                        
                        # Если это комбинация новинок, дополнительно фильтруем по дате создания
                        if is_novelty_combination and "Дата создания" in df.columns and report_end_date is not None:
                            if isinstance(report_end_date, pd.Timestamp):
                                start_date_for_novelties = report_end_date - pd.DateOffset(months=novelty_months)
                            else:
                                start_date_for_novelties = pd.to_datetime(report_end_date) - pd.DateOffset(months=novelty_months)
                            
                            # Дополнительная фильтрация: только товары с датой создания в периоде новинок
                            novelty_date_mask = (
                                (df["Дата создания"] >= start_date_for_novelties) & 
                                (df["Дата создания"] <= report_end_date)
                            )
                            mask = mask & novelty_date_mask
                        
                        filtered_df = df[mask]
                        
                        if not filtered_df.empty:
                            total_revenue = filtered_df["Выручка"].sum() if "Выручка" in filtered_df.columns else 0
                            total_orders = filtered_df["Заказы"].sum() if "Заказы" in filtered_df.columns else 0
                            avg_price = filtered_df["Средняя цена"].mean() if "Средняя цена" in filtered_df.columns else 0
                            lost_revenue = filtered_df["Упущенная выручка"].sum() if "Упущенная выручка" in filtered_df.columns else 0
                            revenue_per_product = total_revenue / len(filtered_df) if len(filtered_df) > 0 else 0
                            lost_revenue_per_product = lost_revenue / len(filtered_df) if len(filtered_df) > 0 else 0
                            avg_position = filtered_df["Позиция в выдаче"].mean() if "Позиция в выдаче" in filtered_df.columns else 0
                            avg_cpm = filtered_df["Стоимость за 1000 показов"].mean() if "Стоимость за 1000 показов" in filtered_df.columns else 0
                            cpm_per_product = avg_cpm / len(filtered_df) if len(filtered_df) > 0 else 0
                            
                            combination_analytics.append({
                                "Комбинация": combination_key,
                                "Общая выручка": total_revenue,
                                "Количество артикулов": len(filtered_df),
                                "Выручка на 1 артикул": revenue_per_product,
                                "Средняя цена без СПП": avg_price,
                                "Упущенная выручка": lost_revenue,
                                "Упущенная выручка на 1 артикул": lost_revenue_per_product,
                                "Позиция в выдаче (средняя)": avg_position,
                                "Стоимость за 1000 показов на 1 артикул": cpm_per_product
                            })
                    
                    if combination_analytics:
                        # Рассчитываем рейтинг для каждой комбинации
                        # Нормализуем значения для расчета рейтинга
                        all_revenues = [c["Выручка на 1 артикул"] for c in combination_analytics]
                        all_prices = [c["Средняя цена без СПП"] for c in combination_analytics]
                        all_lost = [c["Упущенная выручка на 1 артикул"] for c in combination_analytics]
                        
                        def calculate_combination_score(combo):
                            revenue_per_sku = combo["Выручка на 1 артикул"]
                            avg_price = combo["Средняя цена без СПП"]
                            lost_revenue_per_sku = combo["Упущенная выручка на 1 артикул"]
                            
                            # Нормализуем выручку (0-1, где 1 = максимум)
                            if max(all_revenues) > min(all_revenues):
                                norm_revenue = (revenue_per_sku - min(all_revenues)) / (max(all_revenues) - min(all_revenues))
                            else:
                                norm_revenue = 0.5
                            
                            # Нормализуем цену (0-1, где 1 = максимум)
                            if max(all_prices) > min(all_prices):
                                norm_price = (avg_price - min(all_prices)) / (max(all_prices) - min(all_prices))
                            else:
                                norm_price = 0.5
                            
                            # Нормализуем упущенную выручку (0-1, где 1 = минимум, т.е. лучше)
                            if max(all_lost) > min(all_lost):
                                norm_lost = 1 - (lost_revenue_per_sku - min(all_lost)) / (max(all_lost) - min(all_lost))
                            else:
                                norm_lost = 0.5
                            
                            # Взвешенная сумма с приоритетами: Выручка на 1 артикул (90%) → Средняя цена (9%) → Упущенная выручка на 1 артикул (1%)
                            score = (norm_revenue * 0.9) + (norm_price * 0.09) + (norm_lost * 0.01)
                            return score
                        
                        # Разделяем комбинации на новинки и не-новинки, если выбран период для новинок
                        if novelty_months is not None:
                            novelty_combinations = []
                            regular_combinations = []
                            
                            for combo in combination_analytics:
                                if "Период появления:Новинка" in combo["Комбинация"]:
                                    novelty_combinations.append(combo)
                                else:
                                    regular_combinations.append(combo)
                            
                            # Рассчитываем рейтинг отдельно для новинок
                            if novelty_combinations:
                                novelty_revenues = [c["Выручка на 1 артикул"] for c in novelty_combinations]
                                novelty_prices = [c["Средняя цена без СПП"] for c in novelty_combinations]
                                novelty_lost = [c["Упущенная выручка на 1 артикул"] for c in novelty_combinations]
                                
                                def calculate_novelty_score(combo):
                                    revenue_per_sku = combo["Выручка на 1 артикул"]
                                    avg_price = combo["Средняя цена без СПП"]
                                    lost_revenue_per_sku = combo["Упущенная выручка на 1 артикул"]
                                    
                                    if max(novelty_revenues) > min(novelty_revenues):
                                        norm_revenue = (revenue_per_sku - min(novelty_revenues)) / (max(novelty_revenues) - min(novelty_revenues))
                                    else:
                                        norm_revenue = 0.5
                                    
                                    if max(novelty_prices) > min(novelty_prices):
                                        norm_price = (avg_price - min(novelty_prices)) / (max(novelty_prices) - min(novelty_prices))
                                    else:
                                        norm_price = 0.5
                                    
                                    if max(novelty_lost) > min(novelty_lost):
                                        norm_lost = 1 - (lost_revenue_per_sku - min(novelty_lost)) / (max(novelty_lost) - min(novelty_lost))
                                    else:
                                        norm_lost = 0.5
                                    
                                    score = (norm_revenue * 0.9) + (norm_price * 0.09) + (norm_lost * 0.01)
                                    return score
                                
                                for combo in novelty_combinations:
                                    combo["Рейтинг"] = calculate_novelty_score(combo)
                                novelty_combinations.sort(key=lambda x: x["Рейтинг"], reverse=True)
                            
                            # Рассчитываем рейтинг отдельно для не-новинок
                            if regular_combinations:
                                regular_revenues = [c["Выручка на 1 артикул"] for c in regular_combinations]
                                regular_prices = [c["Средняя цена без СПП"] for c in regular_combinations]
                                regular_lost = [c["Упущенная выручка на 1 артикул"] for c in regular_combinations]
                                
                                def calculate_regular_score(combo):
                                    revenue_per_sku = combo["Выручка на 1 артикул"]
                                    avg_price = combo["Средняя цена без СПП"]
                                    lost_revenue_per_sku = combo["Упущенная выручка на 1 артикул"]
                                    
                                    if max(regular_revenues) > min(regular_revenues):
                                        norm_revenue = (revenue_per_sku - min(regular_revenues)) / (max(regular_revenues) - min(regular_revenues))
                                    else:
                                        norm_revenue = 0.5
                                    
                                    if max(regular_prices) > min(regular_prices):
                                        norm_price = (avg_price - min(regular_prices)) / (max(regular_prices) - min(regular_prices))
                                    else:
                                        norm_price = 0.5
                                    
                                    if max(regular_lost) > min(regular_lost):
                                        norm_lost = 1 - (lost_revenue_per_sku - min(regular_lost)) / (max(regular_lost) - min(regular_lost))
                                    else:
                                        norm_lost = 0.5
                                    
                                    score = (norm_revenue * 0.9) + (norm_price * 0.09) + (norm_lost * 0.01)
                                    return score
                                
                                for combo in regular_combinations:
                                    combo["Рейтинг"] = calculate_regular_score(combo)
                                regular_combinations.sort(key=lambda x: x["Рейтинг"], reverse=True)
                            
                            # Сохраняем разделенные комбинации
                            combination_analytics_novelty = novelty_combinations
                            combination_analytics_regular = regular_combinations
                            # Для обратной совместимости оставляем combination_analytics как объединенный список
                            combination_analytics = novelty_combinations + regular_combinations
                        else:
                            # Если фильтр новинок не выбран, рассчитываем рейтинг для всех комбинаций вместе
                            for combo in combination_analytics:
                                combo["Рейтинг"] = calculate_combination_score(combo)
                            
                            # Сортируем по рейтингу (от большего к меньшему)
                            combination_analytics.sort(key=lambda x: x["Рейтинг"], reverse=True)
                            
                            # Разделяем комбинации на новинки и не-новинки
                            if novelty_months is not None:
                                combination_analytics_novelty = [c for c in combination_analytics if "Период появления:Новинка" in c["Комбинация"]]
                                combination_analytics_regular = [c for c in combination_analytics if "Период появления:Новинка" not in c["Комбинация"]]
                            else:
                                combination_analytics_novelty = []
                                combination_analytics_regular = combination_analytics
                        
                        # Загружаем настройки исключения параметров из файла (если еще не загружены)
                        if "excluded_params_loaded" not in st.session_state:
                            load_excluded_params_settings()
                            st.session_state["excluded_params_loaded"] = True
                        
                        # Инициализируем список исключенных значений параметров
                        if "excluded_param_values" not in st.session_state:
                            st.session_state["excluded_param_values"] = {}
                        
                        # Инициализируем список полностью исключенных параметров
                        if "excluded_params" not in st.session_state:
                            st.session_state["excluded_params"] = []
                        
                        # Собираем все уникальные значения для каждого параметра
                        param_values_dict = {}
                        for combo in combination_analytics:
                            combo_str = combo["Комбинация"]
                            # Парсим комбинацию: "Параметр1:Значение1 | Параметр2:Значение2 | ..."
                            parts = [p.strip() for p in combo_str.split("|")]
                            for part in parts:
                                if ":" in part:
                                    param_name, param_value = part.split(":", 1)
                                    param_name = param_name.strip()
                                    param_value = param_value.strip()
                                    if param_name not in param_values_dict:
                                        param_values_dict[param_name] = set()
                                    param_values_dict[param_name].add(param_value)
                        
                        # Дополняем значениями из param_options, чтобы были доступны все варианты параметров
                        # (даже те, которые не присутствуют в текущих комбинациях)
                        param_options = st.session_state.get("param_options", {})
                        for param_name, options_list in param_options.items():
                            if param_name not in param_values_dict:
                                param_values_dict[param_name] = set()
                            # Добавляем все значения из param_options
                            for option_value in options_list:
                                if option_value and str(option_value).strip():
                                    param_values_dict[param_name].add(str(option_value).strip())
                        
                        # Показываем интерфейс для исключения значений параметров и самих параметров
                        with st.expander("🚫 Исключить значения параметров из анализа", expanded=False):
                            st.info("💡 Выберите конкретные значения параметров для исключения. Все комбинации, содержащие эти значения, будут исключены из топ-10.")
                            st.info("💡 Вы также можете полностью исключить параметр из анализа. Товары останутся, но параметр не будет участвовать в формировании комбинаций и рейтинге.")
                            
                            # Раздел 1: Исключение параметров полностью
                            st.markdown("### ❌ Полностью исключить параметры")
                            current_excluded_params = st.session_state.get("excluded_params", [])
                            
                            excluded_params_selected = st.multiselect(
                                "Исключить параметры из анализа (товары останутся, но параметр не будет учитываться):",
                                options=sorted(param_values_dict.keys()),
                                default=current_excluded_params,
                                help="Выберите параметры, которые нужно полностью исключить из формирования комбинаций. Товары не будут исключены, но параметр не будет участвовать в анализе."
                            )
                            
                            # Обновляем список исключенных параметров
                            st.session_state["excluded_params"] = excluded_params_selected
                            # Сохраняем настройки при изменении
                            save_excluded_params_settings()
                            
                            # Раздел 2: Исключение конкретных значений параметров
                            st.markdown("### 🚫 Исключить конкретные значения параметров")
                            
                            # Для каждого параметра создаем мультиселект (только для неисключенных параметров)
                            excluded_selected = {}
                            for param_name in sorted(param_values_dict.keys()):
                                # Пропускаем полностью исключенные параметры
                                if param_name in excluded_params_selected:
                                    continue
                                
                                available_values = sorted(list(param_values_dict[param_name]))
                                current_excluded = st.session_state["excluded_param_values"].get(param_name, [])
                                
                                excluded_values = st.multiselect(
                                    f"Исключить значения параметра '{param_name}':",
                                    options=available_values,
                                    default=current_excluded,
                                    help=f"Выберите значения параметра '{param_name}', которые нужно исключить из анализа"
                                )
                                
                                excluded_selected[param_name] = excluded_values
                            
                            # Обновляем список исключенных значений
                            st.session_state["excluded_param_values"] = excluded_selected
                            # Сохраняем настройки при изменении
                            save_excluded_params_settings()
                            
                            # Кнопка для сброса всех исключений
                            if st.button("🔄 Сбросить все исключения", type="secondary"):
                                st.session_state["excluded_param_values"] = {}
                                st.session_state["excluded_params"] = []
                                # Сохраняем очищенные настройки
                                save_excluded_params_settings()
                                st.rerun()
                            
                            # Показываем статистику исключений
                            total_excluded_params = len(excluded_params_selected)
                            total_excluded_values = sum(len(values) for values in excluded_selected.values())
                            
                            if total_excluded_params > 0 or total_excluded_values > 0:
                                st.markdown("---")
                                st.markdown("**Статистика исключений:**")
                                if total_excluded_params > 0:
                                    st.write(f"  ❌ Полностью исключено параметров: {total_excluded_params}")
                                    for param_name in excluded_params_selected:
                                        st.write(f"    • {param_name}")
                                if total_excluded_values > 0:
                                    excluded_list = []
                                    for param_name, values in excluded_selected.items():
                                        if values:
                                            excluded_list.append(f"{param_name}: {', '.join(values)}")
                                    st.write(f"  🚫 Исключено значений параметров: {total_excluded_values}")
                                    for item in excluded_list:
                                        st.write(f"    • {item}")
                        
                        # Получаем список полностью исключенных параметров
                        excluded_params = set(st.session_state.get("excluded_params", []))
                        
                        # Функция для удаления исключенных параметров из строки комбинации
                        def remove_excluded_params_from_combo(combo_str):
                            """Удаляет исключенные параметры из строки комбинации"""
                            if not combo_str:
                                return combo_str
                            parts = [p.strip() for p in combo_str.split("|")]
                            filtered_parts = []
                            for part in parts:
                                if ":" in part:
                                    param_name, param_value = part.split(":", 1)
                                    param_name = param_name.strip()
                                    # Пропускаем исключенные параметры
                                    if param_name not in excluded_params:
                                        filtered_parts.append(part.strip())
                            return " | ".join(sorted(filtered_parts))
                        
                        # Функция для фильтрации комбинаций по исключенным значениям параметров
                        def should_exclude_combination(combo_str):
                            """Проверяет, нужно ли исключить комбинацию"""
                            parts = [p.strip() for p in combo_str.split("|")]
                            for part in parts:
                                if ":" in part:
                                    param_name, param_value = part.split(":", 1)
                                    param_name = param_name.strip()
                                    param_value = param_value.strip()
                                    # Пропускаем полностью исключенные параметры
                                    if param_name in excluded_params:
                                        continue
                                    # Проверяем, исключено ли это значение параметра
                                    excluded_values = st.session_state["excluded_param_values"].get(param_name, [])
                                    if param_value in excluded_values:
                                        return True
                            return False
                        
                        # Переформируем комбинации: удаляем исключенные параметры и перегруппировываем
                        # Создаем новый словарь комбинаций БЕЗ исключенных параметров
                        regrouped_combinations = {}
                        for combo in combination_analytics:
                            # Удаляем исключенные параметры из строки комбинации
                            cleaned_combo_str = remove_excluded_params_from_combo(combo["Комбинация"])
                            
                            # Пропускаем комбинации с исключенными значениями
                            if should_exclude_combination(combo["Комбинация"]):
                                continue
                            
                            # Пропускаем пустые комбинации
                            if not cleaned_combo_str:
                                continue
                            
                            # Перегруппировываем: собираем все артикулы для очищенной комбинации
                            if cleaned_combo_str not in regrouped_combinations:
                                regrouped_combinations[cleaned_combo_str] = {
                                    "skus": [],
                                    "original_combos": []
                                }
                            
                            # Добавляем артикулы из исходной комбинации
                            combo_skus = all_combinations.get(combo["Комбинация"], [])
                            
                            # Если это комбинация новинок, проверяем дату создания для каждого артикула
                            is_original_novelty_combo = novelty_months is not None and "Период появления:Новинка" in combo["Комбинация"]
                            if is_original_novelty_combo and "Дата создания" in df.columns and report_end_date is not None:
                                if isinstance(report_end_date, pd.Timestamp):
                                    start_date_for_novelties = report_end_date - pd.DateOffset(months=novelty_months)
                                else:
                                    start_date_for_novelties = pd.to_datetime(report_end_date) - pd.DateOffset(months=novelty_months)
                                
                                # Фильтруем артикулы: оставляем только те, у которых дата создания в периоде новинок
                                filtered_skus = []
                                for sku in combo_skus:
                                    sku_mask = df["Артикул"].astype(str).str.replace(".0", "") == sku
                                    sku_rows = df[sku_mask]
                                    
                                    if not sku_rows.empty:
                                        date_mask = (
                                            (sku_rows["Дата создания"] >= start_date_for_novelties) & 
                                            (sku_rows["Дата создания"] <= report_end_date) &
                                            (sku_rows["Дата создания"].notna())
                                        )
                                        if date_mask.any():
                                            filtered_skus.append(sku)
                                
                                combo_skus = filtered_skus
                            
                            regrouped_combinations[cleaned_combo_str]["skus"].extend(combo_skus)
                            regrouped_combinations[cleaned_combo_str]["original_combos"].append(combo)
                        
                        # Пересчитываем метрики для перегруппированных комбинаций
                        regrouped_analytics = []
                        for cleaned_combo_str, combo_data in regrouped_combinations.items():
                            skus = list(set(combo_data["skus"]))  # Убираем дубликаты
                            if not skus:
                                continue
                            
                            # Проверяем, является ли это комбинацией новинок
                            is_novelty_combo = novelty_months is not None and "Период появления:Новинка" in cleaned_combo_str
                            
                            mask = df["Артикул"].astype(str).str.replace(".0", "").isin(skus)
                            
                            # Если это комбинация новинок, дополнительно фильтруем по дате создания
                            if is_novelty_combo and "Дата создания" in df.columns and report_end_date is not None:
                                if isinstance(report_end_date, pd.Timestamp):
                                    start_date_for_novelties = report_end_date - pd.DateOffset(months=novelty_months)
                                else:
                                    start_date_for_novelties = pd.to_datetime(report_end_date) - pd.DateOffset(months=novelty_months)
                                
                                # Дополнительная фильтрация: только товары с датой создания в периоде новинок
                                novelty_date_mask = (
                                    (df["Дата создания"] >= start_date_for_novelties) & 
                                    (df["Дата создания"] <= report_end_date)
                                )
                                mask = mask & novelty_date_mask
                            
                            filtered_df = df[mask]
                            
                            if not filtered_df.empty:
                                total_revenue = filtered_df["Выручка"].sum() if "Выручка" in filtered_df.columns else 0
                                total_orders = filtered_df["Заказы"].sum() if "Заказы" in filtered_df.columns else 0
                                avg_price = filtered_df["Средняя цена"].mean() if "Средняя цена" in filtered_df.columns else 0
                                lost_revenue = filtered_df["Упущенная выручка"].sum() if "Упущенная выручка" in filtered_df.columns else 0
                                revenue_per_product = total_revenue / len(filtered_df) if len(filtered_df) > 0 else 0
                                lost_revenue_per_product = lost_revenue / len(filtered_df) if len(filtered_df) > 0 else 0
                                avg_position = filtered_df["Позиция в выдаче"].mean() if "Позиция в выдаче" in filtered_df.columns else 0
                                avg_cpm = filtered_df["Стоимость за 1000 показов"].mean() if "Стоимость за 1000 показов" in filtered_df.columns else 0
                                cpm_per_product = avg_cpm / len(filtered_df) if len(filtered_df) > 0 else 0
                                
                                # Применяем СПП для расчета цены без СПП
                                spp = st.session_state.get("spp", 15)
                                avg_price_without_spp = avg_price / (1 - float(spp) / 100.0) if float(spp) < 100 else avg_price
                                
                                regrouped_analytics.append({
                                    "Комбинация": cleaned_combo_str,
                                    "Количество артикулов": len(filtered_df),
                                    "Общая выручка": total_revenue,
                                    "Выручка на 1 артикул": revenue_per_product,
                                    "Средняя цена без СПП": avg_price_without_spp,
                                    "Упущенная выручка": lost_revenue,
                                    "Упущенная выручка на 1 артикул": lost_revenue_per_product,
                                    "Позиция в выдаче (средняя)": avg_position,
                                    "Стоимость за 1000 показов на 1 артикул": cpm_per_product
                                })
                        
                        # Пересчитываем рейтинг для перегруппированных комбинаций
                        if regrouped_analytics:
                            all_revenues = [c["Выручка на 1 артикул"] for c in regrouped_analytics]
                            all_prices = [c["Средняя цена без СПП"] for c in regrouped_analytics]
                            all_lost = [c["Упущенная выручка на 1 артикул"] for c in regrouped_analytics]
                            
                            def calculate_combination_score(combo):
                                revenue_per_sku = combo["Выручка на 1 артикул"]
                                avg_price_without_spp = combo["Средняя цена без СПП"]
                                lost_revenue_per_sku = combo["Упущенная выручка на 1 артикул"]
                                
                                # Нормализуем выручку (0-1)
                                if max(all_revenues) > min(all_revenues):
                                    norm_revenue = (revenue_per_sku - min(all_revenues)) / (max(all_revenues) - min(all_revenues))
                                else:
                                    norm_revenue = 0.5
                                
                                # Нормализуем цену (0-1, где 1 = максимум, т.е. лучше)
                                if max(all_prices) > min(all_prices):
                                    norm_price = (avg_price_without_spp - min(all_prices)) / (max(all_prices) - min(all_prices))
                                else:
                                    norm_price = 0.5
                                
                                # Нормализуем упущенную выручку (0-1, где 1 = минимум, т.е. лучше)
                                if max(all_lost) > min(all_lost):
                                    norm_lost = 1 - (lost_revenue_per_sku - min(all_lost)) / (max(all_lost) - min(all_lost))
                                else:
                                    norm_lost = 0.5
                                
                                # Взвешенная сумма с приоритетами: Выручка на 1 артикул (90%) → Средняя цена (9%) → Упущенная выручка на 1 артикул (1%)
                                score = (norm_revenue * 0.9) + (norm_price * 0.09) + (norm_lost * 0.01)
                                return score
                            
                            for combo in regrouped_analytics:
                                combo["Рейтинг"] = calculate_combination_score(combo)
                            
                            regrouped_analytics.sort(key=lambda x: x["Рейтинг"], reverse=True)
                        
                        # Используем перегруппированные комбинации
                        filtered_combinations = regrouped_analytics
                        
                        # Разделяем комбинации на новинки и не-новинки, если выбран период для новинок
                        if novelty_months is not None:
                            filtered_novelty_combinations = [c for c in filtered_combinations if "Период появления:Новинка" in c["Комбинация"]]
                            filtered_regular_combinations = [c for c in filtered_combinations if "Период появления:Новинка" not in c["Комбинация"]]
                            
                            # Рассчитываем рейтинг отдельно для новинок
                            if filtered_novelty_combinations:
                                novelty_revenues = [c["Выручка на 1 артикул"] for c in filtered_novelty_combinations]
                                novelty_prices = [c["Средняя цена без СПП"] for c in filtered_novelty_combinations]
                                novelty_lost = [c["Упущенная выручка на 1 артикул"] for c in filtered_novelty_combinations]
                                
                                def calculate_novelty_score_regrouped(combo):
                                    revenue_per_sku = combo["Выручка на 1 артикул"]
                                    avg_price = combo["Средняя цена без СПП"]
                                    lost_revenue_per_sku = combo["Упущенная выручка на 1 артикул"]
                                    
                                    if max(novelty_revenues) > min(novelty_revenues):
                                        norm_revenue = (revenue_per_sku - min(novelty_revenues)) / (max(novelty_revenues) - min(novelty_revenues))
                                    else:
                                        norm_revenue = 0.5
                                    
                                    if max(novelty_prices) > min(novelty_prices):
                                        norm_price = (avg_price - min(novelty_prices)) / (max(novelty_prices) - min(novelty_prices))
                                    else:
                                        norm_price = 0.5
                                    
                                    if max(novelty_lost) > min(novelty_lost):
                                        norm_lost = 1 - (lost_revenue_per_sku - min(novelty_lost)) / (max(novelty_lost) - min(novelty_lost))
                                    else:
                                        norm_lost = 0.5
                                    
                                    score = (norm_revenue * 0.9) + (norm_price * 0.09) + (norm_lost * 0.01)
                                    return score
                                
                                for combo in filtered_novelty_combinations:
                                    combo["Рейтинг"] = calculate_novelty_score_regrouped(combo)
                                filtered_novelty_combinations.sort(key=lambda x: x["Рейтинг"], reverse=True)
                            
                            # Рассчитываем рейтинг отдельно для не-новинок
                            if filtered_regular_combinations:
                                regular_revenues = [c["Выручка на 1 артикул"] for c in filtered_regular_combinations]
                                regular_prices = [c["Средняя цена без СПП"] for c in filtered_regular_combinations]
                                regular_lost = [c["Упущенная выручка на 1 артикул"] for c in filtered_regular_combinations]
                                
                                def calculate_regular_score_regrouped(combo):
                                    revenue_per_sku = combo["Выручка на 1 артикул"]
                                    avg_price = combo["Средняя цена без СПП"]
                                    lost_revenue_per_sku = combo["Упущенная выручка на 1 артикул"]
                                    
                                    if max(regular_revenues) > min(regular_revenues):
                                        norm_revenue = (revenue_per_sku - min(regular_revenues)) / (max(regular_revenues) - min(regular_revenues))
                                    else:
                                        norm_revenue = 0.5
                                    
                                    if max(regular_prices) > min(regular_prices):
                                        norm_price = (avg_price - min(regular_prices)) / (max(regular_prices) - min(regular_prices))
                                    else:
                                        norm_price = 0.5
                                    
                                    if max(regular_lost) > min(regular_lost):
                                        norm_lost = 1 - (lost_revenue_per_sku - min(regular_lost)) / (max(regular_lost) - min(regular_lost))
                                    else:
                                        norm_lost = 0.5
                                    
                                    score = (norm_revenue * 0.9) + (norm_price * 0.09) + (norm_lost * 0.01)
                                    return score
                                
                                for combo in filtered_regular_combinations:
                                    combo["Рейтинг"] = calculate_regular_score_regrouped(combo)
                                filtered_regular_combinations.sort(key=lambda x: x["Рейтинг"], reverse=True)
                        else:
                            filtered_novelty_combinations = []
                            filtered_regular_combinations = filtered_combinations
                        
                        # Создаем словарь для быстрого доступа к артикулам по очищенной комбинации
                        cleaned_combo_to_skus = {}
                        for cleaned_combo_str, combo_data in regrouped_combinations.items():
                            cleaned_combo_to_skus[cleaned_combo_str] = list(set(combo_data["skus"]))
                        
                        # Показываем топ-10 отдельно для новинок и не-новинок
                        top_10_novelty_combinations = filtered_novelty_combinations[:10] if novelty_months is not None else []
                        top_10_regular_combinations = filtered_regular_combinations[:10]
                        
                        # Для обратной совместимости создаем объединенный список
                        top_10_combinations = top_10_novelty_combinations + top_10_regular_combinations
                        
                        # Функция для форматирования данных комбинаций
                        def format_combinations_display(combinations_list, start_idx=1):
                            display_data = []
                            for idx, combo in enumerate(combinations_list, start_idx):
                                display_data.append({
                                    "Место": f"#{idx}",
                                    "Комбинация": combo["Комбинация"],
                                    "Общая выручка": f"₽{combo['Общая выручка']:,.0f}".replace(",", " "),
                                    "Количество артикулов": combo["Количество артикулов"],
                                    "Выручка на 1 артикул": f"₽{combo['Выручка на 1 артикул']:,.0f}".replace(",", " "),
                                    "Средняя цена без СПП": f"₽{combo['Средняя цена без СПП']:,.0f}".replace(",", " "),
                                    "Упущенная выручка": f"₽{combo['Упущенная выручка']:,.0f}".replace(",", " "),
                                    "Упущенная выручка на 1 артикул": f"₽{combo['Упущенная выручка на 1 артикул']:,.0f}".replace(",", " "),
                                    "Позиция в выдаче (средняя)": f"{combo['Позиция в выдаче (средняя)']:.1f}",
                                    "Стоимость за 1000 показов на 1 артикул": f"₽{combo['Стоимость за 1000 показов на 1 артикул']:,.0f}".replace(",", " "),
                                    "Рейтинг": f"{combo['Рейтинг']:.4f}"
                                })
                            return display_data
                        
                        # Показываем комбинации новинок отдельно, если они есть
                        if novelty_months is not None and top_10_novelty_combinations:
                            st.markdown("### 🆕 Топ-10 комбинаций новинок")
                            st.info(f"💡 Комбинации только для товаров, появившихся за последние {novelty_months} месяцев")
                            
                            display_data_novelty = format_combinations_display(top_10_novelty_combinations, 1)
                            top_10_novelty_df = pd.DataFrame(display_data_novelty)
                            st.dataframe(top_10_novelty_df, use_container_width=True, hide_index=True)
                            
                            if top_10_novelty_combinations:
                                best_novelty_combo = top_10_novelty_combinations[0]
                                st.success(f"🥇 **Лучшая комбинация новинок**: {best_novelty_combo['Комбинация']}")
                                st.markdown(f"**Рейтинг:** {best_novelty_combo['Рейтинг']:.4f} | **Выручка на 1 артикул:** ₽{best_novelty_combo['Выручка на 1 артикул']:,.0f}".replace(",", " "))
                            
                            st.divider()
                        
                        # Показываем комбинации не-новинок отдельно
                        if top_10_regular_combinations:
                            if novelty_months is not None:
                                st.markdown("### 📦 Топ-10 комбинаций остальных товаров")
                                st.info("💡 Комбинации для товаров, которые не являются новинками")
                            else:
                                st.markdown("### 🏆 Топ-10 лучших комбинаций параметров")
                            
                            display_data_regular = format_combinations_display(top_10_regular_combinations, 1)
                            top_10_regular_df = pd.DataFrame(display_data_regular)
                            st.dataframe(top_10_regular_df, use_container_width=True, hide_index=True)
                            
                            if top_10_regular_combinations:
                                best_regular_combo = top_10_regular_combinations[0]
                                st.success(f"🥇 **Лучшая комбинация**: {best_regular_combo['Комбинация']}")
                                st.markdown(f"**Рейтинг:** {best_regular_combo['Рейтинг']:.4f} | **Выручка на 1 артикул:** ₽{best_regular_combo['Выручка на 1 артикул']:,.0f}".replace(",", " "))
                        
                        # Добавляем возможность просмотра товаров для каждой комбинации
                        if top_10_combinations:
                            st.divider()
                            st.subheader("📦 Детали комбинации")
                            
                            # Формируем список всех комбинаций для выбора (новинки + остальные)
                            combo_options = []
                            if novelty_months is not None and top_10_novelty_combinations:
                                for idx, combo in enumerate(top_10_novelty_combinations, 1):
                                    combo_options.append(f"🆕 #{idx} (Новинка) - {combo['Комбинация']}")
                            
                            if top_10_regular_combinations:
                                for idx, combo in enumerate(top_10_regular_combinations, 1):
                                    prefix = "📦" if novelty_months is not None else ""
                                    combo_options.append(f"{prefix} #{idx} - {combo['Комбинация']}")
                            selected_combo_idx = st.selectbox(
                                "Выберите комбинацию для просмотра товаров:",
                                options=combo_options,
                                help="Выберите комбинацию из топ-10, чтобы увидеть все товары с такими параметрами"
                            )
                            
                            if selected_combo_idx:
                                # Определяем, из какого списка выбрана комбинация
                                is_novelty_selection = "🆕" in selected_combo_idx or "(Новинка)" in selected_combo_idx
                                
                                # Извлекаем индекс выбранной комбинации
                                idx_part = selected_combo_idx.split(" - ")[0].replace("#", "").replace("🆕", "").replace("📦", "").replace("(Новинка)", "").strip()
                                selected_idx = int(idx_part) - 1
                                
                                # Выбираем комбинацию из соответствующего списка
                                if is_novelty_selection and novelty_months is not None and top_10_novelty_combinations:
                                    if selected_idx < len(top_10_novelty_combinations):
                                        selected_combo = top_10_novelty_combinations[selected_idx]
                                    else:
                                        selected_combo = None
                                else:
                                    if selected_idx < len(top_10_regular_combinations):
                                        selected_combo = top_10_regular_combinations[selected_idx]
                                    else:
                                        selected_combo = None
                                
                                if selected_combo:
                                    selected_combo_key = selected_combo["Комбинация"]
                                else:
                                    selected_combo_key = None
                                
                                if not selected_combo_key:
                                    st.warning("⚠️ Комбинация не найдена")
                                else:
                                    # Функция для удаления параметра "Период появления" из комбинации
                                    def remove_novelty_param(combo_str):
                                        """Удаляет параметр 'Период появления' из строки комбинации"""
                                        if not combo_str:
                                            return combo_str
                                        parts = [p.strip() for p in combo_str.split("|")]
                                        filtered_parts = [p for p in parts if not p.startswith("Период появления:")]
                                        return " | ".join(sorted(filtered_parts))
                                    
                                    # Удаляем параметр "Период появления" из ключа для поиска в cleaned_combo_to_skus
                                    # (так как в regrouped_combinations ключ может быть без этого параметра)
                                    cleaned_selected_key = remove_novelty_param(selected_combo_key)
                                    
                                    # Парсим выбранную комбинацию на параметры (без параметра "Период появления")
                                    selected_parts = set([p.strip() for p in cleaned_selected_key.split("|") if p.strip()])
                                    
                                    # Находим артикулы для этой комбинации
                                    # Используем перегруппированные комбинации (они уже содержат очищенные комбинации)
                                    combo_skus = None
                                    
                                    # Сначала ищем по очищенному ключу в cleaned_combo_to_skus
                                    if cleaned_selected_key in cleaned_combo_to_skus:
                                        combo_skus = cleaned_combo_to_skus[cleaned_selected_key]
                                    # Если не найдено, пробуем найти по полному ключу
                                    elif selected_combo_key in cleaned_combo_to_skus:
                                        combo_skus = cleaned_combo_to_skus[selected_combo_key]
                                    # Если не найдено, ищем в cleaned_combo_to_skus по совпадению параметров (без учета "Период появления")
                                    else:
                                        for cleaned_key, skus in cleaned_combo_to_skus.items():
                                            # Удаляем параметр "Период появления" из ключа для сравнения
                                            cleaned_key_no_novelty = remove_novelty_param(cleaned_key)
                                            cleaned_key_parts = set([p.strip() for p in cleaned_key_no_novelty.split("|") if p.strip()])
                                            
                                            # Сравниваем параметры (без учета "Период появления")
                                            if selected_parts == cleaned_key_parts:
                                                combo_skus = skus
                                                break
                                    
                                    # Если не найдено, ищем в оригинальных комбинациях
                                    if combo_skus is None:
                                        if selected_combo_key in all_combinations:
                                            combo_skus = all_combinations[selected_combo_key]
                                        # Если это комбинация новинок и артикулы все еще не найдены, ищем в оригинальных комбинациях с параметром "Период появления"
                                        elif is_novelty_selection and novelty_months is not None:
                                            # Ищем оригинальные комбинации новинок, которые могут соответствовать выбранной комбинации
                                            for orig_combo_key, orig_skus in all_combinations.items():
                                                # Проверяем, содержит ли оригинальная комбинация параметр новинок
                                                if "Период появления:Новинка" in orig_combo_key:
                                                    # Удаляем параметр "Период появления" из оригинальной комбинации для сравнения
                                                    orig_cleaned = remove_novelty_param(orig_combo_key)
                                                    orig_parts = set([p.strip() for p in orig_cleaned.split("|") if p.strip()])
                                                    
                                                    # Сравниваем: все параметры из selected_combo_key должны быть в orig_parts
                                                    # (orig_parts может содержать дополнительные параметры, которые были исключены)
                                                    if selected_parts.issubset(orig_parts) or selected_parts == orig_parts:
                                                        combo_skus = orig_skus
                                                        break
                                    
                                    # Если артикулы все еще не найдены, показываем предупреждение
                                    if not combo_skus:
                                        st.warning(f"⚠️ Не удалось найти товары для комбинации: {selected_combo_key}")
                                    
                                    # Фильтруем данные по артикулам
                                    if combo_skus:
                                        mask = df["Артикул"].astype(str).str.replace(".0", "").isin(combo_skus)
                                        combo_products_df = df[mask].copy()
                                        
                                        # Если это комбинация новинок, дополнительно фильтруем по дате создания
                                        # Используем флаг is_novelty_selection вместо проверки ключа
                                        is_novelty_combo = is_novelty_selection and novelty_months is not None
                                        if is_novelty_combo and "Дата создания" in combo_products_df.columns and report_end_date is not None:
                                            if isinstance(report_end_date, pd.Timestamp):
                                                start_date_for_novelties = report_end_date - pd.DateOffset(months=novelty_months)
                                            else:
                                                start_date_for_novelties = pd.to_datetime(report_end_date) - pd.DateOffset(months=novelty_months)
                                            
                                            # Фильтруем только товары с датой создания в периоде новинок
                                            novelty_date_mask = (
                                                (combo_products_df["Дата создания"] >= start_date_for_novelties) & 
                                                (combo_products_df["Дата создания"] <= report_end_date) &
                                                (combo_products_df["Дата создания"].notna())
                                            )
                                            combo_products_df = combo_products_df[novelty_date_mask].copy()
                                        
                                        if not combo_products_df.empty:
                                            st.markdown(f"**Комбинация:** {selected_combo_key}")
                                            st.markdown(f"**Количество товаров:** {len(combo_products_df)}")
                                            
                                            # Получаем конец периода анализа
                                            report_end_date_local = report_end_date
                                            if report_end_date_local is None:
                                                df_raw = st.session_state.get("df_raw")
                                                header_row = st.session_state.get("header_row")
                                                if df_raw is not None and header_row is not None:
                                                    period_info = get_analysis_period(df, df_raw, header_row)
                                                    if period_info and "end_date" in period_info:
                                                        report_end_date_local = period_info["end_date"]
                                                if report_end_date_local is None:
                                                    report_end_date_local = pd.Timestamp.now()
                                            
                                            # Преобразуем в Timestamp если нужно
                                            if not isinstance(report_end_date_local, pd.Timestamp):
                                                report_end_date_local = pd.to_datetime(report_end_date_local)
                                            
                                            # Показываем основные метрики
                                            col1, col2, col3, col4 = st.columns(4)
                                            with col1:
                                                total_rev = combo_products_df["Выручка"].sum() if "Выручка" in combo_products_df.columns else 0
                                                st.metric("Общая выручка", f"₽{total_rev:,.0f}".replace(",", " "))
                                            with col2:
                                                avg_rev = total_rev / len(combo_products_df) if len(combo_products_df) > 0 else 0
                                                st.metric("Выручка на 1 артикул", f"₽{avg_rev:,.0f}".replace(",", " "))
                                            with col3:
                                                avg_price = combo_products_df["Средняя цена"].mean() if "Средняя цена" in combo_products_df.columns else 0
                                                st.metric("Средняя цена", f"₽{avg_price:,.0f}".replace(",", " "))
                                            with col4:
                                                # Средняя заказов в месяц на 1 артикул (с учетом возраста каждого товара)
                                                orders_per_month_list = []
                                                if "Заказы" in combo_products_df.columns and "Дата создания" in combo_products_df.columns:
                                                    for _, row in combo_products_df.iterrows():
                                                        orders = row.get("Заказы", 0)
                                                        if pd.notna(orders) and orders > 0:
                                                            creation_date = row.get("Дата создания")
                                                            if pd.notna(creation_date):
                                                                # Преобразуем дату создания в Timestamp
                                                                if not isinstance(creation_date, pd.Timestamp):
                                                                    creation_date = pd.to_datetime(creation_date, errors='coerce')
                                                                
                                                                if pd.notna(creation_date):
                                                                    # Рассчитываем количество месяцев от даты создания до конца периода
                                                                    months_diff = (report_end_date_local - creation_date).days / 30.0
                                                                    # Минимум 1 месяц, чтобы избежать деления на ноль
                                                                    months_diff = max(months_diff, 1.0)
                                                                    
                                                                    # Заказы в месяц для этого товара
                                                                    orders_per_month = orders / months_diff
                                                                    orders_per_month_list.append(orders_per_month)
                                                
                                                if orders_per_month_list:
                                                    avg_orders_per_month = sum(orders_per_month_list) / len(orders_per_month_list)
                                                else:
                                                    # Fallback: если нет дат создания, используем старый метод
                                                    total_orders = combo_products_df["Заказы"].sum() if "Заказы" in combo_products_df.columns else 0
                                                    avg_orders_per_month = total_orders / len(combo_products_df) if len(combo_products_df) > 0 else 0
                                                
                                                st.metric("Средняя заказов в месяц на 1 артикул", f"{avg_orders_per_month:.1f}")
                                            
                                            # Дополнительные метрики
                                            col5, col6, col7, col8 = st.columns(4)
                                            with col5:
                                                total_lost_rev = combo_products_df["Упущенная выручка"].sum() if "Упущенная выручка" in combo_products_df.columns else 0
                                                st.metric("Упущенная выручка", f"₽{total_lost_rev:,.0f}".replace(",", " "))
                                            with col6:
                                                avg_lost_rev = total_lost_rev / len(combo_products_df) if len(combo_products_df) > 0 else 0
                                                st.metric("Упущенная выручка на артикул", f"₽{avg_lost_rev:,.0f}".replace(",", " "))
                                            with col7:
                                                # Среднее продаж в месяц на 1 артикул (с учетом возраста каждого товара)
                                                sales_per_month_list = []
                                                if "Дата создания" in combo_products_df.columns:
                                                    # Получаем процент выкупа
                                                    buyout_pct = st.session_state.get("buyout_pct", 25)
                                                    buyout_k = float(buyout_pct) / 100.0 if buyout_pct else 0.25
                                                    
                                                    for _, row in combo_products_df.iterrows():
                                                        creation_date = row.get("Дата создания")
                                                        if pd.notna(creation_date):
                                                            # Преобразуем дату создания в Timestamp
                                                            if not isinstance(creation_date, pd.Timestamp):
                                                                creation_date = pd.to_datetime(creation_date, errors='coerce')
                                                            
                                                            if pd.notna(creation_date):
                                                                # Рассчитываем количество месяцев от даты создания до конца периода
                                                                months_diff = (report_end_date_local - creation_date).days / 30.0
                                                                # Минимум 1 месяц, чтобы избежать деления на ноль
                                                                months_diff = max(months_diff, 1.0)
                                                                
                                                                # Получаем продажи (выкупы) для этого товара
                                                                if "Выкупы" in combo_products_df.columns:
                                                                    buyouts = row.get("Выкупы", 0)
                                                                    if pd.notna(buyouts):
                                                                        sales = buyouts
                                                                    else:
                                                                        # Если нет выкупов, рассчитываем через заказы и процент выкупа
                                                                        orders = row.get("Заказы", 0) if "Заказы" in combo_products_df.columns else 0
                                                                        sales = orders * buyout_k if pd.notna(orders) else 0
                                                                else:
                                                                    # Если нет колонки "Выкупы", рассчитываем через заказы и процент выкупа
                                                                    orders = row.get("Заказы", 0) if "Заказы" in combo_products_df.columns else 0
                                                                    sales = orders * buyout_k if pd.notna(orders) else 0
                                                                
                                                                if sales > 0:
                                                                    # Продажи в месяц для этого товара
                                                                    sales_per_month = sales / months_diff
                                                                    sales_per_month_list.append(sales_per_month)
                                                
                                                if sales_per_month_list:
                                                    avg_sales_per_month = sum(sales_per_month_list) / len(sales_per_month_list)
                                                else:
                                                    # Fallback: если нет дат создания, используем старый метод
                                                    if "Выкупы" in combo_products_df.columns:
                                                        total_buyouts = combo_products_df["Выкупы"].sum()
                                                        avg_sales_per_month = total_buyouts / len(combo_products_df) if len(combo_products_df) > 0 else 0
                                                    else:
                                                        buyout_pct = st.session_state.get("buyout_pct", 25)
                                                        buyout_k = float(buyout_pct) / 100.0 if buyout_pct else 0.25
                                                        total_orders_for_sales = combo_products_df["Заказы"].sum() if "Заказы" in combo_products_df.columns else 0
                                                        total_buyouts = total_orders_for_sales * buyout_k
                                                        avg_sales_per_month = total_buyouts / len(combo_products_df) if len(combo_products_df) > 0 else 0
                                                
                                                st.metric("Среднее продаж в месяц на 1 артикул", f"{avg_sales_per_month:.1f}")
                                            with col8:
                                                # Оборачиваемость на единицу товара (на артикул): средняя оборачиваемость одного артикула
                                                total_orders_turnover = combo_products_df["Заказы"].sum() if "Заказы" in combo_products_df.columns else 0
                                                
                                                # Проверяем наличие данных об остатках
                                                has_stock = False
                                                avg_stock = 0
                                                if "Остаток" in combo_products_df.columns:
                                                    avg_stock = combo_products_df["Остаток"].mean()
                                                    has_stock = True
                                                elif "Stock" in combo_products_df.columns:
                                                    avg_stock = combo_products_df["Stock"].mean()
                                                    has_stock = True
                                                
                                                # Рассчитываем оборачиваемость на артикул
                                                if has_stock and avg_stock > 0:
                                                    # Если есть остатки: оборачиваемость на артикул = (Заказы артикула за 30 дней) / Остаток артикула
                                                    # Средняя оборачиваемость = среднее значение оборачиваемости всех артикулов
                                                    turnover_per_sku_list = []
                                                    for _, row in combo_products_df.iterrows():
                                                        sku_orders = row.get("Заказы", 0) if pd.notna(row.get("Заказы", 0)) else 0
                                                        sku_stock = row.get("Остаток", 0) if "Остаток" in combo_products_df.columns and pd.notna(row.get("Остаток", 0)) else row.get("Stock", 0) if "Stock" in combo_products_df.columns and pd.notna(row.get("Stock", 0)) else 0
                                                        if sku_stock > 0:
                                                            # Оборачиваемость на артикул = заказы за 30 дней / остаток
                                                            turnover_per_sku = sku_orders / sku_stock
                                                            turnover_per_sku_list.append(turnover_per_sku)
                                                    
                                                    if turnover_per_sku_list:
                                                        avg_turnover_per_sku = sum(turnover_per_sku_list) / len(turnover_per_sku_list)
                                                    else:
                                                        avg_turnover_per_sku = 0
                                                else:
                                                    # Если нет остатков: оборачиваемость на артикул = средние заказы на артикул
                                                    avg_turnover_per_sku = total_orders_turnover / len(combo_products_df) if len(combo_products_df) > 0 else 0
                                                
                                                turnover_per_sku_display = f"{avg_turnover_per_sku:.2f}" if avg_turnover_per_sku > 0 else "N/A"
                                                st.metric("Оборачиваемость на артикул", turnover_per_sku_display)
                                            
                                            # Показываем таблицу товаров
                                            st.markdown("**Список товаров:**")
                                            
                                            # Создаем копию DataFrame для работы
                                            display_combo_df = combo_products_df.copy()
                                            
                                            # Добавляем изображения
                                            import base64
                                            imgs = []
                                            img_size = 150
                                            for sku in display_combo_df["Артикул"].astype(str):
                                                sku_clean = sku.replace(".0", "")
                                                path = get_cached_image_path(sku_clean)
                                                if path and os.path.exists(path):
                                                    img_bytes = load_image_bytes(path, img_size)
                                                    if img_bytes:
                                                        b64_data = base64.b64encode(img_bytes).decode()
                                                        data_uri = f"data:image/jpeg;base64,{b64_data}"
                                                        imgs.append(data_uri)
                                                    else:
                                                        imgs.append("")
                                                else:
                                                    imgs.append("")
                                            
                                            display_combo_df.insert(1, "Изображение", imgs)
                                            
                                            # Добавляем ссылки на Wildberries
                                            if "Артикул" in display_combo_df.columns:
                                                display_combo_df["Ссылка"] = display_combo_df["Артикул"].astype(str).str.replace(".0", "").apply(
                                                    lambda sku: f"https://www.wildberries.ru/catalog/{sku}/detail.aspx"
                                                )
                                            
                                            # Добавляем параметры в таблицу
                                            for param_name in available_params:
                                                if param_name in param_values:
                                                    display_combo_df[param_name] = display_combo_df["Артикул"].astype(str).str.replace(".0", "").apply(
                                                        lambda sku: param_values[param_name].get(str(sku), "")
                                                    )
                                            
                                            # Форматируем дату создания, если она есть
                                            if "Дата создания" in display_combo_df.columns:
                                                display_combo_df["Дата создания"] = pd.to_datetime(display_combo_df["Дата создания"], errors="coerce")
                                                # Форматируем дату для отображения
                                                display_combo_df["Дата создания"] = display_combo_df["Дата создания"].dt.strftime("%d.%m.%Y")
                                            
                                            # Настраиваем порядок колонок (как в основной таблице)
                                            main_cols = ["Артикул", "Изображение", "Ссылка"]
                                            if "Дата создания" in display_combo_df.columns:
                                                main_cols.append("Дата создания")
                                            if "Выручка" in display_combo_df.columns:
                                                main_cols.append("Выручка")
                                            if "Заказы" in display_combo_df.columns:
                                                main_cols.append("Заказы")
                                            if "Средняя цена" in display_combo_df.columns:
                                                main_cols.append("Средняя цена")
                                            if "Цена (с СПП)" in display_combo_df.columns:
                                                main_cols.append("Цена (с СПП)")
                                            if "Позиция в выдаче" in display_combo_df.columns:
                                                main_cols.append("Позиция в выдаче")
                                            if "Упущенная выручка" in display_combo_df.columns:
                                                main_cols.append("Упущенная выручка")
                                            if "Прибыль" in display_combo_df.columns:
                                                main_cols.append("Прибыль")
                                            if "Стоимость за 1000 показов" in display_combo_df.columns:
                                                main_cols.append("Стоимость за 1000 показов")
                                            
                                            # Добавляем параметры
                                            param_cols = [p for p in available_params if p in display_combo_df.columns]
                                            
                                            # Остальные колонки
                                            other_cols = [c for c in display_combo_df.columns if c not in main_cols and c not in param_cols]
                                            
                                            # Формируем финальный порядок
                                            final_order = [c for c in main_cols + param_cols + other_cols if c in display_combo_df.columns]
                                            display_combo_df = display_combo_df[final_order]
                                            
                                            # Настройка конфигурации столбцов
                                            from streamlit import column_config as cc
                                            col_cfg = {}
                                            
                                            # Конфигурация для изображений
                                            if "Изображение" in display_combo_df.columns:
                                                col_cfg["Изображение"] = cc.ImageColumn("Изображение", width=img_size + 20)
                                            
                                            # Конфигурация для артикула
                                            if "Артикул" in display_combo_df.columns:
                                                col_cfg["Артикул"] = cc.NumberColumn("Артикул", format="%.0f", width=120)
                                            
                                            # Конфигурация для ссылки
                                            if "Ссылка" in display_combo_df.columns:
                                                col_cfg["Ссылка"] = cc.LinkColumn("Ссылка", display_text="🔗", width=60)
                                            
                                            # Конфигурация для числовых столбцов
                                            money_columns = ["Выручка", "Средняя цена", "Цена (с СПП)", "Упущенная выручка", "Прибыль", "Стоимость за 1000 показов"]
                                            for col in money_columns:
                                                if col in display_combo_df.columns:
                                                    col_cfg[col] = cc.NumberColumn(col, format="%.0f", width=120)
                                            
                                            # Конфигурация для позиции
                                            if "Позиция в выдаче" in display_combo_df.columns:
                                                col_cfg["Позиция в выдаче"] = cc.NumberColumn("Позиция в выдаче", format="%.1f", width=100)
                                            
                                            # Конфигурация для параметров
                                            for param_name in param_cols:
                                                if param_name in display_combo_df.columns:
                                                    col_cfg[param_name] = cc.TextColumn(param_name, width=150)
                                            
                                            # Отображаем таблицу
                                            st.dataframe(
                                                display_combo_df,
                                                use_container_width=True,
                                                hide_index=True,
                                                column_config=col_cfg
                                            )
                                            
                                            # Кнопка для анализа товаров в комбинации
                                            st.divider()
                                            st.markdown("### 🤖 ИИ-анализ комбинации")
                                            
                                            if st.button("🔍 Проанализировать товары с помощью ИИ", type="primary", use_container_width=True):
                                                with st.spinner("🔄 Анализирую товары с помощью ИИ... Это может занять некоторое время."):
                                                    # Определяем категорию (можно сделать настраиваемой)
                                                    category = "Рашрашд мужской (компрессионная одежда)"
                                                    
                                                    # Анализируем товары в комбинации
                                                    analysis_result = analyze_combination_products_with_ai(
                                                        combo_products_df, 
                                                        selected_combo_key,
                                                        category=category
                                                    )
                                                    
                                                    if "error" in analysis_result:
                                                        st.error(f"❌ {analysis_result['error']}")
                                                        if "raw_response" in analysis_result:
                                                            with st.expander("📄 Полный ответ ИИ"):
                                                                st.text(analysis_result["raw_response"])
                                                    else:
                                                        # Отображаем результаты анализа
                                                        st.success("✅ Анализ завершен!")
                                                        
                                                        # Общие характеристики
                                                        if "common_characteristics" in analysis_result:
                                                            st.markdown("#### 📊 Общие характеристики")
                                                            st.info(analysis_result["common_characteristics"])
                                                        
                                                        # Сильные стороны
                                                        if "strengths" in analysis_result and analysis_result["strengths"]:
                                                            st.markdown("#### ✅ Сильные стороны")
                                                            for strength in analysis_result["strengths"]:
                                                                st.markdown(f"- {strength}")
                                                        
                                                        # Слабые стороны
                                                        if "weaknesses" in analysis_result and analysis_result["weaknesses"]:
                                                            st.markdown("#### ⚠️ Слабые стороны")
                                                            for weakness in analysis_result["weaknesses"]:
                                                                st.markdown(f"- {weakness}")
                                                        
                                                        # Рекомендации
                                                        if "recommendations" in analysis_result and analysis_result["recommendations"]:
                                                            st.markdown("#### 💡 Рекомендации по улучшению")
                                                            for rec in analysis_result["recommendations"]:
                                                                priority_emoji = "🔴" if rec.get("priority") == "high" else "🟡" if rec.get("priority") == "medium" else "🟢"
                                                                st.markdown(f"**{priority_emoji} {rec.get('category', 'Общее')}**")
                                                                st.markdown(f"- {rec.get('recommendation', '')}")
                                                                if rec.get('expected_impact'):
                                                                    st.caption(f"💭 Ожидаемый эффект: {rec['expected_impact']}")
                                                        
                                                        # Конкурентные преимущества
                                                        if "competitive_advantages" in analysis_result and analysis_result["competitive_advantages"]:
                                                            st.markdown("#### 🏆 Потенциальные конкурентные преимущества")
                                                            for advantage in analysis_result["competitive_advantages"]:
                                                                st.markdown(f"- {advantage}")
                                                        
                                                        # Рыночные инсайты
                                                        if "market_insights" in analysis_result:
                                                            st.markdown("#### 📈 Рыночные инсайты")
                                                            st.info(analysis_result["market_insights"])
                                                        
                                                        # Анализ категории и CAGR
                                                        st.divider()
                                                        st.markdown("### 📊 Анализ категории и CAGR")
                                                        
                                                        cagr_data = get_category_cagr_analysis(category)
                                                        
                                                        col1, col2, col3 = st.columns(3)
                                                        with col1:
                                                            st.metric("CAGR (5 лет)", f"{cagr_data.get('cagr_5_years', 0):.1f}%")
                                                        with col2:
                                                            st.metric("Прогноз роста 2025", f"{cagr_data.get('projected_growth_2025', 0):.1f}%")
                                                        with col3:
                                                            st.metric("Размер рынка 2024", cagr_data.get('market_size_2024', 'N/A'))
                                                        
                                                        st.markdown("#### 🚀 Ключевые драйверы роста")
                                                        for driver in cagr_data.get('key_drivers', []):
                                                            st.markdown(f"- {driver}")
                                                        
                                                        st.markdown("#### 📈 Тренды рынка")
                                                        for trend in cagr_data.get('trends', []):
                                                            st.markdown(f"- {trend}")
                                                        
                                                        st.markdown(f"**Уровень конкуренции:** {cagr_data.get('competition_level', 'N/A')}")
                                                        st.markdown(f"**Зрелость рынка:** {cagr_data.get('market_maturity', 'N/A')}")
                                                        
                                                        # Анализ WGSN
                                                        st.divider()
                                                        st.markdown("### 📚 Анализ трендов WGSN")
                                                        
                                                        with st.spinner("🔄 Анализирую файлы WGSN... Это может занять некоторое время."):
                                                            # Читаем файлы WGSN
                                                            wgsn_content = read_wgsn_files()
                                                            
                                                            if "error" in wgsn_content:
                                                                st.warning(f"⚠️ {wgsn_content['error']}")
                                                            else:
                                                                # Анализируем тренды WGSN с помощью ИИ
                                                                wgsn_analysis = analyze_wgsn_trends_with_ai(
                                                                    wgsn_content,
                                                                    category=category,
                                                                    combination_key=selected_combo_key
                                                                )
                                                                
                                                                if "error" in wgsn_analysis:
                                                                    st.error(f"❌ {wgsn_analysis['error']}")
                                                                    if "raw_response" in wgsn_analysis:
                                                                        with st.expander("📄 Полный ответ ИИ по WGSN"):
                                                                            st.text(wgsn_analysis["raw_response"])
                                                                else:
                                                                    st.success("✅ Анализ WGSN завершен!")
                                                                    
                                                                    # Резюме трендов WGSN
                                                                    if "wgsn_trends_summary" in wgsn_analysis:
                                                                        st.markdown("#### 📋 Резюме трендов WGSN")
                                                                        st.markdown(wgsn_analysis["wgsn_trends_summary"])
                                                                    
                                                                    # Анализ категории
                                                                    if "category_analysis" in wgsn_analysis:
                                                                        st.markdown("#### 📊 Анализ категории на основе WGSN")
                                                                        cat_analysis = wgsn_analysis["category_analysis"]
                                                                        
                                                                        col1, col2 = st.columns(2)
                                                                        with col1:
                                                                            if "current_trends" in cat_analysis:
                                                                                st.markdown("**✅ Актуальные тренды:**")
                                                                                st.info(cat_analysis["current_trends"])
                                                                        
                                                                        with col2:
                                                                            if "emerging_trends" in cat_analysis:
                                                                                st.markdown("**🔮 Появляющиеся тренды:**")
                                                                                st.success(cat_analysis["emerging_trends"])
                                                                        
                                                                        if "declining_trends" in cat_analysis:
                                                                            st.markdown("**⚠️ Устаревающие тренды:**")
                                                                            st.warning(cat_analysis["declining_trends"])
                                                                    
                                                                    # Детальный анализ параметров комбинации
                                                                    if "combination_parameter_analysis" in wgsn_analysis and wgsn_analysis["combination_parameter_analysis"]:
                                                                        st.markdown("#### 🔍 Детальный анализ параметров комбинации")
                                                                        for param_analysis in wgsn_analysis["combination_parameter_analysis"]:
                                                                            status_emoji = "✅" if param_analysis.get("wgsn_trend_status") == "Соответствует" else "⚠️" if param_analysis.get("wgsn_trend_status") == "Частично соответствует" else "❌"
                                                                            
                                                                            with st.expander(f"{status_emoji} **{param_analysis.get('parameter', 'Параметр')}** = {param_analysis.get('current_value', 'N/A')}"):
                                                                                st.markdown(f"**Статус по трендам WGSN:** {param_analysis.get('wgsn_trend_status', 'N/A')}")
                                                                                
                                                                                if param_analysis.get('wgsn_recommendation'):
                                                                                    st.markdown("**💡 Рекомендация WGSN:**")
                                                                                    st.info(param_analysis['wgsn_recommendation'])
                                                                                
                                                                                if param_analysis.get('suggested_values'):
                                                                                    st.markdown("**📌 Рекомендуемые значения:**")
                                                                                    for value in param_analysis['suggested_values']:
                                                                                        st.markdown(f"- {value}")
                                                                                
                                                                                if param_analysis.get('rationale'):
                                                                                    st.markdown("**📝 Обоснование:**")
                                                                                    st.markdown(param_analysis['rationale'])
                                                                    
                                                                    # Релевантные тренды
                                                                    if "relevant_trends" in wgsn_analysis and wgsn_analysis["relevant_trends"]:
                                                                        st.markdown("#### 🎯 Релевантные тренды из WGSN")
                                                                        for trend in wgsn_analysis["relevant_trends"]:
                                                                            trend_type_emoji = {
                                                                "Цвет": "🎨",
                                                                "Материал": "🧵",
                                                                "Дизайн": "✂️",
                                                                "Технология": "⚙️",
                                                                "Стиль": "👔",
                                                                "Функциональность": "⚡"
                                                            }.get(trend.get('trend_type', ''), "📌")
                                                                            
                                                                            with st.expander(f"{trend_type_emoji} **{trend.get('trend_name', 'Тренд')}** ({trend.get('trend_type', 'Общее')}) - из {trend.get('source_file', 'WGSN')}"):
                                                                                st.markdown("**📄 Описание:**")
                                                                                st.markdown(trend.get('description', ''))
                                                                                
                                                                                st.markdown("**🎯 Релевантность для категории:**")
                                                                                st.info(trend.get('relevance_to_category', ''))
                                                                                
                                                                                st.markdown("**🔗 Применение к комбинации:**")
                                                                                st.markdown(trend.get('application_to_combination', ''))
                                                                                
                                                                                if trend.get('implementation_steps'):
                                                                                    st.markdown("**📋 Шаги внедрения:**")
                                                                                    for i, step in enumerate(trend.get('implementation_steps', []), 1):
                                                                                        st.markdown(f"{i}. {step}")
                                                                                
                                                                                if trend.get('expected_market_impact'):
                                                                                    st.markdown("**💼 Ожидаемое влияние на рынок:**")
                                                                                    st.success(trend.get('expected_market_impact', ''))
                                                                    
                                                                    # Рекомендации на основе трендов WGSN
                                                                    if "trend_recommendations" in wgsn_analysis and wgsn_analysis["trend_recommendations"]:
                                                                        st.markdown("#### 💡 Детальные рекомендации на основе трендов WGSN")
                                                                        for rec in wgsn_analysis["trend_recommendations"]:
                                                                            priority_emoji = "🔴" if rec.get("priority") == "high" else "🟡" if rec.get("priority") == "medium" else "🟢"
                                                                            
                                                                            with st.expander(f"{priority_emoji} **{rec.get('trend_category', 'Общее')}** - {rec.get('timeline', 'Срок не указан')}"):
                                                                                st.markdown("**💡 Рекомендация:**")
                                                                                st.markdown(rec.get('recommendation', ''))
                                                                                
                                                                                if rec.get('specific_changes'):
                                                                                    st.markdown("**🔄 Конкретные изменения:**")
                                                                                    st.info(rec.get('specific_changes', ''))
                                                                                
                                                                                if rec.get('expected_impact'):
                                                                                    st.markdown("**💭 Ожидаемый эффект:**")
                                                                                    st.success(rec.get('expected_impact', ''))
                                                                                
                                                                                if rec.get('wgsn_source'):
                                                                                    st.markdown("**📚 Источник WGSN:**")
                                                                                    st.caption(rec.get('wgsn_source', ''))
                                                                    
                                                                    # Предложения новых комбинаций
                                                                    if "new_combination_suggestions" in wgsn_analysis and wgsn_analysis["new_combination_suggestions"]:
                                                                        st.markdown("#### 🆕 Предложения новых комбинаций на основе WGSN")
                                                                        for new_combo in wgsn_analysis["new_combination_suggestions"]:
                                                                            with st.expander(f"**{new_combo.get('combination', 'Новая комбинация')}**"):
                                                                                st.markdown("**📝 Обоснование:**")
                                                                                st.info(new_combo.get('rationale', ''))
                                                                                
                                                                                if new_combo.get('wgsn_trends_used'):
                                                                                    st.markdown("**🎯 Использованные тренды WGSN:**")
                                                                                    for trend in new_combo.get('wgsn_trends_used', []):
                                                                                        st.markdown(f"- {trend}")
                                                                                
                                                                                if new_combo.get('competitive_advantage'):
                                                                                    st.markdown("**🏆 Конкурентное преимущество:**")
                                                                                    st.success(new_combo.get('competitive_advantage', ''))
                                                                    
                                                                    # Будущие тренды
                                                                    if "future_trends" in wgsn_analysis and wgsn_analysis["future_trends"]:
                                                                        st.markdown("#### 🔮 Будущие тренды")
                                                                        for trend in wgsn_analysis["future_trends"]:
                                                                            if isinstance(trend, dict):
                                                                                with st.expander(f"**{trend.get('trend_name', 'Тренд')}** - {trend.get('timeframe', 'Срок не указан')}"):
                                                                                    st.markdown("**📄 Описание:**")
                                                                                    st.markdown(trend.get('description', ''))
                                                                                    
                                                                                    if trend.get('preparation_steps'):
                                                                                        st.markdown("**📋 Шаги подготовки:**")
                                                                                        for i, step in enumerate(trend.get('preparation_steps', []), 1):
                                                                                            st.markdown(f"{i}. {step}")
                                                                            else:
                                                                                st.markdown(f"- {trend}")
                                                                    
                                                                    # Конкурентные преимущества из трендов
                                                                    if "competitive_advantages_from_trends" in wgsn_analysis and wgsn_analysis["competitive_advantages_from_trends"]:
                                                                        st.markdown("#### 🏆 Конкурентные преимущества на основе трендов WGSN")
                                                                        for advantage in wgsn_analysis["competitive_advantages_from_trends"]:
                                                                            if isinstance(advantage, dict):
                                                                                with st.expander(f"**{advantage.get('advantage', 'Преимущество')}**"):
                                                                                    st.markdown("**📚 Основано на тренде WGSN:**")
                                                                                    st.info(advantage.get('wgsn_trend_basis', ''))
                                                                                    
                                                                                    if advantage.get('implementation'):
                                                                                        st.markdown("**⚙️ Реализация:**")
                                                                                        st.markdown(advantage.get('implementation', ''))
                                                                                    
                                                                                    if advantage.get('market_positioning'):
                                                                                        st.markdown("**📍 Позиционирование на рынке:**")
                                                                                        st.success(advantage.get('market_positioning', ''))
                                                                            else:
                                                                                st.markdown(f"- {advantage}")
                                                                    
                                                                    # План действий
                                                                    if "action_plan" in wgsn_analysis:
                                                                        st.markdown("#### 📋 План действий на основе WGSN")
                                                                        action_plan = wgsn_analysis["action_plan"]
                                                                        
                                                                        if action_plan.get("immediate_actions"):
                                                                            st.markdown("**⚡ Немедленные действия:**")
                                                                            for action in action_plan.get("immediate_actions", []):
                                                                                st.markdown(f"- {action}")
                                                                        
                                                                        if action_plan.get("short_term_actions"):
                                                                            st.markdown("**📅 Краткосрочные действия (3 месяца):**")
                                                                            for action in action_plan.get("short_term_actions", []):
                                                                                st.markdown(f"- {action}")
                                                                        
                                                                        if action_plan.get("long_term_actions"):
                                                                            st.markdown("**🗓️ Долгосрочные действия (6-12 месяцев):**")
                                                                            for action in action_plan.get("long_term_actions", []):
                                                                                st.markdown(f"- {action}")
                                                                    
                                                                    # Рыночные инсайты из WGSN
                                                                    if "market_insights" in wgsn_analysis:
                                                                        st.markdown("#### 📈 Рыночные инсайты на основе WGSN")
                                                                        st.markdown(wgsn_analysis["market_insights"])
                                                                    
                                                                    # Источники файлов
                                                                    if "source_files" in wgsn_analysis:
                                                                        st.caption(f"📁 Проанализированные файлы: {', '.join(wgsn_analysis['source_files'])}")
                                                                    
                                                                    # Показываем полный ответ для отладки (опционально)
                                                                    if st.session_state.get('debug_ai_analysis', False):
                                                                        with st.expander("🔍 Отладочная информация (полный ответ ИИ по WGSN)", expanded=False):
                                                                            st.text(wgsn_analysis.get("raw_response", ""))
                                                        
                                                        # Показываем полный ответ для отладки (опционально)
                                                        if st.session_state.get('debug_ai_analysis', False):
                                                            with st.expander("🔍 Отладочная информация (полный ответ ИИ)", expanded=False):
                                                                st.text(analysis_result.get("raw_response", ""))
                                        else:
                                            st.warning("⚠️ Не найдено товаров для выбранной комбинации")
                                    else:
                                        st.warning("⚠️ Комбинация не найдена в данных. Возможно, данные были изменены.")
                            else:
                                st.info("ℹ️ Выберите комбинацию из списка выше")
                        else:
                            st.warning("⚠️ Все комбинации исключены. Сбросьте исключения, чтобы увидеть топ-10.")
                    else:
                        st.info("ℹ️ Не найдено комбинаций с полными данными по всем параметрам")
                else:
                    st.info("ℹ️ Не найдено товаров с заполненными всеми параметрами. Заполните параметры для всех товаров, чтобы увидеть лучшие комбинации.")
        
        # Четвертая вкладка - Анализ сезонности
        if seasonality_available:
            with tab4:
                st.subheader("📅 Анализ сезонности")
                
                # Выбор источника данных
                data_source = st.radio(
                    "Выберите источник данных:",
                    ["📁 Файл sezon.csv", "📤 Загрузить свой файл", "✏️ Ручной ввод"],
                    help="Выберите способ загрузки данных для анализа сезонности"
                )
                
                seasonality_df = None
                
                if data_source == "📁 Файл sezon.csv":
                    # Проверяем наличие файла sezon.csv
                    if not os.path.exists('sezon.csv'):
                        st.error("❌ Файл sezon.csv не найден в текущей директории")
                        st.info("Пожалуйста, убедитесь, что файл sezon.csv находится в той же папке, что и приложение")
                    else:
                        # Загружаем данные сезонности
                        try:
                            seasonality_df = load_seasonality_data()
                            seasonality_df = clean_seasonality_data(seasonality_df)
                            st.success("✅ Данные из sezon.csv успешно загружены")
                        except Exception as e:
                            st.error(f"❌ Ошибка при загрузке данных: {e}")
                
                elif data_source == "📤 Загрузить свой файл":
                    st.info("📋 Поддерживаемые форматы: CSV, Excel")
                    st.info("📋 Обязательные столбцы: запрос, категория, январь, февраль, март, апрель, май, июнь, июль, август, сентябрь, октябрь, ноябрь, декабрь")
                    
                    uploaded_file = st.file_uploader(
                        "Выберите файл с данными сезонности:",
                        type=['csv', 'xlsx'],
                        help="Загрузите CSV или Excel файл с данными сезонности"
                    )
                    
                    if uploaded_file is not None:
                        custom_df, message = load_custom_data(uploaded_file)
                        if custom_df is not None:
                            seasonality_df = clean_seasonality_data(custom_df)
                            st.success(f"✅ {message}")
                        else:
                            st.error(f"❌ {message}")
                
                elif data_source == "✏️ Ручной ввод":
                    st.info("📝 Введите данные для одного запроса")
                    
                    col1, col2 = st.columns(2)
                    with col1:
                        manual_query = st.text_input("Поисковый запрос:", placeholder="например: джинсы женские")
                        manual_category = st.text_input("Категория:", placeholder="например: джинсы")
                    
                    with col2:
                        st.write("**Частотность по месяцам:**")
                    
                    # Создаем поля для ввода частотности по месяцам
                    month_names = ['Январь', 'Февраль', 'Март', 'Апрель', 'Май', 'Июнь',
                                   'Июль', 'Август', 'Сентябрь', 'Октябрь', 'Ноябрь', 'Декабрь']
                    
                    col_months1, col_months2, col_months3, col_months4 = st.columns(4)
                    manual_frequencies = []
                    
                    for i, month in enumerate(month_names):
                        if i < 3:
                            with col_months1:
                                freq = st.number_input(f"{month}:", min_value=0, value=0, step=1, key=f"manual_{i}")
                                manual_frequencies.append(freq)
                        elif i < 6:
                            with col_months2:
                                freq = st.number_input(f"{month}:", min_value=0, value=0, step=1, key=f"manual_{i}")
                                manual_frequencies.append(freq)
                        elif i < 9:
                            with col_months3:
                                freq = st.number_input(f"{month}:", min_value=0, value=0, step=1, key=f"manual_{i}")
                                manual_frequencies.append(freq)
                        else:
                            with col_months4:
                                freq = st.number_input(f"{month}:", min_value=0, value=0, step=1, key=f"manual_{i}")
                                manual_frequencies.append(freq)
                    
                    if st.button("📊 Анализировать данные", type="primary"):
                        if manual_query and manual_category:
                            seasonality_df = create_manual_entry_data(manual_query, manual_category, manual_frequencies)
                            seasonality_df = clean_seasonality_data(seasonality_df)
                            st.success("✅ Данные успешно созданы и готовы для анализа")
                        else:
                            st.error("❌ Пожалуйста, заполните запрос и категорию")
                
                # Продолжаем только если данные загружены
                if seasonality_df is not None and not seasonality_df.empty:
                        
                        # Создаем вкладки для анализа сезонности
                        seasonality_tab1, seasonality_tab2 = st.tabs(["🔍 Анализ запроса", "📅 Анализ по месяцам"])
                        
                        with seasonality_tab1:
                            st.markdown("---")
                            st.subheader("🔍 Выбор товара для анализа")
                            
                            col1, col2 = st.columns([2, 3])
                            with col1:
                                if 'категория' in seasonality_df.columns:
                                    categories = sorted(seasonality_df['категория'].dropna().unique())
                                    selected_category = st.selectbox(
                                        "Выберите категорию:",
                                        categories,
                                        help="Сначала выберите категорию товаров"
                                    )
                                else:
                                    st.error("Столбец 'категория' не найден")
                                    selected_category = None
                            
                            with col2:
                                if selected_category:
                                    category_df = seasonality_df[seasonality_df['категория'] == selected_category]
                                    if 'запрос' in category_df.columns:
                                        queries_in_category = sorted(category_df['запрос'].dropna().unique())
                                        if queries_in_category:
                                            selected_item = st.selectbox(
                                                "Выберите запрос:",
                                                queries_in_category,
                                                help="Выберите поисковый запрос для анализа сезонности"
                                            )
                                            if selected_item:
                                                filtered_df = category_df[category_df['запрос'] == selected_item]
                                            else:
                                                filtered_df = pd.DataFrame()
                                        else:
                                            st.warning("В этой категории нет поисковых запросов")
                                            filtered_df = pd.DataFrame()
                                    else:
                                        st.error("Столбец 'запрос' не найден")
                                        filtered_df = pd.DataFrame()
                                else:
                                    st.info("Выберите категорию для продолжения")
                                    filtered_df = pd.DataFrame()
                            
                            st.markdown("---")
                            
                            if not filtered_df.empty:
                                # Основная информация
                                row = filtered_df.iloc[0]
                                col_info1, col_info2, col_info3 = st.columns(3)
                                
                                with col_info1:
                                    st.metric("Категория", row.get('категория', 'Н/Д'))
                                with col_info2:
                                    st.metric("Товар", row.get('наименование товара', 'Н/Д'))
                                with col_info3:
                                    st.metric("Запрос", row.get('запрос', 'Н/Д'))
                                
                                # График сезонности
                                fig = create_seasonality_graph(filtered_df, selected_item)
                                if fig:
                                    st.plotly_chart(fig, use_container_width=True)
                                
                                # Детальная таблица по месяцам
                                st.subheader("📊 Детальные данные по месяцам")
                                month_columns = ['январь', 'февраль', 'март', 'апрель', 'май', 'июнь',
                                                'июль', 'август', 'сентябрь', 'октябрь', 'ноябрь', 'декабрь']
                                
                                month_data = []
                                for month in month_columns:
                                    if month in row.index:
                                        month_data.append({
                                            'Месяц': month.capitalize(),
                                            'Частота': f"{row[month]:,.0f}",
                                            'Процент от максимума': f"{(row[month] / max([row[m] for m in month_columns if m in row.index])) * 100:.1f}%"
                                        })
                                
                                if month_data:
                                    month_df = pd.DataFrame(month_data)
                                    st.dataframe(month_df, use_container_width=True)
                        
                        with seasonality_tab2:
                            st.subheader("📅 Анализ по месяцам")
                            
                            # Выбор месяца
                            month_columns = ['январь', 'февраль', 'март', 'апрель', 'май', 'июнь',
                                            'июль', 'август', 'сентябрь', 'октябрь', 'ноябрь', 'декабрь']
                            
                            selected_month = st.selectbox(
                                "Выберите месяц для анализа:",
                                [month.capitalize() for month in month_columns],
                                help="Выберите месяц для анализа всех запросов"
                            )
                            
                            # Получаем данные для выбранного месяца
                            month_lower = selected_month.lower()
                            if month_lower in seasonality_df.columns:
                                columns_to_select = ['запрос', 'категория', 'наименование товара'] + month_columns
                                month_data = seasonality_df[columns_to_select].copy()
                                month_data = month_data[month_data[month_lower] > 0]  # Только с данными
                                
                                if not month_data.empty:
                                    st.write(f"**Найдено {len(month_data)} запросов с данными в {selected_month}**")
                                    
                                    # Получаем статистику по статусам
                                    stats, month_data_with_status = get_status_stats(month_data, month_lower)
                                    
                                    # Показываем KPI метрики
                                    col_kpi1, col_kpi2, col_kpi3, col_kpi4, col_kpi5 = st.columns(5)
                                    col_kpi1.metric("Всего", stats.get('Всего', 0))
                                    col_kpi2.metric("Пик max", stats.get('Пик max', 0))
                                    col_kpi3.metric("Пик min", stats.get('Пик min', 0))
                                    col_kpi4.metric("Рост", stats.get('Рост', 0))
                                    col_kpi5.metric("Падение", stats.get('Падение', 0) + stats.get('Большое падение', 0))
                                    
                                    # Фильтрация и сортировка
                                    col1, col2 = st.columns(2)
                                    
                                    with col1:
                                        status_options = ['Все', 'Пик max', 'Пик min', 'Рост', 'Падение', 'Большое падение']
                                        selected_status = st.selectbox(
                                            "Фильтр по статусу:",
                                            status_options,
                                            help="Выберите статус для фильтрации запросов"
                                        )
                                    
                                    with col2:
                                        sort_options = ['По цвету (зеленый → красный)', 'По частотности (высокая → низкая)', 'По алфавиту']
                                        selected_sort = st.selectbox(
                                            "Сортировка:",
                                            sort_options,
                                            help="Выберите способ сортировки запросов"
                                        )
                                    
                                    # Применяем фильтр по статусу
                                    if selected_status != 'Все':
                                        month_data_with_status = month_data_with_status[month_data_with_status['Статус'] == selected_status]
                                    
                                    # Применяем сортировку
                                    if selected_sort == 'По цвету (зеленый → красный)':
                                        def get_color_priority(row):
                                            current_month_value = row[month_lower]
                                            month_values = [row[month] for month in month_columns]
                                            max_val = max(month_values) if month_values else 1
                                            if max_val == 0:
                                                return 5
                                            
                                            intensity = current_month_value / max_val
                                            
                                            if intensity >= 0.9:
                                                return 1
                                            elif intensity >= 0.5:
                                                return 2
                                            elif intensity >= 0.3:
                                                return 3
                                            else:
                                                return 4
                                        
                                        month_data_with_status['sort_key'] = month_data_with_status.apply(get_color_priority, axis=1)
                                        month_data_with_status = month_data_with_status.sort_values('sort_key')
                                        month_data_with_status = month_data_with_status.drop('sort_key', axis=1)
                                        
                                    elif selected_sort == 'По частотности (высокая → низкая)':
                                        month_data_with_status = month_data_with_status.sort_values(month_lower, ascending=False)
                                        
                                    else:  # По алфавиту
                                        month_data_with_status = month_data_with_status.sort_values('запрос')
                                    
                                    # Показываем таблицу
                                    st.subheader("📋 Список запросов")
                                    
                                    # Стилизуем таблицу
                                    styled_df = style_dataframe(month_data_with_status, month_lower)
                                    st.dataframe(styled_df, use_container_width=True)
                                    
                                    # Легенда цветов
                                    st.markdown("---")
                                    st.caption("**Легенда цветов:**")
                                    col_legend1, col_legend2, col_legend3, col_legend4 = st.columns(4)
                                    col_legend1.markdown("🟢 **Зеленый** - 90%+ от максимума")
                                    col_legend2.markdown("🟡 **Желтый** - 50-90% от максимума")
                                    col_legend3.markdown("🟠 **Бледно-желтый** - 30-50% от максимума")
                                    col_legend4.markdown("🔴 **Красный** - <30% от максимума")
                                    
                                else:
                                    st.warning(f"Нет данных для месяца {selected_month}")
                            else:
                                st.error(f"Столбец '{selected_month}' не найден в данных")
                else:
                    st.info("📊 Выберите источник данных и загрузите информацию для анализа сезонности")
        else:
            st.warning("Модуль анализа сезонности недоступен")
        
        # Пятая вкладка - Прогнозирование с Prophet
        if PROPHET_AVAILABLE:
            if seasonality_available:
                with tab5:
                    st.subheader("🔮 Прогнозирование с Prophet")
                    
                    # Проверяем наличие данных
                    if df is not None and not df.empty:
                        # Настройки прогнозирования
                        col_settings1, col_settings2 = st.columns(2)
                        
                        with col_settings1:
                            # Выбор метрики для прогнозирования
                            numeric_columns = df.select_dtypes(include=[np.number]).columns.tolist()
                            if numeric_columns:
                                metric_choice = st.selectbox(
                                    "Выберите метрику для прогнозирования:",
                                    numeric_columns,
                                    key="prophet_metric_choice"
                                )
                            else:
                                st.warning("Нет числовых колонок для прогнозирования")
                                metric_choice = None
                        
                        with col_settings2:
                            # Выбор колонки с датами
                            date_columns = []
                            for col in df.columns:
                                if df[col].dtype == 'datetime64[ns]' or 'дата' in col.lower() or 'date' in col.lower():
                                    date_columns.append(col)
                            
                            if date_columns:
                                date_choice = st.selectbox(
                                    "Выберите колонку с датами (опционально):",
                                    ["Автоматически"] + date_columns,
                                    key="prophet_date_choice"
                                )
                                if date_choice == "Автоматически":
                                    date_choice = None
                            else:
                                date_choice = None
                                st.info("Колонка с датами не найдена, будет создана автоматически")
                        
                        # Дополнительные настройки
                        col_periods, col_seasonality = st.columns(2)
                        
                        with col_periods:
                            forecast_periods = st.number_input(
                                "Период прогнозирования (дни):",
                                min_value=1,
                                max_value=365,
                                value=30,
                                key="prophet_periods"
                            )
                        
                        with col_seasonality:
                            seasonality_mode = st.selectbox(
                                "Режим сезонности:",
                                ["additive", "multiplicative"],
                                key="prophet_seasonality"
                            )
                        
                        # Кнопка создания прогноза
                        if st.button("🔮 Создать прогноз", type="primary", key="create_forecast_btn"):
                            if metric_choice:
                                with st.spinner("Создание прогноза..."):
                                    # Подготавливаем данные
                                    df_prophet = prepare_data_for_prophet(df, metric_choice, date_choice)
                                    
                                    if df_prophet is not None and len(df_prophet) > 1:
                                        # Создаем прогноз
                                        model, forecast, future = create_prophet_forecast(
                                            df_prophet, 
                                            periods=forecast_periods,
                                            seasonality_mode=seasonality_mode
                                        )
                                        
                                        if model and forecast is not None:
                                            # Отображаем основной график прогноза
                                            st.subheader("📈 Прогноз")
                                            fig_forecast = plot_prophet_forecast(
                                                model, 
                                                forecast, 
                                                f"Прогноз {metric_choice}"
                                            )
                                            if fig_forecast:
                                                st.plotly_chart(fig_forecast, use_container_width=True)
                                            
                                            # Отображаем компоненты прогноза
                                            st.subheader("🔍 Компоненты прогноза")
                                            fig_components = plot_prophet_components(
                                                model, 
                                                forecast, 
                                                f"Компоненты прогноза {metric_choice}"
                                            )
                                            if fig_components:
                                                st.plotly_chart(fig_components, use_container_width=True)
                                            
                                            # Статистика прогноза
                                            st.subheader("📊 Статистика прогноза")
                                            
                                            # Получаем последние прогнозные значения
                                            forecast_future = forecast[forecast['ds'] > df_prophet['ds'].max()]
                                            
                                            if not forecast_future.empty:
                                                col_stat1, col_stat2, col_stat3, col_stat4 = st.columns(4)
                                                
                                                with col_stat1:
                                                    mean_forecast = forecast_future['yhat'].mean()
                                                    st.metric("Средний прогноз", f"{mean_forecast:,.0f}")
                                                
                                                with col_stat2:
                                                    max_forecast = forecast_future['yhat'].max()
                                                    st.metric("Максимальный прогноз", f"{max_forecast:,.0f}")
                                                
                                                with col_stat3:
                                                    min_forecast = forecast_future['yhat'].min()
                                                    st.metric("Минимальный прогноз", f"{min_forecast:,.0f}")
                                                
                                                with col_stat4:
                                                    trend = forecast_future['trend'].iloc[-1] - forecast_future['trend'].iloc[0]
                                                    st.metric("Изменение тренда", f"{trend:,.0f}")
                                                
                                                # Таблица с прогнозными значениями
                                                st.subheader("📋 Детальный прогноз")
                                                forecast_display = forecast_future[['ds', 'yhat', 'yhat_lower', 'yhat_upper']].copy()
                                                forecast_display.columns = ['Дата', 'Прогноз', 'Нижняя граница', 'Верхняя граница']
                                                forecast_display['Дата'] = forecast_display['Дата'].dt.strftime('%Y-%m-%d')
                                                
                                                st.dataframe(
                                                    forecast_display,
                                                    use_container_width=True,
                                                    hide_index=True
                                                )
                                                
                                                # Экспорт прогноза
                                                csv_data = forecast_display.to_csv(index=False)
                                                st.download_button(
                                                    label="💾 Скачать прогноз (CSV)",
                                                    data=csv_data,
                                                    file_name=f"prophet_forecast_{metric_choice}.csv",
                                                    mime="text/csv"
                                                )
                                            
                                        else:
                                            st.error("Не удалось создать прогноз. Проверьте данные.")
                                    else:
                                        st.error("Недостаточно данных для создания прогноза. Нужно минимум 2 точки данных.")
                            else:
                                st.warning("Выберите метрику для прогнозирования")
                    
                    else:
                        st.info("📊 Загрузите данные в первой вкладке для создания прогнозов")
            
            else:
                # Если нет вкладки сезонности, используем tab4
                with tab4:
                    st.subheader("🔮 Прогнозирование с Prophet")
                    
                    # Проверяем наличие данных
                    if df is not None and not df.empty:
                        # Настройки прогнозирования
                        col_settings1, col_settings2 = st.columns(2)
                        
                        with col_settings1:
                            # Выбор метрики для прогнозирования
                            numeric_columns = df.select_dtypes(include=[np.number]).columns.tolist()
                            if numeric_columns:
                                metric_choice = st.selectbox(
                                    "Выберите метрику для прогнозирования:",
                                    numeric_columns,
                                    key="prophet_metric_choice"
                                )
                            else:
                                st.warning("Нет числовых колонок для прогнозирования")
                                metric_choice = None
                        
                        with col_settings2:
                            # Выбор колонки с датами
                            date_columns = []
                            for col in df.columns:
                                if df[col].dtype == 'datetime64[ns]' or 'дата' in col.lower() or 'date' in col.lower():
                                    date_columns.append(col)
                            
                            if date_columns:
                                date_choice = st.selectbox(
                                    "Выберите колонку с датами (опционально):",
                                    ["Автоматически"] + date_columns,
                                    key="prophet_date_choice"
                                )
                                if date_choice == "Автоматически":
                                    date_choice = None
                            else:
                                date_choice = None
                                st.info("Колонка с датами не найдена, будет создана автоматически")
                        
                        # Дополнительные настройки
                        col_periods, col_seasonality = st.columns(2)
                        
                        with col_periods:
                            forecast_periods = st.number_input(
                                "Период прогнозирования (дни):",
                                min_value=1,
                                max_value=365,
                                value=30,
                                key="prophet_periods"
                            )
                        
                        with col_seasonality:
                            seasonality_mode = st.selectbox(
                                "Режим сезонности:",
                                ["additive", "multiplicative"],
                                key="prophet_seasonality"
                            )
                        
                        # Кнопка создания прогноза
                        if st.button("🔮 Создать прогноз", type="primary", key="create_forecast_btn"):
                            if metric_choice:
                                with st.spinner("Создание прогноза..."):
                                    # Подготавливаем данные
                                    df_prophet = prepare_data_for_prophet(df, metric_choice, date_choice)
                                    
                                    if df_prophet is not None and len(df_prophet) > 1:
                                        # Создаем прогноз
                                        model, forecast, future = create_prophet_forecast(
                                            df_prophet, 
                                            periods=forecast_periods,
                                            seasonality_mode=seasonality_mode
                                        )
                                        
                                        if model and forecast is not None:
                                            # Отображаем основной график прогноза
                                            st.subheader("📈 Прогноз")
                                            fig_forecast = plot_prophet_forecast(
                                                model, 
                                                forecast, 
                                                f"Прогноз {metric_choice}"
                                            )
                                            if fig_forecast:
                                                st.plotly_chart(fig_forecast, use_container_width=True)
                                            
                                            # Отображаем компоненты прогноза
                                            st.subheader("🔍 Компоненты прогноза")
                                            fig_components = plot_prophet_components(
                                                model, 
                                                forecast, 
                                                f"Компоненты прогноза {metric_choice}"
                                            )
                                            if fig_components:
                                                st.plotly_chart(fig_components, use_container_width=True)
                                            
                                            # Статистика прогноза
                                            st.subheader("📊 Статистика прогноза")
                                            
                                            # Получаем последние прогнозные значения
                                            forecast_future = forecast[forecast['ds'] > df_prophet['ds'].max()]
                                            
                                            if not forecast_future.empty:
                                                col_stat1, col_stat2, col_stat3, col_stat4 = st.columns(4)
                                                
                                                with col_stat1:
                                                    mean_forecast = forecast_future['yhat'].mean()
                                                    st.metric("Средний прогноз", f"{mean_forecast:,.0f}")
                                                
                                                with col_stat2:
                                                    max_forecast = forecast_future['yhat'].max()
                                                    st.metric("Максимальный прогноз", f"{max_forecast:,.0f}")
                                                
                                                with col_stat3:
                                                    min_forecast = forecast_future['yhat'].min()
                                                    st.metric("Минимальный прогноз", f"{min_forecast:,.0f}")
                                                
                                                with col_stat4:
                                                    trend = forecast_future['trend'].iloc[-1] - forecast_future['trend'].iloc[0]
                                                    st.metric("Изменение тренда", f"{trend:,.0f}")
                                                
                                                # Таблица с прогнозными значениями
                                                st.subheader("📋 Детальный прогноз")
                                                forecast_display = forecast_future[['ds', 'yhat', 'yhat_lower', 'yhat_upper']].copy()
                                                forecast_display.columns = ['Дата', 'Прогноз', 'Нижняя граница', 'Верхняя граница']
                                                forecast_display['Дата'] = forecast_display['Дата'].dt.strftime('%Y-%m-%d')
                                                
                                                st.dataframe(
                                                    forecast_display,
                                                    use_container_width=True,
                                                    hide_index=True
                                                )
                                                
                                                # Экспорт прогноза
                                                csv_data = forecast_display.to_csv(index=False)
                                                st.download_button(
                                                    label="💾 Скачать прогноз (CSV)",
                                                    data=csv_data,
                                                    file_name=f"prophet_forecast_{metric_choice}.csv",
                                                    mime="text/csv"
                                                )
                                            
                                        else:
                                            st.error("Не удалось создать прогноз. Проверьте данные.")
                                    else:
                                        st.error("Недостаточно данных для создания прогноза. Нужно минимум 2 точки данных.")
                            else:
                                st.warning("Выберите метрику для прогнозирования")
                    
                    else:
                        st.info("📊 Загрузите данные в первой вкладке для создания прогнозов")
        
        # Вкладка истории изменений параметров
        # Определяем правильную вкладку в зависимости от доступных модулей
        if seasonality_available:
            if PROPHET_AVAILABLE:
                history_tab = tab6
            else:
                history_tab = tab5
        else:
            if PROPHET_AVAILABLE:
                history_tab = tab5
            else:
                history_tab = tab4
        
        with history_tab:
            st.subheader("📜 История изменений параметров")
            st.info("💡 Здесь отображается история всех изменений параметров товаров. При изменении параметра старое значение сохраняется в историю.")
            
            # Загружаем историю при открытии вкладки
            load_param_history_from_file()
            
            # Фильтры для истории
            col_filter1, col_filter2, col_filter3 = st.columns(3)
            
            with col_filter1:
                filter_sku = st.text_input(
                    "🔍 Фильтр по артикулу:",
                    placeholder="Введите артикул",
                    key="history_filter_sku"
                )
            
            with col_filter2:
                filter_param = st.selectbox(
                    "🔍 Фильтр по параметру:",
                    options=["Все"] + list(st.session_state.get("param_options", {}).keys()),
                    key="history_filter_param"
                )
            
            with col_filter3:
                limit_history = st.number_input(
                    "📊 Количество записей:",
                    min_value=10,
                    max_value=1000,
                    value=100,
                    step=10,
                    key="history_limit"
                )
            
            # Получаем историю с фильтрами
            history = get_param_history(
                sku=filter_sku if filter_sku else None,
                param=filter_param if filter_param != "Все" else None
            )
            
            # Ограничиваем количество записей
            history = history[:limit_history]
            
            if history:
                # Создаем DataFrame для отображения
                history_df = pd.DataFrame(history)
                
                # Переименовываем колонки для удобства
                history_df = history_df.rename(columns={
                    "timestamp": "Дата и время",
                    "sku": "Артикул",
                    "param": "Параметр",
                    "old_value": "Старое значение",
                    "new_value": "Новое значение"
                })
                
                # Переупорядочиваем колонки
                history_df = history_df[["Дата и время", "Артикул", "Параметр", "Старое значение", "Новое значение"]]
                
                st.dataframe(history_df, use_container_width=True, height=600)
                
                # Статистика
                st.markdown("### 📊 Статистика изменений")
                col_stat1, col_stat2, col_stat3 = st.columns(3)
                
                with col_stat1:
                    st.metric("Всего изменений", len(history))
                
                with col_stat2:
                    unique_skus = history_df["Артикул"].nunique()
                    st.metric("Уникальных товаров", unique_skus)
                
                with col_stat3:
                    unique_params = history_df["Параметр"].nunique()
                    st.metric("Уникальных параметров", unique_params)
                
                # График изменений по времени
                if len(history) > 1:
                    st.markdown("### 📈 График изменений по времени")
                    history_df["Дата и время"] = pd.to_datetime(history_df["Дата и время"])
                    history_df_sorted = history_df.sort_values("Дата и время")
                    
                    # Группируем по дням
                    history_df_sorted["Дата"] = history_df_sorted["Дата и время"].dt.date
                    daily_changes = history_df_sorted.groupby("Дата").size().reset_index(name="Количество изменений")
                    
                    fig = px.line(daily_changes, x="Дата", y="Количество изменений", 
                                 title="Количество изменений параметров по дням",
                                 markers=True)
                    fig.update_layout(xaxis_title="Дата", yaxis_title="Количество изменений")
                    st.plotly_chart(fig, use_container_width=True)
                
                # Кнопка экспорта
                if st.button("📥 Экспорт истории в CSV", key="export_history_btn"):
                    csv_data = history_df.to_csv(index=False, encoding='utf-8-sig')
                    st.download_button(
                        label="💾 Скачать CSV",
                        data=csv_data,
                        file_name=f"param_history_{pd.Timestamp.now().strftime('%Y%m%d_%H%M%S')}.csv",
                        mime="text/csv",
                        key="download_history_csv"
                    )
                
                # Кнопка очистки истории
                if st.button("🗑️ Очистить историю", key="clear_history_btn", type="secondary"):
                    if st.session_state.get("confirm_clear_history", False):
                        st.session_state["param_history"] = []
                        save_param_history_to_file()
                        st.success("✅ История очищена!")
                        st.session_state["confirm_clear_history"] = False
                        st.rerun()
                    else:
                        st.session_state["confirm_clear_history"] = True
                        st.warning("⚠️ Нажмите еще раз для подтверждения очистки истории")
            else:
                st.info("📝 История изменений пуста. Изменения параметров будут автоматически сохраняться здесь.")

# end of file
