# -*- coding: utf-8 -*-
"""
Модуль для чтения файлов из папки MC (дополнительные исследования)
"""
import os
import re

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    Document = None

try:
    from bs4 import BeautifulSoup
    BS4_AVAILABLE = True
except ImportError:
    BS4_AVAILABLE = False
    BeautifulSoup = None


def read_mc_files(mc_dir: str = None) -> dict:
    """
    Читает все файлы из папки MC и возвращает их содержимое.
    Поддерживает .docx, .txt, .md, .html, .htm файлы.
    Для HTML файлов извлекает текст, удаляя HTML-теги.
    
    Args:
        mc_dir: Путь к папке MC (если None, используется папка MC относительно корня проекта)
    
    Returns:
        Словарь с содержимым файлов: {"filename": "content", ...}
        При ошибках возвращает словарь с ключом "error"
    """
    mc_data = {}
    
    try:
        if mc_dir is None:
            # Определяем путь к папке MC относительно корня проекта
            current_file = os.path.abspath(__file__)
            project_root = os.path.dirname(os.path.dirname(current_file))
            mc_dir = os.path.join(project_root, "MC")
        
        if not os.path.exists(mc_dir):
            return {
                "error": f"Папка MC не найдена: {mc_dir}"
            }
        
        # Читаем все поддерживаемые файлы из папки MC
        for filename in os.listdir(mc_dir):
            file_path = os.path.join(mc_dir, filename)
            
            # Пропускаем директории и скрытые файлы
            if os.path.isdir(file_path) or filename.startswith('.'):
                continue
            
            try:
                # Читаем .docx файлы
                if filename.endswith('.docx'):
                    if not DOCX_AVAILABLE:
                        mc_data[filename] = "Ошибка: библиотека python-docx не установлена"
                        continue
                    
                    doc = Document(file_path)
                    # Извлекаем весь текст из документа
                    full_text = []
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            full_text.append(paragraph.text.strip())
                    
                    # Также извлекаем текст из таблиц
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = []
                            for cell in row.cells:
                                if cell.text.strip():
                                    row_text.append(cell.text.strip())
                            if row_text:
                                full_text.append(" | ".join(row_text))
                    
                    mc_data[filename] = "\n".join(full_text)
                
                # Читаем HTML файлы
                elif filename.endswith(('.html', '.htm')):
                    try:
                        with open(file_path, 'r', encoding='utf-8') as f:
                            html_content = f.read()
                        
                        if BS4_AVAILABLE:
                            # Используем BeautifulSoup для парсинга HTML
                            soup = BeautifulSoup(html_content, 'html.parser')
                            
                            # Удаляем скрипты и стили
                            for script in soup(["script", "style"]):
                                script.decompose()
                            
                            # Извлекаем текст
                            text = soup.get_text()
                            
                            # Очищаем текст от лишних пробелов
                            lines = (line.strip() for line in text.splitlines())
                            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                            text = '\n'.join(chunk for chunk in chunks if chunk)
                            
                            mc_data[filename] = text
                        else:
                            # Если BeautifulSoup не установлен, используем простой парсинг через регулярные выражения
                            # Удаляем теги HTML
                            text = re.sub(r'<[^>]+>', '', html_content)
                            # Удаляем множественные пробелы и переносы строк
                            text = re.sub(r'\s+', ' ', text)
                            text = re.sub(r'\n\s*\n', '\n', text)
                            mc_data[filename] = text.strip()
                    except Exception as html_e:
                        mc_data[filename] = f"Ошибка при чтении HTML файла: {str(html_e)}"
                
                # Читаем текстовые файлы (.txt, .md)
                elif filename.endswith(('.txt', '.md')):
                    with open(file_path, 'r', encoding='utf-8') as f:
                        mc_data[filename] = f.read()
                
            except Exception as e:
                mc_data[filename] = f"Ошибка при чтении файла: {str(e)}"
        
        if not mc_data:
            return {
                "error": "Не найдено поддерживаемых файлов (.docx, .txt, .md, .html, .htm) в папке MC"
            }
        
        return mc_data
    
    except Exception as e:
        return {
            "error": f"Ошибка при чтении файлов MC: {str(e)}"
        }











