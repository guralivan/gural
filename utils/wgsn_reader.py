# -*- coding: utf-8 -*-
"""
Модуль для чтения файлов WGSN
"""
import os

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    Document = None


def read_wgsn_files(wgsn_dir: str = None) -> dict:
    """
    Читает все файлы из папки WGSN и возвращает их содержимое.
    
    Args:
        wgsn_dir: Путь к папке WGSN (если None, используется папка WGSN относительно корня проекта)
    
    Returns:
        Словарь с содержимым файлов: {"filename": "content", ...}
    """
    wgsn_data = {}
    
    if not DOCX_AVAILABLE:
        return {
            "error": "Библиотека python-docx не установлена. Установите её командой: pip install python-docx"
        }
    
    try:
        if wgsn_dir is None:
            # Определяем путь к папке WGSN относительно корня проекта
            current_file = os.path.abspath(__file__)
            project_root = os.path.dirname(os.path.dirname(current_file))
            wgsn_dir = os.path.join(project_root, "WGSN")
        
        if not os.path.exists(wgsn_dir):
            return {
                "error": f"Папка WGSN не найдена: {wgsn_dir}"
            }
        
        # Читаем все .docx файлы из папки WGSN
        for filename in os.listdir(wgsn_dir):
            if filename.endswith('.docx'):
                file_path = os.path.join(wgsn_dir, filename)
                try:
                    doc = Document(file_path)
                    # Извлекаем весь текст из документа
                    full_text = []
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            full_text.append(paragraph.text.strip())
                    
                    # Также извлекаем текст из таблиц
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = []
                            for cell in row.cells:
                                if cell.text.strip():
                                    row_text.append(cell.text.strip())
                            if row_text:
                                full_text.append(" | ".join(row_text))
                    
                    wgsn_data[filename] = "\n".join(full_text)
                except Exception as e:
                    wgsn_data[filename] = f"Ошибка при чтении файла: {str(e)}"
        
        if not wgsn_data:
            return {
                "error": "Не найдено .docx файлов в папке WGSN"
            }
        
        return wgsn_data
        
    except Exception as e:
        return {
            "error": f"Ошибка при чтении файлов WGSN: {str(e)}"
        }



























